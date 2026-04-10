"""
上海海洋大学模板制作 - 分步处理，避免索引偏移。

Step 1: 编辑封面/摘要（body前，索引不变）
Step 2: body_maker 设置正文循环（插入控制段落）
Step 3: 重新定位致谢/参考文献，编辑（索引已偏移）
"""
import os
from editor import TemplateEditor
from body_maker import setup_body_template

SRC = "../副本上海海洋大学毕业论文（设计）模板.docx"
OUT = "output/shou_template.docx"


def main():
    os.makedirs("output", exist_ok=True)

    # ============================================================
    # Step 1: 封面 + 摘要（都在正文之前，索引不受后续影响）
    # ============================================================
    ed = TemplateEditor(SRC)

    # 封面
    ed.replace_run(6, 3, "{{ title_zh }}")
    ed.replace_run(8, 4, "{{ title_en }}")
    ed.replace_run(10, 1, "学 院：{{ college }}")
    ed.replace_run(11, 0, "   班 级：{{ class_name }}")
    ed.replace_run(12, 0, "   姓 名：{{ name }}")
    ed.replace_run(13, 0, "   学 号：{{ student_id }}")
    ed.replace_run(14, 0, "指导教师：{{ advisor }}")
    ed.replace_run(16, 0, "{{ year }}年")
    ed.replace_run(16, 2, "{{ month }}月")

    # 摘要
    ed.clear_para(40)   # (空一行)
    ed.replace_run(41, 1, "{{ abstract_zh }}")
    ed.clear_runs(41, [2, 3, 4, 5, 6, 7, 8, 9, 10, 11])
    ed.clear_para(42)   # (空一行)
    ed.replace_run(43, 1, "：{{ keywords_zh }}")
    ed.clear_runs(43, [2, 3, 4, 5, 6, 7, 8])

    # 英文摘要
    ed.clear_para(45)   # (空两行)
    ed.clear_para(47)   # (空一行)
    ed.replace_run(48, 0, "{{ abstract_en }}")
    ed.clear_runs(48, [1, 2, 3, 4])
    ed.replace_run(49, 2, "{{ keywords_en }}")
    # 删除 ":" 之后除了 keywords_en 变量外的所有 run（避免空格）
    p49 = ed.doc.paragraphs[49]
    for r in reversed(p49.runs[3:]):
        r._r.getparent().remove(r._r)
    # 同样清理中文关键词的空 run
    p43 = ed.doc.paragraphs[43]
    for r in reversed(p43.runs[2:]):
        r._r.getparent().remove(r._r)
    ed.clear_para(51)   # (空一行)
    ed.clear_para(53)   # (空一行)

    ed.save(OUT)
    print("Step 1: 封面/摘要完成")

    # ============================================================
    # Step 2: 正文循环
    # ============================================================
    setup_body_template(OUT, OUT)
    print("Step 2: 正文循环完成")

    # ============================================================
    # Step 3: 致谢/参考文献 - 全部用 XML 元素引用，不用索引
    # ============================================================
    from docx import Document as _Doc
    from refs_maker import setup_refs_template
    doc = _Doc(OUT)
    paras = doc.paragraphs
    _ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    for i, p in enumerate(paras):
        text = (p.text or "").strip()

        # 清空格式说明
        if text in ["（空一行）", "（空两行）"] and i > 60:
            for r in p.runs:
                r.text = ""
        if "宋体小四号" in text and "行间距" in text:
            for r in p.runs:
                r.text = ""
        if "正文中公式" in text:
            for r in p.runs:
                r.text = ""

        # 参考文献 → Jinja2 循环
        if text == "参考文献" and i > 60:
            setup_refs_template(doc, i)

        # 致谢内容
        if "致谢内容致谢内容" in text:
            if p.runs:
                p.runs[0].text = "    {{ acknowledgement }}"
                for r in p.runs[1:]:
                    r.text = ""

    # 清空致谢格式说明残留
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        if "小四号宋体" in text and "行间距" in text:
            for r in p.runs:
                r.text = ""

    doc.save(OUT)
    print("Step 3: 致谢/参考文献完成")

    # ============================================================
    # Step 4: 删除文本框注释（黄色格式说明标签）
    # ============================================================
    import zipfile, re, shutil
    tmp = OUT + ".tmp"
    with zipfile.ZipFile(OUT, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w') as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    xml = data.decode('utf-8')
                    n = xml.count('<mc:AlternateContent')
                    xml = re.sub(
                        r'<mc:AlternateContent>.*?</mc:AlternateContent>',
                        '', xml, flags=re.DOTALL)
                    print(f"Step 4: 删除了 {n} 个文本框注释")
                    data = xml.encode('utf-8')
                zout.writestr(item, data)
    shutil.move(tmp, OUT)
    print(f"\n模板输出: {OUT}")


if __name__ == "__main__":
    main()
