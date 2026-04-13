"""
智学云课论文组装器。
支持文本、图片（InlineImage）、表格。

用法: python build.py <模板.docx> [输出.docx]
"""
import os
import sys
import json
import glob

ROOT = os.path.join(os.path.dirname(__file__), "../..")
sys.path.insert(0, ROOT)

HERE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(HERE, "images")


def _ensure_rgb(img_path):
    """灰度PNG在Word中会偏色，转为RGB PNG"""
    try:
        from PIL import Image
        img = Image.open(img_path)
        if img.mode in ('L', 'LA', 'P', 'PA'):
            canvas = Image.new('RGB', img.size, (255, 255, 255))
            rgba = img.convert('RGBA')
            canvas.paste(rgba, mask=rgba.split()[3])
            canvas.save(img_path, 'PNG')
    except Exception:
        pass
    return img_path


def load_json(filename):
    path = os.path.join(HERE, filename)
    if not os.path.exists(path):
        return None
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def process_content(content_list, doc):
    """
    处理内容列表：
    - 字符串 → 保持不变
    - {"type":"image"} → InlineImage（图片）
    - {"type":"table"} → 转为特殊标记文本（后处理时插入表格）
    """
    from docxtpl import InlineImage
    from docx.shared import Mm

    result = []
    for item in content_list:
        if isinstance(item, str):
            result.append(item)
        elif isinstance(item, dict):
            if item.get("type") == "image":
                img_path = os.path.join(IMG_DIR, item["path"])
                if os.path.exists(img_path):
                    width = Mm(item.get("width", 120))
                    result.append(InlineImage(doc, img_path,
                                             width=width))
                    # 不跳过标注——让标注文本正常输出
                else:
                    result.append(f"[图片缺失: {item['path']}]")
            elif item.get("type") == "table":
                _pending_tables.append(item)
                tid = len(_pending_tables) - 1
                # 如果 table 对象自带 caption，先输出标注行
                if item.get("caption"):
                    result.append(item["caption"])
                result.append(f"__TABLE_PLACEHOLDER_{tid}__")
            else:
                result.append(str(item))
    return result


# 全局表格暂存
_pending_tables = []


def process_chapters(chapters, doc):
    """处理所有章节，转换图片引用"""
    for ch in chapters:
        for sec in ch.get("sections", []):
            sec["content"] = process_content(
                sec.get("content", []), doc)
            for sub in sec.get("subsections", []):
                sub["content"] = process_content(
                    sub.get("content", []), doc)
    return chapters


def build_data(doc=None):
    """组装所有章节数据"""
    meta = load_json("meta.json")
    if not meta:
        raise FileNotFoundError("meta.json not found")

    chapters = []
    for ch_file in sorted(glob.glob(
            os.path.join(HERE, "ch*.json"))):
        ch = load_json(os.path.basename(ch_file))
        if ch:
            chapters.append(ch)
            print(f"  {os.path.basename(ch_file)}"
                  f" -> {ch['title']}")

    if doc:
        chapters = process_chapters(chapters, doc)

    refs = load_json("references.json")
    if isinstance(refs, list):
        references = refs
    elif isinstance(refs, dict):
        references = refs.get("references", [])
    else:
        references = []

    return {**meta, "chapters": chapters,
            "references": references}


def main():
    if len(sys.argv) < 2:
        print("用法: python build.py <模板.docx> [输出.docx]")
        sys.exit(1)

    template = sys.argv[1]
    output = sys.argv[2] if len(sys.argv) > 2 else "output.docx"

    from docxtpl import DocxTemplate
    doc = DocxTemplate(template)

    _pending_tables.clear()
    print("组装论文数据...")
    data = build_data(doc)
    print(f"  {len(data['chapters'])} 章, "
          f"{len(data.get('references', []))} 条参考文献")

    print(f"渲染: {template} -> {output}")
    doc.render(data)
    doc.save(output)

    # 后处理：修正图片居中和图标注格式
    _post_process(output)
    # 校验
    errors = _verify(output)
    if errors:
        print(f"  ⚠ 发现 {len(errors)} 个问题:")
        for e in errors:
            print(f"    {e}")
    else:
        print(f"  ✓ 校验通过")
    print(f"完成: {output}")


def _verify(docx_path):
    """校验输出文档：检查每个表/图标注是否有对应内容"""
    from docx import Document
    import re

    doc = Document(docx_path)
    errors = []
    body = doc.element.body
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # 按 DOM 顺序遍历，记录每个元素的类型
    elements = []  # [(type, text)]
    for elem in body:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            texts = [t.text for t in elem.findall('.//w:t', ns) if t.text]
            text = ''.join(texts).strip()
            has_img = bool(elem.findall('.//w:drawing', ns))
            if has_img:
                elements.append(('img', text))
            else:
                elements.append(('p', text))
        elif tag == 'tbl':
            elements.append(('tbl', ''))

    # 检查每个表标注后面是否有表格
    for i, (etype, text) in enumerate(elements):
        if etype == 'p' and re.match(r'^表\d+-\d+\s', text) and not re.search(r'[。，；]', text) and len(text) < 30:
            # 往后找 3 个元素内是否有表格
            found_tbl = False
            for j in range(i + 1, min(i + 6, len(elements))):
                if elements[j][0] == 'tbl':
                    found_tbl = True
                    break
            if not found_tbl:
                errors.append(f"缺表格: {text}")

    # 检查每个图标注附近是否有图片
    for i, (etype, text) in enumerate(elements):
        if etype == 'p' and re.match(r'^图\d+-\d+\s', text) and len(text) < 50:
            found_img = False
            # 往前找 3 个元素
            for j in range(max(0, i - 3), i):
                if elements[j][0] == 'img':
                    found_img = True
                    break
            if not found_img:
                errors.append(f"缺图片: {text}")

    return errors


def _post_process(docx_path):
    """渲染后：图片居中、标注格式、插入表格"""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

    doc = Document(docx_path)
    fig_pat = re.compile(r'^图\d')
    tbl_cap_pat = re.compile(r'^表\d')
    placeholder_pat = re.compile(r'^__TABLE_PLACEHOLDER_(\d+)__$')

    for p in list(doc.paragraphs):
        t = (p.text or "").strip()

        # Heading 1 → 每章前分页
        if p.style and p.style.name == 'Heading 1':
            p.paragraph_format.page_break_before = True

        # 图片段落 → 居中 + 清除首行缩进 + 清除继承的正文rPr
        has_drawing = bool(p._p.findall(
            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
        if has_drawing:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 清除首行缩进（Normal样式有 firstLineChars="200"，必须显式清零）
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn as _qn2
            pPr = p._p.pPr
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p._p.insert(0, pPr)
            # 删除旧的 ind，创建新的干净的
            old_ind = pPr.find(_qn2('w:ind'))
            if old_ind is not None:
                pPr.remove(old_ind)
            new_ind = OxmlElement('w:ind')
            new_ind.set(_qn2('w:left'), '0')
            new_ind.set(_qn2('w:leftChars'), '0')
            new_ind.set(_qn2('w:firstLine'), '0')
            new_ind.set(_qn2('w:firstLineChars'), '0')
            pPr.append(new_ind)
            # 固定行间距(EXACTLY 20pt)会裁切图片，改为单倍行距
            from docx.enum.text import WD_LINE_SPACING
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            # 清除段落级 rPr（继承的字体格式会影响图片渲染）
            from docx.oxml.ns import qn as _qn
            pPr = p._p.find(_qn('w:pPr'))
            if pPr is not None:
                rPr = pPr.find(_qn('w:rPr'))
                if rPr is not None:
                    pPr.remove(rPr)
            # 清除段落中非图片 run（空文本 run 带字体格式）
            for r in p._p.findall(_qn('w:r')):
                if not r.findall(_qn('w:drawing')) and not r.findall(
                        './/{http://schemas.openxmlformats.org/drawingml/2006/main}graphic'):
                    txt = r.findall(_qn('w:t'))
                    if all((t.text or '').strip() == '' for t in txt):
                        r.getparent().remove(r)

        # 图/表标注 → 居中 + 五号黑体（短文本）
        if (fig_pat.match(t) or tbl_cap_pat.match(t)) and len(t) < 50:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(10.5)
                r.font.name = "黑体"
                # 设置东亚字体
                r._r.rPr.rFonts.set(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia',
                    '黑体')

        # 表格占位符 → 插入真实表格
        m = placeholder_pat.match(t)
        if m:
            tid = int(m.group(1))
            if tid < len(_pending_tables):
                _insert_table(doc, p, _pending_tables[tid])
                # 清空占位文本
                for r in p.runs:
                    r.text = ""

    doc.save(docx_path)
    print(f"  后处理: 图片居中+{len(_pending_tables)}个表格")


def _insert_table(doc, after_para, tbl_data):
    """在段落后插入带边框的表格"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    headers = tbl_data.get("headers", [])
    rows = tbl_data.get("rows", [])
    if not headers:
        return

    ncols = len(headers)
    all_rows = [headers] + rows

    # 直接构建表格 XML
    tbl = OxmlElement('w:tbl')

    # 表格属性：居中 + 自动列宽 + 边框
    tblPr = OxmlElement('w:tblPr')
    # 居中
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblPr.append(jc)
    # 表格宽度 100%
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    # 自动列宽（让 Word 根据内容调整）
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'autofit')
    tblPr.append(tblLayout)
    # 边框（三线表：顶线1.5pt、底线1.5pt，无左右竖线）
    borders = OxmlElement('w:tblBorders')
    for edge, sz in [('top', '12'), ('bottom', '12')]:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    for edge in ['left', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        borders.append(el)
    tblPr.append(borders)
    tbl.append(tblPr)

    # 表格网格 — 按内容长度智能分配列宽
    total_width = 9000
    # 计算每列最大文本长度（中文字符算2）
    col_max_len = [0] * ncols
    for row_data in all_rows:
        for ci in range(ncols):
            cell_text = str(row_data[ci]) if ci < len(row_data) else ""
            # 中文字符算2个宽度单位
            w = sum(2 if ord(c) > 127 else 1 for c in cell_text)
            col_max_len[ci] = max(col_max_len[ci], w)
    # 最小列宽1200 twips（约2cm），确保短文本列也够宽
    min_w = 1200
    total_len = sum(max(l, 4) for l in col_max_len) or 1
    col_widths = [max(min_w, int(total_width * max(l, 4) / total_len))
                  for l in col_max_len]
    # 修正总宽度
    diff = total_width - sum(col_widths)
    if diff > 0:
        # 多余的宽度加到最宽的列
        widest = col_widths.index(max(col_widths))
        col_widths[widest] += diff
    elif diff < 0:
        # 超出的从最宽的列减
        widest = col_widths.index(max(col_widths))
        col_widths[widest] += diff

    tblGrid = OxmlElement('w:tblGrid')
    for w in col_widths:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(w))
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    # 行
    for ri, row_data in enumerate(all_rows):
        tr = OxmlElement('w:tr')
        for ci in range(ncols):
            tc = OxmlElement('w:tc')
            # 单元格属性：宽度 + 垂直居中
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(col_widths[ci] if ci < len(col_widths) else 1800))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)
            # 三线表：表头行底部加细线（0.75pt）
            if ri == 0:
                tcBorders = OxmlElement('w:tcBorders')
                btm = OxmlElement('w:bottom')
                btm.set(qn('w:val'), 'single')
                btm.set(qn('w:sz'), '6')
                btm.set(qn('w:space'), '0')
                btm.set(qn('w:color'), '000000')
                tcBorders.append(btm)
                tcPr.append(tcBorders)
            tc.append(tcPr)
            p = OxmlElement('w:p')
            # 段落属性：水平居中 + 段前段后间距为0 + 单倍行距
            pPr = OxmlElement('w:pPr')
            pJc = OxmlElement('w:jc')
            pJc.set(qn('w:val'), 'center')
            pPr.append(pJc)
            pSpacing = OxmlElement('w:spacing')
            pSpacing.set(qn('w:before'), '0')
            pSpacing.set(qn('w:after'), '0')
            pSpacing.set(qn('w:line'), '240')
            pSpacing.set(qn('w:lineRule'), 'auto')
            pPr.append(pSpacing)
            # 清除所有缩进（firstLine + firstLineChars 都要清）
            pInd = OxmlElement('w:ind')
            pInd.set(qn('w:firstLine'), '0')
            pInd.set(qn('w:firstLineChars'), '0')
            pInd.set(qn('w:left'), '0')
            pInd.set(qn('w:right'), '0')
            pPr.append(pInd)
            p.append(pPr)
            # 文本 — 字体：宋体(中文) + Times New Roman(英文)，五号(10.5pt)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            # 字体
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), '宋体')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rPr.append(rFonts)
            # 字号 五号=10.5pt=21半磅
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '21')
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '21')
            rPr.append(szCs)
            # 表头加粗
            if ri == 0:
                b = OxmlElement('w:b')
                rPr.append(b)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = str(row_data[ci]) if ci < len(row_data) else ''
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    # 插入到段落后面
    after_para._p.addnext(tbl)


if __name__ == "__main__":
    main()
