"""
城市智慧停车系统论文组装器。

数据文件:
  meta.json, ch1.json ... ch7.json, references.json

用法:
  cd /Users/a136/vs/45425/thesis_project/papers/smart-parking
  python build.py [输出.docx]
"""
from __future__ import annotations

import json
import os
import re
import sys
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
HERE = Path(__file__).resolve().parent
TPL = ROOT / "templates" / "shopmall_parking" / "template.docx"
IMG_DIR = HERE / "images"

sys.path.insert(0, str(ROOT))

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docxtpl import DocxTemplate


def load_json(name):
    path = HERE / name
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def split_paras(value, count=3):
    if isinstance(value, list):
        parts = [str(x).strip() for x in value if str(x).strip()]
    else:
        parts = [p.strip() for p in str(value or "").split("\n") if p.strip()]
    parts = parts[:count]
    while len(parts) < count:
        parts.append("")
    return parts


def resolve_images(blocks):
    for block in blocks:
        if isinstance(block, dict) and "image" in block:
            image = Path(block["image"])
            if not image.is_absolute():
                block["image"] = str((HERE / image).resolve())
        elif isinstance(block, dict) and "subsections" in block:
            for sub in block.get("subsections", []):
                resolve_images(sub.get("content", []))


def fmt_run(run, size=12, bold=None):
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if bold is not None:
        run.bold = bold


def add_heading(sd, text, level=1, page_break=False):
    style_by_level = {
        1: "摘要",
        2: "论文一级标题",
        3: "论文二级目录",
    }
    try:
        style = style_by_level.get(level)
        p = sd.add_paragraph(str(text), style=style) if style else sd.add_heading(text, level=level)
    except KeyError:
        # Some thesis templates keep only the first two built-in heading styles.
        # Use a normal paragraph with outline level so Word/WPS TOC can still update.
        p = sd.add_paragraph(str(text))
    pPr = p._p.get_or_add_pPr()
    outline = pPr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        pPr.append(outline)
    outline.set(qn("w:val"), str(max(0, level - 1)))
    p.paragraph_format.line_spacing = 1.5
    if page_break:
        p.paragraph_format.page_break_before = True
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        size = 16
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        size = 14
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        size = 12
    for r in p.runs:
        fmt_run(r, size=size, bold=True)
    return p


def add_text(sd, text):
    p = sd.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.85)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    fmt_run(r, 12)
    return p


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is not None:
        tcPr.remove(borders)
    borders = OxmlElement("w:tcBorders")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        if val:
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), val.get("sz", "4"))
            el.set(qn("w:color"), val.get("color", "000000"))
            el.set(qn("w:space"), "0")
            borders.append(el)
    tcPr.append(borders)


def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.first_child_found_in("w:tcW")
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")


def set_table_grid(table, widths_cm):
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(int(sum(widths_cm) * 567)))
    tblW.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 567)))
        grid.append(col)
    tbl.insert(1, grid)


def add_table(sd, table_data):
    caption = table_data.get("caption", "")
    headers = table_data.get("headers", [])
    rows = table_data.get("rows", [])
    if caption:
        p = sd.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        fmt_run(r, 10.5)

    table = sd.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    default_widths = table_data.get("col_widths")
    if not default_widths:
        if len(headers) == 3:
            default_widths = [3.2, 3.2, 8.0]
        elif len(headers) == 4:
            default_widths = [3.0, 3.0, 3.0, 5.0]
        else:
            default_widths = [13.8 / max(1, len(headers))] * max(1, len(headers))
    set_table_grid(table, default_widths)

    thick = {"val": "single", "sz": "12", "color": "000000"}
    thin = {"val": "single", "sz": "4", "color": "000000"}
    none = {"val": "nil"}

    for j, text in enumerate(headers):
        cell = table.cell(0, j)
        if j < len(default_widths):
            cell.width = Cm(default_widths[j])
            set_cell_width(cell, default_widths[j])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = str(text)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                fmt_run(r, 10.5, True)
        set_cell_border(cell, top=thick, bottom=thin, left=none, right=none)

    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.cell(i + 1, j)
            if j < len(default_widths):
                cell.width = Cm(default_widths[j])
                set_cell_width(cell, default_widths[j])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = str(text)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) < 18 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    fmt_run(r, 10.5)
            set_cell_border(
                cell,
                top=none,
                bottom=thick if i == len(rows) - 1 else none,
                left=none,
                right=none,
            )
    sd.add_paragraph()


def add_image(sd, path, caption, width_cm=12):
    p = sd.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Cm(width_cm))
    if caption:
        cap = sd.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = cap.add_run(caption)
        fmt_run(rr, 10.5)


def add_code(sd, code, caption=None):
    if caption:
        p = sd.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        fmt_run(r, 10.5)
    p = sd.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    for line in str(code).strip("\n").splitlines():
        r = p.add_run(line + "\n")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(9)


def build_blocks(sd, blocks):
    for block in blocks or []:
        if isinstance(block, str):
            add_text(sd, block)
        elif isinstance(block, dict):
            if "text" in block:
                add_text(sd, block["text"])
            elif "table" in block:
                add_table(sd, block["table"])
            elif "image" in block:
                add_image(sd, block["image"], block.get("caption", ""), block.get("width_cm", 12))
            elif "code" in block:
                add_code(sd, block["code"], block.get("caption"))


def build_body(doc, chapters):
    sd = doc.new_subdoc()
    for i, ch in enumerate(chapters):
        add_heading(sd, ch["title"], level=1, page_break=True)
        build_blocks(sd, ch.get("content", []))
        for sec in ch.get("sections", []):
            add_heading(sd, sec["title"], level=2)
            build_blocks(sd, sec.get("content", []))
            for sub in sec.get("subsections", []):
                add_heading(sd, sub["title"], level=3)
                build_blocks(sd, sub.get("content", []))
    return sd


def build_refs(doc, refs):
    sd = doc.new_subdoc()
    for i, ref in enumerate(refs, 1):
        text = re.sub(r"^\[\d+\]\s*", "", str(ref).strip())
        p = sd.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(f"[{i}] {text}")
        fmt_run(r, 12)
    return sd


def patch_update_fields(docx_path):
    with zipfile.ZipFile(docx_path, "a") as zf:
        xml = zf.read("word/settings.xml").decode("utf-8")
        if "w:updateFields" not in xml:
            xml = xml.replace("</w:settings>", '<w:updateFields w:val="true"/></w:settings>')
            zf.writestr("word/settings.xml", xml)


def build_data(doc):
    meta = load_json("meta.json") or {}
    chapters = []
    for i in range(1, 8):
        ch = load_json(f"ch{i}.json")
        if ch:
            print(f"  ch{i}.json -> {ch['title']}")
            for block in ch.get("content", []):
                pass
            resolve_images(ch.get("content", []))
            for sec in ch.get("sections", []):
                resolve_images(sec.get("content", []))
                for sub in sec.get("subsections", []):
                    resolve_images(sub.get("content", []))
            chapters.append(ch)

    zh = split_paras(meta.get("abstract_zh"), 3)
    en = split_paras(meta.get("abstract_en"), 3)
    for i, text in enumerate(zh, 1):
        meta[f"abstract_zh_{i}"] = text
    for i, text in enumerate(en, 1):
        meta[f"abstract_en_{i}"] = text

    refs = load_json("references.json") or []
    if isinstance(refs, dict):
        refs = refs.get("references", [])

    meta["body"] = build_body(doc, chapters)
    meta["references_doc"] = build_refs(doc, refs)
    return meta


def verify_docx(docx_path):
    from docx import Document
    doc = Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    errors = []
    for tag in ("{{", "}}", "{%", "%}"):
        if tag in text:
            errors.append(f"残留模板标记 {tag}")
    if "购物商城" in text:
        errors.append("残留购物商城文本")
    if "城市智慧停车" not in text and "智慧停车" not in text:
        errors.append("未检测到停车系统主题文本")
    return errors


def main():
    output = Path(sys.argv[1]) if len(sys.argv) > 1 else HERE / "城市智慧停车系统论文.docx"
    doc = DocxTemplate(TPL)
    print("组装论文数据...")
    data = build_data(doc)
    print("渲染模板...")
    doc.render(data)
    doc.save(output)
    patch_update_fields(output)
    errors = verify_docx(output)
    if errors:
        print("校验提醒:")
        for e in errors:
            print(" -", e)
    else:
        print("校验通过")
    print(f"完成: {output}")


if __name__ == "__main__":
    main()
