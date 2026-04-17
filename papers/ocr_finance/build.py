"""
基于OCR的智能财务报销系统论文组装器。
支持文本、图片（InlineImage）、表格（后处理三线表）。

用法: python build.py [输出.docx]
"""
import os
import sys
import json
import re

ROOT = os.path.join(os.path.dirname(__file__), "../..")
sys.path.insert(0, ROOT)

HERE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(HERE, "images")
TPL = os.path.join(ROOT, "templates/dslg/template.docx")

_pending_tables = []


def load_json(filename):
    path = os.path.join(HERE, filename)
    if not os.path.exists(path):
        return None
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def process_content(content_list, doc):
    from docxtpl import InlineImage
    from docx.shared import Mm
    result = []
    for item in content_list:
        if isinstance(item, str):
            result.append(item)
        elif isinstance(item, dict):
            if item.get("type") == "image":
                img_path = os.path.join(IMG_DIR, item["path"])
                if os.path.exists(img_path):
                    width = Mm(item.get("width", 148))
                    result.append(InlineImage(doc, img_path, width=width))
                    if item.get("caption"):
                        result.append(item["caption"])
                else:
                    result.append(f"[图片缺失: {item['path']}]")
                    if item.get("caption"):
                        result.append(item["caption"])
            elif item.get("type") == "table":
                _pending_tables.append(item)
                tid = len(_pending_tables) - 1
                if item.get("caption"):
                    result.append(item["caption"])
                result.append(f"__TABLE_PLACEHOLDER_{tid}__")
            else:
                result.append(str(item))
    return result


def process_chapters(chapters, doc):
    for ch in chapters:
        for sec in ch.get("sections", []):
            sec["content"] = process_content(sec.get("content", []), doc)
            for sub in sec.get("subsections", []):
                sub["content"] = process_content(sub.get("content", []), doc)
    return chapters


def build_data(doc=None):
    meta = load_json("meta.json")
    if not meta:
        raise FileNotFoundError("meta.json not found")

    chapters = []
    for idx in range(1, 8):
        ch = load_json(f"ch{idx}.json")
        if ch:
            chapters.append(ch)
            print(f"  ch{idx}.json -> {ch['title']}")

    if doc:
        chapters = process_chapters(chapters, doc)

    refs = load_json("references.json") or []

    # 摘要拆段
    for key, list_key in [("abstract_zh", "abstract_zh_list"),
                          ("abstract_en", "abstract_en_list")]:
        text = meta.get(key, "")
        if isinstance(text, list):
            meta[list_key] = text
        else:
            meta[list_key] = [p.strip() for p in text.split("\n") if p.strip()]

    # 结论拆段
    conclusion = meta.get("conclusion_list", [])
    if isinstance(conclusion, str):
        meta["conclusion_list"] = [p.strip() for p in conclusion.split("\n") if p.strip()]

    return {**meta, "chapters": chapters, "references": refs}


def main():
    output = sys.argv[1] if len(sys.argv) > 1 else None

    from docxtpl import DocxTemplate
    doc = DocxTemplate(TPL)

    _pending_tables.clear()
    print("组装论文数据...")
    data = build_data(doc)
    print(f"  {len(data['chapters'])} 章, "
          f"{len(data.get('references', []))} 条参考文献")

    if not output:
        from datetime import datetime
        ts = datetime.now().strftime("%m%d_%H%M")
        output = os.path.join(HERE,
            f"基于OCR的智能财务报销系统设计与实现_{ts}.docx")

    print(f"渲染: {output}")
    doc.render(data)
    doc.save(output)

    _post_process(output)
    errors = _verify(output)
    if errors:
        print(f"  ⚠ {len(errors)} 个问题:")
        for e in errors:
            print(f"    {e}")
    else:
        print("  ✓ 校验通过")
    print(f"完成: {output}")


def _verify(docx_path):
    from docx import Document
    doc = Document(docx_path)
    errors = []
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    body = doc.element.body
    elements = []
    for elem in body:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            texts = [t.text for t in elem.findall('.//w:t', ns) if t.text]
            text = ''.join(texts).strip()
            has_img = bool(elem.findall('.//w:drawing', ns))
            elements.append(('img' if has_img else 'p', text))
        elif tag == 'tbl':
            elements.append(('tbl', ''))

    for i, (etype, text) in enumerate(elements):
        if '__TABLE_PLACEHOLDER_' in text:
            errors.append(f"未替换占位符: {text[:40]}")

    return errors


def _post_process(docx_path):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    doc = Document(docx_path)
    tbl_cap_pat = re.compile(r'^表\d')
    fig_cap_pat = re.compile(r'^图\d')
    placeholder_pat = re.compile(r'^__TABLE_PLACEHOLDER_(\d+)__$')

    for p in list(doc.paragraphs):
        t = (p.text or "").strip()

        # 图片居中
        has_drawing = bool(p._p.findall(
            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
        if has_drawing:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pPr = p._p.pPr
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p._p.insert(0, pPr)
            old_ind = pPr.find(_qn('w:ind'))
            if old_ind is not None:
                pPr.remove(old_ind)
            new_ind = OxmlElement('w:ind')
            new_ind.set(_qn('w:firstLine'), '0')
            new_ind.set(_qn('w:firstLineChars'), '0')
            pPr.append(new_ind)
            from docx.enum.text import WD_LINE_SPACING
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            rPr = pPr.find(_qn('w:rPr'))
            if rPr is not None:
                pPr.remove(rPr)

        # 图表标注居中
        if (fig_cap_pat.match(t) or tbl_cap_pat.match(t)) and len(t) < 60:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(10.5)

        # 表格占位符
        m = placeholder_pat.match(t)
        if m:
            tid = int(m.group(1))
            if tid < len(_pending_tables):
                _insert_table(doc, p, _pending_tables[tid])
                for r in p.runs:
                    r.text = ""

    # 英文摘要字体 → Times New Roman
    en_abs_started = False
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and t == "Abstract":
            en_abs_started = True
            continue
        if en_abs_started and style == "Heading 1":
            break
        if en_abs_started and t and not t.startswith("Key"):
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                rPr = r._r.get_or_add_rPr()
                rFonts = rPr.find(_qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.insert(0, rFonts)
                rFonts.set(_qn('w:ascii'), 'Times New Roman')
                rFonts.set(_qn('w:hAnsi'), 'Times New Roman')
                rFonts.set(_qn('w:eastAsia'), 'Times New Roman')

    # 中文摘要正文字体 → 宋体 小四
    zh_abs_started = False
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and "摘" in t and "要" in t:
            zh_abs_started = True
            continue
        if zh_abs_started and style == "Heading 1":
            break
        if zh_abs_started and t and not t.startswith("关键词"):
            for r in p.runs:
                r.font.size = Pt(12)

    # 清理模板残留的空段落
    _ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    to_remove = []
    paras = doc.paragraphs

    # 第一遍：定位关键位置
    keywords_zh_idx = None
    abstract_idx = None
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        style = p.style.name if p.style else ""
        if t.startswith("关键词") and keywords_zh_idx is None:
            keywords_zh_idx = i
        if style == "Heading 1" and t == "Abstract":
            abstract_idx = i

    # 第二遍：标记要删除的空段落
    in_ref = False
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        style = p.style.name if p.style else ""
        # 中文关键词与Abstract之间的空段落
        if keywords_zh_idx is not None and abstract_idx is not None:
            if keywords_zh_idx < i < abstract_idx and not t:
                has_sect = p._p.find('.//w:sectPr', _ns) is not None
                if not has_sect:
                    to_remove.append(p)
        # Abstract 后、正文前的空"标题3"段落
        if abstract_idx is not None and i > abstract_idx:
            if style in ("标题3",) and not t:
                to_remove.append(p)
        # 参考文献后的空编号段落
        if '参考文献' in t:
            in_ref = True
        if in_ref:
            has_num = p._p.find('.//w:numPr', _ns) is not None
            if has_num and not t:
                to_remove.append(p)

    for p in to_remove:
        parent = p._p.getparent()
        if parent is not None:
            parent.remove(p._p)
    if to_remove:
        print(f"  清理: {len(to_remove)} 个空段落")

    # 确保目录(TOC sdt)前有分页符
    body = doc.element.body
    _wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for child in body:
        if child.tag == f'{{{_wns}}}sdt':
            gallery = child.find(f'.//{{{_wns}}}docPartGallery')
            if gallery is not None and 'Table of Contents' in (gallery.get(f'{{{_wns}}}val') or ''):
                # 在 sdt 前插入一个带分页符的空段落
                page_break_p = OxmlElement('w:p')
                pPr = OxmlElement('w:pPr')
                pbb = OxmlElement('w:pageBreakBefore')
                pPr.append(pbb)
                page_break_p.append(pPr)
                child.addprevious(page_break_p)
                break

    # 确保摘要和Abstract标题前有分页符
    for p in paras:
        t = (p.text or "").strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and t in ("摘 要", "Abstract", "结 论", "致 谢", "参考文献"):
            existing = p._p.find('.//w:pageBreakBefore',
                {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if existing is None:
                pPr = p._p.get_or_add_pPr()
                pageBreakBefore = OxmlElement('w:pageBreakBefore')
                pPr.insert(0, pageBreakBefore)

    doc.save(docx_path)
    print(f"  后处理: {len(_pending_tables)} 个表格")


def _insert_table(doc, after_para, tbl_data):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    headers = tbl_data.get("headers", [])
    rows = tbl_data.get("rows", [])
    if not headers:
        return

    ncols = len(headers)
    all_rows = [headers] + rows

    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblPr.append(jc)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'autofit')
    tblPr.append(tblLayout)
    borders = OxmlElement('w:tblBorders')
    for edge, sz in [('top', '12'), ('bottom', '12')]:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    for edge in ['left', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        borders.append(el)
    tblPr.append(borders)
    tbl.append(tblPr)

    total_width = 9000
    col_max_len = [0] * ncols
    for row_data in all_rows:
        for ci in range(ncols):
            cell_text = str(row_data[ci]) if ci < len(row_data) else ""
            w = sum(2 if ord(c) > 127 else 1 for c in cell_text)
            col_max_len[ci] = max(col_max_len[ci], w)
    min_w = 1200
    total_len = sum(max(l, 4) for l in col_max_len) or 1
    col_widths = [max(min_w, int(total_width * max(l, 4) / total_len))
                  for l in col_max_len]
    diff = total_width - sum(col_widths)
    widest = col_widths.index(max(col_widths))
    col_widths[widest] += diff

    tblGrid = OxmlElement('w:tblGrid')
    for w in col_widths:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(w))
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    for ri, row_data in enumerate(all_rows):
        tr = OxmlElement('w:tr')
        for ci in range(ncols):
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(col_widths[ci]))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)
            if ri == 0:
                tcBorders = OxmlElement('w:tcBorders')
                btm = OxmlElement('w:bottom')
                btm.set(qn('w:val'), 'single')
                btm.set(qn('w:sz'), '6')
                btm.set(qn('w:space'), '0')
                btm.set(qn('w:color'), '000000')
                tcBorders.append(btm)
                tcPr.append(tcBorders)
            tc.append(tcPr)
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            pJc = OxmlElement('w:jc')
            pJc.set(qn('w:val'), 'center')
            pPr.append(pJc)
            pSpacing = OxmlElement('w:spacing')
            pSpacing.set(qn('w:before'), '0')
            pSpacing.set(qn('w:after'), '0')
            pSpacing.set(qn('w:line'), '240')
            pSpacing.set(qn('w:lineRule'), 'auto')
            pPr.append(pSpacing)
            pInd = OxmlElement('w:ind')
            pInd.set(qn('w:firstLine'), '0')
            pInd.set(qn('w:firstLineChars'), '0')
            pPr.append(pInd)
            p.append(pPr)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), '宋体')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '21')
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '21')
            rPr.append(szCs)
            if ri == 0:
                b = OxmlElement('w:b')
                rPr.append(b)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = str(row_data[ci]) if ci < len(row_data) else ''
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    after_para._p.addnext(tbl)


if __name__ == "__main__":
    main()
