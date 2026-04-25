"""
社区老年食堂订餐系统论文组装器（黑龙江工商学院模板）。
支持文本、图片（InlineImage）、表格（三线表）、代码块。

用法: python build.py [输出.docx]
"""
import os
import sys
import json
import re

ROOT = os.path.join(os.path.dirname(__file__), "../..")
sys.path.insert(0, ROOT)

HERE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(HERE, "images")
TPL = os.path.join(ROOT, "templates/hljgsxy/template.docx")

_pending_tables = []
_pending_code_blocks = []


def load_json(filename):
    path = os.path.join(HERE, filename)
    if not os.path.exists(path):
        return None
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def process_content(content_list, doc):
    from docxtpl import InlineImage
    from docx.shared import Mm
    result = []
    for item in content_list:
        if isinstance(item, str):
            result.append(item)
        elif isinstance(item, dict):
            if item.get("type") == "image":
                img_path = os.path.join(IMG_DIR, item["path"])
                if os.path.exists(img_path):
                    width = Mm(item.get("width", 148))
                    result.append(InlineImage(doc, img_path, width=width))
                    if item.get("caption"):
                        result.append(item["caption"])
                else:
                    result.append(f"[图片缺失: {item['path']}]")
                    if item.get("caption"):
                        result.append(item["caption"])
            elif item.get("type") == "code":
                _pending_code_blocks.append(item)
                cid = len(_pending_code_blocks) - 1
                result.append(f"__CODE_PLACEHOLDER_{cid}__")
            elif item.get("type") == "table":
                _pending_tables.append(item)
                tid = len(_pending_tables) - 1
                if item.get("caption"):
                    result.append(item["caption"])
                result.append(f"__TABLE_PLACEHOLDER_{tid}__")
            else:
                result.append(str(item))
    return result


def process_chapters(chapters, doc):
    for ch in chapters:
        ch["content"] = process_content(ch.get("content", []), doc)
        for sec in ch.get("sections", []):
            sec["content"] = process_content(sec.get("content", []), doc)
            for sub in sec.get("subsections", []):
                sub["content"] = process_content(sub.get("content", []), doc)
    return chapters


def build_data(doc=None):
    meta = load_json("meta.json") or {}
    chapters = []
    for i in range(1, 6):
        ch = load_json(f"ch{i}.json")
        if ch:
            print(f"  ch{i}.json -> {ch['title']}")
            chapters.append(ch)

    if doc:
        chapters = process_chapters(chapters, doc)

    refs = load_json("references.json") or []
    refs = [_format_reference(ref, i + 1) for i, ref in enumerate(refs)]

    # 摘要拆段
    for key, list_key in [("abstract_zh", "abstract_zh_list"),
                          ("abstract_en", "abstract_en_list")]:
        text = meta.get(key, "")
        if isinstance(text, list):
            meta[list_key] = text
        else:
            meta[list_key] = [p.strip() for p in text.split("\n") if p.strip()]

    # 结论拆段
    conclusion = meta.get("conclusion_list", [])
    if isinstance(conclusion, str):
        meta["conclusion_list"] = [p.strip() for p in conclusion.split("\n") if p.strip()]

    # 致谢拆段
    ack = meta.get("acknowledgement", "")
    if isinstance(ack, list):
        meta["acknowledgement_list"] = ack
    else:
        meta["acknowledgement_list"] = [p.strip() for p in ack.split("\n") if p.strip()]

    return {**meta, "chapters": chapters, "references": refs}


def _format_reference(ref, index):
    """Return reference text without numbering (template style adds numbers)."""
    text = str(ref).strip()
    if not text:
        return text
    # Strip existing [N] prefix if present
    text = re.sub(r"^\[\d+\]\s*", "", text)
    return text


def main():
    output = sys.argv[1] if len(sys.argv) > 1 else None

    from docxtpl import DocxTemplate
    doc = DocxTemplate(TPL)

    _pending_tables.clear()
    _pending_code_blocks.clear()
    print("组装论文数据...")
    data = build_data(doc)
    print(f"  {len(data['chapters'])} 章, "
          f"{len(data.get('references', []))} 条参考文献")

    if not output:
        from datetime import datetime
        ts = datetime.now().strftime("%m%d_%H%M")
        output = os.path.join(HERE, f"社区老年食堂订餐系统_{ts}.docx")

    print(f"渲染: {output}")
    doc.render(data)
    doc.save(output)

    _post_process(output)
    errors = _verify(output)
    if errors:
        print(f"  ⚠ {len(errors)} 个问题:")
        for e in errors:
            print(f"    - {e}")
    else:
        print("  ✓ 校验通过")
    print(f"完成: {output}")


def _post_process(docx_path):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    doc = Document(docx_path)
    tbl_cap_pat = re.compile(r'^表\d')
    fig_cap_pat = re.compile(r'^图\d')
    placeholder_pat = re.compile(r'^__TABLE_PLACEHOLDER_(\d+)__$')
    code_placeholder_pat = re.compile(r'^__CODE_PLACEHOLDER_(\d+)__$')

    for p in list(doc.paragraphs):
        t = (p.text or "").strip()

        # 代码块占位符
        cm = code_placeholder_pat.match(t)
        if cm:
            cid = int(cm.group(1))
            if cid < len(_pending_code_blocks):
                _insert_code_block(doc, p, _pending_code_blocks[cid])
                for r in p.runs:
                    r.text = ""

        # 图片居中
        has_drawing = bool(p._p.findall(
            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
        if has_drawing:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pPr = p._p.pPr
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p._p.insert(0, pPr)
            old_ind = pPr.find(_qn('w:ind'))
            if old_ind is not None:
                pPr.remove(old_ind)
            new_ind = OxmlElement('w:ind')
            new_ind.set(_qn('w:firstLine'), '0')
            new_ind.set(_qn('w:firstLineChars'), '0')
            pPr.append(new_ind)
            from docx.enum.text import WD_LINE_SPACING
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 图表标注居中
        if (fig_cap_pat.match(t) or tbl_cap_pat.match(t)) and len(t) < 60:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if t == "表2-2 注册用例描述":
                p.paragraph_format.page_break_before = True
            for r in p.runs:
                r.font.size = Pt(10.5)

        # 表格占位符
        m = placeholder_pat.match(t)
        if m:
            tid = int(m.group(1))
            if tid < len(_pending_tables):
                _insert_table(doc, p, _pending_tables[tid])
                for r in p.runs:
                    r.text = ""

    # ---- 清理空白页：合并连续SECT段，删除多余空段 ----
    _NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    paras = doc.paragraphs
    to_remove = []
    removed_blank = 0

    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        has_sect = p._p.find('.//w:sectPr', _NS) is not None

        if not text and has_sect:
            if i + 1 < len(paras):
                next_text = (paras[i+1].text or "").strip()
                next_sect = paras[i+1]._p.find('.//w:sectPr', _NS) is not None
                if not next_text and next_sect:
                    to_remove.append(p._p)
                    removed_blank += 1
                    continue

        if not text and not has_sect:
            if i > 0 and i + 1 < len(paras):
                prev_sect = paras[i-1]._p.find('.//w:sectPr', _NS) is not None
                next_style = paras[i+1].style.name if paras[i+1].style else ""
                if prev_sect and "Heading" in next_style:
                    to_remove.append(p._p)
                    removed_blank += 1

    for elem in to_remove:
        if elem.getparent() is not None:
            elem.getparent().remove(elem)

    # ---- 章标题前分页 ----
    chapter_pat = re.compile(r'^\d\s')
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ''
        t = (p.text or '').strip()
        if '一级标题' in style_name and chapter_pat.match(t):
            p.paragraph_format.page_break_before = True

    # Remove empty appendix heading left by the structural template.
    for p in list(doc.paragraphs):
        style_name = p.style.name if p.style else ''
        t = (p.text or '').strip()
        if t == '附录' and '附录标题' in style_name:
            p._p.getparent().remove(p._p)

    _fix_keyword_runs(doc)

    doc.save(docx_path)
    print(f"  后处理: {len(_pending_tables)} 个表格, {len(_pending_code_blocks)} 个代码块, 清理{removed_blank}个空白段")


def _fix_keyword_runs(doc):
    """Keep the hljgsxy keyword line split like the source template.

    The source template uses separate runs for:
    label, first keyword, '；second keyword', '；', remaining keywords.
    If docxtpl renders the whole keyword list into a single run, Word shows a
    subtly different character style for the latter terms.
    """
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        style_name = p.style.name if p.style else ""
        if style_name != "论-摘要关键字" or not text.startswith("关键词  "):
            continue
        keywords = text.replace("关键词", "", 1).strip()
        parts = [x.strip() for x in keywords.split("；") if x.strip()]
        if len(parts) < 2 or len(p.runs) < 5:
            continue
        values = [
            "关键词  ",
            parts[0],
            "；" + parts[1],
            "；" if len(parts) > 2 else "",
            "；".join(parts[2:]) if len(parts) > 2 else "",
        ]
        for i, r in enumerate(p.runs):
            r.text = values[i] if i < len(values) else ""


def _insert_code_block(doc, after_para, code_data):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    code_text = code_data.get("content", "")
    lines = code_text.split("\n")

    for line in reversed(lines):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), '0')
        ind.set(qn('w:firstLineChars'), '0')
        pPr.append(ind)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), '240')
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)
        p.append(pPr)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Courier New')
        rFonts.set(qn('w:hAnsi'), 'Courier New')
        rFonts.set(qn('w:cs'), 'Courier New')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '21')
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '21')
        rPr.append(szCs)
        r.append(rPr)
        t_el = OxmlElement('w:t')
        t_el.set(qn('xml:space'), 'preserve')
        t_el.text = line
        r.append(t_el)
        p.append(r)
        after_para._p.addnext(p)


def _insert_table(doc, after_para, tbl_data):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    headers = tbl_data.get("headers", [])
    rows = tbl_data.get("rows", [])
    if not headers:
        return

    all_rows = [headers] + rows
    col_count = len(headers)
    page_width = 9638
    raw_widths = tbl_data.get("widths")
    if raw_widths and len(raw_widths) == col_count:
        total = float(sum(raw_widths)) or 1.0
        grid_widths = [str(int(page_width * float(w) / total)) for w in raw_widths]
    else:
        grid_widths = [str(int(page_width / col_count)) for _ in range(col_count)]
    table_style = tbl_data.get("style", "three_line")
    grid_table = table_style == "grid"

    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblPr.append(jc)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'autofit')
    tblPr.append(tblLayout)
    borders = OxmlElement('w:tblBorders')
    for edge, sz in [('top', '12'), ('bottom', '12')]:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    for edge in ['left', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single' if grid_table else 'none')
        el.set(qn('w:sz'), '6' if grid_table else '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)
    tbl.append(tblPr)

    # Add tblGrid (required for valid table XML)
    tblGrid = OxmlElement('w:tblGrid')
    for col_width in grid_widths:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), col_width)
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    for row_idx, cells in enumerate(all_rows):
        tr = OxmlElement('w:tr')
        for cell_idx, cell_text in enumerate(cells):
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), grid_widths[cell_idx])
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            tcMar = OxmlElement('w:tcMar')
            for side in ['top', 'left', 'bottom', 'right']:
                mar = OxmlElement(f'w:{side}')
                mar.set(qn('w:w'), '90' if grid_table else '60')
                mar.set(qn('w:type'), 'dxa')
                tcMar.append(mar)
            tcPr.append(tcMar)
            if row_idx == 0 or grid_table:
                tcBorders = OxmlElement('w:tcBorders')
                edges = ['bottom'] if not grid_table else ['top', 'left', 'bottom', 'right']
                for edge in edges:
                    border = OxmlElement(f'w:{edge}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '6')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
            tc.append(tcPr)
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            pJc = OxmlElement('w:jc')
            pJc.set(qn('w:val'), 'center' if row_idx == 0 or cell_idx == 0 else 'left')
            pPr.append(pJc)
            p.append(pPr)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '21')
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '21')
            rPr.append(szCs)
            if row_idx == 0:
                b = OxmlElement('w:b')
                rPr.append(b)
            r.append(rPr)
            t_el = OxmlElement('w:t')
            t_el.set(qn('xml:space'), 'preserve')
            t_el.text = str(cell_text)
            r.append(t_el)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    after_para._p.addnext(tbl)


def _verify(docx_path):
    from docx import Document
    doc = Document(docx_path)
    errors = []
    allowed_cover_placeholders = {
        "{{ name }}",
        "{{ student_id }}",
        "{{ advisor }}",
        "{{ advisor_title }}",
    }
    texts = []
    for p in doc.paragraphs:
        text = p.text or ""
        for ph in allowed_cover_placeholders:
            text = text.replace(ph, "")
        texts.append(text)
    full = "\n".join(texts)
    for tag in ["{%", "%}", "{{", "}}"]:
        if tag in full:
            errors.append(f"残留模板标记: {tag}")
    return errors


if __name__ == "__main__":
    main()
