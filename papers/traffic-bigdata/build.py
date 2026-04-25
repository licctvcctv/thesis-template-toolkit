"""
智能出行交通数据可视化分析系统论文组装器。

数据文件:
  meta.json, ch1.json ... ch7.json, references.json

用法:
  cd /Users/a136/vs/45425/thesis_project/papers/traffic-bigdata
  python build.py [输出.docx]
"""
from __future__ import annotations

import json
import os
import re
import sys
import zipfile
from copy import deepcopy
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
HERE = Path(__file__).resolve().parent
TPL = ROOT / "templates" / "tjut_traffic" / "template.docx"
IMG_DIR = HERE / "images"

sys.path.insert(0, str(ROOT))

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docxtpl import DocxTemplate


def load_json(name):
    path = HERE / name
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def split_paras(value, count=3):
    if isinstance(value, list):
        parts = [str(x).strip() for x in value if str(x).strip()]
    else:
        parts = [p.strip() for p in str(value or "").split("\n") if p.strip()]
    parts = parts[:count]
    while len(parts) < count:
        parts.append("")
    return parts


def resolve_images(blocks):
    for block in blocks:
        if isinstance(block, dict) and "image" in block:
            image = Path(block["image"])
            if not image.is_absolute():
                block["image"] = str((HERE / image).resolve())
        elif isinstance(block, dict) and "subsections" in block:
            for sub in block.get("subsections", []):
                resolve_images(sub.get("content", []))


def fmt_run(run, size=12, bold=None):
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if bold is not None:
        run.bold = bold


def fmt_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12, bold=None):
    run.font.size = Pt(size)
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    if bold is not None:
        run.bold = bold


def add_heading(sd, text, level=1, page_break=False):
    style_by_level = {
        1: "Heading 1",
        2: "Heading 2",
        3: "Heading 3",
    }
    try:
        style = style_by_level.get(level)
        p = sd.add_paragraph(str(text), style=style) if style else sd.add_heading(text, level=level)
    except KeyError:
        # Some thesis templates keep only the first two built-in heading styles.
        # Use a normal paragraph with outline level so Word/WPS TOC can still update.
        p = sd.add_paragraph(str(text))
    pPr = p._p.get_or_add_pPr()
    outline = pPr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        pPr.append(outline)
    outline.set(qn("w:val"), str(max(0, level - 1)))
    p.paragraph_format.line_spacing = 1.5
    if page_break:
        p.paragraph_format.page_break_before = True
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        size = 16
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        size = 14
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        size = 12
    for r in p.runs:
        fmt_run(r, size=size, bold=True)
    return p


def add_text(sd, text):
    p = sd.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.85)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    fmt_run(r, 12)
    return p


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is not None:
        tcPr.remove(borders)
    borders = OxmlElement("w:tcBorders")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        if val:
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), val.get("sz", "4"))
            el.set(qn("w:color"), val.get("color", "000000"))
            el.set(qn("w:space"), "0")
            borders.append(el)
    tcPr.append(borders)


def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.first_child_found_in("w:tcW")
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")


def set_table_grid(table, widths_cm):
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(int(sum(widths_cm) * 567)))
    tblW.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 567)))
        grid.append(col)
    tbl.insert(1, grid)


def add_table(sd, table_data):
    caption = table_data.get("caption", "")
    headers = table_data.get("headers", [])
    rows = table_data.get("rows", [])
    if caption:
        p = sd.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        fmt_run(r, 10.5)

    table = sd.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    default_widths = table_data.get("col_widths")
    if not default_widths:
        if len(headers) == 3:
            default_widths = [3.2, 3.2, 8.0]
        elif len(headers) == 4:
            default_widths = [3.0, 3.0, 3.0, 5.0]
        else:
            default_widths = [13.8 / max(1, len(headers))] * max(1, len(headers))
    set_table_grid(table, default_widths)

    thick = {"val": "single", "sz": "12", "color": "000000"}
    thin = {"val": "single", "sz": "4", "color": "000000"}
    none = {"val": "nil"}

    for j, text in enumerate(headers):
        cell = table.cell(0, j)
        if j < len(default_widths):
            cell.width = Cm(default_widths[j])
            set_cell_width(cell, default_widths[j])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = str(text)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                fmt_run(r, 10.5, True)
        set_cell_border(cell, top=thick, bottom=thin, left=none, right=none)

    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.cell(i + 1, j)
            if j < len(default_widths):
                cell.width = Cm(default_widths[j])
                set_cell_width(cell, default_widths[j])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = str(text)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) < 18 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    fmt_run(r, 10.5)
            set_cell_border(
                cell,
                top=none,
                bottom=thick if i == len(rows) - 1 else none,
                left=none,
                right=none,
            )
    sd.add_paragraph()


def add_image(sd, path, caption, width_cm=12):
    p = sd.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Cm(width_cm))
    if caption:
        cap = sd.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = cap.add_run(caption)
        fmt_run(rr, 10.5)


def add_code(sd, code, caption=None):
    if caption:
        p = sd.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        fmt_run(r, 10.5)
    p = sd.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    for line in str(code).strip("\n").splitlines():
        r = p.add_run(line + "\n")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(9)


def build_blocks(sd, blocks):
    for block in blocks or []:
        if isinstance(block, str):
            add_text(sd, block)
        elif isinstance(block, dict):
            if "text" in block:
                add_text(sd, block["text"])
            elif "table" in block:
                add_table(sd, block["table"])
            elif "image" in block:
                add_image(sd, block["image"], block.get("caption", ""), block.get("width_cm", 12))
            elif "code" in block:
                add_code(sd, block["code"], block.get("caption"))


def build_body(doc, chapters):
    sd = doc.new_subdoc()
    for i, ch in enumerate(chapters):
        add_heading(sd, ch["title"], level=1, page_break=True)
        build_blocks(sd, ch.get("content", []))
        for sec in ch.get("sections", []):
            add_heading(sd, sec["title"], level=2)
            build_blocks(sd, sec.get("content", []))
            for sub in sec.get("subsections", []):
                add_heading(sd, sub["title"], level=3)
                build_blocks(sd, sub.get("content", []))
    return sd


def build_refs(doc, refs):
    sd = doc.new_subdoc()
    for i, ref in enumerate(refs, 1):
        text = re.sub(r"^\[\d+\]\s*", "", str(ref).strip())
        p = sd.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(f"[{i}] {text}")
        fmt_run(r, 12)
    return sd


def patch_update_fields(docx_path):
    replacements = {}
    with zipfile.ZipFile(docx_path, "r") as zf:
        for info in zf.infolist():
            if info.filename == "word/settings.xml":
                xml = zf.read(info.filename).decode("utf-8")
                if "w:updateFields" not in xml:
                    xml = xml.replace("</w:settings>", '<w:updateFields w:val="true"/></w:settings>')
                replacements[info.filename] = xml.encode("utf-8")

        tmp_path = Path(str(docx_path) + ".tmp")
        seen = set()
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as out:
            for info in zf.infolist():
                if info.filename in seen:
                    continue
                seen.add(info.filename)
                data = replacements.get(info.filename)
                if data is None:
                    data = zf.read(info.filename)
                out.writestr(info, data)
    tmp_path.replace(docx_path)


def _clear_runs(paragraph):
    for run in paragraph.runs:
        run.text = ""


def _add_page_field(paragraph, number_format=None):
    _clear_runs(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE" if not number_format else f"PAGE \\\\* {number_format}"
    run._r.append(instr)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)
    fmt_run(run, 10.5)


def _set_section_page_start(section, start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def _set_section_page_numbering(section, start=None, fmt=None):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    if start is not None:
        pg_num.set(qn("w:start"), str(start))
    if fmt is not None:
        pg_num.set(qn("w:fmt"), fmt)


def _set_sect_pr_page_numbering(sect_pr, start=None, fmt=None):
    for el in list(sect_pr.findall(qn("w:pgNumType"))):
        sect_pr.remove(el)
    pg_num = OxmlElement("w:pgNumType")
    if start is not None:
        pg_num.set(qn("w:start"), str(start))
    if fmt is not None:
        pg_num.set(qn("w:fmt"), fmt)
    sect_pr.append(pg_num)


def _copy_header_footer_refs(src_sect_pr, dst_sect_pr):
    for tag in ("w:headerReference", "w:footerReference"):
        for el in list(dst_sect_pr.findall(qn(tag))):
            dst_sect_pr.remove(el)
    insert_at = 0
    for tag in ("w:headerReference", "w:footerReference"):
        for ref in src_sect_pr.findall(qn(tag)):
            dst_sect_pr.insert(insert_at, deepcopy(ref))
            insert_at += 1


def _reinforce_toc_body_section_refs(doc):
    sect_prs = [el for el in doc.element.body.iter() if el.tag == qn("w:sectPr")]
    if len(sect_prs) < 4:
        return
    roman_src = sect_prs[2]
    toc_end = sect_prs[-2]
    body_end = sect_prs[-1]
    _copy_header_footer_refs(roman_src, toc_end)
    _set_sect_pr_page_numbering(toc_end, start=1, fmt="upperRoman")
    _set_sect_pr_page_numbering(body_end, start=1, fmt="decimal")


def _clear_header_footer(container):
    for p in container.paragraphs:
        _clear_runs(p)
        p_pr = p._p.pPr
        if p_pr is not None:
            for bdr in list(p_pr.findall(qn("w:pBdr"))):
                p_pr.remove(bdr)


def fix_headers_footers(docx_path, meta):
    from docx import Document

    doc = Document(docx_path)
    header_text = f"天津职业技术师范大学{meta.get('year', '2026')}届本科生毕业设计"
    for idx, section in enumerate(doc.sections):
        if idx == 0:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            _clear_header_footer(section.header)
            _clear_header_footer(section.footer)
            continue
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        hp = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        _clear_runs(hp)
        if idx >= 3:
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hr = hp.add_run(header_text)
            fmt_run(hr, 10.5)

        fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        if idx == 1:
            _clear_runs(fp)
        elif idx == 2:
            _add_page_field(fp, "ROMAN")
            _set_section_page_numbering(section, start=1, fmt="upperRoman")
        else:
            _add_page_field(fp)
            _set_section_page_numbering(section, start=1 if idx == 3 else None, fmt="decimal")
    _reinforce_toc_body_section_refs(doc)
    doc.save(docx_path)


def _replace_runs(paragraph, parts):
    _clear_runs(paragraph)
    for text, east_asia, ascii_font, size, bold in parts:
        run = paragraph.add_run(text)
        fmt_run_font(run, east_asia=east_asia, ascii_font=ascii_font, size=size, bold=bold)


def _remove_page_break_before(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    for el in list(p_pr.findall(qn("w:pageBreakBefore"))):
        p_pr.remove(el)


def fix_abstract_styles(docx_path, meta):
    from docx import Document

    doc = Document(docx_path)
    paragraphs = doc.paragraphs
    zh_idx = en_idx = toc_idx = None
    for i, p in enumerate(paragraphs):
        key = re.sub(r"\s+", "", p.text or "")
        if key == "摘要" and zh_idx is None:
            zh_idx = i
        elif key.upper() == "ABSTRACT" and en_idx is None:
            en_idx = i
        elif key == "目录" and toc_idx is None:
            toc_idx = i
    if zh_idx is None or en_idx is None:
        doc.save(docx_path)
        return

    # 中文摘要：标题黑体三号，正文宋体小四；英文、数字用 Times New Roman。
    p = paragraphs[zh_idx]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.page_break_before = True
    p.paragraph_format.line_spacing = 1.5
    _replace_runs(p, [("摘    要", "黑体", "Times New Roman", 16, True)])
    for p in paragraphs[zh_idx + 1:en_idx]:
        text_value = (p.text or "").strip()
        if not text_value:
            continue
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        if text_value.startswith("关键词"):
            p.paragraph_format.first_line_indent = None
            keywords = meta.get("keywords_zh", "")
            _replace_runs(p, [
                ("关键词：", "宋体", "Times New Roman", 12, True),
                (keywords, "宋体", "Times New Roman", 12, False),
            ])
        else:
            p.paragraph_format.first_line_indent = Cm(0.85)
            _replace_runs(p, [(text_value, "宋体", "Times New Roman", 12, False)])

    # 英文摘要：标题和正文均使用 Times New Roman。
    p = paragraphs[en_idx]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.page_break_before = True
    p.paragraph_format.line_spacing = 1.5
    _replace_runs(p, [("ABSTRACT", "Times New Roman", "Times New Roman", 16, True)])
    if toc_idx is not None:
        _remove_page_break_before(paragraphs[toc_idx])
    end = toc_idx if toc_idx is not None else len(paragraphs)
    for p in paragraphs[en_idx + 1:end]:
        text_value = (p.text or "").strip()
        if not text_value:
            continue
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        if text_value.startswith("Key Words"):
            p.paragraph_format.first_line_indent = None
            keywords = meta.get("keywords_en", "")
            _replace_runs(p, [
                ("Key Words: ", "Times New Roman", "Times New Roman", 12, True),
                (keywords, "Times New Roman", "Times New Roman", 12, False),
            ])
        else:
            p.paragraph_format.first_line_indent = Cm(0.85)
            _replace_runs(p, [(text_value, "Times New Roman", "Times New Roman", 12, False)])
    doc.save(docx_path)


def _remove_paragraph(paragraph):
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def _has_section_break(paragraph):
    return bool(paragraph._p.findall(".//" + qn("w:sectPr")))


def _make_section_pr(doc, fmt=None, start=None, section_type="continuous"):
    sect_pr = deepcopy(doc.sections[-1]._sectPr)
    for tag in ("w:headerReference", "w:footerReference", "w:pgNumType", "w:type"):
        for el in list(sect_pr.findall(qn(tag))):
            sect_pr.remove(el)
    if section_type:
        type_el = OxmlElement("w:type")
        type_el.set(qn("w:val"), section_type)
        sect_pr.insert(0, type_el)
    if fmt or start is not None:
        pg_num = OxmlElement("w:pgNumType")
        if start is not None:
            pg_num.set(qn("w:start"), str(start))
        if fmt:
            pg_num.set(qn("w:fmt"), fmt)
        sect_pr.append(pg_num)
    return sect_pr


def _set_paragraph_section(paragraph, sect_pr):
    p_pr = paragraph._p.get_or_add_pPr()
    old = p_pr.find(qn("w:sectPr"))
    if old is not None:
        p_pr.remove(old)
    p_pr.append(sect_pr)


def _element_text(el):
    return "".join(t.text or "" for t in el.findall(".//" + qn("w:t")))


def _ensure_toc_section_break(doc, toc_title_text="目    录"):
    body = doc.element.body
    children = list(body)
    toc_idx = None
    for i, el in enumerate(children):
        if el.tag == qn("w:p") and _element_text(el).strip().replace(" ", "") == toc_title_text.replace(" ", ""):
            toc_idx = i
            break
    if toc_idx is None:
        return
    body_start = None
    for el in children[toc_idx + 1:]:
        if el.tag == qn("w:p") and _element_text(el).strip().startswith("1 "):
            body_start = el
            break
    if body_start is None:
        return

    prev_el = body_start.getprevious()
    if (
        prev_el is not None
        and prev_el.tag == qn("w:p")
        and not _element_text(prev_el).strip()
        and bool(prev_el.findall(".//" + qn("w:sectPr")))
    ):
        p_el = prev_el
    else:
        p_el = OxmlElement("w:p")
        body_start.addprevious(p_el)
    p_pr = p_el.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        p_el.insert(0, p_pr)
    old = p_pr.find(qn("w:sectPr"))
    if old is not None:
        p_pr.remove(old)
    p_pr.append(_make_section_pr(doc, fmt="upperRoman", start=1, section_type="continuous"))


def _paragraph_has_page_break(p_el):
    return any(br.get(qn("w:type")) == "page" for br in p_el.findall(".//" + qn("w:br")))


def _add_page_break_to_paragraph(p_el):
    if _paragraph_has_page_break(p_el):
        return
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    p_pr = p_el.find(qn("w:pPr"))
    if p_pr is None:
        p_el.insert(0, run)
    else:
        p_pr.addnext(run)


def _fix_toc_page_boundary(doc, toc_title_text="目    录"):
    body = doc.element.body
    children = list(body)
    toc_el = None
    for el in children:
        if el.tag == qn("w:p") and _element_text(el).strip().replace(" ", "") == toc_title_text.replace(" ", ""):
            toc_el = el
            break
    if toc_el is None:
        return

    p_pr = toc_el.find(qn("w:pPr"))
    if p_pr is not None:
        for el in list(p_pr.findall(qn("w:pageBreakBefore"))):
            p_pr.remove(el)

    prev_el = toc_el.getprevious()
    if (
        prev_el is None
        or prev_el.tag != qn("w:p")
        or _element_text(prev_el).strip()
        or not prev_el.findall("./" + qn("w:pPr") + "/" + qn("w:sectPr"))
    ):
        prev_el = OxmlElement("w:p")
        toc_el.addprevious(prev_el)
        prev_pr = OxmlElement("w:pPr")
        prev_el.insert(0, prev_pr)
        prev_pr.append(_make_section_pr(doc, fmt="upperRoman", start=1, section_type="continuous"))
    _add_page_break_to_paragraph(prev_el)


def fix_front_matter(docx_path, meta):
    """Keep a single clean cover and remove the template's duplicate title page."""
    from docx import Document

    doc = Document(docx_path)
    paragraphs = list(doc.paragraphs)

    # The source template contains an outer cover plus an inner title page. For this
    # deliverable, keep the outer cover and remove the duplicate inner title page.
    inner_idx = statement_idx = None
    for i, p in enumerate(paragraphs):
        text_value = (p.text or "").strip()
        if text_value == "天津职业技术师范大学本科生毕业设计" and inner_idx is None:
            inner_idx = i
        elif text_value == "毕业论文（设计）诚信声明" and statement_idx is None:
            statement_idx = i
    if inner_idx is not None and statement_idx is not None and inner_idx < statement_idx:
        start_idx = inner_idx
        while start_idx > 0 and not (paragraphs[start_idx - 1].text or "").strip() and not _has_section_break(paragraphs[start_idx - 1]):
            start_idx -= 1
        for p in paragraphs[start_idx:statement_idx]:
            _remove_paragraph(p)

    paragraphs = list(doc.paragraphs)
    title_zh = str(meta.get("title_zh", "")).strip()
    title_en = str(meta.get("title_en", "")).strip()
    zh_title_idx = en_title_idx = major_idx = None
    for i, p in enumerate(paragraphs):
        text_value = (p.text or "").strip()
        if text_value == title_zh and zh_title_idx is None:
            zh_title_idx = i
        elif text_value == title_en and en_title_idx is None:
            en_title_idx = i
        elif text_value.startswith("专") and "业：" in text_value and major_idx is None:
            major_idx = i

    # Long Chinese/English titles can push the outer-cover metadata onto a second
    # page. Tighten only the outer cover area so it remains one page.
    if en_title_idx is not None:
        for r in paragraphs[en_title_idx].runs:
            if r.text:
                fmt_run_font(r, east_asia="Times New Roman", ascii_font="Times New Roman", size=13, bold=True)
    if zh_title_idx is not None:
        for r in paragraphs[zh_title_idx].runs:
            if r.text:
                fmt_run_font(r, east_asia="华文楷体", ascii_font="Times New Roman", size=15, bold=True)
    if en_title_idx is not None and major_idx is not None and en_title_idx < major_idx:
        blank_paras = [
            p for p in paragraphs[en_title_idx + 1:major_idx]
            if not (p.text or "").strip() and not _has_section_break(p)
        ]
        for p in blank_paras[1:]:
            _remove_paragraph(p)

    paragraphs = list(doc.paragraphs)
    statement_para = next((p for p in paragraphs if (p.text or "").strip() == "毕业论文（设计）诚信声明"), None)
    if statement_para is not None:
        statement_para.paragraph_format.page_break_before = True

    for p in list(doc.paragraphs[:25]):
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    paragraphs = list(doc.paragraphs)
    en_key_para = None
    for p in paragraphs:
        text_value = (p.text or "").strip()
        if text_value.startswith("Key Words"):
            en_key_para = p
            break
    if en_key_para is not None:
        _set_paragraph_section(en_key_para, _make_section_pr(doc, section_type="continuous"))
    _ensure_toc_section_break(doc)
    _fix_toc_page_boundary(doc)

    doc.save(docx_path)


def build_data(doc):
    meta = load_json("meta.json") or {}
    chapters = []
    for i in range(1, 8):
        ch = load_json(f"ch{i}.json")
        if ch:
            print(f"  ch{i}.json -> {ch['title']}")
            for block in ch.get("content", []):
                pass
            resolve_images(ch.get("content", []))
            for sec in ch.get("sections", []):
                resolve_images(sec.get("content", []))
                for sub in sec.get("subsections", []):
                    resolve_images(sub.get("content", []))
            chapters.append(ch)

    zh = split_paras(meta.get("abstract_zh"), 3)
    en = split_paras(meta.get("abstract_en"), 3)
    for i, text in enumerate(zh, 1):
        meta[f"abstract_zh_{i}"] = text
    for i, text in enumerate(en, 1):
        meta[f"abstract_en_{i}"] = text

    refs = load_json("references.json") or []
    if isinstance(refs, dict):
        refs = refs.get("references", [])

    meta["body"] = build_body(doc, chapters)
    meta["references_doc"] = build_refs(doc, refs)
    return meta


def verify_docx(docx_path):
    from docx import Document
    doc = Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    errors = []
    for tag in ("{{", "}}", "{%", "%}"):
        if tag in text:
            errors.append(f"残留模板标记 {tag}")
    for stale in ("购物商城", "智慧停车", "毕业设计的结构", "摘要应具有独立性"):
        if stale in text:
            errors.append(f"残留模板/旧项目文本: {stale}")
    if "智能出行交通数据可视化分析系统" not in text and "交通数据可视化分析系统" not in text:
        errors.append("未检测到交通系统主题文本")
    return errors


def main():
    output = Path(sys.argv[1]) if len(sys.argv) > 1 else HERE / "基于HadoopSparkHive的智能出行交通数据可视化分析系统毕业设计.docx"
    doc = DocxTemplate(TPL)
    print("组装论文数据...")
    data = build_data(doc)
    print("渲染模板...")
    doc.render(data)
    doc.save(output)
    fix_front_matter(output, data)
    fix_abstract_styles(output, data)
    fix_headers_footers(output, data)
    patch_update_fields(output)
    errors = verify_docx(output)
    if errors:
        print("校验提醒:")
        for e in errors:
            print(" -", e)
    else:
        print("校验通过")
    print(f"完成: {output}")


if __name__ == "__main__":
    main()
