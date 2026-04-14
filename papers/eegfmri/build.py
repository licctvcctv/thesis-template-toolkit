"""
智学云课论文组装器。
支持文本、图片（InlineImage）、表格。

用法: python build.py <模板.docx> [输出.docx]
"""
import os
import sys
import json
import glob

ROOT = os.path.join(os.path.dirname(__file__), "../..")
sys.path.insert(0, ROOT)

HERE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(HERE, "images")


def load_json(filename):
    path = os.path.join(HERE, filename)
    if not os.path.exists(path):
        return None
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def process_content(content_list, doc):
    """
    处理内容列表：
    - 字符串 → 保持不变
    - {"type":"image"} → InlineImage（图片）
    - {"type":"table"} → 转为特殊标记文本（后处理时插入表格）
    """
    from docxtpl import InlineImage
    from docx.shared import Mm

    result = []
    for item in content_list:
        if isinstance(item, str):
            result.append(item)
        elif isinstance(item, dict):
            if item.get("type") == "image":
                img_path = os.path.join(IMG_DIR, item["path"])
                if os.path.exists(img_path):
                    width = Mm(item.get("width", 120))
                    result.append(InlineImage(doc, img_path,
                                             width=width))
                    # 不跳过标注——让标注文本正常输出
                else:
                    result.append(f"[图片缺失: {item['path']}]")
            elif item.get("type") == "table":
                _pending_tables.append(item)
                tid = len(_pending_tables) - 1
                # 只放占位标记，不重复生成标注
                # （标注已在 JSON 内容中作为前一个字符串存在）
                result.append(f"__TABLE_PLACEHOLDER_{tid}__")
            else:
                result.append(str(item))
    return result


# 全局表格暂存
_pending_tables = []


def process_chapters(chapters, doc):
    """处理所有章节，转换图片引用"""
    for ch in chapters:
        for sec in ch.get("sections", []):
            sec["content"] = process_content(
                sec.get("content", []), doc)
            for sub in sec.get("subsections", []):
                sub["content"] = process_content(
                    sub.get("content", []), doc)
    return chapters


def build_data(doc=None):
    """组装所有章节数据"""
    meta = load_json("meta.json")
    if not meta:
        raise FileNotFoundError("meta.json not found")

    chapters = []
    for ch_file in sorted(glob.glob(
            os.path.join(HERE, "ch*.json"))):
        ch = load_json(os.path.basename(ch_file))
        if ch:
            chapters.append(ch)
            print(f"  {os.path.basename(ch_file)}"
                  f" -> {ch['title']}")

    if doc:
        chapters = process_chapters(chapters, doc)

    refs = load_json("references.json")
    if isinstance(refs, list):
        references = refs
    elif isinstance(refs, dict):
        references = refs.get("references", [])
    else:
        references = []

    # 给参考文献加 [n] 编号
    references = [f"[{i+1}] {r}" if not r.startswith("[") else r
                  for i, r in enumerate(references)]

    return {**meta, "chapters": chapters,
            "references": references}


def main():
    if len(sys.argv) < 2:
        print("用法: python build.py <模板.docx> [输出.docx]")
        sys.exit(1)

    template = sys.argv[1]
    output = sys.argv[2] if len(sys.argv) > 2 else "output.docx"

    from docxtpl import DocxTemplate
    doc = DocxTemplate(template)

    _pending_tables.clear()
    print("组装论文数据...")
    data = build_data(doc)
    print(f"  {len(data['chapters'])} 章, "
          f"{len(data.get('references', []))} 条参考文献")

    print(f"渲染: {template} -> {output}")
    doc.render(data)
    doc.save(output)

    # 后处理：修正图片居中和图标注格式
    _post_process(output)
    # 校验
    errors = _verify(output)
    if errors:
        print(f"  ⚠ 发现 {len(errors)} 个问题:")
        for e in errors:
            print(f"    {e}")
    else:
        print(f"  ✓ 校验通过")
    print(f"完成: {output}")


def _verify(docx_path):
    """校验输出文档：检查每个表/图标注是否有对应内容"""
    from docx import Document
    import re

    doc = Document(docx_path)
    errors = []
    body = doc.element.body
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # 按 DOM 顺序遍历，记录每个元素的类型
    elements = []  # [(type, text)]
    for elem in body:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            texts = [t.text for t in elem.findall('.//w:t', ns) if t.text]
            text = ''.join(texts).strip()
            has_img = bool(elem.findall('.//w:drawing', ns))
            if has_img:
                elements.append(('img', text))
            else:
                elements.append(('p', text))
        elif tag == 'tbl':
            elements.append(('tbl', ''))

    # 检查每个表标注后面是否有表格
    for i, (etype, text) in enumerate(elements):
        if etype == 'p' and re.match(r'^表\d+-\d+\s', text) and not re.search(r'[。，；]', text) and len(text) < 30:
            # 往后找 3 个元素内是否有表格
            found_tbl = False
            for j in range(i + 1, min(i + 6, len(elements))):
                if elements[j][0] == 'tbl':
                    found_tbl = True
                    break
            if not found_tbl:
                errors.append(f"缺表格: {text}")

    # 检查每个图标注附近是否有图片
    for i, (etype, text) in enumerate(elements):
        if etype == 'p' and re.match(r'^图\d+-\d+\s', text) and len(text) < 50:
            found_img = False
            # 往前找 3 个元素
            for j in range(max(0, i - 3), i):
                if elements[j][0] == 'img':
                    found_img = True
                    break
            if not found_img:
                errors.append(f"缺图片: {text}")

    return errors


def _post_process(docx_path):
    """渲染后：图片居中、标注格式、插入表格"""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

    doc = Document(docx_path)
    fig_pat = re.compile(r'^图\d')
    tbl_cap_pat = re.compile(r'^表\d')
    placeholder_pat = re.compile(r'^__TABLE_PLACEHOLDER_(\d+)__$')

    for p in list(doc.paragraphs):
        t = (p.text or "").strip()

        # 图片段落 → 居中 + 单倍行距（防止固定行距裁剪图片）
        has_drawing = bool(p._p.findall(
            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
        if has_drawing:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = 0
            # 设置 lineRule="auto" 防止图片被裁剪为一行
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn as _qn
            pPr = p._p.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
            if pPr is not None:
                spacing = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing')
                if spacing is None:
                    spacing = OxmlElement('w:spacing')
                    pPr.append(spacing)
                spacing.set(_qn('w:line'), '240')
                spacing.set(_qn('w:lineRule'), 'auto')

        # 图/表标注 → 居中 + 五号（短文本）
        if (fig_pat.match(t) or tbl_cap_pat.match(t)) and len(t) < 50:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(10.5)

        # 表格占位符 → 插入真实表格
        m = placeholder_pat.match(t)
        if m:
            tid = int(m.group(1))
            if tid < len(_pending_tables):
                _insert_table(doc, p, _pending_tables[tid])
                # 清空占位文本
                for r in p.runs:
                    r.text = ""

    # 模板末尾存在静态“附录”占位时，若未实际生成附录内容则移除。
    last_nonempty = None
    for p in reversed(doc.paragraphs):
        if (p.text or "").strip():
            last_nonempty = p
            break
    if last_nonempty is not None and (last_nonempty.text or "").strip() == "附录":
        last_nonempty._element.getparent().remove(
            last_nonempty._element
        )

    doc.save(docx_path)
    print(f"  后处理: 图片居中+{len(_pending_tables)}个表格")


def _insert_table(doc, after_para, tbl_data):
    """在段落后插入三线表"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    headers = tbl_data.get("headers", [])
    rows = tbl_data.get("rows", [])
    if not headers:
        return

    ncols = len(headers)
    all_rows = [headers] + rows

    tbl = OxmlElement('w:tbl')

    # 表格属性：居中 + 自动列宽
    tblPr = OxmlElement('w:tblPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblPr.append(jc)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'autofit')
    tblPr.append(tblLayout)
    # 三线表边框：顶线1.5pt、底线1.5pt，无左右竖线
    borders = OxmlElement('w:tblBorders')
    for edge, sz in [('top', '12'), ('bottom', '12')]:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    for edge in ['left', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        borders.append(el)
    tblPr.append(borders)
    tbl.append(tblPr)

    # 智能列宽
    total_width = 9000
    col_max_len = [0] * ncols
    for row_data in all_rows:
        for ci in range(ncols):
            cell_text = str(row_data[ci]) if ci < len(row_data) else ""
            w = sum(2 if ord(c) > 127 else 1 for c in cell_text)
            col_max_len[ci] = max(col_max_len[ci], w)
    min_w = 1200
    total_len = sum(max(l, 4) for l in col_max_len) or 1
    col_widths = [max(min_w, int(total_width * max(l, 4) / total_len))
                  for l in col_max_len]
    diff = total_width - sum(col_widths)
    widest = col_widths.index(max(col_widths))
    col_widths[widest] += diff

    tblGrid = OxmlElement('w:tblGrid')
    for w in col_widths:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(w))
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    # 行
    for ri, row_data in enumerate(all_rows):
        tr = OxmlElement('w:tr')
        for ci in range(ncols):
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(col_widths[ci] if ci < len(col_widths) else 1800))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)
            # 三线表：表头行底部加细线(0.75pt)
            if ri == 0:
                tcBorders = OxmlElement('w:tcBorders')
                btm = OxmlElement('w:bottom')
                btm.set(qn('w:val'), 'single')
                btm.set(qn('w:sz'), '6')
                btm.set(qn('w:space'), '0')
                btm.set(qn('w:color'), '000000')
                tcBorders.append(btm)
                tcPr.append(tcBorders)
            tc.append(tcPr)
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            pJc = OxmlElement('w:jc')
            pJc.set(qn('w:val'), 'center')
            pPr.append(pJc)
            pSpacing = OxmlElement('w:spacing')
            pSpacing.set(qn('w:before'), '0')
            pSpacing.set(qn('w:after'), '0')
            pSpacing.set(qn('w:line'), '240')
            pSpacing.set(qn('w:lineRule'), 'auto')
            pPr.append(pSpacing)
            pInd = OxmlElement('w:ind')
            pInd.set(qn('w:firstLine'), '0')
            pInd.set(qn('w:firstLineChars'), '0')
            pPr.append(pInd)
            p.append(pPr)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), '宋体')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '21')
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '21')
            rPr.append(szCs)
            if ri == 0:
                b = OxmlElement('w:b')
                rPr.append(b)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = str(row_data[ci]) if ci < len(row_data) else ''
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    # 插入到段落后面
    after_para._p.addnext(tbl)


if __name__ == "__main__":
    main()
