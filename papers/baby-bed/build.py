"""
基于STM32的智能婴儿床论文组装器（景德镇学院模板）。

用法:
    cd thesis_project
    python papers/baby-bed/build.py templates/jdzxy/template.docx [输出.docx]
"""
from __future__ import annotations

import json
import os
import re
import sys
import zipfile
import tempfile
import shutil
from pathlib import Path


ROOT = Path(__file__).resolve().parents[2]
HERE = Path(__file__).resolve().parent
IMG_DIR = HERE / "images"
WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


_pending_tables = []


def load_json(name):
    path = HERE / name
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def load_appendix():
    appendix = load_json("appendix.json")
    if not appendix:
        return None
    if not appendix.get("sections"):
        return None
    return appendix


def process_content(items, doc):
    from docxtpl import InlineImage
    from docx.shared import Mm
    from PIL import Image

    result = []
    for item in items:
        if isinstance(item, str):
            result.append(item)
        elif isinstance(item, dict) and item.get("type") == "image":
            img = IMG_DIR / item["path"]
            if img.exists():
                width_mm = fit_image_width_mm(img, item.get("width", 115))
                result.append(InlineImage(doc, str(img), width=Mm(width_mm)))
            else:
                result.append(f"[图片缺失: {item['path']}]")
            if item.get("caption"):
                result.append(item["caption"])
        elif isinstance(item, dict) and item.get("type") == "table":
            _pending_tables.append(item)
            idx = len(_pending_tables) - 1
            if item.get("caption"):
                result.append(item["caption"])
            result.append(f"__TABLE_PLACEHOLDER_{idx}__")
        elif isinstance(item, dict) and item.get("text"):
            result.append(item["text"])
    return result


def fit_image_width_mm(img_path, requested_width):
    from PIL import Image

    max_width = 105
    max_height = 112
    width = min(float(requested_width), max_width)
    try:
        with Image.open(img_path) as im:
            px_w, px_h = im.size
    except Exception:
        return width
    if not px_w or not px_h:
        return width
    height = width * px_h / px_w
    if height > max_height:
        width = max_height * px_w / px_h
    return max(45, round(width, 1))


def load_data(doc):
    meta = load_json("meta.json")
    if not meta:
        raise FileNotFoundError("meta.json missing")

    chapters = []
    for idx in range(1, 7):
        ch = load_json(f"ch{idx}.json")
        if not ch:
            continue
        for sec in ch.get("sections", []):
            sec["content"] = process_content(sec.get("content", []), doc)
            for sub in sec.get("subsections", []):
                sub["content"] = process_content(sub.get("content", []), doc)
        chapters.append(ch)

    refs = load_json("references.json") or []
    if isinstance(refs, dict):
        refs = refs.get("references", [])
    references = [{"idx": i, "text": ref} for i, ref in enumerate(refs, 1)]
    return {**meta, "chapters": chapters, "references": references}


def make_toc_paragraph(title, page, level, begin=False, end=False):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    style_map = {1: "9", 2: "10", 3: "11"}
    p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    pstyle = OxmlElement("w:pStyle")
    pstyle.set(qn("w:val"), style_map.get(level, "10"))
    ppr.append(pstyle)
    p.append(ppr)
    if begin:
        r = OxmlElement("w:r")
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        fld.set(qn("w:dirty"), "true")
        r.append(fld)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = 'TOC \\o "1-3" \\h \\u'
        r.append(instr)
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        r.append(sep)
        p.append(r)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = title
    r.append(t)
    r.append(OxmlElement("w:tab"))
    t = OxmlElement("w:t")
    t.text = str(page)
    r.append(t)
    p.append(r)
    if end:
        r = OxmlElement("w:r")
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "end")
        r.append(fld)
        p.append(r)
    return p


def insert_table_after(paragraph, table_data):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    headers = table_data.get("headers", [])
    rows = table_data.get("rows", [])
    if not headers:
        return
    ncols = len(headers)
    all_rows = [headers] + rows

    tbl = OxmlElement("w:tbl")
    tbl_pr = OxmlElement("w:tblPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tbl_pr.append(jc)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")
    tbl_pr.append(tbl_w)
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "autofit")
    tbl_pr.append(tbl_layout)
    borders = OxmlElement("w:tblBorders")
    for edge, val, sz in [
        ("top", "single", "12"),
        ("bottom", "single", "12"),
        ("left", "nil", "0"),
        ("right", "nil", "0"),
        ("insideH", "nil", "0"),
        ("insideV", "nil", "0"),
    ]:
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), val)
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tbl_pr.append(borders)
    tbl.append(tbl_pr)

    total_width = 9000
    col_max = [0] * ncols
    for row_data in all_rows:
        for ci in range(ncols):
            cell_text = str(row_data[ci]) if ci < len(row_data) else ""
            visual_len = sum(2 if ord(ch) > 127 else 1 for ch in cell_text)
            col_max[ci] = max(col_max[ci], visual_len, 4)
    total_len = sum(col_max) or 1
    col_widths = [max(1200, int(total_width * v / total_len)) for v in col_max]
    diff = total_width - sum(col_widths)
    if col_widths:
        col_widths[col_widths.index(max(col_widths))] += diff

    grid = OxmlElement("w:tblGrid")
    for width in col_widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.append(grid)

    for ri, row_data in enumerate(all_rows):
        tr = OxmlElement("w:tr")
        for ci in range(ncols):
            cell_text = row_data[ci] if ci < len(row_data) else ""
            tc = OxmlElement("w:tc")
            tc_pr = OxmlElement("w:tcPr")
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), str(col_widths[ci]))
            tc_w.set(qn("w:type"), "dxa")
            tc_pr.append(tc_w)
            v_align = OxmlElement("w:vAlign")
            v_align.set(qn("w:val"), "center")
            tc_pr.append(v_align)
            if ri == 0:
                tc_borders = OxmlElement("w:tcBorders")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "6")
                bottom.set(qn("w:space"), "0")
                bottom.set(qn("w:color"), "000000")
                tc_borders.append(bottom)
                tc_pr.append(tc_borders)
            tc.append(tc_pr)
            p = OxmlElement("w:p")
            p_pr = OxmlElement("w:pPr")
            p_jc = OxmlElement("w:jc")
            p_jc.set(qn("w:val"), "center")
            p_pr.append(p_jc)
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:before"), "0")
            spacing.set(qn("w:after"), "0")
            spacing.set(qn("w:line"), "240")
            spacing.set(qn("w:lineRule"), "auto")
            p_pr.append(spacing)
            ind = OxmlElement("w:ind")
            ind.set(qn("w:firstLine"), "0")
            ind.set(qn("w:firstLineChars"), "0")
            ind.set(qn("w:left"), "0")
            ind.set(qn("w:right"), "0")
            p_pr.append(ind)
            p.append(p_pr)
            r = OxmlElement("w:r")
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:eastAsia"), "宋体")
            fonts.set(qn("w:ascii"), "Times New Roman")
            fonts.set(qn("w:hAnsi"), "Times New Roman")
            r_pr.append(fonts)
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "21")
            r_pr.append(sz)
            if ri == 0:
                bold = OxmlElement("w:b")
                r_pr.append(bold)
            r.append(r_pr)
            t = OxmlElement("w:t")
            t.text = str(cell_text)
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)
    paragraph._p.addnext(tbl)


def post_process(docx_path, toc_entries):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document(docx_path)

    for p in list(doc.paragraphs):
        text = (p.text or "").strip()
        if text == "__TOC_PLACEHOLDER__":
            parent = p._p.getparent()
            idx = parent.index(p._p)
            parent.remove(p._p)
            for n, (title, page, level) in enumerate(toc_entries):
                parent.insert(idx + n, make_toc_paragraph(
                    title, page, int(level), begin=n == 0,
                    end=n == len(toc_entries) - 1
                ))
            continue

        m = re.fullmatch(r"__TABLE_PLACEHOLDER_(\d+)__", text)
        if m:
            tid = int(m.group(1))
            if tid < len(_pending_tables):
                insert_table_after(p, _pending_tables[tid])
            for r in p.runs:
                r.text = ""
            continue

        if re.match(r"^图\d+\.\d+", text) or re.match(r"^表\d+\.\d+", text):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.page_break_before = False
            for r in p.runs:
                r.font.size = Pt(10.5)
                r.font.name = "黑体"
                if r._r.rPr is not None:
                    r._r.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

        if p._p.findall(f".//{{{WNS}}}drawing"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.page_break_before = False
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.line_spacing = 1.0
            p_pr = p._p.pPr
            if p_pr is None:
                p_pr = OxmlElement("w:pPr")
                p._p.insert(0, p_pr)
            old_rpr = p_pr.find(qn("w:rPr"))
            if old_rpr is not None:
                p_pr.remove(old_rpr)
            old_ind = p_pr.find(qn("w:ind"))
            if old_ind is not None:
                p_pr.remove(old_ind)
            ind = OxmlElement("w:ind")
            ind.set(qn("w:firstLine"), "0")
            ind.set(qn("w:left"), "0")
            p_pr.append(ind)

        if re.match(r"^[1-6]\s+", text) or text == "参考文献":
            p.paragraph_format.page_break_before = True

        if text.replace(" ", "").replace("　", "") == "致谢":
            p._p.getparent().remove(p._p)

    doc.save(docx_path)
    format_references(docx_path)
    insert_appendix(docx_path, load_appendix())
    remove_disabled_page_breaks(docx_path)
    add_update_fields(docx_path)
    remove_ref_fields(docx_path)


def insert_appendix(docx_path, appendix):
    """Append source-code listings after references."""
    if not appendix:
        return

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from docx.oxml.ns import qn

    doc = Document(docx_path)
    title = appendix.get("title", "附件 程序代码")

    heading = doc.add_paragraph(title, style="Heading 1")
    heading.paragraph_format.page_break_before = True
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = appendix.get("intro", "")
    if intro:
        p = doc.add_paragraph(intro, style=doc.styles["Normal"])
        p.paragraph_format.first_line_indent = Pt(21)

    for idx, section in enumerate(appendix.get("sections", []), 1):
        sub_title = section.get("title") or f"程序清单{idx}"
        p = doc.add_paragraph(sub_title)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        for r in p.runs:
            r.font.size = Pt(10.5)
            r.font.name = "黑体"
            if r._r.rPr is not None:
                r._r.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

        note = section.get("note", "")
        if note:
            p = doc.add_paragraph(note, style=doc.styles["Normal"])
            p.paragraph_format.first_line_indent = Pt(21)

        for line in section.get("code", "").splitlines():
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(14)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(line if line else " ")
            r.font.size = Pt(7.5)
            r.font.name = "Courier New"
            if r._r.rPr is not None:
                r._r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.save(docx_path)


def remove_disabled_page_breaks(docx_path):
    from lxml import etree

    tmp = tempfile.mktemp(suffix=".docx")
    ns = {"w": WNS}
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                root = etree.fromstring(data)
                for pbr in list(root.xpath(".//w:pageBreakBefore[@w:val='0' or @w:val='false']", namespaces=ns)):
                    pbr.getparent().remove(pbr)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(item, data)
    shutil.move(tmp, docx_path)


def format_references(docx_path):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn

    doc = Document(docx_path)
    in_refs = False
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text == "参考文献":
            in_refs = True
            p.paragraph_format.page_break_before = True
            continue
        if in_refs and text.replace(" ", "").replace("　", "") in {"致谢", "附件", "附录"}:
            in_refs = False
            continue
        if in_refs and re.match(r"^\[\d+\]", text):
            p.style = doc.styles["Normal"]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = p.paragraph_format
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(0)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.page_break_before = False
            pf.line_spacing = 1.5
            for r in p.runs:
                r.font.size = Pt(10.5)
                r.font.name = "Times New Roman"
                if r._r.rPr is not None:
                    r._r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.save(docx_path)


def add_update_fields(docx_path):
    tmp = tempfile.mktemp(suffix=".docx")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml":
                from lxml import etree
                root = etree.fromstring(data)
                if not root.xpath("./w:updateFields", namespaces={"w": WNS}):
                    node = etree.SubElement(root, f"{{{WNS}}}updateFields")
                    node.set(f"{{{WNS}}}val", "true")
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(item, data)
    shutil.move(tmp, docx_path)


def remove_ref_fields(docx_path):
    """Delete broken REF fields inherited from the source template."""
    from lxml import etree

    tmp = tempfile.mktemp(suffix=".docx")
    ns = {"w": WNS}
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                root = etree.fromstring(data)
                for p in root.xpath(".//w:p", namespaces=ns):
                    runs = list(p.xpath("./w:r", namespaces=ns))
                    i = 0
                    while i < len(runs):
                        run = runs[i]
                        if not run.xpath(".//w:fldChar[@w:fldCharType='begin']", namespaces=ns):
                            i += 1
                            continue
                        field_runs = [run]
                        is_ref = bool(run.xpath(".//w:instrText[contains(., 'REF')]", namespaces=ns))
                        j = i + 1
                        ended = bool(run.xpath(".//w:fldChar[@w:fldCharType='end']", namespaces=ns))
                        while j < len(runs) and not ended:
                            field_runs.append(runs[j])
                            if runs[j].xpath(".//w:instrText[contains(., 'REF')]", namespaces=ns):
                                is_ref = True
                            if runs[j].xpath(".//w:fldChar[@w:fldCharType='end']", namespaces=ns):
                                ended = True
                            j += 1
                        if is_ref:
                            for fr in field_runs:
                                if fr.getparent() is not None:
                                    fr.getparent().remove(fr)
                            runs = list(p.xpath("./w:r", namespaces=ns))
                            continue
                        i = j
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(item, data)
    shutil.move(tmp, docx_path)


def verify(docx_path):
    from docx import Document
    doc = Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    errors = []
    for token in ["{{", "{%", "__TABLE_PLACEHOLDER", "__TOC_PLACEHOLDER"]:
        if token in text:
            errors.append(f"残留模板标记: {token}")
    for token in ["指纹", "打卡", "AS608", "DS1307", "MAX30102", "HC-05", "VS1053", "白噪音", "翻身", "心率"]:
        if token in text:
            errors.append(f"非本项目功能残留: {token}")
    return errors


def main():
    if len(sys.argv) < 2:
        print("用法: python build.py <模板.docx> [输出.docx]")
        sys.exit(1)
    template = sys.argv[1]
    output = sys.argv[2] if len(sys.argv) > 2 else str(HERE / "output.docx")

    from docxtpl import DocxTemplate

    doc = DocxTemplate(template)
    _pending_tables.clear()
    data = load_data(doc)
    print(f"组装数据: {len(data['chapters'])}章, {len(data['references'])}条参考文献, {len(_pending_tables)}个表格")
    doc.render(data)
    doc.save(output)
    post_process(output, data.get("toc_entries", []))
    errors = verify(output)
    if errors:
        print(f"发现 {len(errors)} 个问题:")
        for e in errors:
            print(f"  - {e}")
    else:
        print("校验通过")
    print(f"完成: {output}")


if __name__ == "__main__":
    main()
