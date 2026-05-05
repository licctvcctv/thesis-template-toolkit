"""
校园导航系统论文组装器（哈尔滨商业大学模板）。
支持文本、图片（InlineImage）、表格（三线表）、代码块。

用法: python build.py [输出.docx]
"""
import os
import sys
import json
import re

ROOT = os.path.join(os.path.dirname(__file__), "../..")
sys.path.insert(0, ROOT)

HERE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(HERE, "images")
TPL = os.path.join(ROOT, "templates/hrbcu/template.docx")

_pending_tables = []
_pending_code_blocks = []


def load_json(filename):
    path = os.path.join(HERE, filename)
    if not os.path.exists(path):
        return None
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def process_content(content_list, doc):
    from docxtpl import InlineImage
    from docx.shared import Mm
    result = []
    for item in content_list:
        if isinstance(item, str):
            result.append(item)
        elif isinstance(item, dict):
            if item.get("type") == "image":
                img_path = os.path.join(IMG_DIR, item["path"])
                if os.path.exists(img_path):
                    width = Mm(item.get("width", 148))
                    result.append(InlineImage(doc, img_path, width=width))
                    if item.get("caption"):
                        result.append(item["caption"])
                else:
                    result.append(f"[图片缺失: {item['path']}]")
                    if item.get("caption"):
                        result.append(item["caption"])
            elif item.get("type") == "code":
                _pending_code_blocks.append(item)
                cid = len(_pending_code_blocks) - 1
                result.append(f"__CODE_PLACEHOLDER_{cid}__")
            elif item.get("type") == "table":
                _pending_tables.append(item)
                tid = len(_pending_tables) - 1
                if item.get("caption"):
                    result.append(item["caption"])
                result.append(f"__TABLE_PLACEHOLDER_{tid}__")
            else:
                result.append(str(item))
    return result


def process_chapters(chapters, doc):
    for ch in chapters:
        ch["content"] = process_content(ch.get("content", []), doc)
        for sec in ch.get("sections", []):
            sec["content"] = process_content(sec.get("content", []), doc)
            for sub in sec.get("subsections", []):
                sub["content"] = process_content(sub.get("content", []), doc)
    return chapters


def build_data(doc=None):
    meta = load_json("meta.json") or {}
    chapters = []
    for i in range(1, 6):
        ch = load_json(f"ch{i}.json")
        if ch:
            print(f"  ch{i}.json -> {ch['title']}")
            chapters.append(ch)

    if doc:
        chapters = process_chapters(chapters, doc)

    refs = load_json("references.json") or []
    refs = [f"[{i+1}] {ref}" if not ref.startswith("[") else ref
            for i, ref in enumerate(refs)]

    # 摘要拆段
    for key, list_key in [("abstract_zh", "abstract_zh_list"),
                          ("abstract_en", "abstract_en_list")]:
        text = meta.get(key, "")
        if isinstance(text, list):
            meta[list_key] = text
        else:
            meta[list_key] = [p.strip() for p in text.split("\n") if p.strip()]

    # 结论拆段
    conclusion = meta.get("conclusion_list", [])
    if isinstance(conclusion, str):
        meta["conclusion_list"] = [p.strip() for p in conclusion.split("\n") if p.strip()]

    return {**meta, "chapters": chapters, "references": refs}


def main():
    output = sys.argv[1] if len(sys.argv) > 1 else None

    from docxtpl import DocxTemplate
    doc = DocxTemplate(TPL)

    _pending_tables.clear()
    _pending_code_blocks.clear()
    print("组装论文数据...")
    data = build_data(doc)
    print(f"  {len(data['chapters'])} 章, "
          f"{len(data.get('references', []))} 条参考文献")

    if not output:
        from datetime import datetime
        ts = datetime.now().strftime("%m%d_%H%M")
        output = os.path.join(HERE, f"校园导航系统_{ts}.docx")

    print(f"渲染: {output}")
    doc.render(data)
    doc.save(output)

    _post_process(output)
    errors = _verify(output)
    if errors:
        print(f"  ⚠ {len(errors)} 个问题:")
        for e in errors:
            print(f"    - {e}")
    else:
        print("  ✓ 校验通过")
    print(f"完成: {output}")


def _post_process(docx_path):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    doc = Document(docx_path)
    tbl_cap_pat = re.compile(r'^表\d')
    fig_cap_pat = re.compile(r'^图\d')
    placeholder_pat = re.compile(r'^__TABLE_PLACEHOLDER_(\d+)__$')
    code_placeholder_pat = re.compile(r'^__CODE_PLACEHOLDER_(\d+)__$')

    for p in list(doc.paragraphs):
        t = (p.text or "").strip()

        # 代码块占位符
        cm = code_placeholder_pat.match(t)
        if cm:
            cid = int(cm.group(1))
            if cid < len(_pending_code_blocks):
                _insert_code_block(doc, p, _pending_code_blocks[cid])
                for r in p.runs:
                    r.text = ""

        # 图片居中
        has_drawing = bool(p._p.findall(
            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
        if has_drawing:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            pPr = p._p.pPr
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p._p.insert(0, pPr)
            old_ind = pPr.find(_qn('w:ind'))
            if old_ind is not None:
                pPr.remove(old_ind)
            new_ind = OxmlElement('w:ind')
            new_ind.set(_qn('w:firstLine'), '0')
            new_ind.set(_qn('w:firstLineChars'), '0')
            pPr.append(new_ind)
            from docx.enum.text import WD_LINE_SPACING
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 图表标注居中
        if (fig_cap_pat.match(t) or tbl_cap_pat.match(t)) and len(t) < 60:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                _set_run_font(r, 10.5)

        # 表格占位符
        m = placeholder_pat.match(t)
        if m:
            tid = int(m.group(1))
            if tid < len(_pending_tables):
                _insert_table(doc, p, _pending_tables[tid])
                _remove_paragraph(p)
                continue

    # ---- 清理空白页：合并连续SECT段，删除多余空段 ----
    _NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    paras = doc.paragraphs
    to_remove = []
    removed_blank = 0

    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        has_sect = p._p.find('.//w:sectPr', _NS) is not None

        if not text and has_sect:
            # 如果下一段也是空SECT段，删掉当前这个（保留最后一个SECT）
            if i + 1 < len(paras):
                next_text = (paras[i+1].text or "").strip()
                next_sect = paras[i+1]._p.find('.//w:sectPr', _NS) is not None
                if not next_text and next_sect:
                    to_remove.append(p._p)
                    removed_blank += 1
                    continue

        # 空段（无SECT）夹在SECT段和Heading之间 → 删
        if not text and not has_sect:
            if i > 0 and i + 1 < len(paras):
                prev_sect = paras[i-1]._p.find('.//w:sectPr', _NS) is not None
                next_style = paras[i+1].style.name if paras[i+1].style else ""
                if prev_sect and "Heading" in next_style:
                    to_remove.append(p._p)
                    removed_blank += 1

    for elem in to_remove:
        if elem.getparent() is not None:
            elem.getparent().remove(elem)

    _replace_static_toc(doc)
    _keep_figures_with_captions(doc, fig_cap_pat)
    _format_references(doc)
    _superscript_citations(doc)

    doc.save(docx_path)
    print(f"  后处理: {len(_pending_tables)} 个表格, {len(_pending_code_blocks)} 个代码块, 清理{removed_blank}个空白段")


def _remove_paragraph(p):
    parent = p._p.getparent()
    if parent is not None:
        parent.remove(p._p)


def _set_run_font(run, size_pt=None, bold=None):
    from docx.shared import Pt
    from docx.oxml.ns import qn

    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run._element.rPr.rFonts.set(qn('w:cs'), 'Times New Roman')
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def _append_rpr_fonts(rPr):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rPr.append(rFonts)


def _keep_figures_with_captions(doc, fig_cap_pat):
    from docx.shared import Pt

    drawing_xpath = './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
    paragraphs = doc.paragraphs
    for idx, p in enumerate(paragraphs[:-1]):
        has_drawing = bool(p._p.findall(drawing_xpath))
        if not has_drawing:
            continue
        caption = (paragraphs[idx + 1].text or "").strip()
        if fig_cap_pat.match(caption):
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            paragraphs[idx + 1].paragraph_format.keep_together = True
            paragraphs[idx + 1].paragraph_format.space_before = Pt(0)


def _replace_static_toc(doc):
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_LINE_SPACING
    from docx.shared import Cm, Pt

    toc_entries = [
        (1, "摘  要", "I"),
        (1, "Abstract", "II"),
        (1, "1 绪论", "1"),
        (2, "1.1 课题背景与意义", "1"),
        (2, "1.2 国内外研究现状", "1"),
        (2, "1.3 研究目标", "2"),
        (2, "1.4 论文组织结构", "3"),
        (1, "2 系统需求分析与关键技术", "4"),
        (2, "2.1 可行性分析", "4"),
        (3, "2.1.1 技术可行性", "4"),
        (3, "2.1.2 经济可行性", "4"),
        (3, "2.1.3 操作可行性", "4"),
        (2, "2.2 关键技术介绍", "4"),
        (3, "2.2.1 Vue3框架", "4"),
        (3, "2.2.2 Express框架与SQLite数据库", "5"),
        (3, "2.2.3 MapLibre GL地图库", "5"),
        (3, "2.2.4 可视图算法与Dijkstra算法", "5"),
        (3, "2.2.5 JWT认证", "6"),
        (2, "2.3 功能需求分析", "6"),
        (3, "2.3.1 用户端功能需求", "7"),
        (3, "2.3.2 管理端功能需求", "7"),
        (2, "2.4 本章小结", "7"),
        (1, "3 系统设计", "8"),
        (2, "3.1 系统功能设计", "8"),
        (2, "3.2 系统架构设计", "8"),
        (2, "3.3 业务流程设计", "9"),
        (2, "3.4 数据库设计", "10"),
        (3, "3.4.1 概念结构设计", "10"),
        (3, "3.4.2 逻辑结构设计", "14"),
        (2, "3.5 本章小结", "17"),
        (1, "4 系统功能的实现", "18"),
        (2, "4.1 用户功能模块实现", "18"),
        (3, "4.1.1 登录与注册实现", "18"),
        (3, "4.1.2 地图浏览与视角切换", "19"),
        (3, "4.1.3 兴趣点搜索与标记交互", "20"),
        (3, "4.1.4 路径规划实现", "21"),
        (3, "4.1.5 统计图表页面", "23"),
        (3, "4.1.6 游客浏览模式", "24"),
        (2, "4.2 管理员功能模块实现", "24"),
        (3, "4.2.1 管理员统计仪表盘", "25"),
        (3, "4.2.2 用户管理实现", "25"),
        (3, "4.2.3 兴趣点管理实现", "26"),
        (3, "4.2.4 公告管理实现", "27"),
        (2, "4.3 本章小结", "28"),
        (1, "5 系统测试", "29"),
        (2, "5.1 测试概述", "29"),
        (2, "5.2 测试方法", "29"),
        (3, "5.2.1 功能测试方法", "29"),
        (3, "5.2.2 性能测试方法", "29"),
        (3, "5.2.3 安全测试方法", "29"),
        (2, "5.3 功能测试", "29"),
        (3, "5.3.1 用户端功能测试", "30"),
        (3, "5.3.2 管理端功能测试", "30"),
        (2, "5.4 测试结果与分析", "31"),
        (3, "5.4.1 功能测试结果", "31"),
        (3, "5.4.2 性能测试结果", "31"),
        (3, "5.4.3 安全测试结果", "32"),
        (2, "5.5 本章小结", "32"),
        (1, "结  论", "33"),
        (1, "参考文献", "34"),
        (1, "致  谢", "36"),
        (1, "附  录", "37"),
    ]

    paragraphs = doc.paragraphs
    toc_title_idx = None
    for idx, p in enumerate(paragraphs):
        if (p.text or "").strip() == "目  录":
            toc_title_idx = idx
            break
    if toc_title_idx is None:
        return

    body_idx = None
    for idx in range(toc_title_idx + 1, len(paragraphs)):
        p = paragraphs[idx]
        if (p.text or "").strip() == "1 绪论" and p.style and "Heading" in p.style.name:
            body_idx = idx
            break
    if body_idx is None:
        return

    sect_break = None
    for p in paragraphs[toc_title_idx + 1:body_idx]:
        if p._p.find('.//w:sectPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}) is not None:
            sect_break = p
    if sect_break is None:
        return

    for p in reversed(paragraphs[toc_title_idx + 1:body_idx]):
        if p is sect_break:
            continue
        _remove_paragraph(p)

    for level, title, page in toc_entries:
        p = sect_break.insert_paragraph_before()
        try:
            p.style = f"toc {level}"
        except KeyError:
            pass
        p.paragraph_format.left_indent = Cm(0.56 * (level - 1))
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.keep_with_next = False
        p.paragraph_format.tab_stops.clear_all()
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(14.7),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )
        run = p.add_run(f"{title}\t{page}")
        _set_run_font(run, 10.5, bold=(level == 1))


def _insert_code_block(doc, after_para, code_data):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    code_text = code_data.get("content", "")
    lines = code_text.split("\n")

    for line in reversed(lines):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), '0')
        ind.set(qn('w:firstLineChars'), '0')
        pPr.append(ind)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), '240')
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)
        p.append(pPr)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Courier New')
        rFonts.set(qn('w:hAnsi'), 'Courier New')
        rFonts.set(qn('w:cs'), 'Courier New')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '21')
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '21')
        rPr.append(szCs)
        r.append(rPr)
        t_el = OxmlElement('w:t')
        t_el.set(qn('xml:space'), 'preserve')
        t_el.text = line
        r.append(t_el)
        p.append(r)
        after_para._p.addnext(p)


def _insert_table(doc, after_para, tbl_data):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    headers = tbl_data.get("headers", [])
    rows = tbl_data.get("rows", [])
    if not headers:
        return

    all_rows = [headers] + rows

    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblPr.append(jc)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'autofit')
    tblPr.append(tblLayout)
    borders = OxmlElement('w:tblBorders')
    for edge, sz in [('top', '6'), ('bottom', '6')]:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    for edge in ['left', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)
    tbl.append(tblPr)

    for row_idx, cells in enumerate(all_rows):
        tr = OxmlElement('w:tr')
        for cell_text in cells:
            tc = OxmlElement('w:tc')
            if row_idx == 0:
                tcPr = OxmlElement('w:tcPr')
                tcBorders = OxmlElement('w:tcBorders')
                btm = OxmlElement('w:bottom')
                btm.set(qn('w:val'), 'single')
                btm.set(qn('w:sz'), '6')
                btm.set(qn('w:space'), '0')
                btm.set(qn('w:color'), '000000')
                tcBorders.append(btm)
                tcPr.append(tcBorders)
                tc.append(tcPr)
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            pJc = OxmlElement('w:jc')
            pJc.set(qn('w:val'), 'center')
            pPr.append(pJc)
            p.append(pPr)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            _append_rpr_fonts(rPr)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '21')
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '21')
            rPr.append(szCs)
            if row_idx == 0:
                b = OxmlElement('w:b')
                rPr.append(b)
            r.append(rPr)
            t_el = OxmlElement('w:t')
            t_el.set(qn('xml:space'), 'preserve')
            t_el.text = str(cell_text)
            r.append(t_el)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    after_para._p.addnext(tbl)


def _format_references(doc):
    from docx.shared import Cm

    ref_started = False
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text == "参考文献":
            ref_started = True
            continue
        if not ref_started:
            continue
        if re.match(r'^\[\d+\]', text):
            p.paragraph_format.first_line_indent = Cm(-0.74)
            p.paragraph_format.left_indent = Cm(0.74)


def _superscript_citations(doc):
    from docx.shared import Pt

    citation_pat = re.compile(r'(\[\d+(?:,\s*\d+)*\](?:\[\d+(?:,\s*\d+)*\])*)')
    ref_started = False
    for p in doc.paragraphs:
        text = p.text or ""
        if text.strip() == "参考文献":
            ref_started = True
            continue
        if ref_started or not citation_pat.search(text):
            continue

        style = p.style
        alignment = p.alignment
        parts = citation_pat.split(text)
        for r in list(p.runs):
            r.text = ""
        for part in parts:
            if not part:
                continue
            run = p.add_run(part)
            _set_run_font(run, 12)
            if citation_pat.fullmatch(part):
                run.font.superscript = True
                run.font.size = Pt(9)
        p.style = style
        p.alignment = alignment


def _verify(docx_path):
    from docx import Document
    doc = Document(docx_path)
    errors = []
    texts = [p.text or "" for p in doc.paragraphs]
    full = "\n".join(texts)
    for tag in ["{%", "%}", "{{", "}}"]:
        if tag in full:
            errors.append(f"残留模板标记: {tag}")
    return errors


if __name__ == "__main__":
    main()
