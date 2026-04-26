"""
学习伴侣系统论文组装器。

项目规范：
- 模板由 templates/hrbust_software/make.py 从学校 Word 模板清理生成；
- 论文内容放在 meta.json、ch*.json、references.json 等数据文件中；
- build.py 只负责读取 JSON、插入图片/表格/代码并渲染最终 DOCX。

用法:
  cd /Users/a136/vs/45425/thesis_project/papers/study-buddy
  python build.py [输出.docx]
"""
from __future__ import annotations

import json
import re
import sys
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
HERE = Path(__file__).resolve().parent
TPL = ROOT / "templates" / "hrbust_software" / "template.docx"

sys.path.insert(0, str(ROOT))

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docxtpl import DocxTemplate


def load_json(name: str, default=None):
    path = HERE / name
    if not path.exists():
        return default
    return json.loads(path.read_text(encoding="utf-8"))


def split_paras(value, count=3):
    if isinstance(value, list):
        parts = [str(x).strip() for x in value if str(x).strip()]
    else:
        parts = [p.strip() for p in str(value or "").split("\n") if p.strip()]
    parts = parts[:count]
    while len(parts) < count:
        parts.append("")
    return parts


def fmt_run(run, size=12, bold=None, east="宋体", ascii_font="Times New Roman"):
    run.font.size = Pt(size)
    run.font.name = ascii_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), east)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    if bold is not None:
        run.bold = bold


def ensure_outline(p, level: int):
    p_pr = p._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))


def add_heading(sd, text, level=1, page_break=False):
    # Do not use the source Heading styles directly: this school template binds
    # automatic numbering to Heading 1/2/3, while the thesis data already carries
    # visible numbers such as "第1章" and "1.1".
    p = sd.add_paragraph(str(text))
    ensure_outline(p, max(0, level - 1))
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if page_break:
        p.paragraph_format.page_break_before = True
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size, east, bold = 16, "黑体", True
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size, east, bold = 14, "黑体", True
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size, east, bold = 12, "宋体", True
    for r in p.runs:
        fmt_run(r, size=size, east=east, bold=bold)
    return p


def add_toc_line(sd, title, page, level=1):
    p = sd.add_paragraph()
    p.paragraph_format.left_indent = Cm({1: 0, 2: 0.42, 3: 0.82}.get(level, 0))
    p.paragraph_format.tab_stops.add_tab_stop(Cm(13.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    r = p.add_run(f"{title}\t{page}")
    fmt_run(r, 10.5, bold=(level == 1))


def build_toc(doc, chapters):
    sd = doc.new_subdoc()
    add_toc_line(sd, "摘要", "I", 1)
    add_toc_line(sd, "Abstract", "II", 1)
    page = 1
    chapter_pages = {
        "第1章 绪论": 1,
        "第2章 系统需求分析": 7,
        "第3章 系统设计": 16,
        "第4章 系统实现与测试": 39,
    }
    section_offsets = {
        "第1章 绪论": [0, 1, 2, 3],
        "第2章 系统需求分析": [0, 1, 3],
        "第3章 系统设计": [0, 4, 12, 22],
        "第4章 系统实现与测试": [0, 3, 8, 12],
    }
    for ch in chapters:
        page = chapter_pages.get(ch["title"], page)
        add_toc_line(sd, ch["title"], page, 1)
        sec_idx = 0
        for sec in ch.get("sections", []):
            sec_page = page + (section_offsets.get(ch["title"], [0])[sec_idx] if sec_idx < len(section_offsets.get(ch["title"], [])) else sec_idx)
            if sec.get("title"):
                add_toc_line(sd, sec["title"], sec_page, 2)
            for sub_idx, sub in enumerate(sec.get("subsections", [])):
                add_toc_line(sd, sub["title"], sec_page, 3)
            sec_idx += 1
    add_toc_line(sd, "结论", 52, 1)
    add_toc_line(sd, "致谢", 53, 1)
    add_toc_line(sd, "参考文献", 54, 1)
    add_toc_line(sd, "附录", 56, 1)
    return sd


def add_text(sd, text):
    try:
        p = sd.add_paragraph(style="Body Text First Indent")
    except Exception:
        p = sd.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    parts = re.split(r"(\[[0-9,，\\-—]+\])", str(text))
    for part in parts:
        if not part:
            continue
        r = p.add_run(part)
        fmt_run(r, size=12)
        if re.fullmatch(r"\[[0-9,，\\-—]+\]", part):
            r.font.superscript = True
    return p


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.first_child_found_in("w:tcBorders")
    if old is not None:
        tc_pr.remove(old)
    borders = OxmlElement("w:tcBorders")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        if val is None:
            continue
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val.get("val", "single"))
        el.set(qn("w:sz"), val.get("sz", "4"))
        el.set(qn("w:color"), val.get("color", "000000"))
        el.set(qn("w:space"), "0")
        borders.append(el)
    tc_pr.append(borders)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=70, bottom=70, left=90, right=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.first_child_found_in("w:tcMar")
    if old is not None:
        tc_pr.remove(old)
    mar = OxmlElement("w:tcMar")
    for edge, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tc_pr.append(mar)


def set_table_grid(table, widths_cm):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths_cm) * 567)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 567)))
        grid.append(col)
    tbl.insert(1, grid)


def add_table(sd, data):
    caption = data.get("caption", "")
    headers = data.get("headers", [])
    rows = data.get("rows", [])
    if not headers:
        return
    if caption:
        p = sd.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        fmt_run(r, 10.5, True)

    widths = data.get("col_widths") or [13.2 / len(headers)] * len(headers)
    max_width = 13.2
    total = sum(widths)
    if total > max_width:
        widths = [w * max_width / total for w in widths]

    table = sd.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    set_table_grid(table, widths)

    thick = {"val": "single", "sz": "12", "color": "000000"}
    thin = {"val": "single", "sz": "4", "color": "000000"}
    none = {"val": "nil"}

    for r_idx, row in enumerate([headers] + rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_width(cell, widths[c_idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(
                cell,
                top=thick if r_idx == 0 else none,
                bottom=thin if r_idx == 0 else (thick if r_idx == len(rows) else none),
                left=none,
                right=none,
            )
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx < 2 or len(str(value)) < 14 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = Pt(16)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            fmt_run(r, 9 if len(headers) >= 4 else 10.5, bold=(r_idx == 0))
    sd.add_paragraph()


def add_image(sd, image, caption="", width_cm=12.0):
    path = Path(image)
    if not path.is_absolute():
        path = HERE / path
    if not path.exists():
        add_text(sd, f"[图片缺失：{image}]")
        return
    p = sd.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    if caption:
        cap = sd.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = cap.add_run(caption)
        fmt_run(rr, 10.5, True)


def add_code(sd, code, caption=None):
    if caption:
        p = sd.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        fmt_run(r, 10.5, True)
    p = sd.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Cm(0.5)
    for line in str(code).strip("\n").splitlines():
        r = p.add_run(line + "\n")
        fmt_run(r, 8.5, east="Consolas", ascii_font="Consolas")


def build_blocks(sd, blocks):
    for block in blocks or []:
        if isinstance(block, str):
            add_text(sd, block)
        elif isinstance(block, dict):
            if "text" in block:
                add_text(sd, block["text"])
            elif "table" in block:
                add_table(sd, block["table"])
            elif "image" in block:
                add_image(sd, block["image"], block.get("caption", ""), block.get("width_cm", 12.0))
            elif "code" in block:
                add_code(sd, block["code"], block.get("caption"))


def build_body(doc, chapters):
    sd = doc.new_subdoc()
    for idx, ch in enumerate(chapters):
        add_heading(sd, ch["title"], level=1, page_break=idx > 0)
        build_blocks(sd, ch.get("content", []))
        for sec in ch.get("sections", []):
            if sec.get("title"):
                add_heading(sd, sec["title"], level=2)
            build_blocks(sd, sec.get("content", []))
            for sub in sec.get("subsections", []):
                add_heading(sd, sub["title"], level=3)
                build_blocks(sd, sub.get("content", []))
    return sd


def build_para_doc(doc, paras):
    sd = doc.new_subdoc()
    for para in paras or []:
        add_text(sd, para)
    return sd


def build_appendix_doc(doc, blocks):
    sd = doc.new_subdoc()
    build_blocks(sd, blocks)
    return sd


def build_refs(doc, refs):
    sd = doc.new_subdoc()
    for i, ref in enumerate(refs or [], 1):
        text = re.sub(r"^\[\d+\]\s*", "", str(ref).strip())
        p = sd.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.line_spacing = Pt(20)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        r = p.add_run(f"[{i}] {text}")
        fmt_run(r, 10.5)
    return sd


def patch_update_fields(docx_path: Path):
    tmp = Path(str(docx_path) + ".tmp")
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "word/settings.xml":
                    xml = data.decode("utf-8")
                    if "w:updateFields" not in xml:
                        xml = xml.replace("</w:settings>", '<w:updateFields w:val="true"/></w:settings>')
                    data = xml.encode("utf-8")
                zout.writestr(info, data)
    tmp.replace(docx_path)


def build_data(doc):
    meta = load_json("meta.json", {})
    zh = split_paras(meta.get("abstract_zh"), 3)
    en = split_paras(meta.get("abstract_en"), 3)
    for i in range(3):
        meta[f"abstract_zh_{i + 1}"] = zh[i]
        meta[f"abstract_en_{i + 1}"] = en[i]

    chapters = []
    for idx in range(1, 10):
        ch = load_json(f"ch{idx}.json")
        if ch:
            print(f"  ch{idx}.json -> {ch['title']}")
            chapters.append(ch)

    meta["toc_doc"] = build_toc(doc, chapters)
    conclusion = load_json("conclusion.json", [])
    acknowledgement = load_json("acknowledgement.json", [])
    appendix = load_json("appendix.json", [])
    references = load_json("references.json", [])
    if isinstance(references, dict):
        references = references.get("references", [])

    meta["body"] = build_body(doc, chapters)
    meta["conclusion_doc"] = build_para_doc(doc, conclusion)
    meta["acknowledgement_doc"] = build_para_doc(doc, acknowledgement)
    meta["references_doc"] = build_refs(doc, references)
    meta["appendix_doc"] = build_appendix_doc(doc, appendix)
    return meta


def verify_docx(docx_path: Path):
    from docx import Document

    doc = Document(str(docx_path))
    text = "\n".join(p.text for p in doc.paragraphs)
    errors = []
    # Do not treat literal JSX style={{...}} snippets in appendix code as
    # template residue. Real docxtpl control tags are checked explicitly.
    for token in ("{%p", "%}", "{{p "):
        if token in text:
            errors.append(f"残留模板标记 {token}")
    for old in ("停车场管理系统", "物联网综合在线实验平台", "基于SSM", "香菇", "YOLO"):
        if old in text:
            errors.append(f"残留旧模板/旧项目文本：{old}")
    if "学习伴侣" not in text:
        errors.append("未检测到学习伴侣主题文本")
    return errors


def main():
    output = Path(sys.argv[1]) if len(sys.argv) > 1 else HERE / "姓名-学号-基于微信小程序的学习伴侣系统设计与实现.docx"
    if not TPL.exists():
        raise FileNotFoundError(f"模板不存在，请先运行 templates/hrbust_software/make.py: {TPL}")

    doc = DocxTemplate(str(TPL))
    print("组装论文数据...")
    data = build_data(doc)
    print("渲染模板...")
    doc.render(data)
    doc.save(output)
    patch_update_fields(output)
    errors = verify_docx(output)
    if errors:
        print("校验提醒:")
        for e in errors:
            print(" -", e)
    else:
        print("校验通过")
    print(f"完成: {output}")


if __name__ == "__main__":
    main()
