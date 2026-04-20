"""
基于Flask+Bootstrap的企业官网CMS论文组装器（青岛科技大学模板）。
支持文本、图片（InlineImage）、表格（后处理三线表）。

用法: python build.py [输出.docx]
"""
import os
import sys
import json
import re
import shutil
import tempfile
import zipfile
from xml.etree import ElementTree as ET

ROOT = os.path.join(os.path.dirname(__file__), "../..")
sys.path.insert(0, ROOT)

HERE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(HERE, "images")
TPL = os.path.join(ROOT, "templates/qdkjdx/template.docx")

_pending_tables = []
_pending_code_blocks = []


def load_json(filename):
    path = os.path.join(HERE, filename)
    if not os.path.exists(path):
        return None
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def process_content(content_list, doc):
    from docxtpl import InlineImage
    from docx.shared import Mm
    result = []
    for item in content_list:
        if isinstance(item, str):
            result.append(item)
        elif isinstance(item, dict):
            if item.get("type") == "image":
                img_path = os.path.join(IMG_DIR, item["path"])
                if os.path.exists(img_path):
                    width = Mm(item.get("width", 148))
                    result.append(InlineImage(doc, img_path, width=width))
                    if item.get("caption"):
                        result.append(item["caption"])
                else:
                    result.append(f"[图片缺失: {item['path']}]")
                    if item.get("caption"):
                        result.append(item["caption"])
            elif item.get("type") == "code":
                _pending_code_blocks.append(item)
                cid = len(_pending_code_blocks) - 1
                result.append(f"__CODE_PLACEHOLDER_{cid}__")
            elif item.get("type") == "table":
                _pending_tables.append(item)
                tid = len(_pending_tables) - 1
                if item.get("caption"):
                    result.append(item["caption"])
                result.append(f"__TABLE_PLACEHOLDER_{tid}__")
            else:
                result.append(str(item))
    return result


def process_chapters(chapters, doc):
    for ch in chapters:
        ch["content"] = process_content(ch.get("content", []), doc)
        for sec in ch.get("sections", []):
            sec["content"] = process_content(sec.get("content", []), doc)
            for sub in sec.get("subsections", []):
                sub["content"] = process_content(sub.get("content", []), doc)
    return chapters


def build_data(doc=None):
    meta = load_json("meta.json") or {}
    chapters = []
    for i in range(1, 7):
        ch = load_json(f"ch{i}.json")
        if ch:
            print(f"  ch{i}.json -> {ch['title']}")
            chapters.append(ch)

    if doc:
        chapters = process_chapters(chapters, doc)

    refs = load_json("references.json") or []
    refs = [f"[{i+1}]{ref}" if not ref.startswith("[") else ref
            for i, ref in enumerate(refs)]

    return {**meta, "chapters": chapters, "references": refs}


def main():
    output = sys.argv[1] if len(sys.argv) > 1 else None

    from docxtpl import DocxTemplate
    doc = DocxTemplate(TPL)

    _pending_tables.clear()
    _pending_code_blocks.clear()
    print("组装论文数据...")
    data = build_data(doc)
    print(f"  {len(data['chapters'])} 章, "
          f"{len(data.get('references', []))} 条参考文献")

    if not output:
        from datetime import datetime
        ts = datetime.now().strftime("%m%d_%H%M")
        output = os.path.join(HERE,
            f"基于Flask的企业官网CMS_{ts}.docx")

    print(f"渲染: {output}")
    doc.render(data)
    doc.save(output)

    _post_process(output)
    _enable_update_fields(output)
    errors = _verify(output)
    if errors:
        print(f"  ⚠ {len(errors)} 个问题:")
        for e in errors:
            print(f"    - {e}")
    else:
        print("  ✓ 校验通过")
    print(f"完成: {output}")


def _post_process(docx_path):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    doc = Document(docx_path)
    tbl_cap_pat = re.compile(r'^表\d')
    fig_cap_pat = re.compile(r'^图\d')
    placeholder_pat = re.compile(r'^__TABLE_PLACEHOLDER_(\d+)__$')
    code_placeholder_pat = re.compile(r'^__CODE_PLACEHOLDER_(\d+)__$')

    for p in list(doc.paragraphs):
        t = (p.text or "").strip()

        # 代码块占位符
        cm = code_placeholder_pat.match(t)
        if cm:
            cid = int(cm.group(1))
            if cid < len(_pending_code_blocks):
                _insert_code_block(doc, p, _pending_code_blocks[cid])
                for r in p.runs:
                    r.text = ""

        # 章标题前分页（每章另起一页），跳过"摘 要"和"前言"
        if p.style and p.style.name == '论文正文章标题' and re.match(r'^\d', t):
            p.paragraph_format.page_break_before = True

        # 图片居中
        has_drawing = bool(p._p.findall(
            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
        if has_drawing:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pPr = p._p.pPr
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p._p.insert(0, pPr)
            old_ind = pPr.find(_qn('w:ind'))
            if old_ind is not None:
                pPr.remove(old_ind)
            new_ind = OxmlElement('w:ind')
            new_ind.set(_qn('w:firstLine'), '0')
            new_ind.set(_qn('w:firstLineChars'), '0')
            pPr.append(new_ind)
            from docx.enum.text import WD_LINE_SPACING
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 图表标注居中
        if (fig_cap_pat.match(t) or tbl_cap_pat.match(t)) and len(t) < 60:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(10.5)

        # 表格占位符
        m = placeholder_pat.match(t)
        if m:
            tid = int(m.group(1))
            if tid < len(_pending_tables):
                _insert_table(doc, p, _pending_tables[tid])
                for r in p.runs:
                    r.text = ""

    doc.save(docx_path)
    print(f"  后处理: {len(_pending_tables)} 个表格, {len(_pending_code_blocks)} 个代码块")


def _insert_code_block(doc, after_para, code_data):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    code_text = code_data.get("content", "")
    lines = code_text.split("\n")

    for line in reversed(lines):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), '0')
        ind.set(qn('w:firstLineChars'), '0')
        pPr.append(ind)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), '240')
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)
        p.append(pPr)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Courier New')
        rFonts.set(qn('w:hAnsi'), 'Courier New')
        rFonts.set(qn('w:cs'), 'Courier New')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '21')
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '21')
        rPr.append(szCs)
        r.append(rPr)
        t_el = OxmlElement('w:t')
        t_el.set(qn('xml:space'), 'preserve')
        t_el.text = line
        r.append(t_el)
        p.append(r)
        after_para._p.addnext(p)


def _insert_table(doc, after_para, tbl_data):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    headers = tbl_data.get("headers", [])
    rows = tbl_data.get("rows", [])
    if not headers:
        return

    all_rows = [headers] + rows

    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblPr.append(jc)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'autofit')
    tblPr.append(tblLayout)
    borders = OxmlElement('w:tblBorders')
    for edge, sz in [('top', '12'), ('bottom', '12')]:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    for edge in ['left', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)
    tbl.append(tblPr)

    for row_idx, cells in enumerate(all_rows):
        tr = OxmlElement('w:tr')
        for cell_text in cells:
            tc = OxmlElement('w:tc')
            if row_idx == 0:
                tcPr = OxmlElement('w:tcPr')
                tcBorders = OxmlElement('w:tcBorders')
                btm = OxmlElement('w:bottom')
                btm.set(qn('w:val'), 'single')
                btm.set(qn('w:sz'), '6')
                btm.set(qn('w:space'), '0')
                btm.set(qn('w:color'), '000000')
                tcBorders.append(btm)
                tcPr.append(tcBorders)
                tc.append(tcPr)
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            pJc = OxmlElement('w:jc')
            pJc.set(qn('w:val'), 'center')
            pPr.append(pJc)
            p.append(pPr)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '21')
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '21')
            rPr.append(szCs)
            if row_idx == 0:
                b = OxmlElement('w:b')
                rPr.append(b)
            r.append(rPr)
            t_el = OxmlElement('w:t')
            t_el.set(qn('xml:space'), 'preserve')
            t_el.text = str(cell_text)
            r.append(t_el)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    after_para._p.addnext(tbl)


def _verify(docx_path):
    from docx import Document
    doc = Document(docx_path)
    errors = []
    texts = [p.text or "" for p in doc.paragraphs]
    full = "\n".join(texts)
    for tag in ["{%", "%}", "{{", "}}"]:
        if tag in full:
            errors.append(f"残留模板标记: {tag}")
    return errors


def _enable_update_fields(docx_path):
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ET.register_namespace('w', ns)

    with zipfile.ZipFile(docx_path, 'r') as zin:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp_path = tmp.name

        with zipfile.ZipFile(tmp_path, 'w') as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/settings.xml":
                    root = ET.fromstring(data)
                    node = root.find(f'{{{ns}}}updateFields')
                    if node is None:
                        node = ET.Element(f'{{{ns}}}updateFields')
                        root.append(node)
                    node.set(f'{{{ns}}}val', 'true')
                    data = ET.tostring(root, encoding='utf-8', xml_declaration=True)
                zout.writestr(item, data)

    shutil.move(tmp_path, docx_path)


if __name__ == "__main__":
    main()
