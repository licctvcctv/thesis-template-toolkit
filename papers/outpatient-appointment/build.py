"""
门诊预约挂号系统论文组装器。
支持文本、图片（InlineImage）、表格（三线表）。

用法:
    python build.py [输出.docx]
    python build.py <模板.docx> <输出.docx>
"""
import os
import sys
import json
import re

ROOT = os.path.join(os.path.dirname(__file__), "../..")
sys.path.insert(0, ROOT)

HERE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(HERE, "images")
TPL = os.path.join(ROOT, "templates/outpatient_original/template.docx")
DEFAULT_OUTPUT = os.path.join(HERE, "基于Spring Boot的门诊预约挂号系统设计与实现_修改版.docx")


def _ensure_rgb(img_path):
    """灰度PNG在Word中会偏色，转为RGB PNG"""
    try:
        from PIL import Image
        img = Image.open(img_path)
        if img.mode in ('L', 'LA', 'P', 'PA'):
            canvas = Image.new('RGB', img.size, (255, 255, 255))
            rgba = img.convert('RGBA')
            canvas.paste(rgba, mask=rgba.split()[3])
            canvas.save(img_path, 'PNG')
    except Exception:
        pass
    return img_path


def load_json(filename):
    path = os.path.join(HERE, filename)
    if not os.path.exists(path):
        return None
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def process_content(content_list, doc):
    """
    处理内容列表：
    - 字符串 → 保持不变
    - {"type":"image"} → InlineImage（图片）
    - {"type":"table"} → 转为特殊标记文本（后处理时插入表格）
    """
    from docxtpl import InlineImage
    from docx.shared import Mm

    result = []
    for item in content_list:
        if isinstance(item, str):
            result.append(item)
        elif isinstance(item, dict):
            if item.get("type") == "image":
                img_path = os.path.join(IMG_DIR, item["path"])
                if os.path.exists(img_path):
                    width = Mm(item.get("width", 120))
                    result.append(InlineImage(doc, img_path,
                                             width=width))
                    if item.get("caption"):
                        result.append(item["caption"])
                else:
                    result.append(f"[图片缺失: {item['path']}]")
                    if item.get("caption"):
                        result.append(item["caption"])
            elif item.get("type") == "table":
                _pending_tables.append(item)
                tid = len(_pending_tables) - 1
                # 如果 table 对象自带 caption，先输出标注行
                if item.get("caption"):
                    result.append(item["caption"])
                result.append(f"__TABLE_PLACEHOLDER_{tid}__")
            else:
                result.append(str(item))
    return result


# 全局表格暂存
_pending_tables = []


def process_chapters(chapters, doc):
    """处理所有章节，转换图片引用"""
    for ch in chapters:
        ch["content"] = process_content(ch.get("content", []), doc)
        for sec in ch.get("sections", []):
            sec["content"] = process_content(
                sec.get("content", []), doc)
            for sub in sec.get("subsections", []):
                sub["content"] = process_content(
                    sub.get("content", []), doc)
    return chapters


def build_data(doc=None):
    """组装所有章节数据"""
    meta = load_json("meta.json")
    if not meta:
        raise FileNotFoundError("meta.json not found")

    chapters = []
    for idx in range(1, 7):
        filename = f"ch{idx}.json"
        ch = load_json(filename)
        if ch:
            chapters.append(ch)
            print(f"  {filename} -> {ch['title']}")

    if doc:
        chapters = process_chapters(chapters, doc)

    refs = load_json("references.json")
    if isinstance(refs, list):
        references = refs
    elif isinstance(refs, dict):
        references = refs.get("references", [])
    else:
        references = []

    return {**meta, "chapters": chapters,
            "references": references}


def main():
    if len(sys.argv) >= 3:
        template = sys.argv[1]
        output = sys.argv[2]
    elif len(sys.argv) == 2:
        template = TPL
        output = sys.argv[1]
    else:
        template = TPL
        output = DEFAULT_OUTPUT

    from docxtpl import DocxTemplate
    doc = DocxTemplate(template)

    _pending_tables.clear()
    print("组装论文数据...")
    data = build_data(doc)
    print(f"  {len(data['chapters'])} 章, "
          f"{len(data.get('references', []))} 条参考文献")

    print(f"渲染: {template} -> {output}")
    doc.render(data)
    doc.save(output)

    # 摘要中使用空行分段，渲染后拆成真正的 Word 段落。
    _split_paragraph_breaks(output)
    # 删除 Word 批注
    _remove_comments(output)
    # 后处理：修正图片居中和图标注格式
    _post_process(output)
    _patch_auto_toc(output)
    _patch_section_headers(output)
    # 校验
    errors = _verify(output)
    errors.extend(_verify_citations(output))
    if errors:
        print(f"  ⚠ 发现 {len(errors)} 个问题:")
        for e in errors:
            print(f"    {e}")
    else:
        print(f"  ✓ 校验通过")
    print(f"完成: {output}")


def _split_paragraph_breaks(docx_path):
    """把模板字段中的双换行转换为真正段落，避免摘要挤成一整段。"""
    from copy import deepcopy
    from docx import Document
    from docx.text.paragraph import Paragraph

    doc = Document(docx_path)
    changed = False

    def clear_para_xml(p_xml):
        for child in list(p_xml):
            if child.tag.endswith('}pPr'):
                continue
            p_xml.remove(child)

    for para in list(doc.paragraphs):
        if '\n\n' not in para.text:
            continue
        parts = [part.strip() for part in para.text.split('\n\n')
                 if part.strip()]
        if len(parts) <= 1:
            continue

        clear_para_xml(para._p)
        para.add_run(parts[0])
        current = para._p
        for part in parts[1:]:
            new_p = deepcopy(para._p)
            clear_para_xml(new_p)
            current.addnext(new_p)
            new_para = Paragraph(new_p, para._parent)
            new_para.add_run(part)
            current = new_p
        changed = True

    if changed:
        doc.save(docx_path)


def _remove_comments(docx_path):
    """删除 Word 文档中的所有批注（comments）"""
    import zipfile
    import shutil
    import tempfile
    from lxml import etree

    WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    tmp = tempfile.mktemp(suffix='.docx')
    with zipfile.ZipFile(docx_path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w') as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                # 跳过 comments 部分文件
                if 'comments' in item.filename.lower():
                    continue
                # 从 document.xml 中删除批注标记
                if item.filename == 'word/document.xml':
                    root = etree.fromstring(data)
                    for tag in ['commentRangeStart', 'commentRangeEnd',
                                'commentReference']:
                        for el in root.findall(f'.//{{{WNS}}}{tag}'):
                            el.getparent().remove(el)
                    data = etree.tostring(root, xml_declaration=True,
                                          encoding='UTF-8', standalone=True)
                # 从 [Content_Types].xml 中移除 comments 引用
                if item.filename == '[Content_Types].xml':
                    root = etree.fromstring(data)
                    for override in root.findall(
                            '{http://schemas.openxmlformats.org/package/2006/content-types}Override'):
                        pn = override.get('PartName', '')
                        if 'comments' in pn.lower():
                            root.remove(override)
                    data = etree.tostring(root, xml_declaration=True,
                                          encoding='UTF-8', standalone=True)
                # 从 relationships 中移除 comments 引用
                if item.filename == 'word/_rels/document.xml.rels':
                    root = etree.fromstring(data)
                    for rel in list(root):
                        target = rel.get('Target', '')
                        if 'comments' in target.lower():
                            root.remove(rel)
                    data = etree.tostring(root, xml_declaration=True,
                                          encoding='UTF-8', standalone=True)
                zout.writestr(item, data)
    shutil.move(tmp, docx_path)
    print("  批注已删除")


def _verify(docx_path):
    """校验输出文档：检查每个表/图标注是否有对应内容"""
    from docx import Document
    import re

    doc = Document(docx_path)
    errors = []
    body = doc.element.body
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # 按 DOM 顺序遍历，记录每个元素的类型
    elements = []  # [(type, text)]
    for elem in body:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            texts = [t.text for t in elem.findall('.//w:t', ns) if t.text]
            text = ''.join(texts).strip()
            has_img = bool(elem.findall('.//w:drawing', ns))
            if has_img:
                elements.append(('img', text))
            else:
                elements.append(('p', text))
        elif tag == 'tbl':
            elements.append(('tbl', ''))

    # 检查每个表标注后面是否有表格
    for i, (etype, text) in enumerate(elements):
        if etype == 'p' and re.match(r'^表\d+-\d+\s', text) and not re.search(r'[。，；]', text) and len(text) < 30:
            # 往后找 3 个元素内是否有表格
            found_tbl = False
            for j in range(i + 1, min(i + 6, len(elements))):
                if elements[j][0] == 'tbl':
                    found_tbl = True
                    break
            if not found_tbl:
                errors.append(f"缺表格: {text}")

    # 检查每个图标注附近是否有图片
    for i, (etype, text) in enumerate(elements):
        if etype == 'p' and re.match(r'^图\d+-\d+\s', text) and len(text) < 50:
            found_img = False
            # 往前找 3 个元素
            for j in range(max(0, i - 3), i):
                if elements[j][0] == 'img':
                    found_img = True
                    break
            if not found_img:
                errors.append(f"缺图片: {text}")

    return errors


def _post_process(docx_path):
    """渲染后：图片居中、标注格式、插入表格"""
    import copy
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import re

    doc = Document(docx_path)
    fig_pat = re.compile(r'^图\d')
    tbl_cap_pat = re.compile(r'^表\d')
    placeholder_pat = re.compile(r'^__TABLE_PLACEHOLDER_(\d+)__$')
    citation_pat = re.compile(r'\[\d+(?:[-,，]\d+)*\]')

    def _set_or_append(parent, tag):
        child = parent.find(qn(tag))
        if child is None:
            child = OxmlElement(tag)
            parent.append(child)
        return child

    def _fix_cover_info_table():
        """封面专业/姓名/导师信息表保持一行显示，避免专业名被挤断。"""
        for table in doc.tables:
            table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
            if "专" not in table_text or "指导教师" not in table_text:
                continue
            tbl_pr = _set_or_append(table._tbl, 'w:tblPr')
            tbl_w = _set_or_append(tbl_pr, 'w:tblW')
            tbl_w.set(qn('w:w'), '4800')
            tbl_w.set(qn('w:type'), 'dxa')
            layout = _set_or_append(tbl_pr, 'w:tblLayout')
            layout.set(qn('w:type'), 'fixed')
            grid = table._tbl.tblGrid
            if grid is not None:
                for col in grid.gridCol_lst:
                    col.set(qn('w:w'), '4800')
            for row in table.rows:
                for cell in row.cells:
                    tc_pr = cell._tc.get_or_add_tcPr()
                    tc_w = tc_pr.tcW
                    if tc_w is None:
                        tc_w = OxmlElement('w:tcW')
                        tc_pr.append(tc_w)
                    tc_w.set(qn('w:w'), '4800')
                    tc_w.set(qn('w:type'), 'dxa')
                    no_wrap = tc_pr.find(qn('w:noWrap'))
                    if no_wrap is None:
                        no_wrap = OxmlElement('w:noWrap')
                        tc_pr.append(no_wrap)
                    no_wrap.set(qn('w:val'), '1')
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(13)
            break

    def _set_outline_level(paragraph, level):
        p_pr = paragraph._p.pPr
        if p_pr is None:
            p_pr = OxmlElement('w:pPr')
            paragraph._p.insert(0, p_pr)
        old = p_pr.find(qn('w:outlineLvl'))
        if old is not None:
            p_pr.remove(old)
        outline = OxmlElement('w:outlineLvl')
        outline.set(qn('w:val'), str(level))
        r_pr = p_pr.find(qn('w:rPr'))
        if r_pr is not None:
            p_pr.insert(list(p_pr).index(r_pr), outline)
        else:
            p_pr.append(outline)

    def _toc_level(text):
        compact = text.replace(" ", "").replace("　", "")
        if compact in {"摘要", "Abstract", "总结与展望", "结论", "致谢", "参考文献"}:
            return 0
        if re.fullmatch(r'\d+\s+\S.*', text):
            return 0
        if re.fullmatch(r'\d+\.\d+\s+\S.*', text):
            return 1
        if re.fullmatch(r'\d+\.\d+\.\d+\s+\S.*', text):
            return 2
        return None

    def _format_body_citations():
        ref_seen = False
        for p in list(doc.paragraphs):
            text = p.text or ""
            compact = text.strip().replace(" ", "").replace("　", "")
            if compact == "参考文献":
                ref_seen = True
                continue
            if ref_seen or not citation_pat.search(text):
                continue

            for run in list(p.runs):
                run_text = run.text or ""
                if not citation_pat.search(run_text):
                    continue
                parent = run._r.getparent()
                insert_at = parent.index(run._r)
                original_rpr = copy.deepcopy(run._r.rPr) if run._r.rPr is not None else None
                parent.remove(run._r)

                pos = 0
                for match in citation_pat.finditer(run_text):
                    if match.start() > pos:
                        insert_at = _insert_text_run(
                            parent, insert_at, run_text[pos:match.start()], original_rpr, False
                        )
                    insert_at = _insert_text_run(
                        parent, insert_at, match.group(0), original_rpr, True
                    )
                    pos = match.end()
                if pos < len(run_text):
                    _insert_text_run(parent, insert_at, run_text[pos:], original_rpr, False)

    def _insert_text_run(parent, insert_at, text, original_rpr, is_citation):
        new_r = OxmlElement('w:r')
        if original_rpr is not None:
            new_r.append(copy.deepcopy(original_rpr))
        r_pr = new_r.rPr
        if r_pr is None:
            r_pr = OxmlElement('w:rPr')
            new_r.insert(0, r_pr)
        if is_citation:
            for tag in ['w:vertAlign', 'w:sz', 'w:szCs']:
                old = r_pr.find(qn(tag))
                if old is not None:
                    r_pr.remove(old)
            vert = OxmlElement('w:vertAlign')
            vert.set(qn('w:val'), 'superscript')
            r_pr.append(vert)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '16')
            r_pr.append(sz)
            sz_cs = OxmlElement('w:szCs')
            sz_cs.set(qn('w:val'), '16')
            r_pr.append(sz_cs)
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        new_r.append(t)
        parent.insert(insert_at, new_r)
        return insert_at + 1

    _fix_cover_info_table()

    # 删除标题前后的多余空段落（Jinja 控制标签渲染残留）
    paras_list = list(doc.paragraphs)
    for i, p in enumerate(paras_list):
        t = (p.text or "").strip()
        has_break_or_section = bool(p._p.findall(
            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'
        )) or bool(p._p.findall(
            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr'
        ))
        if has_break_or_section:
            continue
        if not t and i > 0 and i < len(paras_list) - 1:
            prev_t = (paras_list[i - 1].text or "").strip()
            next_t = (paras_list[i + 1].text or "").strip()
            # 空段落夹在两个非空段落之间，且前后至少一个看起来是标题
            is_after_title = bool(re.match(r'^\d+(\.\d+)*\s', prev_t))
            is_before_title = bool(re.match(r'^\d+(\.\d+)*\s', next_t))
            if is_after_title or is_before_title:
                p._p.getparent().remove(p._p)
                continue

    for p in list(doc.paragraphs):
        t = (p.text or "").strip()
        compact_t = t.replace(" ", "").replace("　", "")

        if compact_t == "目录":
            p.paragraph_format.page_break_before = True

        toc_level = _toc_level(t)
        if toc_level is not None:
            _set_outline_level(p, toc_level)
            if toc_level == 0:
                p.paragraph_format.page_break_before = True

        # 一级标题与结尾部分统一新开一页。
        if (p.style and p.style.name == 'Heading 1') or re.match(r'^\d+\s+\S', t):
            p.paragraph_format.page_break_before = True

        # 图片段落 → 居中 + 清除首行缩进 + 清除继承的正文rPr
        has_drawing = bool(p._p.findall(
            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
        if has_drawing:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 清除首行缩进（Normal样式有 firstLineChars="200"，必须显式清零）
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn as _qn2
            pPr = p._p.pPr
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p._p.insert(0, pPr)
            # 删除旧的 ind，创建新的干净的
            old_ind = pPr.find(_qn2('w:ind'))
            if old_ind is not None:
                pPr.remove(old_ind)
            new_ind = OxmlElement('w:ind')
            new_ind.set(_qn2('w:left'), '0')
            new_ind.set(_qn2('w:leftChars'), '0')
            new_ind.set(_qn2('w:firstLine'), '0')
            new_ind.set(_qn2('w:firstLineChars'), '0')
            pPr.append(new_ind)
            # 固定行间距(EXACTLY 20pt)会裁切图片，改为单倍行距
            from docx.enum.text import WD_LINE_SPACING
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            # 清除段落级 rPr（继承的字体格式会影响图片渲染）
            from docx.oxml.ns import qn as _qn
            pPr = p._p.find(_qn('w:pPr'))
            if pPr is not None:
                rPr = pPr.find(_qn('w:rPr'))
                if rPr is not None:
                    pPr.remove(rPr)
            # 清除段落中非图片 run（空文本 run 带字体格式）
            for r in p._p.findall(_qn('w:r')):
                if not r.findall(_qn('w:drawing')) and not r.findall(
                        './/{http://schemas.openxmlformats.org/drawingml/2006/main}graphic'):
                    txt = r.findall(_qn('w:t'))
                    if all((t.text or '').strip() == '' for t in txt):
                        r.getparent().remove(r)

        # 图/表标注 → 居中 + 五号宋体（短文本）
        if (fig_pat.match(t) or tbl_cap_pat.match(t)) and len(t) < 50:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if tbl_cap_pat.match(t):
                p.paragraph_format.keep_with_next = True
                if t in {"表4-3 医师表"}:
                    p.paragraph_format.page_break_before = True
            for r in p.runs:
                r.font.size = Pt(10.5)
                r.font.name = "宋体"
                r.font.bold = False
                # 设置东亚字体
                r._r.rPr.rFonts.set(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia',
                    '宋体')

        # 表格占位符 → 插入真实表格
        m = placeholder_pat.match(t)
        if m:
            tid = int(m.group(1))
            if tid < len(_pending_tables):
                _insert_table(doc, p, _pending_tables[tid])
                # 删除占位段，避免表题和表格之间被空段落隔开。
                p._p.getparent().remove(p._p)

    _format_body_citations()

    meta = load_json("meta.json") or {}
    conclusion = (meta.get("conclusion") or "").strip()
    paras = list(doc.paragraphs)
    ack_heading = next((p for p in paras if (p.text or "").strip() == "致谢"), None)
    ref_heading = next((p for p in paras if (p.text or "").strip() == "参考文献"), None)
    has_conclusion = any((p.text or "").strip() == "结论" for p in paras)

    if conclusion and ack_heading and not has_conclusion:
        body_style = None
        ack_seen = False
        for p in paras:
            if p is ack_heading:
                ack_seen = True
                continue
            if ack_seen and (p.text or "").strip():
                body_style = p.style
                break
        if body_style is None:
            body_style = ack_heading.style

        heading = ack_heading.insert_paragraph_before("结论", style=ack_heading.style)
        heading.paragraph_format.page_break_before = True
        for line in [x.strip() for x in conclusion.splitlines() if x.strip()]:
            ack_heading.insert_paragraph_before(line, style=body_style)

    # 参考参考稿结构：结论后接参考文献，再接致谢
    paras = list(doc.paragraphs)
    ack_heading = next((p for p in paras if (p.text or "").strip() == "致谢"), None)
    ref_heading = next((p for p in paras if (p.text or "").strip() == "参考文献"), None)
    if ack_heading is not None and ref_heading is not None:
        body = doc.element.body
        children = list(body)
        ack_idx = children.index(ack_heading._p)
        ref_idx = children.index(ref_heading._p)
        if ack_idx < ref_idx:
            moving = []
            for el in children[ref_idx:]:
                if el.tag.endswith('}sectPr'):
                    continue
                moving.append(el)
            for el in moving:
                body.remove(el)
            ack_idx = list(body).index(ack_heading._p)
            for offset, el in enumerate(moving):
                body.insert(ack_idx + offset, el)

    # 清理模板残留的空参考文献编号，如单独一行的 [20]、[21]
    ref_seen = False
    for p in list(doc.paragraphs):
        t = (p.text or "").strip()
        compact_t = t.replace(" ", "").replace("　", "")
        if compact_t == "参考文献":
            ref_seen = True
            continue
        if re.fullmatch(r'\[\d+\]', t):
            p._p.getparent().remove(p._p)
            continue
        ppr = p._p.pPr
        num_pr = ppr.numPr if ppr is not None and ppr.numPr is not None else None
        if ref_seen and not t and num_pr is not None:
            p._p.getparent().remove(p._p)

    doc.save(docx_path)
    print(f"  后处理: 图片居中+{len(_pending_tables)}个表格")


def _insert_table(doc, after_para, tbl_data):
    """在段落后插入带边框的表格"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    headers = tbl_data.get("headers", [])
    rows = tbl_data.get("rows", [])
    if not headers:
        return

    ncols = len(headers)
    all_rows = [headers] + rows

    # 直接构建表格 XML
    tbl = OxmlElement('w:tbl')

    # 表格属性：居中 + 固定列宽 + 三线表边框
    tblPr = OxmlElement('w:tblPr')
    # 居中
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblPr.append(jc)
    # 参考论文中的表格不铺满整页，固定宽度后更接近正文表的视觉比例
    total_width = 7800
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(total_width))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    # 边框（三线表：顶线1.5pt、底线1.5pt，无左右竖线）
    borders = OxmlElement('w:tblBorders')
    for edge, sz in [('top', '12'), ('bottom', '12')]:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    for edge in ['left', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        borders.append(el)
    tblPr.append(borders)
    tbl.append(tblPr)

    # 表格网格：优先使用 JSON 中的列宽比例；没有配置时再按内容长度估算
    width_weights = tbl_data.get("widths") or []
    if len(width_weights) == ncols and sum(float(w) for w in width_weights) > 0:
        weight_total = sum(float(w) for w in width_weights)
        col_widths = [int(total_width * float(w) / weight_total) for w in width_weights]
    else:
        col_max_len = [0] * ncols
        for row_data in all_rows:
            for ci in range(ncols):
                cell_text = str(row_data[ci]) if ci < len(row_data) else ""
                text_width = sum(2 if ord(c) > 127 else 1 for c in cell_text)
                col_max_len[ci] = max(col_max_len[ci], text_width)
        total_len = sum(max(l, 4) for l in col_max_len) or 1
        col_widths = [int(total_width * max(l, 4) / total_len) for l in col_max_len]
    diff = total_width - sum(col_widths)
    if col_widths:
        col_widths[-1] += diff

    tblGrid = OxmlElement('w:tblGrid')
    for w in col_widths:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(w))
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    # 行
    for ri, row_data in enumerate(all_rows):
        tr = OxmlElement('w:tr')
        for ci in range(ncols):
            tc = OxmlElement('w:tc')
            # 单元格属性：宽度 + 垂直居中
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(col_widths[ci] if ci < len(col_widths) else 1800))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)
            tcMar = OxmlElement('w:tcMar')
            for edge, value in [('top', '45'), ('bottom', '45'), ('left', '80'), ('right', '80')]:
                mar = OxmlElement(f'w:{edge}')
                mar.set(qn('w:w'), value)
                mar.set(qn('w:type'), 'dxa')
                tcMar.append(mar)
            tcPr.append(tcMar)
            # 三线表：表头行底部加细线（0.75pt）
            if ri == 0:
                tcBorders = OxmlElement('w:tcBorders')
                btm = OxmlElement('w:bottom')
                btm.set(qn('w:val'), 'single')
                btm.set(qn('w:sz'), '6')
                btm.set(qn('w:space'), '0')
                btm.set(qn('w:color'), '000000')
                tcBorders.append(btm)
                tcPr.append(tcBorders)
            tc.append(tcPr)
            p = OxmlElement('w:p')
            # 段落属性：水平居中 + 段前段后间距为0 + 单倍行距
            pPr = OxmlElement('w:pPr')
            pJc = OxmlElement('w:jc')
            pJc.set(qn('w:val'), 'center')
            pPr.append(pJc)
            pSpacing = OxmlElement('w:spacing')
            pSpacing.set(qn('w:before'), '0')
            pSpacing.set(qn('w:after'), '0')
            pSpacing.set(qn('w:line'), '220')
            pSpacing.set(qn('w:lineRule'), 'auto')
            pPr.append(pSpacing)
            # 清除所有缩进（firstLine + firstLineChars 都要清）
            pInd = OxmlElement('w:ind')
            pInd.set(qn('w:firstLine'), '0')
            pInd.set(qn('w:firstLineChars'), '0')
            pInd.set(qn('w:left'), '0')
            pInd.set(qn('w:right'), '0')
            pPr.append(pInd)
            p.append(pPr)
            # 文本 — 字体：宋体(中文) + Times New Roman(英文)，小五号(9pt)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            # 字体
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), '宋体')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rPr.append(rFonts)
            # 字号 小五号=9pt=18半磅
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '18')
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '18')
            rPr.append(szCs)
            # 表头加粗
            if ri == 0:
                b = OxmlElement('w:b')
                rPr.append(b)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = str(row_data[ci]) if ci < len(row_data) else ''
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    # 插入到段落后面
    after_para._p.addnext(tbl)


def _verify_citations(docx_path):
    """检查正文引用序号连续、位置符合句末标注习惯。"""
    from docx import Document

    doc = Document(docx_path)
    body_parts = []
    in_refs = False
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        compact = text.replace(" ", "").replace("　", "")
        if compact == "参考文献":
            in_refs = True
            continue
        if not in_refs:
            body_parts.append(text)
    body_text = "\n".join(body_parts)
    nums = [int(x) for x in re.findall(r"\[(\d+)\]", body_text)]
    errors = []
    if nums:
        first_seen = []
        for n in nums:
            if n not in first_seen:
                first_seen.append(n)
        expected = list(range(1, max(first_seen) + 1))
        if first_seen != expected:
            errors.append(f"正文参考文献首次出现顺序不连续: {first_seen}")
    bad_positions = re.findall(r"[。；，,.;!?]\s*\[\d+\]", body_text)
    if bad_positions:
        errors.append("存在标点后引用标注，应调整为引用标注在句号或逗号前")
    if "可行性分析" in body_text:
        errors.append("正文仍包含“可行性分析”")
    return errors


def _patch_auto_toc(docx_path):
    """把模板中的静态目录替换为真正的 Word TOC 域，并刷新其显示结果。"""
    import shutil
    import socket
    import subprocess
    import tempfile
    import time
    import zipfile
    from lxml import etree

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    NS = {"w": W}

    def qn(name):
        return f"{{{W}}}{name.split(':', 1)[1] if ':' in name else name}"

    def para_text(p):
        return "".join(p.xpath(".//w:t/text()", namespaces=NS)).strip()

    def make_toc_field_para():
        p = etree.Element(qn("w:p"))
        r = etree.SubElement(p, qn("w:r"))
        fld_begin = etree.SubElement(r, qn("w:fldChar"))
        fld_begin.set(qn("w:fldCharType"), "begin")
        fld_begin.set(qn("w:dirty"), "true")

        r = etree.SubElement(p, qn("w:r"))
        instr = etree.SubElement(r, qn("w:instrText"))
        instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        instr.text = ' TOC \\o "1-3" \\h \\z \\u '

        r = etree.SubElement(p, qn("w:r"))
        fld_sep = etree.SubElement(r, qn("w:fldChar"))
        fld_sep.set(qn("w:fldCharType"), "separate")

        r = etree.SubElement(p, qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.text = "目录更新中"

        r = etree.SubElement(p, qn("w:r"))
        fld_end = etree.SubElement(r, qn("w:fldChar"))
        fld_end.set(qn("w:fldCharType"), "end")
        return p

    def clone_section_for_front_matter(root):
        sects = root.xpath(".//w:sectPr", namespaces=NS)
        if not sects:
            return etree.Element(qn("w:sectPr"))
        return etree.fromstring(etree.tostring(sects[-1]))

    def set_pg_num_type(sect, fmt=None, start=None):
        pg = sect.find("w:pgNumType", NS)
        if pg is None:
            pg = etree.Element(qn("w:pgNumType"))
            sect.insert(0, pg)
        if fmt:
            pg.set(qn("w:fmt"), fmt)
        elif qn("w:fmt") in pg.attrib:
            del pg.attrib[qn("w:fmt")]
        if start is not None:
            pg.set(qn("w:start"), str(start))

    def make_section_break_para(front_sect):
        p = etree.Element(qn("w:p"))
        ppr = etree.SubElement(p, qn("w:pPr"))
        front_sect = etree.fromstring(etree.tostring(front_sect))
        sect_type = front_sect.find("w:type", NS)
        if sect_type is None:
            sect_type = etree.Element(qn("w:type"))
            front_sect.insert(0, sect_type)
        sect_type.set(qn("w:val"), "nextPage")
        ppr.append(front_sect)
        return p

    def refresh_toc_with_libreoffice(path):
        soffice = shutil.which("soffice") or "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        lo_python_candidates = [
            "/Applications/LibreOffice.app/Contents/Resources/python",
            shutil.which("python3"),
        ]
        lo_python = next((p for p in lo_python_candidates if p and os.path.exists(p)), None)
        if not soffice or not os.path.exists(soffice) or not lo_python:
            print("  目录域已插入；未找到 LibreOffice 字段刷新环境")
            return False

        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            s.bind(("127.0.0.1", 0))
            port = s.getsockname()[1]

        profile = tempfile.mkdtemp(prefix="lo_toc_profile_")
        proc = subprocess.Popen([
            soffice,
            f"-env:UserInstallation=file://{profile}",
            "--headless",
            "--invisible",
            "--norestore",
            f"--accept=socket,host=127.0.0.1,port={port};urp;StarOffice.ComponentContext",
        ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)

        try:
            time.sleep(1.5)
            script = f"""
import sys
import time
import uno
from com.sun.star.beans import PropertyValue
from com.sun.star.connection import NoConnectException

def prop(name, value):
    p = PropertyValue()
    p.Name = name
    p.Value = value
    return p

local_ctx = uno.getComponentContext()
resolver = local_ctx.ServiceManager.createInstanceWithContext(
    'com.sun.star.bridge.UnoUrlResolver', local_ctx
)
for _ in range(30):
    try:
        ctx = resolver.resolve(
            'uno:socket,host=127.0.0.1,port={port};urp;StarOffice.ComponentContext'
        )
        break
    except NoConnectException:
        time.sleep(0.5)
else:
    raise RuntimeError('LibreOffice UNO connection failed')

smgr = ctx.ServiceManager
desktop = smgr.createInstanceWithContext('com.sun.star.frame.Desktop', ctx)
url = uno.systemPathToFileUrl({os.path.abspath(path)!r})
doc = desktop.loadComponentFromURL(url, '_blank', 0, (
    prop('Hidden', True),
    prop('ReadOnly', False),
))
indexes = doc.getDocumentIndexes()
for i in range(indexes.getCount()):
    indexes.getByIndex(i).update()
doc.getTextFields().refresh()
doc.refresh()
doc.store()
doc.close(True)
"""
            run = subprocess.run([lo_python, "-c", script],
                                 stdout=subprocess.PIPE,
                                 stderr=subprocess.PIPE,
                                 text=True,
                                 timeout=90)
            if run.returncode != 0:
                print("  目录域已插入；LibreOffice 刷新失败")
                if run.stderr.strip():
                    print(run.stderr.strip())
                return False
            print("  目录已替换为 Word TOC 域并刷新")
            return True
        finally:
            proc.terminate()
            try:
                proc.wait(timeout=5)
            except subprocess.TimeoutExpired:
                proc.kill()
            shutil.rmtree(profile, ignore_errors=True)

    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {info.filename: zin.read(info.filename) for info in zin.infolist()}

    root = etree.fromstring(files["word/document.xml"])
    body = root.find("w:body", NS)
    children = list(body)
    paras = [el for el in children if el.tag == qn("w:p")]

    toc_title = None
    body_start = None
    for p in paras:
        text = para_text(p).replace(" ", "").replace("　", "")
        if text == "目录" and toc_title is None:
            toc_title = p
        elif toc_title is not None and re.match(r"^1绪论$", text):
            body_start = p
            break
    if toc_title is None or body_start is None:
        return

    front_sect = clone_section_for_front_matter(root)
    set_pg_num_type(front_sect, fmt="upperRoman", start=1)

    final_sect = body.find("w:sectPr", NS)
    if final_sect is not None:
        set_pg_num_type(final_sect, fmt="decimal", start=1)

    children = list(body)
    start = children.index(toc_title) + 1
    stop = children.index(body_start)
    for elem in children[start:stop]:
        body.remove(elem)
    body.insert(start, make_toc_field_para())
    body.insert(start + 1, make_section_break_para(front_sect))

    settings = etree.fromstring(files["word/settings.xml"])
    update = settings.find("w:updateFields", NS)
    if update is None:
        update = etree.Element(qn("w:updateFields"))
        settings.insert(0, update)
    update.set(qn("w:val"), "true")

    files["word/document.xml"] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    files["word/settings.xml"] = etree.tostring(settings, xml_declaration=True, encoding="UTF-8", standalone=True)

    tmp = tempfile.mktemp(suffix=".docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)
    shutil.move(tmp, docx_path)
    refresh_toc_with_libreoffice(docx_path)


def _patch_section_headers(docx_path):
    """为摘要、目录、各章和尾部部分建立独立页眉。"""
    import shutil
    import socket
    import subprocess
    import tempfile
    import time
    import zipfile
    from lxml import etree

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    REL = "http://schemas.openxmlformats.org/package/2006/relationships"
    CT = "http://schemas.openxmlformats.org/package/2006/content-types"
    NS = {"w": W, "r": R}

    def qn(name):
        return f"{{{W}}}{name.split(':', 1)[1] if ':' in name else name}"

    def rel_qn(name):
        return f"{{{REL}}}{name}"

    def ct_qn(name):
        return f"{{{CT}}}{name}"

    def para_text(p):
        return "".join(p.xpath(".//w:t/text()", namespaces=NS)).strip()

    def compact(text):
        return (text or "").strip().replace(" ", "").replace("　", "")

    def header_title(text):
        c = compact(text)
        mapping = {
            "摘要": "摘要",
            "Abstract": "Abstract",
            "目录": "目录",
            "总结": "总结",
            "总结与展望": "总结与展望",
            "结论": "结论",
            "致谢": "致谢",
            "参考文献": "参考文献",
        }
        return mapping.get(c, (text or "").strip())

    def is_section_start(text):
        c = compact(text)
        if c in {"摘要", "Abstract", "目录", "总结", "总结与展望", "结论", "致谢", "参考文献"}:
            return True
        return bool(re.fullmatch(r"[1-6]\s+\S.*", (text or "").strip()))

    def remove_page_break_before(p):
        ppr = p.find("w:pPr", NS)
        if ppr is None:
            return
        for el in list(ppr.findall("w:pageBreakBefore", NS)):
            ppr.remove(el)

    def next_rel_id(rels_root):
        max_id = 0
        for rel in rels_root:
            rid = rel.get("Id", "")
            m = re.fullmatch(r"rId(\d+)", rid)
            if m:
                max_id = max(max_id, int(m.group(1)))
        max_id += 1
        return f"rId{max_id}"

    def next_header_name(files):
        max_num = 0
        for name in files:
            m = re.fullmatch(r"word/header(\d+)\.xml", name)
            if m:
                max_num = max(max_num, int(m.group(1)))
        return f"word/header{max_num + 1}.xml"

    def next_footer_name(files):
        max_num = 0
        for name in files:
            m = re.fullmatch(r"word/footer(\d+)\.xml", name)
            if m:
                max_num = max(max_num, int(m.group(1)))
        return f"word/footer{max_num + 1}.xml"

    def make_header_xml(template_xml, title):
        root = etree.fromstring(template_xml)
        p = root.find("w:p", NS)
        if p is None:
            p = etree.SubElement(root, qn("w:p"))
        sample_rpr = None
        for r in p.findall("w:r", NS):
            text = "".join(r.xpath(".//w:t/text()", namespaces=NS)).strip()
            rpr = r.find("w:rPr", NS)
            if text and rpr is not None:
                sample_rpr = etree.fromstring(etree.tostring(rpr))
                break
        for r in list(p.findall("w:r", NS)):
            p.remove(r)
        r = etree.SubElement(p, qn("w:r"))
        if sample_rpr is not None:
            rpr = etree.fromstring(etree.tostring(sample_rpr))
        else:
            rpr = etree.SubElement(r, qn("w:rPr"))
        rfonts = rpr.find("w:rFonts", NS)
        if rfonts is None:
            rfonts = etree.Element(qn("w:rFonts"))
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:eastAsia"), "宋体")
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        r.append(rpr)
        t = etree.SubElement(r, qn("w:t"))
        t.text = title
        return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

    def ensure_header_part(files, rels_root, ct_root, template_xml, title, cache):
        if title in cache:
            return cache[title]
        header_name = next_header_name(files)
        files[header_name] = make_header_xml(template_xml, title)

        rid = next_rel_id(rels_root)
        rel = etree.SubElement(rels_root, rel_qn("Relationship"))
        rel.set("Id", rid)
        rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/header")
        rel.set("Target", header_name.replace("word/", ""))

        part_name = "/" + header_name
        exists = any(
            el.get("PartName") == part_name
            for el in ct_root.findall(f"{{{CT}}}Override")
        )
        if not exists:
            override = etree.SubElement(ct_root, ct_qn("Override"))
            override.set("PartName", part_name)
            override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml")

        cache[title] = rid
        return rid

    def make_page_footer_xml():
        ftr = etree.Element(qn("w:ftr"), nsmap={"w": W, "r": R})
        p = etree.SubElement(ftr, qn("w:p"))
        ppr = etree.SubElement(p, qn("w:pPr"))
        jc = etree.SubElement(ppr, qn("w:jc"))
        jc.set(qn("w:val"), "center")
        rpr = etree.SubElement(ppr, qn("w:rPr"))
        fonts = etree.SubElement(rpr, qn("w:rFonts"))
        fonts.set(qn("w:ascii"), "Times New Roman")
        fonts.set(qn("w:hAnsi"), "Times New Roman")
        fonts.set(qn("w:eastAsia"), "宋体")
        sz = etree.SubElement(rpr, qn("w:sz"))
        sz.set(qn("w:val"), "20")
        szcs = etree.SubElement(rpr, qn("w:szCs"))
        szcs.set(qn("w:val"), "20")

        for kind, value in [
            ("begin", None),
            ("instr", " PAGE "),
            ("separate", None),
            ("text", "1"),
            ("end", None),
        ]:
            r = etree.SubElement(p, qn("w:r"))
            rrpr = etree.SubElement(r, qn("w:rPr"))
            rfonts = etree.SubElement(rrpr, qn("w:rFonts"))
            rfonts.set(qn("w:ascii"), "Times New Roman")
            rfonts.set(qn("w:hAnsi"), "Times New Roman")
            rfonts.set(qn("w:eastAsia"), "宋体")
            rsz = etree.SubElement(rrpr, qn("w:sz"))
            rsz.set(qn("w:val"), "20")
            rszcs = etree.SubElement(rrpr, qn("w:szCs"))
            rszcs.set(qn("w:val"), "20")
            if kind in {"begin", "separate", "end"}:
                fld = etree.SubElement(r, qn("w:fldChar"))
                fld.set(qn("w:fldCharType"), kind)
            elif kind == "instr":
                instr = etree.SubElement(r, qn("w:instrText"))
                instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                instr.text = value
            else:
                t = etree.SubElement(r, qn("w:t"))
                t.text = value
        return etree.tostring(ftr, xml_declaration=True, encoding="UTF-8", standalone=True)

    def ensure_page_footer_part(files, rels_root, ct_root):
        footer_name = next_footer_name(files)
        files[footer_name] = make_page_footer_xml()

        rid = next_rel_id(rels_root)
        rel = etree.SubElement(rels_root, rel_qn("Relationship"))
        rel.set("Id", rid)
        rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer")
        rel.set("Target", footer_name.replace("word/", ""))

        part_name = "/" + footer_name
        exists = any(
            el.get("PartName") == part_name
            for el in ct_root.findall(f"{{{CT}}}Override")
        )
        if not exists:
            override = etree.SubElement(ct_root, ct_qn("Override"))
            override.set("PartName", part_name)
            override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml")
        return rid

    def footer_refs_for_rid(rid):
        refs = []
        for ref_type in ["even", "default", "first"]:
            fr = etree.Element(qn("w:footerReference"))
            fr.set(qn("w:type"), ref_type)
            fr.set(f"{{{R}}}id", rid)
            refs.append(fr)
        return refs

    def footer_refs_from(sect):
        return [
            etree.fromstring(etree.tostring(el))
            for el in sect.findall("w:footerReference", NS)
        ]

    def clean_sect_clone(sect):
        new_sect = etree.fromstring(etree.tostring(sect))
        for el in list(new_sect):
            if el.tag in {
                qn("w:headerReference"),
                qn("w:footerReference"),
                qn("w:type"),
                qn("w:pgNumType"),
            }:
                new_sect.remove(el)
        return new_sect

    def make_sect(base_sect, footer_refs, header_rid, fmt, start=None):
        sect = clean_sect_clone(base_sect)
        pg_mar = sect.find("w:pgMar", NS)
        if pg_mar is not None and int(pg_mar.get(qn("w:footer"), "0") or 0) < 567:
            pg_mar.set(qn("w:footer"), "567")
        insert_at = 0
        for ref_type in ["even", "default", "first"]:
            hr = etree.Element(qn("w:headerReference"))
            hr.set(qn("w:type"), ref_type)
            hr.set(f"{{{R}}}id", header_rid)
            sect.insert(insert_at, hr)
            insert_at += 1
        for fr in footer_refs:
            sect.insert(insert_at, etree.fromstring(etree.tostring(fr)))
            insert_at += 1
        stype = etree.Element(qn("w:type"))
        stype.set(qn("w:val"), "nextPage")
        sect.insert(insert_at, stype)
        insert_at += 1
        pg = etree.Element(qn("w:pgNumType"))
        if fmt:
            pg.set(qn("w:fmt"), fmt)
        if start is not None:
            pg.set(qn("w:start"), str(start))
        sect.insert(insert_at, pg)
        return sect

    def section_break_para(sect):
        p = etree.Element(qn("w:p"))
        ppr = etree.SubElement(p, qn("w:pPr"))
        ppr.append(sect)
        return p

    def has_sect_pr(p):
        return p is not None and p.tag == qn("w:p") and p.find("w:pPr/w:sectPr", NS) is not None

    def last_sdt_paragraph(sdt):
        if sdt is None or sdt.tag != qn("w:sdt"):
            return None
        paras = sdt.findall(".//w:sdtContent/w:p", NS)
        return paras[-1] if paras else None

    def attach_sect_to_paragraph(p, sect):
        ppr = p.find("w:pPr", NS)
        if ppr is None:
            ppr = etree.Element(qn("w:pPr"))
            p.insert(0, ppr)
        old = ppr.find("w:sectPr", NS)
        if old is not None:
            ppr.remove(old)
        ppr.append(sect)

    def set_break_sect(before_para, sect):
        body = before_para.getparent()
        children = list(body)
        idx = children.index(before_para)
        prev = children[idx - 1] if idx > 0 else None
        prev2 = children[idx - 2] if idx > 1 else None
        if has_sect_pr(prev) and not para_text(prev):
            sdt_para = last_sdt_paragraph(prev2)
            if sdt_para is not None:
                body.remove(prev)
                attach_sect_to_paragraph(sdt_para, sect)
                return
            attach_sect_to_paragraph(prev, sect)
            return
        sdt_para = last_sdt_paragraph(prev)
        if sdt_para is not None:
            attach_sect_to_paragraph(sdt_para, sect)
            return
        if prev is not None and prev.tag == qn("w:p"):
            attach_sect_to_paragraph(prev, sect)
            return
        body.insert(idx, section_break_para(sect))

    def refresh_fields(path):
        soffice = shutil.which("soffice") or "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        lo_python_candidates = [
            "/Applications/LibreOffice.app/Contents/Resources/python",
            shutil.which("python3"),
        ]
        lo_python = next((p for p in lo_python_candidates if p and os.path.exists(p)), None)
        if not soffice or not os.path.exists(soffice) or not lo_python:
            return False
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            s.bind(("127.0.0.1", 0))
            port = s.getsockname()[1]
        profile = tempfile.mkdtemp(prefix="lo_header_profile_")
        proc = subprocess.Popen([
            soffice,
            f"-env:UserInstallation=file://{profile}",
            "--headless",
            "--invisible",
            "--norestore",
            f"--accept=socket,host=127.0.0.1,port={port};urp;StarOffice.ComponentContext",
        ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        try:
            time.sleep(1.5)
            script = f"""
import time
import uno
from com.sun.star.beans import PropertyValue
from com.sun.star.connection import NoConnectException

def prop(name, value):
    p = PropertyValue()
    p.Name = name
    p.Value = value
    return p

local_ctx = uno.getComponentContext()
resolver = local_ctx.ServiceManager.createInstanceWithContext(
    'com.sun.star.bridge.UnoUrlResolver', local_ctx
)
for _ in range(30):
    try:
        ctx = resolver.resolve(
            'uno:socket,host=127.0.0.1,port={port};urp;StarOffice.ComponentContext'
        )
        break
    except NoConnectException:
        time.sleep(0.5)
else:
    raise RuntimeError('LibreOffice UNO connection failed')

smgr = ctx.ServiceManager
desktop = smgr.createInstanceWithContext('com.sun.star.frame.Desktop', ctx)
url = uno.systemPathToFileUrl({os.path.abspath(path)!r})
doc = desktop.loadComponentFromURL(url, '_blank', 0, (
    prop('Hidden', True),
    prop('ReadOnly', False),
))
indexes = doc.getDocumentIndexes()
for i in range(indexes.getCount()):
    indexes.getByIndex(i).update()
doc.getTextFields().refresh()
doc.refresh()
doc.store()
doc.close(True)
"""
            run = subprocess.run([lo_python, "-c", script],
                                 stdout=subprocess.PIPE,
                                 stderr=subprocess.PIPE,
                                 text=True,
                                 timeout=90)
            return run.returncode == 0
        finally:
            proc.terminate()
            try:
                proc.wait(timeout=5)
            except subprocess.TimeoutExpired:
                proc.kill()
            shutil.rmtree(profile, ignore_errors=True)

    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {info.filename: zin.read(info.filename) for info in zin.infolist()}

    root = etree.fromstring(files["word/document.xml"])
    rels_root = etree.fromstring(files["word/_rels/document.xml.rels"])
    ct_root = etree.fromstring(files["[Content_Types].xml"])
    body = root.find("w:body", NS)

    header_template_name = next(
        (name for name, data in files.items()
         if re.fullmatch(r"word/header\d+\.xml", name)
         and "本科毕业设计".encode("utf-8") in data),
        "word/header2.xml",
    )
    header_template = files[header_template_name]
    sects = root.xpath(".//w:sectPr", namespaces=NS)
    base_front = next(
        (s for s in sects
         if (s.find("w:pgNumType", NS) is not None
             and s.find("w:pgNumType", NS).get(qn("w:fmt")) == "upperRoman")),
        sects[-1],
    )
    base_body = next(
        (s for s in sects
         if (s.find("w:pgNumType", NS) is not None
             and s.find("w:pgNumType", NS).get(qn("w:fmt")) == "decimal"
             and s.find("w:headerReference", NS) is not None)),
        sects[-1],
    )
    footer_base = next(
        (s for s in sects if s.findall("w:footerReference", NS)),
        base_body,
    )
    page_footers = footer_refs_from(footer_base)
    if not page_footers:
        page_footers = footer_refs_for_rid(
            ensure_page_footer_part(files, rels_root, ct_root)
        )

    starts = []
    for p in body.findall("w:p", NS):
        text = para_text(p)
        if is_section_start(text):
            starts.append((p, header_title(text), compact(text)))
            remove_page_break_before(p)

    if len(starts) < 4:
        return

    header_cache = {}
    for idx, (start_p, title, key) in enumerate(starts):
        is_front = key in {"摘要", "Abstract", "目录"}
        base = base_front if is_front else base_body
        footers = page_footers
        fmt = "upperRoman" if is_front else "decimal"
        start_num = None
        if key == "摘要":
            start_num = 1
        elif key == "1绪论":
            start_num = 1
        rid = ensure_header_part(files, rels_root, ct_root, header_template, title, header_cache)
        sect = make_sect(base, footers, rid, fmt, start_num)
        if idx + 1 < len(starts):
            set_break_sect(starts[idx + 1][0], sect)
        else:
            final_sect = body.find("w:sectPr", NS)
            if final_sect is not None:
                parent = final_sect.getparent()
                parent.remove(final_sect)
                parent.append(sect)

    files["word/document.xml"] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    files["word/_rels/document.xml.rels"] = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    files["[Content_Types].xml"] = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    tmp = tempfile.mktemp(suffix=".docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)
    shutil.move(tmp, docx_path)
    refresh_fields(docx_path)
    print("  页眉已按各部分更新")


if __name__ == "__main__":
    main()
