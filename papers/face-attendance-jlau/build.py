from __future__ import annotations

import json
import re
import shutil
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

HERE = Path(__file__).resolve().parent
SRC_TEMPLATE = Path("/Users/a136/vs/45425/要求66/吉林农业大学本科生毕业论文参考模板.docx")
OUTPUT = HERE / "基于百度AI人脸识别的考勤系统设计与实现.docx"


def load_json(name: str):
    return json.loads((HERE / name).read_text(encoding="utf-8"))


def set_run_font(run, east="宋体", ascii_font="Times New Roman", size=12, bold=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_exact_20(p):
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def set_outline_level(p, level: int):
    p_pr = p._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))


def clear_runs(p):
    for r in p.runs:
        r.text = ""


def remove_all_body_content(doc: Document):
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def remove_all_but_cover(doc: Document):
    body = doc.element.body
    kept_first_table = False
    for child in list(body):
        if child.tag == qn("w:tbl") and not kept_first_table:
            kept_first_table = True
            continue
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def replace_cover_field(paragraph, label: str, value: str):
    found = False
    wrote = False
    for r in paragraph.runs:
        if label in r.text:
            found = True
        elif found and r.font.underline and not wrote:
            r.text = value
            wrote = True
        elif found and r.font.underline:
            r.text = ""


def fill_cover(doc: Document, meta: dict):
    cell = doc.tables[0].cell(0, 0)
    cps = cell.paragraphs
    replace_cover_field(cps[10], "目：", meta["title_zh"])
    replace_cover_field(cps[12], "名：", meta.get("name", ""))
    replace_cover_field(cps[12], "号：", meta.get("student_id", ""))
    replace_cover_field(cps[14], "院", meta.get("college", ""))
    replace_cover_field(cps[14], "业", meta.get("major", ""))
    replace_cover_field(cps[16], "教师：", meta.get("advisor", ""))
    replace_cover_field(cps[16], "称：", meta.get("advisor_title", ""))
    if len(cps) > 23 and cps[23].runs:
        cps[23].runs[0].text = f"{meta.get('year','2026')}年    {meta.get('month','4')}月    {meta.get('day','25')}日"
        for r in cps[23].runs[1:]:
            r.text = ""


def set_margins(section):
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)


def _add_field(paragraph, instr_text: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    run._r.append(instr)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)
    text = OxmlElement("w:t")
    text.text = "1"
    run._r.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)
    return run


def add_page_number(section, roman=False):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    clear_runs(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = _add_field(p, "PAGE \\* ROMAN" if roman else "PAGE")
    set_run_font(run, size=10.5)
    pg_num = section._sectPr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        section._sectPr.append(pg_num)
    pg_num.set(qn("w:start"), "1")
    pg_num.set(qn("w:fmt"), "upperRoman" if roman else "decimal")


def add_new_page_section(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(doc.sections[-1])
    pg_num = doc.sections[-1]._sectPr.find(qn("w:pgNumType"))
    if pg_num is not None:
        doc.sections[-1]._sectPr.remove(pg_num)


def add_title(doc, text, level=1, page_break=False):
    if page_break:
        add_new_page_section(doc)
    p = doc.add_paragraph()
    try:
        p.style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(level, "Normal")
    except Exception:
        pass
    p.paragraph_format.page_break_before = False
    set_exact_20(p)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, "黑体", "Times New Roman", 16, True)
        set_outline_level(p, 0)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_run_font(r, "黑体", "Times New Roman", 14, True)
        set_outline_level(p, 1)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_run_font(r, "宋体", "Times New Roman", 12, True)
        set_outline_level(p, 2)
    return p


def add_para(doc, text, indent=True):
    p = doc.add_paragraph()
    set_exact_20(p)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    parts = re.split(r"(\[[0-9,，\-—]+\])", str(text))
    for part in parts:
        if not part:
            continue
        r = p.add_run(part)
        set_run_font(r, "宋体", "Times New Roman", 12)
        if re.fullmatch(r"\[[0-9,，\-—]+\]", part):
            r.font.superscript = True
    return p


def add_center_para(doc, text, east="宋体", size=12, bold=False):
    p = doc.add_paragraph()
    set_exact_20(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, east, "Times New Roman", size, bold)
    return p


def add_picture_para(doc):
    """Pictures need natural line height, otherwise Word clips inline images."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def add_toc_entry(doc, title, page, indent_cm=0.0, bold=False):
    p = doc.add_paragraph()
    set_exact_20(p)
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(f"{title}\t{page}")
    set_run_font(r, "宋体", "Times New Roman", 12, bold)


def add_toc(doc, chapters):
    add_new_page_section(doc)
    p = add_center_para(doc, "目  录", "黑体", 16, True)
    p.paragraph_format.space_after = Pt(8)
    estimated = {
        "第1章 前言": 1,
        "第2章 相关技术介绍": 5,
        "第3章 系统需求分析": 8,
        "第4章 系统总体设计": 13,
        "第5章 系统详细设计与实现": 23,
        "第6章 系统测试": 32,
        "第7章 总结与展望": 36,
        "参考文献": 38,
        "附  录": 39,
        "致  谢": 40,
    }
    section_offsets = {
        "第1章 前言": [0, 0, 1, 1, 2, 3],
        "第2章 相关技术介绍": [0, 0, 0, 0, 1, 2],
        "第3章 系统需求分析": [0, 0, 2, 2, 3, 4],
        "第4章 系统总体设计": [0, 1, 1, 2, 3, 4, 8, 9],
        "第5章 系统详细设计与实现": [0, 1, 2, 3, 3, 4, 5, 5, 6, 8],
        "第6章 系统测试": [0, 0, 1, 2, 2, 2, 3],
        "第7章 总结与展望": [0, 0, 0],
    }
    for chapter in chapters:
        base = estimated.get(chapter["title"], 1)
        add_toc_entry(doc, chapter["title"], base, bold=True)
        h2_idx = 0
        offsets = section_offsets.get(chapter["title"], [])
        for block in chapter.get("blocks", []):
            if block.get("type") == "h2":
                page = base + (offsets[h2_idx] if h2_idx < len(offsets) else h2_idx)
                add_toc_entry(doc, block["text"], page, indent_cm=0.55)
                h2_idx += 1
    add_toc_entry(doc, "参考文献", estimated["参考文献"], bold=True)
    add_toc_entry(doc, "附  录", estimated["附  录"], bold=True)
    add_toc_entry(doc, "致  谢", estimated["致  谢"], bold=True)


def _add_toc_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run._r.append(sep)
    t = OxmlElement("w:t")
    t.text = "目录更新后显示"
    run._r.append(t)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.first_child_found_in("w:tcBorders")
    if old is not None:
        tc_pr.remove(old)
    borders = OxmlElement("w:tcBorders")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        if val is None:
            continue
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val.get("val", "single"))
        el.set(qn("w:sz"), val.get("sz", "4"))
        el.set(qn("w:color"), val.get("color", "000000"))
        el.set(qn("w:space"), "0")
        borders.append(el)
    tc_pr.append(borders)


def set_cell_margins(cell, top=80, bottom=80, left=80, right=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.first_child_found_in("w:tcMar")
    if old is not None:
        tc_pr.remove(old)
    mar = OxmlElement("w:tcMar")
    for edge, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tc_pr.append(mar)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_grid(table, widths_cm):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths_cm) * 567)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 567)))
        grid.append(col)
    tbl.insert(1, grid)


def add_table(doc, block):
    caption = block["caption"]
    headers = block["headers"]
    rows = block["rows"]
    widths = block.get("widths") or ([13.8 / len(headers)] * len(headers))
    add_center_para(doc, caption, size=10.5, bold=True)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_grid(table, widths)
    table.autofit = False
    line_thick = {"val": "single", "sz": "12", "color": "000000"}
    line_thin = {"val": "single", "sz": "6", "color": "000000"}

    for r_idx, row in enumerate([headers] + rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[c_idx])
            top = line_thick if r_idx == 0 else None
            bottom = line_thin if r_idx == 0 else (line_thick if r_idx == len(rows) else None)
            set_cell_border(cell, top=top, bottom=bottom)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx < 2 or len(str(value)) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = Pt(16)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, "宋体", "Times New Roman", 9.5, bold=(r_idx == 0))
    add_para(doc, "", indent=False)


def render_table_image(caption, headers, rows, widths):
    from PIL import Image, ImageDraw, ImageFont

    out_dir = HERE / "generated_tables"
    out_dir.mkdir(exist_ok=True)
    idx = len(list(out_dir.glob("table_*.png"))) + 1
    path = out_dir / f"table_{idx:02d}.png"
    total = sum(widths)
    px_width = 1800
    margin_x = 40
    table_width = px_width - margin_x * 2
    col_widths = [max(120, int(table_width * (w / total))) for w in widths]
    delta = table_width - sum(col_widths)
    col_widths[-1] += delta

    font_path = "/System/Library/Fonts/Supplemental/Songti.ttc"
    bold_path = "/System/Library/Fonts/STHeiti Medium.ttc"
    font = ImageFont.truetype(font_path, 34)
    header_font = ImageFont.truetype(bold_path, 34)

    def wrap(text, width_px, fnt):
        s = str(text)
        if not s:
            return [""]
        lines, buf = [], ""
        for ch in s:
            trial = buf + ch
            if ImageDraw.Draw(Image.new("RGB", (1, 1))).textbbox((0, 0), trial, font=fnt)[2] <= width_px - 28:
                buf = trial
            else:
                if buf:
                    lines.append(buf)
                buf = ch
        if buf:
            lines.append(buf)
        return lines or [""]

    all_rows = [headers] + rows
    row_lines = []
    row_heights = []
    for r_idx, row in enumerate(all_rows):
        fnt = header_font if r_idx == 0 else font
        wrapped = [wrap(cell, col_widths[i], fnt) for i, cell in enumerate(row)]
        row_lines.append(wrapped)
        max_lines = max(len(x) for x in wrapped)
        row_heights.append(max(62, 20 + max_lines * 42))

    height = 36 + sum(row_heights) + 36
    im = Image.new("RGB", (px_width, height), "white")
    draw = ImageDraw.Draw(im)
    x0 = margin_x
    y = 24
    draw.line((x0, y, x0 + table_width, y), fill="black", width=5)
    y += 12

    for r_idx, wrapped in enumerate(row_lines):
        h = row_heights[r_idx]
        x = x0
        fnt = header_font if r_idx == 0 else font
        for c_idx, lines in enumerate(wrapped):
            line_h = 42
            text_h = len(lines) * line_h
            yy = y + max(0, (h - text_h) // 2)
            for line in lines:
                bbox = draw.textbbox((0, 0), line, font=fnt)
                tw = bbox[2] - bbox[0]
                if c_idx == len(wrapped) - 1 and len(line) > 9:
                    xx = x + 16
                else:
                    xx = x + max(12, (col_widths[c_idx] - tw) // 2)
                draw.text((xx, yy), line, fill="black", font=fnt)
                yy += line_h
            x += col_widths[c_idx]
        y += h
        if r_idx == 0:
            draw.line((x0, y, x0 + table_width, y), fill="black", width=2)
    draw.line((x0, y + 6, x0 + table_width, y + 6), fill="black", width=5)
    im.save(path, "PNG")
    return path


def add_image(doc, block):
    path = Path(block["path"])
    if not path.is_absolute():
        path = HERE / path
    if not path.exists():
        return
    p = add_picture_para(doc)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(block.get("width_cm", 12.0)))
    add_center_para(doc, block["caption"], size=10.5, bold=True)
    if block.get("caption_en"):
        add_center_para(doc, block["caption_en"], size=10.5)


def add_code(doc, block):
    if block.get("caption"):
        add_center_para(doc, block["caption"], size=10.5, bold=True)
    for line in block["code"].splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(14)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(line)
        set_run_font(r, "Consolas", "Consolas", 9)


def patch_update_fields(docx_path):
    replacements = {}
    with zipfile.ZipFile(docx_path, "r") as zf:
        for info in zf.infolist():
            if info.filename == "word/settings.xml":
                xml = zf.read(info.filename).decode("utf-8")
                if "w:updateFields" not in xml:
                    xml = xml.replace("</w:settings>", '<w:updateFields w:val="true"/></w:settings>')
                replacements[info.filename] = xml.encode("utf-8")
        tmp = Path(str(docx_path) + ".tmp")
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as out:
            seen = set()
            for info in zf.infolist():
                if info.filename in seen:
                    continue
                seen.add(info.filename)
                out.writestr(info, replacements.get(info.filename, zf.read(info.filename)))
    tmp.replace(docx_path)


def repair_table_compat(docx_path):
    """Normalize newer OOXML tokens that some renderers do not accept.

    LibreOffice may write w:jc='start' and table cell margins as
    w:start/w:end. Word/WPS handle them, but artifact-tool expects the older
    left/right token set, so normalize those after every build/update pass.
    """
    replacements = {}
    with zipfile.ZipFile(docx_path, "r") as zf:
        for info in zf.infolist():
            if info.filename == "word/document.xml":
                xml = zf.read(info.filename).decode("utf-8")
                xml = xml.replace('w:val="start"', 'w:val="left"')
                xml = xml.replace("<w:start ", "<w:left ")
                xml = xml.replace("</w:start>", "</w:left>")
                xml = xml.replace("<w:end ", "<w:right ")
                xml = xml.replace("</w:end>", "</w:right>")
                replacements[info.filename] = xml.encode("utf-8")
        tmp = Path(str(docx_path) + ".tmp")
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as out:
            seen = set()
            for info in zf.infolist():
                if info.filename in seen:
                    continue
                seen.add(info.filename)
                out.writestr(info, replacements.get(info.filename, zf.read(info.filename)))
    tmp.replace(docx_path)


def add_front_matter(doc, meta, chapters):
    doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(doc.sections[-1])
    add_page_number(doc.sections[-1], roman=True)
    add_center_para(doc, meta["title_zh"], "黑体", 16, True)
    add_center_para(doc, "摘  要", "黑体", 14, True)
    for para in meta["abstract_zh"]:
        add_para(doc, para)
    add_para(doc, "关键词：" + "；".join(meta["keywords_zh"]), indent=False)
    add_new_page_section(doc)
    add_center_para(doc, meta["title_en"], "Times New Roman", 16, True)
    add_center_para(doc, "ABSTRACT", "Times New Roman", 14, True)
    for para in meta["abstract_en"]:
        add_para(doc, para)
    add_para(doc, "KEY WORDS: " + "; ".join(meta["keywords_en"]), indent=False)
    add_toc(doc, chapters)


def add_body(doc, chapters, references, acknowledgement, appendix_blocks):
    doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(doc.sections[-1])
    add_page_number(doc.sections[-1], roman=False)
    for chapter_idx, chapter in enumerate(chapters):
        add_title(doc, chapter["title"], 1, page_break=chapter_idx > 0)
        for block in chapter.get("blocks", []):
            kind = block.get("type")
            if kind == "h2":
                add_title(doc, block["text"], 2)
            elif kind == "h3":
                add_title(doc, block["text"], 3)
            elif kind == "p":
                add_para(doc, block["text"])
            elif kind == "table":
                add_table(doc, block)
            elif kind == "image":
                add_image(doc, block)
            elif kind == "code":
                add_code(doc, block)
    add_title(doc, "参考文献", 1, page_break=True)
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        set_exact_20(p)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.left_indent = Cm(0.74)
        r = p.add_run(f"[{i}] {ref}")
        set_run_font(r, "宋体", "Times New Roman", 10.5)
    add_title(doc, "附  录", 1, page_break=True)
    for block in appendix_blocks:
        if block["type"] == "p":
            add_para(doc, block["text"])
        elif block["type"] == "code":
            add_code(doc, block)
    add_title(doc, "致  谢", 1, page_break=True)
    for para in acknowledgement:
        add_para(doc, para)


def validate_docx_text(path):
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    bad = ["YOLO", "香菇", "绝缘子", "×××", "{{", "}}"]
    found = [x for x in bad if x in text]
    if found:
        raise RuntimeError("残留异常文本: " + ", ".join(found))


def main():
    meta = load_json("meta.json")
    chapters = load_json("chapters.json")
    references = load_json("references.json")
    acknowledgement = load_json("acknowledgement.json")
    appendix = load_json("appendix.json")
    doc = Document(str(SRC_TEMPLATE))
    remove_all_but_cover(doc)
    for idx, section in enumerate(doc.sections):
        if idx > 0:
            set_margins(section)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.footer.paragraphs:
            clear_runs(p)
    fill_cover(doc, meta)
    add_front_matter(doc, meta, chapters)
    add_body(doc, chapters, references, acknowledgement, appendix)
    doc.save(OUTPUT)
    patch_update_fields(OUTPUT)
    repair_table_compat(OUTPUT)
    validate_docx_text(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
