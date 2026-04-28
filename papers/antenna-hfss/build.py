"""
多频段手机天线 HFSS 仿真论文组装器。

项目规范：
- 模板由 templates/mjxy_antenna/make.py 从源论文格式样例生成；
- 论文内容放在 meta.json、ch*.json、references.json、conclusion.json 等数据文件中；
- build.py 只负责读取 JSON、插入图片/表格并渲染最终 DOCX。

用法：
  cd /Users/a136/vs/45425/thesis_project/papers/antenna-hfss
  python build.py [输出.docx]
"""
from __future__ import annotations

import json
import hashlib
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path


ROOT = Path(__file__).resolve().parents[2]
HERE = Path(__file__).resolve().parent
IMG_DIR = HERE / "images"
TPL = ROOT / "templates" / "mjxy_antenna" / "template.docx"
WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

sys.path.insert(0, str(ROOT))

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docxtpl import DocxTemplate


def load_json(name: str, default=None):
    path = HERE / name
    if not path.exists():
        return default
    return json.loads(path.read_text(encoding="utf-8"))


def split_paras(value, count=3):
    if isinstance(value, list):
        parts = [str(x).strip() for x in value if str(x).strip()]
    else:
        parts = [p.strip() for p in str(value or "").split("\n") if p.strip()]
    parts = parts[:count]
    while len(parts) < count:
        parts.append("")
    return parts


def fmt_run(run, size=12, bold=None, east="宋体", ascii_font="宋体"):
    run.font.size = Pt(size)
    run.font.name = ascii_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), east)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    if bold is not None:
        run.bold = bold


def ensure_outline(p, level: int):
    p_pr = p._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))


def add_heading(sd, text, level=1, page_break=False):
    p = sd.add_paragraph(str(text))
    try:
        p.style = f"Heading {level}"
    except Exception:
        pass
    ensure_outline(p, max(0, level - 1))
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.space_before = Pt(6 if level > 1 else 0)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 3)
    if page_break:
        p.paragraph_format.page_break_before = True
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size, east, bold = 16, "宋体", True
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size, east, bold = 14, "宋体", True
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size, east, bold = 12, "宋体", True
    for r in p.runs:
        fmt_run(r, size=size, east=east, bold=bold)
    return p


def add_toc_line(sd, title, page, level=1):
    p = sd.add_paragraph()
    p.paragraph_format.left_indent = Cm({1: 0, 2: 0.45, 3: 0.85}.get(level, 0))
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    r = p.add_run(f"{title}\t{page}")
    fmt_run(r, 10.5, bold=(level == 1), east="宋体")


def iter_toc_titles(chapters):
    for ch in chapters:
        yield ch["title"], 1
        for sec in ch.get("sections", []):
            if sec.get("title"):
                yield sec["title"], 2
            for sub in sec.get("subsections", []):
                if sub.get("title"):
                    yield sub["title"], 3
    yield "结论", 1
    yield "致谢", 1
    yield "参考文献", 1


def default_page_for(title, order):
    defaults = {
        "1 绪论": 1,
        "2 手机天线基本理论与仿真方法": 6,
        "3 n78手机PIFA天线设计与建模": 11,
        "4 仿真结果与优化分析": 16,
        "结论": 23,
        "致谢": 24,
        "参考文献": 25,
    }
    return defaults.get(title, max(1, 1 + order // 3))


def build_toc(doc, chapters, toc_pages=None):
    toc_pages = toc_pages or {}
    sd = doc.new_subdoc()
    add_toc_line(sd, "摘 要", "I", 1)
    add_toc_line(sd, "Abstract", "II", 1)
    for order, (title, level) in enumerate(iter_toc_titles(chapters), 1):
        page = toc_pages.get(title, default_page_for(title, order))
        add_toc_line(sd, title, page, level)
    return sd


def add_text(sd, text):
    p = sd.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    parts = re.split(r"(\[[0-9,，,\-—–]+\])", str(text))
    for part in parts:
        if not part:
            continue
        r = p.add_run(part)
        fmt_run(r, size=12)
        if re.fullmatch(r"\[[0-9,，,\-—–]+\]", part):
            r.font.superscript = True
    return p


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.first_child_found_in("w:tcBorders")
    if old is not None:
        tc_pr.remove(old)
    borders = OxmlElement("w:tcBorders")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        if val is None:
            continue
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val.get("val", "single"))
        el.set(qn("w:sz"), val.get("sz", "4"))
        el.set(qn("w:color"), val.get("color", "000000"))
        el.set(qn("w:space"), "0")
        borders.append(el)
    tc_pr.append(borders)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_width_pct(cell, ratio):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(max(1, int(5000 * ratio))))
    tc_w.set(qn("w:type"), "pct")


def set_cell_margins(cell, top=70, bottom=70, left=90, right=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.first_child_found_in("w:tcMar")
    if old is not None:
        tc_pr.remove(old)
    mar = OxmlElement("w:tcMar")
    for edge, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tc_pr.append(mar)


def set_table_grid(table, widths_cm):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "autofit")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")
    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    total = sum(widths_cm) or 1
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(7938 * width / total)))
        grid.append(col)
    tbl.insert(1, grid)


def add_table(sd, data):
    caption = data.get("caption", "")
    if caption:
        p = sd.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = Pt(16)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(caption)
        fmt_run(r, 10.5, False, east="宋体")

    image_path = render_table_image(data)
    p = sd.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(13.2))
    sd.add_paragraph()


def get_table_font(size=34):
    from PIL import ImageFont

    candidates = [
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/Library/Fonts/Songti.ttc",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def wrap_cell_text(draw, text, font, max_width):
    text = str(text)
    if not text:
        return [""]
    lines, current = [], ""
    for ch in text:
        test = current + ch
        bbox = draw.textbbox((0, 0), test, font=font)
        if bbox[2] - bbox[0] <= max_width or not current:
            current = test
        else:
            lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def render_table_image(data):
    from PIL import Image, ImageDraw

    headers = [str(x) for x in data.get("headers", [])]
    rows = [[str(x) for x in row] for row in data.get("rows", [])]
    if not headers:
        raise ValueError("表格缺少 headers")

    table_dir = IMG_DIR / "_generated_tables"
    table_dir.mkdir(parents=True, exist_ok=True)
    digest = hashlib.md5(json.dumps(data, ensure_ascii=False, sort_keys=True).encode("utf-8")).hexdigest()[:12]
    out = table_dir / f"table_{digest}.png"

    widths = data.get("col_widths") or [1] * len(headers)
    if len(widths) != len(headers):
        widths = [1] * len(headers)
    total = float(sum(widths)) or 1.0
    canvas_w = 1600
    margin_x = 30
    usable_w = canvas_w - margin_x * 2
    col_ws = [max(120, int(usable_w * w / total)) for w in widths]
    delta = usable_w - sum(col_ws)
    col_ws[-1] += delta

    font = get_table_font(34)
    line_h = 48
    pad_x = 18
    pad_y = 16
    probe = Image.new("RGB", (canvas_w, 200), "white")
    draw = ImageDraw.Draw(probe)

    wrapped_rows = []
    for row in [headers] + rows:
        wrapped = []
        for value, width in zip(row, col_ws):
            wrapped.append(wrap_cell_text(draw, value, font, width - pad_x * 2))
        wrapped_rows.append(wrapped)

    row_heights = [max(line_h + pad_y * 2, max(len(cell) for cell in row) * line_h + pad_y * 2)
                   for row in wrapped_rows]
    canvas_h = sum(row_heights) + 36
    image = Image.new("RGB", (canvas_w, canvas_h), "white")
    draw = ImageDraw.Draw(image)

    x_positions = [margin_x]
    for width in col_ws[:-1]:
        x_positions.append(x_positions[-1] + width)
    table_top = 18
    table_bottom = table_top + sum(row_heights)

    draw.line((margin_x, table_top, margin_x + usable_w, table_top), fill="black", width=4)
    header_bottom = table_top + row_heights[0]
    draw.line((margin_x, header_bottom, margin_x + usable_w, header_bottom), fill="black", width=2)
    draw.line((margin_x, table_bottom, margin_x + usable_w, table_bottom), fill="black", width=4)

    y = table_top
    for r_idx, (row, row_h) in enumerate(zip(wrapped_rows, row_heights)):
        for c_idx, lines in enumerate(row):
            x = x_positions[c_idx]
            col_w = col_ws[c_idx]
            text_h = len(lines) * line_h
            y0 = y + max(pad_y, (row_h - text_h) // 2)
            short_text = len("".join(lines)) <= 14 or r_idx == 0
            for line in lines:
                bbox = draw.textbbox((0, 0), line, font=font)
                text_w = bbox[2] - bbox[0]
                if short_text:
                    tx = x + (col_w - text_w) // 2
                else:
                    tx = x + pad_x
                draw.text((tx, y0), line, fill="black", font=font)
                y0 += line_h
        y += row_h

    image.save(out)
    return out


def add_word_table(sd, data):
    caption = data.get("caption", "")
    headers = data.get("headers", [])
    rows = data.get("rows", [])
    if not headers:
        return
    if caption:
        p = sd.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        fmt_run(r, 10.5, False, east="宋体")

    widths = data.get("col_widths") or [14.0 / len(headers)] * len(headers)
    total = sum(widths)
    if total > 14.0:
        widths = [w * 14.0 / total for w in widths]

    table = sd.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_grid(table, widths)
    width_total = sum(widths) or 1

    thick = {"val": "single", "sz": "12", "color": "000000"}
    thin = {"val": "single", "sz": "4", "color": "000000"}
    none = {"val": "nil"}

    for r_idx, row in enumerate([headers] + rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_width_pct(cell, widths[c_idx] / width_total)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(
                cell,
                top=thick if r_idx == 0 else none,
                bottom=thin if r_idx == 0 else (thick if r_idx == len(rows) else none),
                left=none,
                right=none,
            )
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.right_indent = Cm(0)
            p.paragraph_format.line_spacing = Pt(16)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for run in list(p.runs):
                p._p.remove(run._r)
            r = p.add_run(str(value))
            fmt_run(r, 9.5 if len(headers) >= 5 else 10.5, bold=(r_idx == 0))
    sd.add_paragraph()


def fitted_width_cm(path: Path, requested=13.0, max_height=10.5):
    try:
        from PIL import Image

        with Image.open(path) as im:
            w, h = im.size
        width = min(float(requested), 14.0)
        height = width * h / w
        if height > max_height:
            width = max_height * w / h
        return max(6.0, round(width, 2))
    except Exception:
        return float(requested)


def add_image(sd, image, caption="", width_cm=13.0):
    path = Path(image)
    if not path.is_absolute():
        path = IMG_DIR / path
    if not path.exists():
        add_text(sd, f"[图片缺失：{image}]")
        return
    width_cm = fitted_width_cm(path, requested=width_cm)
    p = sd.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    if caption:
        cap = sd.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.line_spacing = Pt(16)
        cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        cap.paragraph_format.space_before = Pt(0)
        cap.paragraph_format.space_after = Pt(3)
        rr = cap.add_run(caption)
        fmt_run(rr, 10.5, False, east="宋体")


def build_blocks(sd, blocks):
    for block in blocks or []:
        if isinstance(block, str):
            add_text(sd, block)
        elif isinstance(block, dict):
            if "text" in block:
                add_text(sd, block["text"])
            elif "table" in block:
                add_table(sd, block["table"])
            elif "image" in block:
                add_image(sd, block["image"], block.get("caption", ""), block.get("width_cm", 13.0))


def build_body(doc, chapters):
    sd = doc.new_subdoc()
    page = sd.add_paragraph()
    page.add_run().add_break(WD_BREAK.PAGE)
    for idx, ch in enumerate(chapters):
        if idx > 0:
            p = sd.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)
        add_heading(sd, ch["title"], level=1)
        build_blocks(sd, ch.get("content", []))
        for sec in ch.get("sections", []):
            if sec.get("title"):
                add_heading(sd, sec["title"], level=2)
            build_blocks(sd, sec.get("content", []))
            for sub in sec.get("subsections", []):
                add_heading(sd, sub["title"], level=3)
                build_blocks(sd, sub.get("content", []))
    return sd


def build_para_doc(doc, paras):
    sd = doc.new_subdoc()
    build_blocks(sd, paras)
    return sd


def build_refs(doc, refs):
    sd = doc.new_subdoc()
    for i, ref in enumerate(refs or [], 1):
        text = re.sub(r"^\[\d+\]\s*", "", str(ref).strip())
        p = sd.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.line_spacing = Pt(20)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"[{i}] {text}")
        fmt_run(r, 10.5)
    return sd


def load_chapters():
    chapters = []
    for idx in range(1, 10):
        ch = load_json(f"ch{idx}.json")
        if ch:
            print(f"  ch{idx}.json -> {ch['title']}")
            chapters.append(ch)
    return chapters


def build_data(doc, toc_pages=None):
    meta = load_json("meta.json", {})
    zh = split_paras(meta.get("abstract_zh"), 3)
    en = split_paras(meta.get("abstract_en"), 3)
    for i in range(3):
        meta[f"abstract_zh_{i + 1}"] = zh[i]
        meta[f"abstract_en_{i + 1}"] = en[i]

    chapters = load_chapters()
    references = load_json("references.json", [])
    if isinstance(references, dict):
        references = references.get("references", [])

    meta["toc_doc"] = build_toc(doc, chapters, toc_pages=toc_pages)
    meta["body"] = build_body(doc, chapters)
    meta["conclusion_doc"] = build_para_doc(doc, load_json("conclusion.json", []))
    meta["acknowledgement_doc"] = build_para_doc(doc, load_json("acknowledgement.json", []))
    meta["references_doc"] = build_refs(doc, references)
    return meta, chapters


def patch_update_fields(docx_path: Path):
    tmp = Path(str(docx_path) + ".tmp")
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "word/settings.xml":
                    xml = data.decode("utf-8")
                    if "w:updateFields" not in xml:
                        xml = xml.replace("</w:settings>", '<w:updateFields w:val="true"/></w:settings>')
                    data = xml.encode("utf-8")
                zout.writestr(info, data)
    tmp.replace(docx_path)


def clear_runs(p):
    for run in list(p.runs):
        p._p.remove(run._r)


def replace_runs(p, parts):
    clear_runs(p)
    for text, size, bold, east, ascii_font in parts:
        r = p.add_run(text)
        fmt_run(r, size=size, bold=bold, east=east, ascii_font=ascii_font)


def set_para_rule(p, line_pt=22, first_line="keep", align=None, page_break_before=None):
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    if first_line != "keep":
        pf.first_line_indent = first_line
    if line_pt is not None:
        pf.line_spacing = Pt(line_pt)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if page_break_before is not None:
        pf.page_break_before = page_break_before


def keyword_content(text, label):
    content = str(text or "").strip()
    if content.startswith(label):
        content = content[len(label):]
    return content.lstrip("：: ").strip()


def fix_front_matter(doc):
    paragraphs = doc.paragraphs
    zh_idx = en_idx = toc_idx = None
    for i, p in enumerate(paragraphs):
        compact = normalize_text(p.text)
        if compact == "摘要" and zh_idx is None:
            zh_idx = i
        elif compact == "Abstract" and en_idx is None:
            en_idx = i
        elif compact == "目录" and toc_idx is None:
            toc_idx = i
    if zh_idx is None or en_idx is None:
        return

    p = paragraphs[zh_idx]
    set_para_rule(p, line_pt=22, first_line=Cm(0), align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_runs(p, [("摘 要", 15, False, "黑体", "黑体")])

    for p in paragraphs[zh_idx + 1:en_idx]:
        text = (p.text or "").strip()
        if not text:
            continue
        if text.startswith("关键词"):
            content = keyword_content(text, "关键词")
            set_para_rule(p, line_pt=22, first_line=None, align=WD_ALIGN_PARAGRAPH.LEFT)
            replace_runs(p, [
                ("关键词", 12, False, "黑体", "黑体"),
                ("：", 14, False, "黑体", "黑体"),
                (content, 12, False, "楷体", "楷体"),
            ])
        else:
            set_para_rule(p, line_pt=22, first_line=Cm(0.74), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            replace_runs(p, [(text, 12, False, "楷体", "楷体")])

    p = paragraphs[en_idx]
    set_para_rule(p, line_pt=22, first_line=Cm(0), align=WD_ALIGN_PARAGRAPH.CENTER, page_break_before=True)
    replace_runs(p, [("Abstract", 15, False, "楷体", "Arial Black")])

    end = toc_idx if toc_idx is not None else len(paragraphs)
    for p in paragraphs[en_idx + 1:end]:
        text = (p.text or "").strip()
        if not text:
            continue
        if text.startswith("Key words"):
            content = keyword_content(text, "Key words")
            set_para_rule(p, line_pt=22, first_line=None, align=WD_ALIGN_PARAGRAPH.LEFT)
            replace_runs(p, [
                ("Key words", 12, False, "楷体", "Arial Black"),
                (": ", 12, False, "楷体", "Times New Roman"),
                (content, 12, False, "楷体", "Times New Roman"),
            ])
        else:
            set_para_rule(p, line_pt=22, first_line=Cm(0.74), align=WD_ALIGN_PARAGRAPH.LEFT)
            replace_runs(p, [(text, 12, False, "楷体", "Times New Roman")])

    if toc_idx is not None:
        p = paragraphs[toc_idx]
        set_para_rule(p, line_pt=22, first_line=Cm(0), align=WD_ALIGN_PARAGRAPH.CENTER, page_break_before=True)
        replace_runs(p, [("目    录", 15, False, "黑体", "黑体")])


def is_drawing_paragraph(p):
    return bool(p._p.findall(f".//{{{WNS}}}drawing"))


def fix_generated_body(doc):
    fig_cap_pat = re.compile(r"^[图表]\d+(?:[-.．]\d+)?")
    h1_pat = re.compile(r"^\d+\s+\S")
    h2_pat = re.compile(r"^\d+\.\d+\s+\S")
    h3_pat = re.compile(r"^\d+\.\d+\.\d+\s+\S")
    body_started = False
    ref_started = False

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text == "1 绪论":
            body_started = True
        if not body_started:
            continue
        if not text:
            continue
        if is_drawing_paragraph(p):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            continue

        if text == "参考文献":
            ref_started = True

        if h1_pat.match(text) or text in {"结论", "致谢", "参考文献"}:
            try:
                p.style = "Heading 1"
            except Exception:
                pass
            set_para_rule(p, line_pt=22, first_line=None, align=WD_ALIGN_PARAGRAPH.CENTER)
            replace_runs(p, [(text, 16, True, "宋体", "宋体")])
        elif h2_pat.match(text):
            try:
                p.style = "Heading 2"
            except Exception:
                pass
            set_para_rule(p, line_pt=22, first_line=None, align=WD_ALIGN_PARAGRAPH.LEFT)
            replace_runs(p, [(text, 14, True, "宋体", "宋体")])
        elif h3_pat.match(text):
            try:
                p.style = "Heading 3"
            except Exception:
                pass
            set_para_rule(p, line_pt=22, first_line=None, align=WD_ALIGN_PARAGRAPH.LEFT)
            replace_runs(p, [(text, 12, True, "宋体", "宋体")])
        elif fig_cap_pat.match(text) and len(text) < 80:
            set_para_rule(p, line_pt=16, first_line=Cm(0), align=WD_ALIGN_PARAGRAPH.CENTER)
            replace_runs(p, [(text, 10.5, False, "宋体", "宋体")])
        else:
            size = 10.5 if ref_started else 12
            for run in p.runs:
                superscript = bool(run.font.superscript)
                fmt_run(run, size=size, bold=run.bold, east="宋体", ascii_font="宋体")
                run.font.superscript = superscript


def fix_table_fonts(doc):
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_para_rule(p, line_pt=16, first_line=Cm(0), align=p.alignment)
                    for run in p.runs:
                        fmt_run(run, size=10.5, bold=run.bold if run.bold is not None else (row_idx == 0),
                                east="宋体", ascii_font="宋体")


def post_process_docx(docx_path: Path):
    from docx import Document

    doc = Document(str(docx_path))
    fix_front_matter(doc)
    fix_generated_body(doc)
    fix_table_fonts(doc)
    doc.save(str(docx_path))


def render(output: Path, toc_pages=None):
    doc = DocxTemplate(str(TPL))
    data, chapters = build_data(doc, toc_pages=toc_pages)
    doc.render(data)
    doc.save(output)
    post_process_docx(output)
    patch_update_fields(output)
    return chapters


def normalize_text(value: str):
    return re.sub(r"\s+", "", value or "")


def infer_toc_pages(docx_path: Path, chapters):
    if not shutil.which("soffice") or not shutil.which("pdftotext"):
        return {}
    titles = [title for title, _ in iter_toc_titles(chapters)]
    if not titles:
        return {}
    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(tmpdir_path), str(docx_path)],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=120,
            check=False,
        )
        if result.returncode != 0:
            return {}
        pdf = tmpdir_path / (docx_path.stem + ".pdf")
        if not pdf.exists():
            return {}
        text = subprocess.run(
            ["pdftotext", "-layout", str(pdf), "-"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=120,
            check=False,
        ).stdout
    pages = text.split("\f")
    norm_pages = [normalize_text(page) for page in pages]
    hit_pages: dict[str, list[int]] = {}
    for title in titles:
        nt = normalize_text(title)
        hits = [idx for idx, page in enumerate(norm_pages, 1) if nt and nt in page]
        if hits:
            hit_pages[title] = hits
    first_ch = chapters[0]["title"] if chapters else None
    if not first_ch or first_ch not in hit_pages:
        return {}
    body_start = max(hit_pages[first_ch])
    toc_pages = {}
    for title in titles:
        hits = hit_pages.get(title)
        if not hits:
            continue
        abs_page = max(hits)
        if abs_page >= body_start:
            toc_pages[title] = str(abs_page - body_start + 1)
    return toc_pages


def verify_docx(docx_path: Path):
    from docx import Document

    doc = Document(str(docx_path))
    text = "\n".join(p.text for p in doc.paragraphs)
    errors = []
    for token in ("{{", "{%", "__TABLE_PLACEHOLDER", "__TOC_PLACEHOLDER"):
        if token in text:
            errors.append(f"残留模板标记：{token}")
    for old in ("购物商城", "学习伴侣", "停车场管理系统", "校园导航系统"):
        if old in text:
            errors.append(f"残留旧项目文本：{old}")
    if "n78" not in text or "PIFA" not in text or "-10.77" not in text:
        errors.append("未检测到关键仿真主题或结果数据")
    missing = [p.text for p in doc.paragraphs if "图片缺失" in p.text]
    errors.extend(missing)
    return errors


def main():
    output = Path(sys.argv[1]) if len(sys.argv) > 1 else HERE / "王艺微-3222607247-多频段手机天线的仿真研究与设计.docx"
    if not TPL.exists():
        raise FileNotFoundError(f"模板不存在，请先运行 templates/mjxy_antenna/make.py: {TPL}")

    print("第一次渲染论文数据...")
    chapters = render(output)
    toc_pages = infer_toc_pages(output, chapters)
    if toc_pages:
        print(f"回扫到 {len(toc_pages)} 个目录页码，二次渲染目录...")
        render(output, toc_pages=toc_pages)
    else:
        print("未能自动回扫目录页码，保留预估目录。")

    errors = verify_docx(output)
    if errors:
        print("校验提醒:")
        for err in errors:
            print(" -", err)
    else:
        print("校验通过")
    print(f"完成: {output}")


if __name__ == "__main__":
    main()
