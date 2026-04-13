#!/usr/bin/env python3
"""
Extract content from the Douban movie thesis .docx into JSON data files.
Produces: meta.json, ch1.json-ch6.json, references.json, and images/fig_XXX.png
"""

import json
import os
import re
from pathlib import Path
from docx import Document
from lxml import etree

# ── paths ──────────────────────────────────────────────────────────────
DOCX = "/Users/a136/vs/45425/Analysis_and_Implementation_of_Classification_Algorithms_Based_on_Douban_Movie_Data_Thesis (1).docx"
OUT  = Path(__file__).resolve().parent
IMG_DIR = OUT / "images"
IMG_DIR.mkdir(exist_ok=True)

NS = {
    'w':  'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a':  'http://schemas.openxmlformats.org/drawingml/2006/main',
    'r':  'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'pic':'http://schemas.openxmlformats.org/drawingml/2006/picture',
}

doc = Document(DOCX)
paragraphs = doc.paragraphs

# ── 1. Extract images ─────────────────────────────────────────────────
# Build a mapping: paragraph-index -> image filename
img_counter = 0
para_img_map = {}   # {para_idx: "fig_001.png"}

for idx, p in enumerate(paragraphs):
    drawings = p._element.findall('.//w:drawing', NS)
    if not drawings:
        continue
    for drawing in drawings:
        blips = drawing.findall('.//a:blip', NS)
        for blip in blips:
            rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if not rId:
                continue
            # Get the image part via the relationship
            part = doc.part.related_parts.get(rId)
            if part is None:
                continue
            img_counter += 1
            # Determine extension from content type
            ct = part.content_type
            ext = 'png'
            if 'jpeg' in ct or 'jpg' in ct:
                ext = 'jpeg'
            elif 'gif' in ct:
                ext = 'gif'
            elif 'emf' in ct:
                ext = 'emf'
            elif 'wmf' in ct:
                ext = 'wmf'
            fname = f"fig_{img_counter:03d}.{ext}"
            with open(IMG_DIR / fname, 'wb') as f:
                f.write(part.blob)
            para_img_map[idx] = fname
            # Get width in mm (from EMU)
            extents = drawing.findall('.//wp:extent', NS)
            width_mm = None
            if extents:
                cx = int(extents[0].get('cx', 0))
                width_mm = round(cx / 914400 * 25.4)  # EMU -> inches -> mm

            para_img_map[idx] = {"path": fname, "width": width_mm or 130}

print(f"Extracted {img_counter} images")

# ── 2. Build interleaved body sequence (paragraphs + tables) ──────────
# Walk through body XML children to get the correct ordering
body = doc.element.body
body_items = []   # list of ('p', para_idx) or ('t', table_idx)
p_idx = 0
t_idx = 0
for child in body:
    tag = etree.QName(child.tag).localname
    if tag == 'p':
        body_items.append(('p', p_idx))
        p_idx += 1
    elif tag == 'tbl':
        body_items.append(('t', t_idx))
        t_idx += 1

# ── helpers ────────────────────────────────────────────────────────────
def get_para_style(p):
    return p.style.name if p.style else "None"

def extract_table(tbl):
    """Return a dict with headers and rows."""
    rows = tbl.rows
    if len(rows) == 0:
        return {"type": "table", "headers": [], "rows": []}
    headers = [cell.text.strip() for cell in rows[0].cells]
    data_rows = []
    for row in rows[1:]:
        data_rows.append([cell.text.strip() for cell in row.cells])
    return {"type": "table", "headers": headers, "rows": data_rows}

def is_subsection_number(num_str):
    """Return True if number has pattern X.Y.Z (three parts)."""
    parts = num_str.strip().split('.')
    return len(parts) >= 3

def parse_heading2_number(text):
    """Extract the number prefix from a section heading like '1.1 研究背景'."""
    m = re.match(r'^(\d+(?:\.\d+)+)', text.strip())
    return m.group(1) if m else None

# ── 3. Build meta.json ────────────────────────────────────────────────
abstract_zh = "\n".join(paragraphs[i].text for i in range(1, 5))
abstract_en = "\n".join(paragraphs[i].text for i in range(8, 12))

# Keywords
kw_zh_text = paragraphs[5].text
kw_zh = kw_zh_text.replace("关键词：", "").replace("关键词:", "").strip()

kw_en_text = paragraphs[12].text
kw_en = kw_en_text.replace("Keywords:", "").replace("Keywords：", "").strip()

# Acknowledgement: paragraphs after p311 ("致谢") until p317 (empty) or p318 (参考文献)
ack_parts = []
for i in range(312, len(paragraphs)):
    style = get_para_style(paragraphs[i])
    if style == 'Thesis Heading 1':
        break
    txt = paragraphs[i].text.strip()
    if txt:
        ack_parts.append(txt)
acknowledgement = "\n".join(ack_parts)

meta = {
    "title_zh": "基于豆瓣电影数据的分类算法分析与实现",
    "title_en": "Analysis and Implementation of Classification Algorithms Based on Douban Movie Data",
    "college": "大数据学院",
    "major": "数据科学与大数据技术",
    "class_name": "",
    "name": "",
    "student_id": "",
    "advisor": "",
    "abstract_zh": abstract_zh,
    "abstract_en": abstract_en,
    "keywords_zh": kw_zh,
    "keywords_en": kw_en,
    "acknowledgement": acknowledgement,
}

with open(OUT / "meta.json", "w", encoding="utf-8") as f:
    json.dump(meta, f, ensure_ascii=False, indent=2)
print("Wrote meta.json")

# ── 4. Build chapter JSON files ───────────────────────────────────────
# Identify chapter boundaries using Thesis Heading 1 in body content (p67+)
# Chapter headings: p67, p85, p131, p171, p240, p290
# Special: p311 = 致谢, p318 = 参考文献

chapter_starts = []  # (para_idx, title_text)
for i, p in enumerate(paragraphs):
    if i < 67:
        continue
    if get_para_style(p) == 'Thesis Heading 1':
        chapter_starts.append((i, p.text.strip()))

# Filter to only actual chapters (第X章)
chapter_headings = [(i, t) for i, t in chapter_starts if t.startswith("第")]
# Also note special sections
ack_start = next(i for i, t in chapter_starts if t == "致谢")
ref_start = next(i for i, t in chapter_starts if t == "参考文献")

print(f"Found {len(chapter_headings)} chapters, ack at p{ack_start}, ref at p{ref_start}")

# For each chapter, determine the end boundary (next Thesis Heading 1)
all_h1_indices = [i for i, t in chapter_starts]

# Build a position map: para_idx -> position in body_items, and also allow lookup of tables
# We need to iterate body_items for the range of each chapter

def body_items_in_range(start_para, end_para):
    """Yield body items (paragraphs and tables) between start_para (exclusive) and end_para (exclusive).
    start_para is the chapter heading paragraph index.
    """
    in_range = False
    for item_type, item_idx in body_items:
        if item_type == 'p' and item_idx == start_para:
            in_range = True
            continue  # skip the heading itself
        if item_type == 'p' and item_idx >= end_para:
            break
        if in_range:
            yield item_type, item_idx

def build_chapter(ch_para_idx, ch_title, end_para_idx):
    """Build a chapter dict from paragraphs/tables between ch_para_idx and end_para_idx."""
    # Parse chapter number from title like "第1章 绪论"
    m = re.match(r'第(\d+)章\s*(.+)', ch_title)
    ch_num = int(m.group(1))
    ch_name = m.group(2).strip()

    sections = []
    current_section = None
    current_subsection = None

    for item_type, item_idx in body_items_in_range(ch_para_idx, end_para_idx):
        if item_type == 'p':
            p = paragraphs[item_idx]
            style = get_para_style(p)
            text = p.text.strip()

            if style == 'Thesis Heading 2':
                num = parse_heading2_number(text)
                if num and is_subsection_number(num):
                    # It's a subsection (X.Y.Z)
                    current_subsection = {
                        "title": text,
                        "number": num,
                        "content": []
                    }
                    if current_section is not None:
                        current_section["subsections"].append(current_subsection)
                else:
                    # It's a section (X.Y)
                    current_subsection = None
                    current_section = {
                        "title": text,
                        "number": num or "",
                        "content": [],
                        "subsections": []
                    }
                    sections.append(current_section)
            else:
                # Regular paragraph or image
                target = current_subsection if current_subsection else current_section
                if target is None:
                    # Content before any section heading (shouldn't happen normally)
                    continue

                # Check if this paragraph has an image
                if item_idx in para_img_map:
                    img_info = para_img_map[item_idx]
                    target["content"].append({
                        "type": "image",
                        "path": img_info["path"],
                        "width": img_info["width"]
                    })
                elif text:
                    target["content"].append(text)

        elif item_type == 't':
            tbl = doc.tables[item_idx]
            table_data = extract_table(tbl)
            target = current_subsection if current_subsection else current_section
            if target is not None:
                target["content"].append(table_data)

    return {
        "title": ch_name,
        "chapter_number": ch_num,
        "sections": sections
    }


for ci, (ch_idx, ch_title) in enumerate(chapter_headings):
    # End boundary is next heading-1
    if ci + 1 < len(chapter_headings):
        end_idx = chapter_headings[ci + 1][0]
    else:
        end_idx = ack_start  # last chapter ends at 致谢

    ch_data = build_chapter(ch_idx, ch_title, end_idx)
    ch_file = OUT / f"ch{ch_data['chapter_number']}.json"
    with open(ch_file, "w", encoding="utf-8") as f:
        json.dump(ch_data, f, ensure_ascii=False, indent=2)
    print(f"Wrote {ch_file.name}: {ch_data['title']} ({len(ch_data['sections'])} sections)")

# ── 5. References ─────────────────────────────────────────────────────
refs = []
for i in range(ref_start + 1, len(paragraphs)):
    text = paragraphs[i].text.strip()
    if text:
        refs.append(text)

with open(OUT / "references.json", "w", encoding="utf-8") as f:
    json.dump(refs, f, ensure_ascii=False, indent=2)
print(f"Wrote references.json ({len(refs)} references)")

print("\nDone!")
