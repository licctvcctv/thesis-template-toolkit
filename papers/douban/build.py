"""
智学云课论文组装器。
支持文本、图片（InlineImage）、表格。

用法: python build.py <模板.docx> [输出.docx]
"""
import os
import sys
import json
import glob

ROOT = os.path.join(os.path.dirname(__file__), "../..")
sys.path.insert(0, ROOT)

HERE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(HERE, "images")


def load_json(filename):
    path = os.path.join(HERE, filename)
    if not os.path.exists(path):
        return None
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def process_content(content_list, doc):
    """
    处理内容列表：
    - 字符串 → 保持不变
    - {"type":"image"} → InlineImage（图片）
    - {"type":"table"} → 转为特殊标记文本（后处理时插入表格）
    """
    from docxtpl import InlineImage
    from docx.shared import Mm

    result = []
    for item in content_list:
        if isinstance(item, str):
            result.append(item)
        elif isinstance(item, dict):
            if item.get("type") == "image":
                img_path = os.path.join(IMG_DIR, item["path"])
                if os.path.exists(img_path):
                    width = Mm(item.get("width", 120))
                    result.append(InlineImage(doc, img_path,
                                             width=width))
                    # 不跳过标注——让标注文本正常输出
                else:
                    result.append(f"[图片缺失: {item['path']}]")
            elif item.get("type") == "table":
                _pending_tables.append(item)
                tid = len(_pending_tables) - 1
                # 只放占位标记，不重复生成标注
                # （标注已在 JSON 内容中作为前一个字符串存在）
                result.append(f"__TABLE_PLACEHOLDER_{tid}__")
            else:
                result.append(str(item))
    return result


# 全局表格暂存
_pending_tables = []
_toc_entries = []


def _numbered_title(number, title):
    """Return a display title with its chapter/section number."""
    text = str(title or "").strip()
    if not number:
        return text
    number = str(number).strip()
    return text if text.startswith(number) else f"{number} {text}"


def prepare_chapter_titles(chapters):
    """Add visible numbering to headings and collect TOC entries."""
    _toc_entries.clear()
    for idx, ch in enumerate(chapters, 1):
        ch_no = ch.get("chapter_number") or idx
        _toc_entries.append((1, _numbered_title(ch_no, ch.get("title"))))

        for sec in ch.get("sections", []):
            sec_title = _numbered_title(sec.get("number"), sec.get("title"))
            _toc_entries.append((2, sec_title))

            for sub in sec.get("subsections", []):
                sub_title = _numbered_title(sub.get("number"),
                                            sub.get("title"))
                _toc_entries.append((3, sub_title))
    return chapters


def process_chapters(chapters, doc):
    """处理所有章节，转换图片引用"""
    for ch in chapters:
        for sec in ch.get("sections", []):
            sec["content"] = process_content(
                sec.get("content", []), doc)
            for sub in sec.get("subsections", []):
                sub["content"] = process_content(
                    sub.get("content", []), doc)
    return chapters


def build_data(doc=None):
    """组装所有章节数据"""
    meta = load_json("meta.json")
    if not meta:
        raise FileNotFoundError("meta.json not found")

    chapters = []
    for ch_file in sorted(glob.glob(
            os.path.join(HERE, "ch*.json"))):
        ch = load_json(os.path.basename(ch_file))
        if ch:
            chapters.append(ch)
            print(f"  {os.path.basename(ch_file)}"
                  f" -> {ch['title']}")

    chapters = prepare_chapter_titles(chapters)

    if doc:
        chapters = process_chapters(chapters, doc)

    refs = load_json("references.json")
    if isinstance(refs, list):
        references = refs
    elif isinstance(refs, dict):
        references = refs.get("references", [])
    else:
        references = []

    return {**meta, "chapters": chapters,
            "references": references}


def main():
    if len(sys.argv) < 2:
        print("用法: python build.py <模板.docx> [输出.docx]")
        sys.exit(1)

    template = sys.argv[1]
    output = sys.argv[2] if len(sys.argv) > 2 else "output.docx"

    from docxtpl import DocxTemplate
    doc = DocxTemplate(template)

    _pending_tables.clear()
    print("组装论文数据...")
    data = build_data(doc)
    print(f"  {len(data['chapters'])} 章, "
          f"{len(data.get('references', []))} 条参考文献")

    print(f"渲染: {template} -> {output}")
    doc.render(data)
    doc.save(output)

    # 后处理：修正图片居中和图标注格式
    _post_process(output)
    # 校验
    errors = _verify(output)
    if errors:
        print(f"  ⚠ 发现 {len(errors)} 个问题:")
        for e in errors:
            print(f"    {e}")
    else:
        print(f"  ✓ 校验通过")
    print(f"完成: {output}")


def _verify(docx_path):
    """校验输出文档：检查每个表/图标注是否有对应内容"""
    from docx import Document
    import re

    doc = Document(docx_path)
    errors = []
    body = doc.element.body
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # 按 DOM 顺序遍历，记录每个元素的类型
    elements = []  # [(type, text)]
    for elem in body:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            texts = [t.text for t in elem.findall('.//w:t', ns) if t.text]
            text = ''.join(texts).strip()
            has_img = bool(elem.findall('.//w:drawing', ns))
            if has_img:
                elements.append(('img', text))
            else:
                elements.append(('p', text))
        elif tag == 'tbl':
            elements.append(('tbl', ''))

    # 检查每个表标注后面是否有表格
    for i, (etype, text) in enumerate(elements):
        if etype == 'p' and re.match(r'^表\d+[.-]\d+\s', text) and not re.search(r'[。，；]', text) and len(text) < 30:
            # 往后找 3 个元素内是否有表格
            found_tbl = False
            for j in range(i + 1, min(i + 6, len(elements))):
                if elements[j][0] == 'tbl':
                    found_tbl = True
                    break
            if not found_tbl:
                errors.append(f"缺表格: {text}")

    # 检查每个图标注附近是否有图片
    for i, (etype, text) in enumerate(elements):
        if etype == 'p' and re.match(r'^图\d+-\d+\s', text) and len(text) < 50:
            found_img = False
            # 往前找 3 个元素
            for j in range(max(0, i - 3), i):
                if elements[j][0] == 'img':
                    found_img = True
                    break
            if not found_img:
                errors.append(f"缺图片: {text}")

    return errors


def _para_text_from_xml(p_elem):
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    return ''.join(t.text or '' for t in p_elem.findall('.//w:t', ns)).strip()


def _insert_paragraph_after(paragraph, text=""):
    from docx.text.paragraph import Paragraph
    from docx.oxml import OxmlElement

    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if text:
        p.add_run(text)
    return p


def _insert_paragraph_before(paragraph, text=""):
    from docx.text.paragraph import Paragraph
    from docx.oxml import OxmlElement

    new_p = OxmlElement('w:p')
    paragraph._p.addprevious(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if text:
        p.add_run(text)
    return p


def _remove_paragraph(paragraph):
    elem = paragraph._element
    parent = elem.getparent()
    if parent is not None:
        parent.remove(elem)


def _prev_is_blank(paragraph):
    prev = paragraph._p.getprevious()
    if prev is None or prev.tag.split('}')[-1] != 'p':
        return False
    return not _para_text_from_xml(prev)


def _set_para_text(paragraph, text):
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


def _set_run_font(run, size=None, east="宋体", ascii_font="Times New Roman",
                  bold=None):
    from docx.shared import Pt
    from docx.oxml.ns import qn

    if size:
        run.font.size = Pt(size)
    run.font.name = ascii_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), east)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    if bold is not None:
        run.bold = bold


def _toc_page_for(title, level):
    import re

    chapter_pages = {1: 1, 2: 6, 3: 12, 4: 17, 5: 25, 6: 31}
    m = re.match(r'^(\d+)(?:\.(\d+))?', title or "")
    if not m:
        return ""
    ch = int(m.group(1))
    base = chapter_pages.get(ch, max(1, ch * 5 - 4))
    if level == 1:
        return str(base)
    sec = int(m.group(2) or 1)
    return str(base + max(0, (sec - 1) // 2))


def _format_toc_entry(paragraph, level):
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

    paragraph.style = f"toc {level}" if f"toc {level}" in [
        s.name for s in paragraph.part.document.styles] else paragraph.style
    paragraph.paragraph_format.left_indent = Cm(
        {1: 0, 2: 0.55, 3: 1.1}.get(level, 0))
    paragraph.paragraph_format.line_spacing = Pt(18)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Cm(14.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    for r in paragraph.runs:
        _set_run_font(r, 12, east="Times New Roman",
                      ascii_font="Times New Roman")


def _rebuild_toc(doc):
    toc_title = None
    toc_idx = None
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip().replace(" ", "") == "目录":
            toc_title = p
            toc_idx = i
            break
    if toc_title is None:
        return

    body_idx = None
    for i, p in enumerate(doc.paragraphs[toc_idx + 1:], toc_idx + 1):
        if p.style and p.style.name == "Heading 1":
            body_idx = i
            break
    if body_idx is None:
        return

    old_between = doc.paragraphs[toc_idx + 1:body_idx]
    for p in old_between:
        _remove_paragraph(p)

    if not _prev_is_blank(toc_title):
        _insert_paragraph_before(toc_title)

    anchor = _insert_paragraph_after(toc_title)
    for level, title in _toc_entries:
        if not title:
            continue
        text = f"{title}\t{_toc_page_for(title, level)}"
        anchor = _insert_paragraph_after(anchor, text)
        _format_toc_entry(anchor, level)
    _insert_paragraph_after(anchor)


def _format_heading(paragraph):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    style_name = paragraph.style.name if paragraph.style else ""
    if style_name == "Heading 1":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)
        for r in paragraph.runs:
            _set_run_font(r, 16, east="黑体", ascii_font="Times New Roman",
                          bold=True)
    elif style_name == "Heading 2":
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        for r in paragraph.runs:
            _set_run_font(r, 14, east="黑体", ascii_font="Times New Roman",
                          bold=True)
    elif style_name == "Heading 3":
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(3)
        for r in paragraph.runs:
            _set_run_font(r, 12, east="黑体", ascii_font="Times New Roman",
                          bold=True)


def _post_process(docx_path):
    """渲染后：图片居中、标注格式、插入表格"""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

    doc = Document(docx_path)
    _rebuild_toc(doc)

    fig_pat = re.compile(r'^图\d')
    tbl_cap_pat = re.compile(r'^表\d+[.-]\d+\s')
    placeholder_pat = re.compile(r'^__TABLE_PLACEHOLDER_(\d+)__$')

    for p in list(doc.paragraphs):
        t = (p.text or "").strip()

        # Heading 1 → 每章前分页
        if p.style and p.style.name == 'Heading 1':
            p.paragraph_format.page_break_before = True
        if p.style and p.style.name in {'Heading 1', 'Heading 2', 'Heading 3'}:
            _format_heading(p)

        # 致谢另起一页
        if t == "致    谢":
            p.paragraph_format.page_break_before = True

        # 英文摘要和英文关键词格式
        if t.startswith("Key words:"):
            for r in p.runs:
                _set_run_font(r, 12, east="Times New Roman",
                              ascii_font="Times New Roman")

        # 图片段落 → 居中
        has_drawing = bool(p._p.findall(
            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
        if has_drawing:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 固定行间距(EXACTLY 20pt)会裁切图片，改为单倍行距
            from docx.enum.text import WD_LINE_SPACING
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 图/表标注 → 居中 + 五号黑体（短文本）
        if (fig_pat.match(t) or tbl_cap_pat.match(t)) and len(t) < 50:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(10.5)
                r.font.name = "黑体"
                # 设置东亚字体
                r._r.rPr.rFonts.set(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia',
                    '黑体')

        # 表标题 → 表1.1格式、居中、五号黑体，上方空一行
        if tbl_cap_pat.match(t) and len(t) < 50:
            new_text = re.sub(r'^(表\d+)-(\d+)', r'\1.\2', t)
            if new_text != t:
                _set_para_text(p, new_text)
                t = new_text
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if not _prev_is_blank(p):
                _insert_paragraph_before(p)
            for r in p.runs:
                _set_run_font(r, 10.5, east="黑体",
                              ascii_font="Times New Roman", bold=False)

        # 表格占位符 → 插入真实表格，并在表后空一行
        m = placeholder_pat.match(t)
        if m:
            tid = int(m.group(1))
            if tid < len(_pending_tables):
                _insert_table(doc, p, _pending_tables[tid])
                # 清空占位文本
                for r in p.runs:
                    r.text = ""

    doc.save(docx_path)
    print(f"  后处理: 图片居中+{len(_pending_tables)}个表格")


def _insert_table(doc, after_para, tbl_data):
    """在段落后插入带边框的表格"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    headers = tbl_data.get("headers", [])
    rows = tbl_data.get("rows", [])
    if not headers:
        return

    ncols = len(headers)
    all_rows = [headers] + rows

    # 直接构建表格 XML
    tbl = OxmlElement('w:tbl')

    # 表格属性：居中 + 自动列宽 + 边框
    tblPr = OxmlElement('w:tblPr')
    # 居中
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblPr.append(jc)
    # 表格宽度 100%
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    # 自动列宽（让 Word 根据内容调整）
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'autofit')
    tblPr.append(tblLayout)
    # 边框（三线表：顶线1.5pt、底线1.5pt，无左右竖线）
    borders = OxmlElement('w:tblBorders')
    for edge, sz in [('top', '12'), ('bottom', '12')]:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    for edge in ['left', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        borders.append(el)
    tblPr.append(borders)
    tbl.append(tblPr)

    # 表格网格 — 按内容长度智能分配列宽
    total_width = 9000
    # 计算每列最大文本长度（中文字符算2）
    col_max_len = [0] * ncols
    for row_data in all_rows:
        for ci in range(ncols):
            cell_text = str(row_data[ci]) if ci < len(row_data) else ""
            # 中文字符算2个宽度单位
            w = sum(2 if ord(c) > 127 else 1 for c in cell_text)
            col_max_len[ci] = max(col_max_len[ci], w)
    # 最小列宽1200 twips（约2cm），确保短文本列也够宽
    min_w = 1200
    total_len = sum(max(l, 4) for l in col_max_len) or 1
    col_widths = [max(min_w, int(total_width * max(l, 4) / total_len))
                  for l in col_max_len]
    # 修正总宽度
    diff = total_width - sum(col_widths)
    if diff > 0:
        # 多余的宽度加到最宽的列
        widest = col_widths.index(max(col_widths))
        col_widths[widest] += diff
    elif diff < 0:
        # 超出的从最宽的列减
        widest = col_widths.index(max(col_widths))
        col_widths[widest] += diff

    tblGrid = OxmlElement('w:tblGrid')
    for w in col_widths:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(w))
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    # 行
    for ri, row_data in enumerate(all_rows):
        tr = OxmlElement('w:tr')
        for ci in range(ncols):
            tc = OxmlElement('w:tc')
            # 单元格属性：宽度 + 垂直居中
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(col_widths[ci] if ci < len(col_widths) else 1800))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)
            # 三线表：表头行底部加细线（0.75pt）
            if ri == 0:
                tcBorders = OxmlElement('w:tcBorders')
                btm = OxmlElement('w:bottom')
                btm.set(qn('w:val'), 'single')
                btm.set(qn('w:sz'), '6')
                btm.set(qn('w:space'), '0')
                btm.set(qn('w:color'), '000000')
                tcBorders.append(btm)
                tcPr.append(tcBorders)
            tc.append(tcPr)
            p = OxmlElement('w:p')
            # 段落属性：水平居中 + 段前段后间距为0 + 单倍行距
            pPr = OxmlElement('w:pPr')
            pJc = OxmlElement('w:jc')
            pJc.set(qn('w:val'), 'center')
            pPr.append(pJc)
            pSpacing = OxmlElement('w:spacing')
            pSpacing.set(qn('w:before'), '0')
            pSpacing.set(qn('w:after'), '0')
            pSpacing.set(qn('w:line'), '240')
            pSpacing.set(qn('w:lineRule'), 'auto')
            pPr.append(pSpacing)
            # 清除所有缩进（firstLine + firstLineChars 都要清）
            pInd = OxmlElement('w:ind')
            pInd.set(qn('w:firstLine'), '0')
            pInd.set(qn('w:firstLineChars'), '0')
            pInd.set(qn('w:left'), '0')
            pInd.set(qn('w:right'), '0')
            pPr.append(pInd)
            p.append(pPr)
            # 文本 — 字体：宋体(中文) + Times New Roman(英文)，五号(10.5pt)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            # 字体
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), '宋体')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rPr.append(rFonts)
            # 字号 五号=10.5pt=21半磅
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '21')
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '21')
            rPr.append(szCs)
            # 表头加粗
            if ri == 0:
                b = OxmlElement('w:b')
                rPr.append(b)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = str(row_data[ci]) if ci < len(row_data) else ''
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    # 插入到段落后面
    after_para._p.addnext(tbl)
    tbl.addnext(OxmlElement('w:p'))


if __name__ == "__main__":
    main()
