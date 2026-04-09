"""
段落级扫描 - 提取每个段落的详细 run 信息。
供 AI 查看具体内容和格式，决定哪些 run 需要替换。
"""
from docx import Document
from docx.oxml.ns import qn

NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def scan_paragraphs(doc_path, start=0, end=None, detail=False):
    """
    扫描指定范围段落的结构信息。

    参数:
        doc_path: docx 文件路径
        start: 起始段落索引
        end: 结束段落索引（不含）
        detail: 是否输出每个 run 的详细信息

    返回:
        [{
            "idx": 段落索引,
            "text": 段落文本,
            "style": 段落样式名,
            "has_sectPr": 是否含分节符,
            "has_numPr": 是否有自动编号,
            "alignment": 对齐方式,
            "runs": [{"idx": run索引, "text": ...,
                      "font": ..., "size": ...,
                      "bold": ..., "underline": ...}]
        }]
    """
    doc = Document(doc_path)
    paras = doc.paragraphs
    end = end or len(paras)
    results = []

    for i in range(start, min(end, len(paras))):
        p = paras[i]
        text = (p.text or "").strip()

        # 跳过纯空段落（除非 detail 模式）
        if not text and not detail:
            continue

        info = {
            "idx": i,
            "text": text[:100],
            "style": p.style.name if p.style else "None",
            "has_sectPr": p._p.find(
                './/w:sectPr', NS) is not None,
            "has_numPr": p._p.find(
                './/w:numPr', NS) is not None,
            "alignment": _get_alignment(p),
        }

        if detail:
            info["runs"] = _get_runs(p)

        results.append(info)

    return results


def scan_para_runs(doc_path, para_idx):
    """
    扫描单个段落的所有 run 详情。
    用于 AI 精确定位要替换的 run。

    返回:
        {
            "idx": 段落索引,
            "text": 完整文本,
            "runs": [{"idx": 0, "text": "关键词：",
                      "font": "黑体", "size": 152400,
                      "bold": True, "underline": False}]
        }
    """
    doc = Document(doc_path)
    p = doc.paragraphs[para_idx]
    return {
        "idx": para_idx,
        "text": p.text or "",
        "style": p.style.name if p.style else "None",
        "alignment": _get_alignment(p),
        "runs": _get_runs(p),
    }


def _get_alignment(para):
    """获取段落对齐方式"""
    a = para.paragraph_format.alignment
    if a is None:
        return None
    return str(a)


def _get_runs(para):
    """提取段落所有 run 的信息"""
    runs = []
    for j, r in enumerate(para.runs):
        runs.append({
            "idx": j,
            "text": r.text,
            "font": r.font.name,
            "size": r.font.size,
            "bold": r.font.bold,
            "italic": r.font.italic,
            "underline": r.font.underline,
        })
    return runs
