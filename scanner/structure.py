"""
文档结构扫描 - 识别文档的逻辑分区（封面/声明/摘要/目录/正文等）。
"""
import re
from docx import Document

SECTION_MARKERS = {
    "cover": ["届本科生毕业", "本科毕业论文", "毕业设计（论文）"],
    "declaration": ["原创性声明", "原创性申明", "学术诚信"],
    "abstract_zh": ["摘  要", "摘 要"],
    "abstract_en": ["ABSTRACT", "Abstract"],
    "toc": ["目  录", "目 录"],
    "conclusion": ["结  论", "结 论", "结论"],
    "acknowledgement": ["致  谢", "致 谢", "致谢"],
    "references": ["参考文献"],
    "appendix": ["附  录", "附 录", "附录"],
}


def scan_structure(doc_or_path):
    """
    扫描文档分区结构。接受文件路径或 Document 对象。

    返回 dict: total_paragraphs, total_tables, total_sections,
               styles_used, parts, body_start, body_end
    """
    doc = doc_or_path if not isinstance(doc_or_path, str) \
        else Document(doc_or_path)
    paras = doc.paragraphs
    result = {
        "total_paragraphs": len(paras),
        "total_tables": len(doc.tables),
        "total_sections": len(doc.sections),
        "styles_used": {},
        "parts": {},
    }

    # Count styles
    for p in paras:
        name = p.style.name if p.style else "None"
        result["styles_used"][name] = \
            result["styles_used"].get(name, 0) + 1

    # Single-pass section detection
    # Skip TOC entries (contain tab = page number separator)
    toc_start = None
    found = set()
    for i, p in enumerate(paras):
        text = p.text or ""
        # Detect TOC start
        for kw in SECTION_MARKERS["toc"]:
            if kw in text:
                toc_start = i
                break
        for name, keywords in SECTION_MARKERS.items():
            if name in found:
                continue
            for kw in keywords:
                if kw in text:
                    # Skip TOC entries (tab = page number)
                    is_toc_entry = ('\t' in text and toc_start
                                    and i > toc_start and name != "toc")
                    if is_toc_entry:
                        continue
                    result["parts"][name] = {
                        "start": i,
                        "text": text.strip()[:60],
                    }
                    found.add(name)
                    break

    # Body start: first chapter heading after TOC, before acknowledgement
    toc_end = result["parts"].get("toc", {}).get("start", 0)
    ack_start = result["parts"].get(
        "acknowledgement", {}).get("start", len(paras))
    for i in range(toc_end + 1, ack_start):
        text = (paras[i].text or "").strip()
        if '\t' in text:
            continue
        if (re.match(r'^\d+[\s　]+', text)
                or re.match(r'^第[一二三四五六七八九十]+章', text)):
            result["body_start"] = i
            break

    result["body_end"] = ack_start
    return result
