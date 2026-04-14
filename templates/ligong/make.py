"""
理工类论文模板制作（自动生成目录版）。
用法: cd thesis_project && python -m templates.ligong.make <原始模板.docx>

模板特点:
- 封面: P4 本科毕业设计 + P5 届次 + table[0] 学生信息(8行2列)
- P12-P17: 诚信承诺页
- P20-P25: 中文摘要 + 关键词
- P37-P42: 英文摘要 + Key words
- P45-P74: 目录（自动生成）
- Heading 1/2/3: 三级标题
- Normal / Body Text Indent 2 / 章节题目: 正文
- 致谢 / 参考文献 / 附录: 末尾区域
"""
import os
import sys
import re
import copy
import zipfile
import shutil

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../.."))
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def make(src_path, out_path):
    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc = Document(src_path)

    # ========== Step 1: 封面 ==========
    # P5: "（2025届）" → "（{{ year }}届）"
    p5 = doc.paragraphs[5]
    if p5.runs:
        # runs: [（20] [2] [5] [届）]  →  合并为一个
        p5.runs[0].text = "（{{ year }}届）"
        for r in p5.runs[1:]:
            r.text = ""

    # table[0]: 封面信息表
    tbl = doc.tables[0]
    cover_map = {
        0: "{{ title_zh }}",       # 题目
        1: "{{ college }}",        # 学院
        2: "{{ major }}",          # 专业
        3: "{{ class_name }}",     # 班级
        4: "{{ student_id }}",     # 学号
        5: "{{ name }}",           # 学生姓名
        6: "{{ advisor }}",        # 指导教师
        7: "{{ finish_date }}",    # 完成日期
    }
    for row_idx, var in cover_map.items():
        cell = tbl.cell(row_idx, 1)
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = var
            for r in p.runs[1:]:
                r.text = ""
        else:
            # 没有 run，添加一个
            from docx.oxml import OxmlElement
            r_elem = OxmlElement('w:r')
            t_elem = OxmlElement('w:t')
            t_elem.set(qn('xml:space'), 'preserve')
            t_elem.text = var
            r_elem.append(t_elem)
            p._p.append(r_elem)

    print("  Step 1: 封面")

    # ========== Step 2: 诚信承诺 ==========
    # P14: "《XXXXXXXX》" → "《{{ title_zh }}》"
    p14 = doc.paragraphs[14]
    if p14.runs:
        text = p14.runs[0].text
        p14.runs[0].text = text.replace("XXXXXXXX", "{{ title_zh }}")

    # P17: 承诺日期 → {{ finish_date }}
    p17 = doc.paragraphs[17]
    if p17.runs:
        # runs: [2025][年 ][4][ 月 ][05][ 日] → 合并
        for j, r in enumerate(p17.runs):
            if r.text.strip() and any(c.isdigit() for c in r.text):
                r.text = "{{ finish_date }}"
                for rr in p17.runs[j + 1:]:
                    rr.text = ""
                break

    print("  Step 2: 诚信承诺")

    # ========== Step 3: 中文摘要 ==========
    # P21-P23: 摘要正文 → 合并到 P21
    p21 = doc.paragraphs[21]
    if p21.runs:
        p21.runs[0].text = "{{ abstract_zh }}"
        for r in p21.runs[1:]:
            r.text = ""
    # 清空 P22, P23
    for idx in [22, 23]:
        for r in doc.paragraphs[idx].runs:
            r.text = ""

    # P25: "关键词：..." → 保留标签，替换值
    p25 = doc.paragraphs[25]
    if p25.runs:
        full = p25.text
        if "：" in full:
            # 整个在一个 run 里
            prefix = full[:full.index("：") + 1]
            p25.runs[0].text = prefix + "{{ keywords_zh }}"
            for r in p25.runs[1:]:
                r.text = ""

    print("  Step 3: 中文摘要")

    # ========== Step 4: 英文摘要 ==========
    # P38-P40: 英文摘要正文 → 合并到 P38
    p38 = doc.paragraphs[38]
    if p38.runs:
        p38.runs[0].text = "{{ abstract_en }}"
        for r in p38.runs[1:]:
            r.text = ""
    for idx in [39, 40]:
        for r in doc.paragraphs[idx].runs:
            r.text = ""

    # P42: "Key words：..." → runs: [Key][ ][words：][value...]
    # 保留 Key words： 标签(run0-2)，替换 run3 为模板变量，清 run4+
    p42 = doc.paragraphs[42]
    if p42.runs and len(p42.runs) > 3:
        p42.runs[3].text = "{{ keywords_en }}"
        for r in p42.runs[4:]:
            r.text = ""
    elif p42.runs:
        # fallback: 全部替换
        p42.runs[0].text = "Key words：{{ keywords_en }}"
        for r in p42.runs[1:]:
            r.text = ""
    else:
        # 无 runs - 添加一个
        from docx.oxml import OxmlElement as _OE
        r_elem = _OE('w:r')
        rPr = _OE('w:rPr')
        b = _OE('w:b')
        rPr.append(b)
        sz = _OE('w:sz')
        sz.set(qn('w:val'), '24')  # 小四
        rPr.append(sz)
        r_elem.append(rPr)
        t_elem = _OE('w:t')
        t_elem.set(qn('xml:space'), 'preserve')
        t_elem.text = "Key words："
        r_elem.append(t_elem)
        p42._p.append(r_elem)
        # 值
        r2 = _OE('w:r')
        t2 = _OE('w:t')
        t2.set(qn('xml:space'), 'preserve')
        t2.text = "{{ keywords_en }}"
        r2.append(t2)
        p42._p.append(r2)

    print("  Step 4: 英文摘要")

    # ========== Step 5: 致谢 ==========
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and text == "致谢":
            for j in range(i + 1, min(i + 10, len(doc.paragraphs))):
                pj = doc.paragraphs[j]
                sj = pj.style.name if pj.style else ""
                if "Heading" in sj:
                    break
                if pj.runs and pj.text.strip():
                    pj.runs[0].text = "    {{ acknowledgement }}"
                    for r in pj.runs[1:]:
                        r.text = ""
                    # 清后续段落到下一个 Heading
                    for k in range(j + 1, min(j + 20, len(doc.paragraphs))):
                        pk = doc.paragraphs[k]
                        pk_style = pk.style.name if pk.style else ""
                        if "Heading" in pk_style:
                            break
                        for r in pk.runs:
                            r.text = ""
                    break
            break

    print("  Step 5: 致谢")

    # ========== Step 6: 参考文献 ==========
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and "参考文献" in text:
            for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                pj = doc.paragraphs[j]
                pt = (pj.text or "").strip()
                if pj.runs and pt and "[" in pt:
                    ed_p = pj._p
                    # for 标签
                    ed_p.addprevious(_mk_ctrl("{%p for ref in references %}"))
                    pj.runs[0].text = "{{ ref }}"
                    for r in pj.runs[1:]:
                        r.text = ""
                    ed_p.addnext(_mk_ctrl("{%p endfor %}"))
                    # 清掉后续参考文献条目
                    cleaning = False
                    for k in range(j + 1, min(j + 30, len(doc.paragraphs))):
                        pk = doc.paragraphs[k]
                        pk_style = pk.style.name if pk.style else ""
                        if "Heading" in pk_style:
                            break
                        pk_text = (pk.text or "").strip()
                        if "{%p endfor" in pk_text:
                            cleaning = True
                            continue
                        if cleaning:
                            for r in pk.runs:
                                r.text = ""
                    break
            break

    print("  Step 6: 参考文献")

    doc.save(out_path)

    # ========== Step 7: 正文循环 ==========
    _setup_body(out_path)
    print("  Step 7: 正文循环")

    # ========== Step 8: 删除文本框注释 ==========
    tmp = out_path + ".tmp"
    with zipfile.ZipFile(out_path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w') as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    xml = data.decode('utf-8')
                    n = xml.count('<mc:AlternateContent')
                    xml = re.sub(
                        r'<mc:AlternateContent>.*?</mc:AlternateContent>',
                        '', xml, flags=re.DOTALL)
                    print(f"  Step 8: 删除 {n} 个文本框注释")
                    data = xml.encode('utf-8')
                zout.writestr(item, data)
    shutil.move(tmp, out_path)


def _setup_body(doc_path):
    """手动设置正文循环。"""
    doc = Document(doc_path)
    paras = doc.paragraphs

    # 定位样本段落 —— 分两轮：先找 h1_idx，再从 h1 之后找其余样本
    h1_idx = h2_idx = h3_idx = body_idx = None
    body_end = len(paras)

    # 第一轮：找章标题和 body_end
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        if style == "Heading 1" and i > 40:
            if any(kw in text for kw in ["致", "参考", "附录"]):
                body_end = i
                break
            if h1_idx is None:
                h1_idx = i

    if not h1_idx:
        print("  警告: 未找到 Heading 1")
        return

    # 第二轮：从 h1_idx 之后找 h2/h3/body 样本
    for i in range(h1_idx + 1, body_end):
        p = paras[i]
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()

        if style == "Heading 2" and h2_idx is None:
            h2_idx = i

        if style == "Heading 3" and h3_idx is None:
            h3_idx = i

        # 正文段落: Normal / Body Text Indent 2 / 章节题目
        if body_idx is None and text and len(text) > 20:
            if style in ("Normal", "Body Text Indent 2", "章节题目"):
                body_idx = i

    if not h1_idx or not body_idx:
        print(f"  警告: 样本不足 h1={h1_idx} body={body_idx}")
        return

    keep = {h1_idx, h2_idx, h3_idx, body_idx} - {None}

    # 清空非样本段落
    for i in range(h1_idx, body_end):
        if i in keep:
            continue
        p = paras[i]
        if p._p.find('.//w:sectPr', _NS) is not None:
            continue
        for r in p.runs:
            r.text = ""
        style = p.style.name if p.style else ""
        if style.startswith("Heading"):
            p.style = doc.styles['Normal']

    # 删除正文区域示例表格
    body_elem = doc.element.body
    tables_rm = []
    pc = 0
    for elem in body_elem:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            pc += 1
        elif tag == 'tbl' and h1_idx <= pc <= body_end:
            tables_rm.append(elem)
    for tbl in tables_rm:
        body_elem.remove(tbl)

    # 替换样本文本
    def _set_text(para, text):
        if para.runs:
            para.runs[0].text = text
            for r in para.runs[1:]:
                r.text = ""

    paras[h1_idx].paragraph_format.page_break_before = True
    _set_text(paras[h1_idx], "{{ ch.title }}")
    if h2_idx:
        _set_text(paras[h2_idx], "{{ sec.title }}")
    if h3_idx:
        _set_text(paras[h3_idx], "{{ sub.title }}")
    _set_text(paras[body_idx], "{{ para }}")

    h1_p = paras[h1_idx]._p
    h2_p = paras[h2_idx]._p if h2_idx else None
    h3_p = paras[h3_idx]._p if h3_idx else None
    body_p = paras[body_idx]._p
    body_copy = copy.deepcopy(body_p)

    def _mk_body(text):
        p = copy.deepcopy(body_copy)
        for t_elem in p.findall('.//w:t', _NS):
            t_elem.text = ""
        t_elems = p.findall('.//w:t', _NS)
        if t_elems:
            t_elems[0].text = text
        return p

    # 移除样本，按正确顺序重新插入
    parent = h1_p.getparent()
    anchor = h1_p.getprevious()

    for elem in [h1_p, h2_p, h3_p, body_p]:
        if elem is not None and elem.getparent() is not None:
            parent.remove(elem)

    cursor = anchor

    def _after(cursor, elem):
        cursor.addnext(elem)
        return elem

    cursor = _after(cursor, _mk_ctrl("{%p for ch in chapters %}"))
    cursor = _after(cursor, h1_p)

    if h2_p is not None:
        cursor = _after(cursor, _mk_ctrl("{%p for sec in ch.sections %}"))
        cursor = _after(cursor, h2_p)
        cursor = _after(cursor, _mk_ctrl("{%p for para in sec.content %}"))
        cursor = _after(cursor, body_p)
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))  # content

        if h3_p is not None:
            cursor = _after(cursor, _mk_ctrl("{%p for sub in sec.subsections %}"))
            cursor = _after(cursor, h3_p)
            cursor = _after(cursor, _mk_ctrl("{%p for p2 in sub.content %}"))
            cursor = _after(cursor, _mk_body("{{ p2 }}"))
            cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))  # sub.content
            cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))  # subsections

        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))  # sections
    else:
        cursor = _after(cursor, _mk_ctrl("{%p for sec in ch.sections %}"))
        cursor = _after(cursor, _mk_ctrl("{%p for para in sec.content %}"))
        cursor = _after(cursor, body_p)
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))

    cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))  # chapters

    # 清正文末尾到致谢之间的残留段落
    in_refs = False
    for i in range(body_end, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        if style == "Heading 1":
            in_refs = "参考" in text
            continue
        if "acknowledgement" in text:
            continue
        if in_refs:
            if "{{ ref" in text or "{%p" in text:
                continue
        for r in p.runs:
            r.text = ""

    # 删除循环区域和致谢/参考文献/附录之间的空段落
    _remove_empty_between(doc, "{%p endfor %}", "致谢")
    _remove_empty_between(doc, "{%p endfor %}", "附录", start_from_last=True)

    doc.save(doc_path)


def _remove_empty_between(doc, start_marker, end_keyword, start_from_last=False):
    """删除 start_marker 和 end_keyword(Heading 1) 之间的空段落"""
    paras = doc.paragraphs
    # 找 start_marker 位置
    marker_idx = None
    for i, p in enumerate(paras):
        if start_marker in (p.text or ""):
            if start_from_last:
                marker_idx = i  # 取最后一个
            elif marker_idx is None:
                marker_idx = i

    # 找 end_keyword Heading 1 位置
    end_idx = None
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        if style == "Heading 1" and end_keyword in text and i > (marker_idx or 0):
            end_idx = i
            break

    if marker_idx is None or end_idx is None:
        return

    to_remove = []
    for i in range(marker_idx + 1, end_idx):
        p = paras[i]
        text = (p.text or "").strip()
        style = p.style.name if p.style else ""
        # 保留 Heading、循环标签、有内容的段落、sectPr
        if style.startswith("Heading"):
            continue
        if "{%p" in text or "{{" in text:
            continue
        if p._p.find('.//w:sectPr', _NS) is not None:
            continue
        if not text:
            to_remove.append(p._p)

    for elem in to_remove:
        if elem.getparent() is not None:
            elem.getparent().remove(elem)
    if to_remove:
        print(f"    删除 {len(to_remove)} 个空段落 ({start_marker[:20]}→{end_keyword})")


def _mk_ctrl(text):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python make.py <原始模板.docx> [输出路径]")
        sys.exit(1)
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else "template.docx"
    print(f"制作模板: {src} -> {out}")
    make(src, out)
    print(f"完成: {out}")
