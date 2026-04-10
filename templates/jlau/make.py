"""
吉林农业大学本科生毕业论文模板制作。
用法: cd thesis_project && python -m templates.jlau.make <原始模板.docx>

模板特点:
- 封面在单元格表格(table[0])内，字段用下划线空格占位
- 有三种目录格式（农理工/毕设/经管社科），只保留农理工类
- 章标题用 Heading 1 + "第X章" 格式
- 正文小四号宋体，首行缩进两字符，行间距固定值20磅
"""
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../.."))
from body_maker import setup_body_template
from refs_maker import setup_refs_template
from docx import Document


def _replace_cover_field(cell_para, label, jinja_var):
    """替换封面表格内的字段：清空下划线空格 run，写入 Jinja2 变量"""
    runs = cell_para.runs
    found_label = False
    for r in runs:
        if label in r.text:
            found_label = True
        elif found_label and r.font.underline:
            # 第一个下划线 run 写变量，其余清空
            if jinja_var:
                r.text = jinja_var
                jinja_var = None  # 只写一次
            else:
                r.text = ""


def _clear_instruction_paras(doc, start, end):
    """清空指定范围内的格式说明段落（含"号""磅""宋体"等关键字）"""
    inst_keywords = ["号", "磅", "宋体", "黑体", "加粗", "行间距",
                     "罗马体", "居中", "居左", "缩进", "标点",
                     "标题", "参考模板", "可参考"]
    for i in range(start, min(end, len(doc.paragraphs))):
        text = (doc.paragraphs[i].text or "").strip()
        if any(kw in text for kw in inst_keywords):
            for r in doc.paragraphs[i].runs:
                r.text = ""


def _setup_jlau_body(doc_path):
    """手动设置正文循环。

    吉林农大模板格式规范：
    - H1: Heading 1 样式, 三号(16pt)黑体加粗居中
    - H2: 四号(14pt)黑体居左, 无缩进
    - H3: 小四号(12pt)宋体居左, 无缩进
    - 正文: 小四号(12pt)宋体, 首行缩进2字符, 行间距固定20磅
    """
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, Emu
    from scanner import scan_structure, scan_body
    import copy

    _NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    doc = Document(doc_path)
    paras = doc.paragraphs
    s = scan_structure(doc)
    bs = s.get('body_start')
    be = s.get('body_end', len(paras))

    if not bs:
        print("  警告: 未找到正文区域")
        return

    body = scan_body(doc, bs, be)
    h1_idx = body['h1_samples'][0]['idx'] if body['h1_samples'] else None

    # 找真正的正文段落（有首行缩进342900 + 宋体 + 小四号）
    body_idx = None
    for sample in body.get('body_samples', []):
        idx = sample['idx']
        p = paras[idx]
        indent = p.paragraph_format.first_line_indent
        if indent and indent > 300000:  # ~342900 = 2字符缩进
            body_idx = idx
            break
    if body_idx is None and body['body_samples']:
        body_idx = body['body_samples'][0]['idx']

    if not h1_idx or not body_idx:
        print(f"  警告: 样本不足 h1={h1_idx} body={body_idx}")
        return

    keep = {h1_idx, body_idx}

    # 清空非样本段落
    _keep_kw = ['参考文献', '致谢', '致 谢', '致  谢', '附录', '附 录']
    for i in range(bs, be):
        if i in keep:
            continue
        p = paras[i]
        if p._p.find('.//w:sectPr', _NS) is not None:
            continue
        text = (p.text or "").strip()
        if any(kw in text for kw in _keep_kw):
            continue
        if p._p.find('.//w:numPr', _NS) is not None and len(text) > 5:
            continue
        for r in p.runs:
            r.text = ""

    # 删除正文区域的示例表格
    body_elem = doc.element.body
    tables_rm = []
    pc = 0
    for elem in body_elem:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            pc += 1
        elif tag == 'tbl' and bs <= pc <= be:
            tables_rm.append(elem)
    for tbl in tables_rm:
        body_elem.remove(tbl)

    # 替换样本文本
    def _set_text(para, text):
        if para.runs:
            para.runs[0].text = text
            for r in para.runs[1:]:
                r.text = ""

    paras[h1_idx].paragraph_format.page_break_before = True
    _set_text(paras[h1_idx], "{{ ch.title }}")
    _set_text(paras[body_idx], "{{ para }}")

    h1_p = paras[h1_idx]._p
    body_p = paras[body_idx]._p

    # 先保存 body_p 的完整 XML 副本（在移动前复制）
    body_copy = copy.deepcopy(body_p)

    # 确保 body 在 h1 后面
    h1_p.addnext(body_p)

    # 创建带格式的段落辅助函数
    def _mk_ctrl(text):
        """创建控制标签段落（无格式）"""
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        return p

    def _mk_heading(text, font_name, font_size_pt, bold=False):
        """创建标题段落：指定字体、字号、左对齐、无缩进"""
        p = OxmlElement('w:p')
        # 段落属性：左对齐，无缩进
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'left')
        pPr.append(jc)
        # 行间距固定20磅
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), '400')  # 20pt = 400 twips
        spacing.set(qn('w:lineRule'), 'exact')
        pPr.append(spacing)
        p.append(pPr)
        # Run
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        # 字体
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rPr.append(rFonts)
        # 字号（半磅单位）
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(font_size_pt * 2))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(font_size_pt * 2))
        rPr.append(szCs)
        # 加粗
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        return p

    def _mk_body(text):
        """创建正文段落：复制 body 样本格式"""
        p = copy.deepcopy(body_copy)
        for t_elem in p.findall('.//w:t', _NS):
            t_elem.text = ""
        t_elems = p.findall('.//w:t', _NS)
        if t_elems:
            t_elems[0].text = text
        return p

    # 构建完整循环结构（从下往上 addnext 插入）
    # 最终顺序:
    # {%p for ch %}
    # {{ ch.title }}          ← H1 (已有)
    # {%p for sec %}
    # {{ sec.title }}         ← H2: 四号(14pt)黑体
    # {%p for para %}
    # {{ para }}              ← body (已有)
    # {%p endfor %}           content
    # {%p for sub %}
    # {{ sub.title }}         ← H3: 小四号(12pt)宋体
    # {%p for p2 %}
    # {{ p2 }}                ← body 复制
    # {%p endfor %}           sub.content
    # {%p endfor %}           subsections
    # {%p endfor %}           sections
    # {%p endfor %}           chapters

    # 在 H1 前插 for ch
    h1_p.addprevious(_mk_ctrl("{%p for ch in chapters %}"))

    # 在 body_p 后面按顺序插入
    anchor = body_p

    # body_p 前面插 for sec + sec.title + for para
    # 需要反过来在 h1_p 后面从上到下构建
    # 先移除 body_p，重新按正确顺序插入所有元素
    parent = body_p.getparent()
    parent.remove(body_p)

    # h1_p 后面依次插入
    cursor = h1_p

    p_for_sec = _mk_ctrl("{%p for sec in ch.sections %}")
    cursor.addnext(p_for_sec)
    cursor = p_for_sec

    # H2: 四号黑体居左
    p_sec_title = _mk_heading("{{ sec.title }}", "黑体", 14, bold=False)
    cursor.addnext(p_sec_title)
    cursor = p_sec_title

    p_for_para = _mk_ctrl("{%p for para in sec.content %}")
    cursor.addnext(p_for_para)
    cursor = p_for_para

    # 正文段落（用原始 body 样本）
    cursor.addnext(body_p)
    cursor = body_p

    p_endfor_c = _mk_ctrl("{%p endfor %}")
    cursor.addnext(p_endfor_c)
    cursor = p_endfor_c

    p_for_sub = _mk_ctrl("{%p for sub in sec.subsections %}")
    cursor.addnext(p_for_sub)
    cursor = p_for_sub

    # H3: 小四号宋体居左
    p_sub_title = _mk_heading("{{ sub.title }}", "宋体", 12, bold=False)
    cursor.addnext(p_sub_title)
    cursor = p_sub_title

    p_for_p2 = _mk_ctrl("{%p for p2 in sub.content %}")
    cursor.addnext(p_for_p2)
    cursor = p_for_p2

    # p2 正文（复制 body 格式）
    p_p2 = _mk_body("{{ p2 }}")
    cursor.addnext(p_p2)
    cursor = p_p2

    # endfor sub.content
    cursor.addnext(_mk_ctrl("{%p endfor %}"))
    cursor = cursor.getnext()
    # endfor subsections
    cursor.addnext(_mk_ctrl("{%p endfor %}"))
    cursor = cursor.getnext()
    # endfor sections
    cursor.addnext(_mk_ctrl("{%p endfor %}"))
    cursor = cursor.getnext()
    # endfor chapters
    cursor.addnext(_mk_ctrl("{%p endfor %}"))

    doc.save(doc_path)
    print(f"  正文循环已设置: {doc_path}")


def make(src_path, out_path):
    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc = Document(src_path)

    # ========== Step 1: 封面 ==========
    cover_cell = doc.tables[0].cell(0, 0)
    cps = cover_cell.paragraphs

    # cp[10]: 题目
    _replace_cover_field(cps[10], "目：", "{{ title_zh }}")
    # 清空题目第二行（cp[11] 可能是续行）
    for r in cps[11].runs:
        r.text = ""

    # cp[12]: 姓名 + 学号
    _replace_cover_field(cps[12], "名：", "{{ name }}")
    _replace_cover_field(cps[12], "号：", "{{ student_id }}")

    # cp[14]: 学院 + 专业
    _replace_cover_field(cps[14], "院", "{{ college }}")
    _replace_cover_field(cps[14], "业", "{{ major }}")

    # cp[16]: 指导教师 + 职称
    _replace_cover_field(cps[16], "教师：", "{{ advisor }}")
    _replace_cover_field(cps[16], "称：", "{{ advisor_title }}")

    # cp[23]: 日期
    if cps[23].runs:
        cps[23].runs[0].text = "{{ year }}年  {{ month }}月  {{ day }}日"
        for r in cps[23].runs[1:]:
            r.text = ""

    print("  Step 1: 封面")

    # ========== Step 2: 中文摘要 ==========
    # p[2]: 论文题目
    p2 = doc.paragraphs[2]
    if p2.runs:
        p2.runs[0].text = "{{ title_zh }}"
        for r in p2.runs[1:]:
            r.text = ""

    # p[4]: "摘  要" 标题 - 清掉格式说明和括号
    p4 = doc.paragraphs[4]
    for r in p4.runs:
        if "号" in r.text or "黑体" in r.text or r.text.strip() in ("（", "）", "（）"):
            r.text = ""

    # p[6]: 摘要正文
    p6 = doc.paragraphs[6]
    if p6.runs:
        p6.runs[0].text = "{{ abstract_zh }}"
        for r in p6.runs[1:]:
            r.text = ""
    # 清空摘要格式说明段落 p[7]-p[13]
    for i in range(7, 14):
        for r in doc.paragraphs[i].runs:
            r.text = ""

    # p[14]: 关键词
    p14 = doc.paragraphs[14]
    if p14.runs:
        p14.runs[0].text = "关键词："
        if len(p14.runs) > 1:
            p14.runs[1].text = "{{ keywords_zh }}"
        for r in p14.runs[2:]:
            r.text = ""

    print("  Step 2: 中文摘要")

    # ========== Step 3: 英文摘要 ==========
    # p[17]: 英文题目
    p17 = doc.paragraphs[17]
    if p17.runs:
        p17.runs[0].text = "{{ title_en }}"
        for r in p17.runs[1:]:
            r.text = ""

    # p[19]: ABSTRACT 标题 - 清掉说明和括号
    p19 = doc.paragraphs[19]
    for r in p19.runs:
        if "号" in r.text or "Roman" in r.text or "加粗" in r.text:
            r.text = ""
        if r.text.strip() in ("（", "）", "（）"):
            r.text = ""

    # p[21]: 英文摘要正文
    p21 = doc.paragraphs[21]
    if p21.runs:
        p21.runs[0].text = "{{ abstract_en }}"
        for r in p21.runs[1:]:
            r.text = ""
    # 清空英文摘要说明 p[22]-p[24]
    for i in range(22, 25):
        for r in doc.paragraphs[i].runs:
            r.text = ""

    # 清除目录标题的残留括号
    p29 = doc.paragraphs[29]
    for r in p29.runs:
        if r.text.strip() in ("（", "）"):
            r.text = ""

    # p[25]: KEY WORDS
    p25 = doc.paragraphs[25]
    if p25.runs:
        p25.runs[0].text = "KEY WORDS: "
        if len(p25.runs) > 1:
            p25.runs[1].text = "{{ keywords_en }}"
        for r in p25.runs[2:]:
            r.text = ""

    print("  Step 3: 英文摘要")

    # ========== Step 4: 目录 - 只保留第一种（农理工），删除后两种 ==========
    # 第一种目录: p[29]-p[45]（保留，清掉标题里的说明文字）
    p29 = doc.paragraphs[29]
    for r in p29.runs:
        if "（" in r.text or "可参考" in r.text:
            r.text = ""

    # 第二种目录: p[49]-p[65]（毕业设计）-> 清空
    # 第三种目录: p[68]-p[84]（经管社科）-> 清空
    for i in range(46, 86):
        for r in doc.paragraphs[i].runs:
            r.text = ""

    print("  Step 4: 目录（保留农理工类）")

    # ========== Step 5: 正文区域格式说明清理 ==========
    # Heading 1 标题段落：删除格式说明 run，保留标题文字
    for i in [88, 99, 135, 143, 175, 213]:
        if i >= len(doc.paragraphs):
            continue
        p = doc.paragraphs[i]
        # 先拼出完整文本判断，再逐 run 清理
        full = p.text or ""
        for r in p.runs:
            t = r.text
            # 删掉独立的括号和格式说明 run
            if t.strip() in ("（", "）", "（正文）"):
                r.text = ""
            elif any(kw in t for kw in
                     ["号", "黑体", "加粗", "Roman", "Times"]):
                r.text = ""
            # "附录名称" → "附录"
            if "名称" in t:
                r.text = t.replace("名称", "")

    # 清空正文区域的格式说明段落（非标题非样本）
    # 保留 p[90] 作为正文格式样本（首行缩进+宋体+小四号）
    heading_indices = {88, 90, 99, 135, 143, 175, 213, 219}
    for i in range(89, 145):
        if i in heading_indices:
            continue
        p = doc.paragraphs[i]
        style = p.style.name if p.style else ""
        if style == "Heading 1":
            continue
        text = (p.text or "").strip()
        # 保留样本段落：表X-1、图X-1、(X-1) 等
        if text and ("表X" in text or "图X" in text or "（X-" in text
                     or "续表" in text):
            continue
        # 清掉说明文字段落
        if any(kw in text for kw in
               ["号", "磅", "宋体", "黑体", "加粗", "行间距",
                "居中", "居左", "缩进", "标点", "参考模板",
                "文献综述", "实验", "结论是对", "结论要严格",
                "展望或建议", "人文社", "最后一章"]):
            for r in p.runs:
                r.text = ""

    # 清空参考文献格式说明和示例 p[145]-p[172]，保留 p[150] 作为样本
    for i in range(145, 173):
        if i == 150:
            continue  # 保留第一条参考文献示例作为格式样本
        for r in doc.paragraphs[i].runs:
            r.text = ""

    # 清空附录/学术成果/致谢的说明段落
    for i in range(176, 225):
        if i >= len(doc.paragraphs) or i in heading_indices:
            continue
        text = (doc.paragraphs[i].text or "").strip()
        if any(kw in text for kw in
               ["号", "磅", "宋体", "黑体", "行间距", "致谢",
                "字数", "感谢"]):
            for r in doc.paragraphs[i].runs:
                r.text = ""

    print("  Step 5: 清理格式说明")

    # 保存中间结果
    doc.save(out_path)

    # ========== Step 6: 正文 Jinja2 循环 ==========
    # 吉林农大模板没有 H2/H3 样本，手动设置循环结构
    _setup_jlau_body(out_path)
    print("  Step 6: 正文循环")

    # ========== Step 7: 参考文献 + 致谢 ==========
    doc = Document(out_path)
    refs_done = False
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if text == "参考文献" and i > 60 and not refs_done:
            # 找参考文献标题后第一个有 runs 的段落作为样本
            for j in range(i + 1, min(i + 20, len(doc.paragraphs))):
                pj = doc.paragraphs[j]
                if pj.runs:
                    # 用这个段落做参考文献模板
                    pj.runs[0].text = "{{ ref }}"
                    for r in pj.runs[1:]:
                        r.text = ""
                    # 在它前面插 for 循环
                    from editor import TemplateEditor
                    ed_p = pj._p
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn as _qn
                    for_p = OxmlElement('w:p')
                    for_r = OxmlElement('w:r')
                    for_t = OxmlElement('w:t')
                    for_t.set(_qn('xml:space'), 'preserve')
                    for_t.text = "{%p for ref in references %}"
                    for_r.append(for_t)
                    for_p.append(for_r)
                    ed_p.addprevious(for_p)
                    # 在它后面插 endfor
                    end_p = OxmlElement('w:p')
                    end_r = OxmlElement('w:r')
                    end_t = OxmlElement('w:t')
                    end_t.set(_qn('xml:space'), 'preserve')
                    end_t.text = "{%p endfor %}"
                    end_r.append(end_t)
                    end_p.append(end_r)
                    ed_p.addnext(end_p)
                    refs_done = True
                    break
            if not refs_done:
                # 没有非空段落，尝试 refs_maker
                setup_refs_template(doc, i)
                refs_done = True

        if text in ("致  谢", "致谢"):
            for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                pt = (doc.paragraphs[j].text or "").strip()
                if pt or doc.paragraphs[j].runs:
                    if doc.paragraphs[j].runs:
                        doc.paragraphs[j].runs[0].text = \
                            "    {{ acknowledgement }}"
                        for r in doc.paragraphs[j].runs[1:]:
                            r.text = ""
                    break
    # 压缩摘要区域多余空行的行间距
    from docx.shared import Pt as _Pt
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if not text and not p.runs:
            # 检查是否在摘要区域（前30个段落内）
            if i < 30:
                pf = p.paragraph_format
                pf.space_before = _Pt(0)
                pf.space_after = _Pt(0)
                pf.line_spacing = _Pt(1)

    doc.save(out_path)
    print("  Step 7: 参考文献/致谢")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python make.py <原始模板.docx> [输出路径]")
        sys.exit(1)
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else "template.docx"
    print(f"制作模板: {src} -> {out}")
    make(src, out)
    print(f"完成: {out}")
