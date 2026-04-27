"""
信科院 2024 本科毕业论文 docxtpl 模板制作脚本。

项目约定：
- source.docx 保存用户提供的学院参考模板原件；
- template.docx 是可重复渲染的 docxtpl 模板；
- 论文内容不写在模板脚本里，由 papers/<paper>/JSON + build.py 注入。
"""
from __future__ import annotations

import os
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


HERE = Path(__file__).resolve().parent
DEFAULT_SRC = HERE / "source.docx"
DEFAULT_OUT = HERE / "template.docx"


def set_run_font(run, size=12, east="宋体", ascii_font="Times New Roman", bold=False):
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.bold = bold
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), east)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)


def set_para_line(p, line_pt=23, first_line=False):
    p.paragraph_format.line_spacing = Pt(line_pt)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)


def clear_body(doc: Document):
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_margins(doc: Document):
    for section in doc.sections:
        section.top_margin = Cm(3.4)
        section.bottom_margin = Cm(2.8)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.6)


def add_center(doc: Document, text: str, size=12, east="宋体", bold=False, line_pt=23):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_line(p, line_pt=line_pt)
    r = p.add_run(text)
    set_run_font(r, size=size, east=east, bold=bold)
    return p


def add_body_para(doc: Document, text: str, first_line=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_line(p, first_line=first_line)
    r = p.add_run(text)
    set_run_font(r, size=12, east="宋体")
    return p


def add_title(doc: Document, text: str, page_break_before=False):
    p = doc.add_paragraph()
    if page_break_before:
        p.paragraph_format.page_break_before = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_line(p, line_pt=23)
    r = p.add_run(text)
    set_run_font(r, size=16, east="黑体", bold=True)
    return p


def add_cover_field(doc: Document, label: str, placeholder: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}：{placeholder}")
    set_run_font(r, size=16, east="黑体", bold=True)


def add_toc_placeholder(doc: Document):
    add_title(doc, "目 录", page_break_before=True)
    p = doc.add_paragraph("{{p toc_doc}}")
    set_para_line(p)


def add_subdoc_section(doc: Document, title: str, tag: str, page_break_before=True):
    add_title(doc, title, page_break_before=page_break_before)
    p = doc.add_paragraph(f"{{{{p {tag}}}}}")
    set_para_line(p)


def clear_paragraph(p):
    for run in list(p.runs):
        p._p.remove(run._r)


def set_page_number(section, fmt="decimal", start=1, enabled=True):
    section.footer.is_linked_to_previous = False
    footer = section.footer
    for p in footer.paragraphs:
        clear_paragraph(p)
    if not footer.paragraphs:
        p = footer.add_paragraph()
    else:
        p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    if enabled:
        run = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)
        set_run_font(run, size=10.5)

    sect_pr = section._sectPr
    old = sect_pr.find(qn("w:pgNumType"))
    if old is not None:
        sect_pr.remove(old)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), str(start))
    pg.set(qn("w:fmt"), fmt)
    sect_pr.append(pg)


def configure_sections(doc: Document):
    set_margins(doc)
    for idx, section in enumerate(doc.sections):
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            clear_paragraph(p)
        if idx == 0:
            set_page_number(section, enabled=False)
        elif idx == 1:
            set_page_number(section, fmt="upperRoman", start=1, enabled=True)
        else:
            set_page_number(section, fmt="decimal", start=1, enabled=True)


def patch_update_fields(docx_path: Path):
    tmp = Path(str(docx_path) + ".tmp")
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "word/settings.xml":
                    xml = data.decode("utf-8")
                    if "w:updateFields" not in xml:
                        xml = xml.replace("</w:settings>", '<w:updateFields w:val="true"/></w:settings>')
                    data = xml.encode("utf-8")
                zout.writestr(info, data)
    tmp.replace(docx_path)


def make(src_path: str | os.PathLike = DEFAULT_SRC, out_path: str | os.PathLike = DEFAULT_OUT):
    src_path = Path(src_path)
    out_path = Path(out_path)
    if not src_path.exists():
        raise FileNotFoundError(f"未找到源模板: {src_path}")

    doc = Document(str(src_path))
    clear_body(doc)
    set_margins(doc)

    # 封面。保留学院模板的版心和页边距，字段由论文 meta.json 渲染。
    for _ in range(4):
        doc.add_paragraph()
    add_center(doc, "{{ title_zh }}", size=26, east="黑体", bold=True, line_pt=42)
    for _ in range(5):
        doc.add_paragraph()
    add_cover_field(doc, "学院", "{{ college }}")
    add_cover_field(doc, "专业班级", "{{ major_class }}")
    add_cover_field(doc, "姓名", "{{ author }}")
    add_cover_field(doc, "学号", "{{ student_id }}")
    add_cover_field(doc, "指导教师", "{{ advisor }}")
    for _ in range(2):
        doc.add_paragraph()
    add_center(doc, "{{ year }}年{{ month }}月", size=16, east="黑体", bold=True)

    # 中英文题名页。
    add_center(doc, "{{ title_zh }}", size=16, east="黑体", bold=True).paragraph_format.page_break_before = True
    add_center(doc, "{{ title_en }}", size=14, east="Times New Roman", bold=True)

    # 中文摘要。
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_title(doc, "摘 要", page_break_before=False)
    add_body_para(doc, "{{ abstract_zh_1 }}")
    add_body_para(doc, "{{ abstract_zh_2 }}")
    add_body_para(doc, "{{ abstract_zh_3 }}")
    p = add_body_para(doc, "关键词：{{ keywords_zh_text }}", first_line=False)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 英文摘要。
    add_title(doc, "Abstract", page_break_before=True)
    add_body_para(doc, "{{ abstract_en_1 }}")
    add_body_para(doc, "{{ abstract_en_2 }}")
    add_body_para(doc, "{{ abstract_en_3 }}")
    p = add_body_para(doc, "Key words: {{ keywords_en_text }}", first_line=False)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_toc_placeholder(doc)
    doc.add_section(WD_SECTION.NEW_PAGE)
    p = doc.add_paragraph("{{p body}}")
    set_para_line(p)
    add_subdoc_section(doc, "结论", "conclusion_doc", page_break_before=True)
    add_subdoc_section(doc, "致谢", "acknowledgement_doc", page_break_before=True)
    add_subdoc_section(doc, "参考文献", "references_doc", page_break_before=True)
    add_subdoc_section(doc, "附录", "appendix_doc", page_break_before=True)

    configure_sections(doc)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    patch_update_fields(out_path)

    text = "\n".join(p.text for p in Document(str(out_path)).paragraphs)
    for tag in ("{{p toc_doc}}", "{{p body}}", "{{p conclusion_doc}}", "{{p references_doc}}"):
        if tag not in text:
            raise RuntimeError(f"模板缺少占位符: {tag}")
    print(f"模板已生成: {out_path}")


def main():
    src = sys.argv[1] if len(sys.argv) > 1 else str(DEFAULT_SRC)
    out = sys.argv[2] if len(sys.argv) > 2 else str(DEFAULT_OUT)
    make(src, out)


if __name__ == "__main__":
    main()
