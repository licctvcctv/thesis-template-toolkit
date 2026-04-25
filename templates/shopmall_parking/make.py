"""
购物商城论文模板清理脚本。

用法:
    cd /Users/a136/vs/45425/thesis_project
    python templates/shopmall_parking/make.py \
        templates/shopmall_parking/source.docx \
        templates/shopmall_parking/template.docx

脚本会保留原文档的声明、摘要、页眉页脚、页面设置和样式，清理原正文、
旧参考文献、空白目录页，并插入 docxtpl 占位符。
"""
from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def para_text(p):
    return re.sub(r"\s+", "", p.text or "")


def find_para(paragraphs, predicate, label):
    for i, p in enumerate(paragraphs):
        if predicate(p.text or "", p):
            return i, p
    raise RuntimeError(f"未找到段落: {label}")


def remove_para(p):
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def remove_block_range(start_el, end_el=None, include_start=True, include_end=False):
    body = start_el.getparent()
    children = list(body)
    start = children.index(start_el)
    if end_el is None:
        end = len(children)
    else:
        end = children.index(end_el)
        if include_end:
            end += 1
    if not include_start:
        start += 1
    for el in children[start:end]:
        if el.tag == qn("w:sectPr"):
            continue
        body.remove(el)


def clear_para(p, text):
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def set_single_run(p, text, size=None, bold=None):
    clear_para(p, text)
    for r in p.runs:
        if size is not None:
            r.font.size = Pt(size)
        if bold is not None:
            r.bold = bold


def strip_shading(p):
    """Remove inherited gray background from reused abstract paragraphs."""
    for shd in list(p._p.xpath(".//w:shd")):
        parent = shd.getparent()
        if parent is not None:
            parent.remove(shd)


def ensure_page_break_before(p):
    p.paragraph_format.page_break_before = True


def add_toc_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(instr)

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run._r.append(sep)

    t = OxmlElement("w:t")
    t.text = "目录更新后显示"
    run._r.append(t)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def patch_update_fields(docx_path: Path):
    with zipfile.ZipFile(docx_path, "a") as zf:
        try:
            xml = zf.read("word/settings.xml").decode("utf-8")
        except KeyError:
            xml = (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                f'<w:settings xmlns:w="{W_NS}"></w:settings>'
            )
        if "w:updateFields" not in xml:
            xml = xml.replace(
                "</w:settings>",
                '<w:updateFields w:val="true"/></w:settings>',
            )
        zf.writestr("word/settings.xml", xml)


def main():
    if len(sys.argv) < 2:
        raise SystemExit("用法: python make.py <source.docx> [template.docx]")

    source = Path(sys.argv[1])
    output = Path(sys.argv[2]) if len(sys.argv) > 2 else Path(__file__).with_name("template.docx")

    doc = Document(source)
    paragraphs = list(doc.paragraphs)

    decl_idx, decl_title = find_para(paragraphs, lambda t, p: para_text(p) == "郑重声明", "声明标题")
    zh_title_idx, zh_title = find_para(paragraphs, lambda t, p: para_text(p) == "摘要", "中文摘要标题")
    zh_key_idx, zh_key = find_para(paragraphs, lambda t, p: t.strip().startswith("关键字"), "中文关键词")
    en_title_idx, en_title = find_para(paragraphs, lambda t, p: t.strip() == "Abstract", "英文摘要标题")
    en_key_idx, en_key = find_para(paragraphs, lambda t, p: t.strip().startswith("Keywords"), "英文关键词")
    body_idx, body_start = find_para(paragraphs, lambda t, p: para_text(p).startswith("1前言"), "正文起点")

    print(f"摘要: P{zh_title_idx}, Abstract: P{en_title_idx}, 正文: P{body_idx}")

    # 清掉源模板开头残留的空白节/空白段，避免办公套件渲染出多张空白页。
    first_el = list(doc.element.body)[0]
    if first_el is not decl_title._element:
        remove_block_range(first_el, decl_title._element, include_start=True, include_end=False)

    _, signature = find_para(paragraphs, lambda t, p: t.strip().startswith("本人签名"), "声明签名")
    remove_block_range(signature._element, zh_title._element, include_start=False, include_end=False)
    ensure_page_break_before(zh_title)
    remove_block_range(zh_key._element, en_title._element, include_start=False, include_end=False)

    # 中文摘要：保留三段段落格式，删除多余旧摘要段。
    zh_abs = [p for p in paragraphs[zh_title_idx + 1:zh_key_idx] if (p.text or "").strip()]
    for i, p in enumerate(zh_abs):
        if i < 3:
            set_single_run(p, "{{ abstract_zh_%d }}" % (i + 1))
            strip_shading(p)
        else:
            remove_para(p)
    set_single_run(zh_key, "关键字：{{ keywords_zh }}")
    strip_shading(zh_key)

    # 英文摘要：同样保留三段。
    en_abs = [p for p in paragraphs[en_title_idx + 1:en_key_idx] if (p.text or "").strip()]
    for i, p in enumerate(en_abs):
        if i < 3:
            set_single_run(p, "{{ abstract_en_%d }}" % (i + 1))
            strip_shading(p)
        else:
            remove_para(p)
    set_single_run(en_key, "Keywords: {{ keywords_en }}")
    strip_shading(en_key)
    ensure_page_break_before(en_title)

    # 删除英文关键词之后到正文之前的旧目录、空白页和表格遗留。
    paragraphs = list(doc.paragraphs)
    _, en_key = find_para(paragraphs, lambda t, p: t.strip().startswith("Keywords"), "英文关键词")
    _, body_start = find_para(paragraphs, lambda t, p: para_text(p).startswith("1前言"), "正文起点")
    remove_block_range(en_key._element, body_start._element, include_start=False, include_end=False)

    # 清掉旧正文、旧参考文献，以及其中的表格等块级元素。
    paragraphs = list(doc.paragraphs)
    _, body_start = find_para(paragraphs, lambda t, p: para_text(p).startswith("1前言"), "正文起点")
    remove_block_range(body_start._element)

    # 目录页：使用 Word/WPS 可更新目录域。
    toc_title = doc.add_paragraph()
    ensure_page_break_before(toc_title)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = toc_title.add_run("目 录")
    r.bold = True
    r.font.size = Pt(16)

    toc_hint = doc.add_paragraph()
    add_toc_field(toc_hint)

    # 正文和参考文献占位由论文 build.py 生成子文档填入。
    body_p = doc.add_paragraph("{{p body}}")
    ensure_page_break_before(body_p)

    ref_title = doc.add_paragraph()
    ensure_page_break_before(ref_title)
    ref_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = ref_title.add_run("参考文献")
    rr.bold = True
    rr.font.size = Pt(16)

    doc.add_paragraph("{{p references_doc}}")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    patch_update_fields(output)
    print(f"模板已生成: {output}")


if __name__ == "__main__":
    main()
