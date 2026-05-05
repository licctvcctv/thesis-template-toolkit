"""
Convert 薛艳.docx into a reusable docxtpl thesis template.

The source document is used as a formatting reference for the hospital-style
paper: cover, abstract pages, TOC page, headers/footers, heading typography,
caption style, body spacing, references, and acknowledgement order are kept.

Usage:
    python make.py [source.docx] [template.docx]
"""
import copy
import os
import re
import shutil
import sys
import tempfile
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


HERE = os.path.dirname(os.path.abspath(__file__))
DEFAULT_SOURCE = os.path.join(HERE, "source.docx")
DEFAULT_OUTPUT = os.path.join(HERE, "template.docx")


def main():
    src = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_SOURCE
    out = sys.argv[2] if len(sys.argv) > 2 else DEFAULT_OUTPUT

    doc = Document(src)
    anchors = _find_anchors(doc)
    _replace_cover(doc)
    _move_toc_after_abstracts(doc, anchors)
    anchors = _find_anchors(doc)
    _replace_front_matter(doc, anchors)
    anchors = _find_anchors(doc)
    _remove_empty_between(doc, doc.paragraphs[21], doc.paragraphs[anchors["abstract_zh"]])
    anchors = _find_anchors(doc)
    _replace_toc(doc, anchors)
    anchors = _find_anchors(doc)
    _replace_body(doc, anchors)
    anchors = _find_anchors(doc)
    _replace_tail(doc, anchors)
    _set_page_breaks(doc)

    os.makedirs(os.path.dirname(os.path.abspath(out)) or ".", exist_ok=True)
    doc.save(out)
    _mark_fields_update_on_open(out)
    _verify(out)
    print(f"Saved: {out}")


def _norm(text):
    return (text or "").strip().replace(" ", "").replace("\u3000", "")


def _find_anchors(doc):
    anchors = {}
    seen_toc = False
    seen_abstract = False
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        compact = _norm(text)
        style = p.style.name if p.style else ""

        if compact == "摘要" and "abstract_zh" not in anchors:
            anchors["abstract_zh"] = i
            seen_abstract = True
        elif compact == "Abstract" and "abstract_en" not in anchors:
            anchors["abstract_en"] = i
        elif compact == "目录" and "toc" not in anchors:
            anchors["toc"] = i
            seen_toc = True
        elif seen_toc and style == "Heading 1" and "body_start" not in anchors:
            anchors["body_start"] = i
        elif "body_start" in anchors and compact in {"总结与展望", "总结", "结论"} and "conclusion" not in anchors:
            anchors["conclusion"] = i
        elif compact == "参考文献" and "refs" not in anchors:
            anchors["refs"] = i
        elif compact == "致谢" and "ack" not in anchors:
            anchors["ack"] = i

    required = ["abstract_zh", "abstract_en", "toc", "body_start", "conclusion", "refs", "ack"]
    missing = [k for k in required if k not in anchors]
    if missing:
        raise RuntimeError(f"Missing anchors: {missing}; found={anchors}")
    return anchors


def _replace_cover(doc):
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        compact = _norm(text)
        if text.startswith("学") and "号" in text and "2018206035" in text:
            _replace_para_text(p, "学  号：{{ student_id }}")
        elif text == "基于B/S结构的医院科研信息管理系统的设计与实现":
            _replace_para_text(p, "{{ title_zh }}")
        elif text == "The Design and Implementation of Hospital Scientific Research Information Management System Based on B / S Structure":
            _replace_para_text(p, "{{ title_en }}")
        elif compact == "二O二二年六月":
            _replace_para_text(p, "{{ year_month }}")

    if doc.tables:
        cell = doc.tables[0].cell(0, 0)
        for p in cell.paragraphs:
            text = (p.text or "").strip()
            if text.startswith("专"):
                _replace_para_text(p, "专    业：{{ major }}")
            elif text.startswith("学"):
                _replace_para_text(p, "学    生：{{ name }}")
            elif text.startswith("指导教师"):
                _replace_para_text(p, "指导教师：{{ advisor }}")


def _move_toc_after_abstracts(doc, anchors):
    """The source places TOC before abstracts; the generator expects abstracts before TOC."""
    body = doc.element.body
    children = list(body)
    toc = doc.paragraphs[anchors["toc"]]._p
    abstract_zh = doc.paragraphs[anchors["abstract_zh"]]._p
    body_start = doc.paragraphs[anchors["body_start"]]._p

    toc_i = children.index(toc)
    abstract_i = children.index(abstract_zh)
    if toc_i > abstract_i:
        return

    moving = children[toc_i:abstract_i]
    for elem in moving:
        body.remove(elem)
    insert_at = list(body).index(body_start)
    for offset, elem in enumerate(moving):
        body.insert(insert_at + offset, elem)


def _replace_front_matter(doc, anchors):
    paras = doc.paragraphs
    zh_title = paras[anchors["abstract_zh"]]
    en_title = paras[anchors["abstract_en"]]
    toc_title = paras[anchors["toc"]]

    zh_body = _first_after_index(doc, anchors["abstract_zh"], "Normal")
    zh_kw = _find_between_text(doc, zh_title, en_title, "关键词")
    en_body = _first_after_index(doc, anchors["abstract_en"], "Normal")
    en_kw = _find_between_text(doc, en_title, toc_title, "Keywords")

    _replace_range_with_blocks(
        doc,
        zh_title,
        en_title,
        [
            _clone_with_text(zh_body, "{{ abstract_zh }}"),
            _clone_with_text(zh_kw or zh_body, "关键词：{{ keywords_zh }}"),
        ],
    )

    anchors = _find_anchors(doc)
    paras = doc.paragraphs
    en_title = paras[anchors["abstract_en"]]
    toc_title = paras[anchors["toc"]]
    en_body = _first_after_index(doc, anchors["abstract_en"], "Normal")
    en_kw = _find_between_text(doc, en_title, toc_title, "Keywords")
    _replace_range_with_blocks(
        doc,
        en_title,
        toc_title,
        [
            _clone_with_text(en_body, "{{ abstract_en }}"),
            _clone_with_text(en_kw or en_body, "Keywords: {{ keywords_en }}"),
        ],
    )


def _replace_toc(doc, anchors):
    paras = doc.paragraphs
    toc_title = paras[anchors["toc"]]
    body_start = paras[anchors["body_start"]]
    sample = paras[anchors["toc"] + 1]
    _replace_range_with_blocks(
        doc,
        toc_title,
        body_start,
        [_clone_with_text(sample, "{{ toc_placeholder }}")],
    )


def _replace_body(doc, anchors):
    paras = doc.paragraphs
    body_start = paras[anchors["body_start"]]
    conclusion = paras[anchors["conclusion"]]
    p_h1 = body_start
    p_h2 = _first_after(doc, anchors["body_start"], "Heading 2")
    p_h3 = _first_after(doc, anchors["body_start"], "Heading 3")
    p_body = _first_after(doc, anchors["body_start"], "Normal")

    blocks = [
        _clone_with_text(p_body, "{%p for ch in chapters %}"),
        _clone_with_text(p_h1, "{{ ch.title }}", no_number=True),
        _clone_with_text(p_body, "{%p for item in ch.content %}"),
        _clone_with_text(p_body, "{{ item }}"),
        _clone_with_text(p_body, "{%p endfor %}"),
        _clone_with_text(p_body, "{%p for sec in ch.sections %}"),
        _clone_with_text(p_h2, "{{ sec.title }}", no_number=True),
        _clone_with_text(p_body, "{%p for item in sec.content %}"),
        _clone_with_text(p_body, "{{ item }}"),
        _clone_with_text(p_body, "{%p endfor %}"),
        _clone_with_text(p_body, "{%p for sub in sec.subsections %}"),
        _clone_with_text(p_h3, "{{ sub.title }}", no_number=True),
        _clone_with_text(p_body, "{%p for item in sub.content %}"),
        _clone_with_text(p_body, "{{ item }}"),
        _clone_with_text(p_body, "{%p endfor %}"),
        _clone_with_text(p_body, "{%p endfor %}"),
        _clone_with_text(p_body, "{%p endfor %}"),
        _clone_with_text(p_body, "{%p endfor %}"),
    ]
    _replace_range_with_blocks(doc, body_start, conclusion, blocks, include_start=True)


def _replace_tail(doc, anchors):
    paras = doc.paragraphs
    conclusion = paras[anchors["conclusion"]]
    refs = paras[anchors["refs"]]
    ack = paras[anchors["ack"]]
    _disable_numbering(conclusion)
    body_sample = _first_between(doc, conclusion, refs) or _first_after(doc, anchors["body_start"], "Normal")
    ref_sample = _first_between(doc, refs, ack) or body_sample
    ack_sample = _first_after_index(doc, anchors["ack"], "Normal") or body_sample

    _replace_range_with_blocks(
        doc,
        conclusion,
        refs,
        [
            _clone_with_text(body_sample, "{%p for para in conclusion_list %}"),
            _clone_with_text(body_sample, "{{ para }}"),
            _clone_with_text(body_sample, "{%p endfor %}"),
        ],
    )

    anchors = _find_anchors(doc)
    refs = doc.paragraphs[anchors["refs"]]
    ack = doc.paragraphs[anchors["ack"]]
    _replace_range_with_blocks(
        doc,
        refs,
        ack,
        [
            _clone_with_text(ref_sample, "{%p for ref in references %}"),
            _clone_with_text(ref_sample, "{{ ref }}"),
            _clone_with_text(ref_sample, "{%p endfor %}"),
        ],
    )

    anchors = _find_anchors(doc)
    ack = doc.paragraphs[anchors["ack"]]
    body = doc.element.body
    children = list(body)
    start = children.index(ack._p) + 1
    tail_sect = children[-1] if children and children[-1].tag == qn("w:sectPr") else None
    for elem in children[start:]:
        if elem is not tail_sect:
            body.remove(elem)
    body.insert(start, _clone_with_text(ack_sample, "{{ acknowledgement }}"))


def _replace_range_with_blocks(doc, start_para, stop_para, blocks, include_start=False):
    body = doc.element.body
    children = list(body)
    start = children.index(start_para._p) + (0 if include_start else 1)
    stop = children.index(stop_para._p)
    for elem in children[start:stop]:
        body.remove(elem)
    for offset, elem in enumerate(blocks):
        body.insert(start + offset, elem)


def _clone_with_text(sample_para, text, no_number=False):
    elem = OxmlElement("w:p")
    ppr = sample_para._p.pPr
    if ppr is not None:
        ppr2 = copy.deepcopy(ppr)
        for tag in ["w:numPr", "w:pageBreakBefore"]:
            node = ppr2.find(qn(tag))
            if node is not None:
                ppr2.remove(node)
        if no_number:
            num_pr = OxmlElement("w:numPr")
            num_id = OxmlElement("w:numId")
            num_id.set(qn("w:val"), "0")
            num_pr.append(num_id)
            ppr2.append(num_pr)
        elem.append(ppr2)

    run = OxmlElement("w:r")
    if sample_para.runs:
        rpr = sample_para.runs[0]._r.rPr
        if rpr is not None:
            run.append(copy.deepcopy(rpr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    elem.append(run)
    return elem


def _replace_para_text(p, text):
    if not p.runs:
        p.add_run("")
    p.runs[0].text = text
    for r in p.runs[1:]:
        r.text = ""


def _disable_numbering(p):
    ppr = p._p.pPr
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        p._p.insert(0, ppr)
    old = ppr.find(qn("w:numPr"))
    if old is not None:
        ppr.remove(old)
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "0")
    num_pr.append(num_id)
    ppr.append(num_pr)


def _remove_empty_between(doc, start_para, stop_para):
    body = doc.element.body
    children = list(body)
    start = children.index(start_para._p) + 1
    stop = children.index(stop_para._p)
    for elem in list(children[start:stop]):
        if elem.tag != qn("w:p"):
            continue
        texts = [t.text or "" for t in elem.findall(".//" + qn("w:t"))]
        has_drawing = bool(elem.findall(".//" + qn("w:drawing")))
        has_sect = elem.find("w:pPr/w:sectPr", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}) is not None
        if not "".join(texts).strip() and not has_drawing and not has_sect:
            body.remove(elem)


def _first_after(doc, idx, style_name):
    found = _first_after_index(doc, idx, style_name)
    if found is None:
        raise RuntimeError(f"No paragraph with style {style_name} after {idx}")
    return found


def _first_after_index(doc, idx, style_name):
    for p in doc.paragraphs[idx + 1:]:
        if p.style and p.style.name == style_name and p.text.strip():
            return p
    return None


def _find_between_text(doc, start_para, stop_para, needle):
    paras = doc.paragraphs
    start = _para_index(doc, start_para)
    stop = _para_index(doc, stop_para)
    for p in paras[start + 1:stop]:
        if needle in (p.text or ""):
            return p
    return None


def _first_between(doc, start_para, stop_para):
    paras = doc.paragraphs
    start = _para_index(doc, start_para)
    stop = _para_index(doc, stop_para)
    for p in paras[start + 1:stop]:
        if p.text.strip():
            return p
    return None


def _para_index(doc, target_para):
    for i, p in enumerate(doc.paragraphs):
        if p._p is target_para._p:
            return i
    raise ValueError("paragraph not found")


def _set_page_breaks(doc):
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text in {"{{ ch.title }}", "总结与展望", "总结", "结论", "参考文献", "致    谢", "致谢"}:
            p.paragraph_format.page_break_before = True


def _verify(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    text = "\n".join(parts)
    required = [
        "{{ student_id }}",
        "{{ title_zh }}",
        "{{ title_en }}",
        "{{ major }}",
        "{{ name }}",
        "{{ advisor }}",
        "{{ abstract_zh }}",
        "{{ keywords_zh }}",
        "{{ abstract_en }}",
        "{{ keywords_en }}",
        "{{ toc_placeholder }}",
        "{%p for ch in chapters %}",
        "{{ ch.title }}",
        "{%p for sec in ch.sections %}",
        "{{ sec.title }}",
        "{%p for sub in sec.subsections %}",
        "{{ sub.title }}",
        "{%p for para in conclusion_list %}",
        "{{ acknowledgement }}",
        "{%p for ref in references %}",
        "{{ ref }}",
    ]
    missing = [tag for tag in required if tag not in text]
    if missing:
        raise RuntimeError(f"Template verification failed, missing: {missing}")
    print("Verification passed")


def _mark_fields_update_on_open(path):
    from lxml import etree

    w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    tmp = tempfile.mktemp(suffix=".docx")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml":
                root = etree.fromstring(data)
                update = root.find(f"{{{w}}}updateFields")
                if update is None:
                    update = etree.Element(f"{{{w}}}updateFields")
                    root.insert(0, update)
                update.set(f"{{{w}}}val", "true")
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(item, data)
    shutil.move(tmp, path)


if __name__ == "__main__":
    main()
