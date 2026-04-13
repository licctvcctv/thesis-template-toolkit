"""
宁德师范学院本科毕业论文模板制作。
用法: cd thesis_project && python -m templates.ndnu.make <原始模板.docx>

模板特点:
- 封面 table[0] 放题目，table[1] 放学生信息（7行2列）
- 摘要/Abstract 用 Normal 样式
- Heading 1: 章标题（引言、结论、致谢、参考文献、附录）
- Heading 2: 二级标题（1.1 xxx）
- Heading 3: 三级标题（1.2.1 xxx）
- 正文: Normal, indent=304800
"""
import os
import sys
import re
import zipfile
import shutil

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../.."))
from body_maker import setup_body_template
from docx import Document


def _setup_ndnu_body(doc_path):
    """手动设置正文循环。宁德师范模板有完整 Heading 1/2/3 样式。"""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import copy

    _NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    doc = Document(doc_path)
    paras = doc.paragraphs

    # 定位 body 区域：第一个 Heading 1（致谢/参考文献之前）到 致谢
    h1_idx = h2_idx = h3_idx = body_idx = None
    body_end = len(paras)

    for i, p in enumerate(paras):
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()

        if style == "Heading 1" and i > 40:
            if "致" in text or "参考" in text or "附录" in text:
                body_end = i
                break
            if h1_idx is None:
                h1_idx = i

        if style == "Heading 2" and h2_idx is None and i > 40:
            h2_idx = i

        if style == "Heading 3" and h3_idx is None and i > 40:
            h3_idx = i

        if style == "Normal" and body_idx is None and i > 40:
            indent = p.paragraph_format.first_line_indent
            if indent and indent > 250000 and text and len(text) > 20:
                body_idx = i

    if not h1_idx or not body_idx:
        print(f"  警告: 样本不足 h1={h1_idx} body={body_idx}")
        return

    keep = {h1_idx, h2_idx, h3_idx, body_idx} - {None}

    # 清空非样本段落
    for i in range(h1_idx, body_end):
        if i in keep:
            continue
        p = paras[i]
        if p._p.find('.//w:sectPr', _NS) is not None:
            continue
        for r in p.runs:
            r.text = ""
        # 空 Heading 段落会显示自动编号数字，降级为 Normal 消除
        style = p.style.name if p.style else ""
        if style.startswith("Heading"):
            p.style = doc.styles['Normal']

    # 删除正文区域示例表格
    body_elem = doc.element.body
    tables_rm = []
    pc = 0
    for elem in body_elem:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            pc += 1
        elif tag == 'tbl' and h1_idx <= pc <= body_end:
            tables_rm.append(elem)
    for tbl in tables_rm:
        body_elem.remove(tbl)

    # 替换样本文本
    def _set_text(para, text):
        if para.runs:
            para.runs[0].text = text
            for r in para.runs[1:]:
                r.text = ""

    paras[h1_idx].paragraph_format.page_break_before = True
    _set_text(paras[h1_idx], "{{ ch.title }}")
    if h2_idx:
        _set_text(paras[h2_idx], "{{ sec.title }}")
    if h3_idx:
        _set_text(paras[h3_idx], "{{ sub.title }}")
    _set_text(paras[body_idx], "{{ para }}")

    h1_p = paras[h1_idx]._p
    h2_p = paras[h2_idx]._p if h2_idx else None
    h3_p = paras[h3_idx]._p if h3_idx else None
    body_p = paras[body_idx]._p
    body_copy = copy.deepcopy(body_p)

    def _mk_ctrl(text):
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        return p

    def _mk_body(text):
        p = copy.deepcopy(body_copy)
        for t_elem in p.findall('.//w:t', _NS):
            t_elem.text = ""
        t_elems = p.findall('.//w:t', _NS)
        if t_elems:
            t_elems[0].text = text
        return p

    # 移除所有样本段落，然后按正确顺序重新插入
    parent = h1_p.getparent()
    # 找到 h1 在 DOM 中的位置
    anchor_prev = h1_p.getprevious()

    for elem in [h1_p, h2_p, h3_p, body_p]:
        if elem is not None and elem.getparent() is not None:
            parent.remove(elem)

    # 从 anchor 后面按顺序插入
    cursor = anchor_prev

    def _insert_after(cursor, elem):
        cursor.addnext(elem)
        return elem

    cursor = _insert_after(cursor, _mk_ctrl("{%p for ch in chapters %}"))
    cursor = _insert_after(cursor, h1_p)

    if h2_p:
        cursor = _insert_after(cursor, _mk_ctrl("{%p for sec in ch.sections %}"))
        cursor = _insert_after(cursor, h2_p)
        cursor = _insert_after(cursor, _mk_ctrl("{%p for para in sec.content %}"))
        cursor = _insert_after(cursor, body_p)
        cursor = _insert_after(cursor, _mk_ctrl("{%p endfor %}"))  # content

        if h3_p:
            cursor = _insert_after(cursor, _mk_ctrl("{%p for sub in sec.subsections %}"))
            cursor = _insert_after(cursor, h3_p)
            cursor = _insert_after(cursor, _mk_ctrl("{%p for p2 in sub.content %}"))
            cursor = _insert_after(cursor, _mk_body("{{ p2 }}"))
            cursor = _insert_after(cursor, _mk_ctrl("{%p endfor %}"))  # sub.content
            cursor = _insert_after(cursor, _mk_ctrl("{%p endfor %}"))  # subsections

        cursor = _insert_after(cursor, _mk_ctrl("{%p endfor %}"))  # sections
    else:
        # 没有 H2，直接用 body
        cursor = _insert_after(cursor, _mk_ctrl("{%p for sec in ch.sections %}"))
        cursor = _insert_after(cursor, _mk_ctrl("{%p for para in sec.content %}"))
        cursor = _insert_after(cursor, body_p)
        cursor = _insert_after(cursor, _mk_ctrl("{%p endfor %}"))
        cursor = _insert_after(cursor, _mk_ctrl("{%p endfor %}"))

    cursor = _insert_after(cursor, _mk_ctrl("{%p endfor %}"))  # chapters

    # 清掉结论正文、附录内容等残留（body_end 之后到文档末尾）
    # 跳过参考文献区域（Step 6 已处理，保留原始 run 结构给 refs_maker 用）
    in_refs_zone = False
    for i in range(body_end, len(paras)):
        p = paras[i]
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        if style == "Heading 1":
            if "参考" in text or "参 考" in text:
                in_refs_zone = True
            else:
                in_refs_zone = False
            continue
        if "acknowledgement" in text:
            continue
        if in_refs_zone:
            # 保留 for/ref/endfor 循环段落
            if "{{ ref" in text or "{%p" in text:
                continue
            # 其余空段落：删掉 numPr 防止显示残留编号 [2][3][4]
            _wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            for numPr in p._p.findall(f'.//{{{_wns}}}numPr'):
                numPr.getparent().remove(numPr)
            continue
        for r in p.runs:
            r.text = ""

    doc.save(doc_path)
    print(f"  正文循环已设置: {doc_path}")


def make(src_path, out_path):
    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc = Document(src_path)

    # ========== Step 1: 封面 ==========
    # table[0]: 题目（单元格里的段落）
    title_cell = doc.tables[0].cell(0, 0)
    for p in title_cell.paragraphs:
        if p.runs:
            p.runs[0].text = "{{ title_zh }}"
            for r in p.runs[1:]:
                r.text = ""
            break

    # 清掉第二行"一行排不下可排两行"
    for p in title_cell.paragraphs[1:]:
        for r in p.runs:
            r.text = ""

    # table[1]: 学生信息
    info_table = doc.tables[1]
    field_map = {
        "姓": "{{ name }}",
        "学": "{{ student_id }}",
        "专业": "{{ major }}",
        "学院": "{{ college }}",  # 注意：学院的"学"也匹配，用"院"区分
        "指": "{{ advisor }}",
        "职": "{{ advisor_title }}",
        "完": "{{ finish_date }}",
    }
    for row in info_table.rows:
        label = row.cells[0].text.strip()
        value_cell = row.cells[1]
        for key, var in field_map.items():
            if label.startswith(key) or key in label:
                # 避免"学号"的"学"匹配到"学院"
                if key == "学" and "院" in label:
                    continue
                if value_cell.paragraphs[0].runs:
                    value_cell.paragraphs[0].runs[0].text = var
                    for r in value_cell.paragraphs[0].runs[1:]:
                        r.text = ""
                else:
                    # 没有 run，直接设文本
                    value_cell.paragraphs[0].text = var
                break

    print("  Step 1: 封面")

    # ========== Step 2: 中文摘要 ==========
    # p[26]: "摘  要" 标题
    # p[27]: 摘要正文
    # p[28]: 关键词
    p27 = doc.paragraphs[27]
    if p27.runs:
        p27.runs[0].text = "{{ abstract_zh }}"
        for r in p27.runs[1:]:
            r.text = ""

    # 关键词：保留原始 run 结构（run[0]=空黑体, run[1]="关键词：", run[2]=宋体值）
    p28 = doc.paragraphs[28]
    label_found = False
    for j, r in enumerate(p28.runs):
        if "关键词：" in r.text or "关键词:" in r.text:
            label_found = True
            continue
        if label_found and r.text.strip():
            r.text = "{{ keywords_zh }}"
            for rr in p28.runs[j + 1:]:
                rr.text = ""
            break

    # 删除 p[29] p[30] 空段落（避免中英文摘要间多出空白页）
    for idx in [30, 29]:
        doc.paragraphs[idx]._p.getparent().remove(doc.paragraphs[idx]._p)

    print("  Step 2: 中文摘要")

    # ========== Step 3: 英文摘要 ==========
    # p[31]: "Abstract"
    # p[32]: 英文摘要正文
    # p[33]: Key words
    # 英文摘要另起一页：原始模板用 <w:br type="page"/>，不用 page_break_before
    # p[29] 是 Abstract 标题，它的第一个 run 里已经有分页符，保持不动即可
    # 如果被清掉了，确保它存在
    p29 = doc.paragraphs[29]
    from docx.oxml.ns import qn as _qn2
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    has_br = bool(p29._p.findall(f'.//{{{ns_w}}}br[@{{{ns_w}}}type="page"]'))
    if not has_br and p29.runs:
        from docx.oxml import OxmlElement as _OE2
        br = _OE2('w:br')
        br.set(_qn2('w:type'), 'page')
        p29.runs[0]._r.insert(0, br)
    # 确保没有 page_break_before（会导致多空一页）
    p29.paragraph_format.page_break_before = False

    # 英文摘要正文（原 p[32] → 现 p[30]）
    p30 = doc.paragraphs[30]
    if p30.runs:
        p30.runs[0].text = "{{ abstract_en }}"
        p30.runs[0].font.name = "Times New Roman"
        from docx.shared import Pt as _Pt3
        p30.runs[0].font.size = _Pt3(12)
        p30.runs[0].font.bold = False
        for r in p30.runs[1:]:
            r.text = ""

    # 英文关键词（原 p[33] → 现 p[31]）
    p31 = doc.paragraphs[31]
    for j, r in enumerate(p31.runs):
        if "Key" in r.text:
            pass
        elif j >= 2 and r.text.strip():
            r.text = "{{ keywords_en }}"
            for rr in p31.runs[j + 1:]:
                rr.text = ""
            break

    # 清掉格式说明（原 p[34] → 现 p[32]）
    for r in doc.paragraphs[32].runs:
        r.text = ""

    print("  Step 3: 英文摘要")

    # ========== Step 4: 清理格式说明 ==========
    inst_kw = ["号", "磅", "宋体", "黑体", "仿宋", "行距", "居中",
               "缩进", "标点", "字体", "格式见", "独占一页", "空格"]
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        style = p.style.name if p.style else ""
        # 跳过标题和目录
        if style.startswith("Heading") or style.startswith("toc"):
            continue
        # 跳过正文内容段落（有实际内容且不是格式说明）
        if i < 26 or i > 135:
            continue
        # 清掉括号里的格式说明
        if text.startswith("（") and text.endswith("）"):
            for r in p.runs:
                r.text = ""

    # 清掉正文标题页残留（"某某安全方案分析与设计"等）
    for i in range(50, 60):
        if i >= len(doc.paragraphs):
            break
        p = doc.paragraphs[i]
        style = p.style.name if p.style else ""
        if style.startswith("Heading"):
            continue
        text = (p.text or "").strip()
        if text and "某某" in text or ("方案" in text and "设计" in text):
            for r in p.runs:
                r.text = ""

    print("  Step 4: 清理格式说明")

    # ========== Step 5: 致谢 ==========
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if text in ("致    谢", "致谢", "致  谢") and p.style and \
                p.style.name == "Heading 1":
            # 找后面的 Normal 段落
            for j in range(i + 1, min(i + 10, len(doc.paragraphs))):
                pj = doc.paragraphs[j]
                sj = pj.style.name if pj.style else ""
                if "Heading" in sj:
                    continue
                if pj.runs:
                    pj.runs[0].text = "    {{ acknowledgement }}"
                    for r in pj.runs[1:]:
                        r.text = ""
                    # 清掉后面的占位段落
                    for k in range(j + 1, min(j + 5, len(doc.paragraphs))):
                        pk = doc.paragraphs[k]
                        pk_style = pk.style.name if pk.style else ""
                        if "Heading" in pk_style:
                            break
                        for r in pk.runs:
                            r.text = ""
                    break
            break

    print("  Step 5: 致谢")

    # ========== Step 6: 参考文献 ==========
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if ("参考文献" in text or "参 考 文 献" in text) and p.style and p.style.name == "Heading 1":
            # 找后面第一个有内容的 Normal 段落作为样本
            for j in range(i + 1, min(i + 10, len(doc.paragraphs))):
                pj = doc.paragraphs[j]
                pt = (pj.text or "").strip()
                sj = pj.style.name if pj.style else ""
                if "Heading" in sj:
                    break
                if pj.runs and pt:
                    # 替换为循环
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    ed_p = pj._p
                    # for 标签
                    for_p = OxmlElement('w:p')
                    for_r = OxmlElement('w:r')
                    for_t = OxmlElement('w:t')
                    for_t.set(qn('xml:space'), 'preserve')
                    for_t.text = "{%p for ref in references %}"
                    for_r.append(for_t)
                    for_p.append(for_r)
                    ed_p.addprevious(for_p)
                    # 替换文本
                    pj.runs[0].text = "{{ ref }}"
                    for r in pj.runs[1:]:
                        r.text = ""
                    # endfor 标签
                    end_p = OxmlElement('w:p')
                    end_r = OxmlElement('w:r')
                    end_t = OxmlElement('w:t')
                    end_t.set(qn('xml:space'), 'preserve')
                    end_t.text = "{%p endfor %}"
                    end_r.append(end_t)
                    end_p.append(end_r)
                    ed_p.addnext(end_p)
                    # 清掉 endfor 之后到下一个 Heading 之间的剩余示例
                    cleaning = False
                    for k in range(j + 1, min(j + 25, len(doc.paragraphs))):
                        pk = doc.paragraphs[k]
                        pk_style = pk.style.name if pk.style else ""
                        if "Heading" in pk_style:
                            break
                        pk_text = (pk.text or "").strip()
                        if "{%p endfor" in pk_text:
                            cleaning = True
                            continue
                        if cleaning:
                            for r in pk.runs:
                                r.text = ""
                    break
            break

    print("  Step 6: 参考文献")

    # 保存中间结果
    doc.save(out_path)

    # ========== Step 7: 正文循环 ==========
    _setup_ndnu_body(out_path)
    print("  Step 7: 正文循环")

    # ========== Step 8: 删除文本框注释 ==========
    tmp = out_path + ".tmp"
    with zipfile.ZipFile(out_path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w') as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    xml = data.decode('utf-8')
                    n = xml.count('<mc:AlternateContent')
                    xml = re.sub(
                        r'<mc:AlternateContent>.*?</mc:AlternateContent>',
                        '', xml, flags=re.DOTALL)
                    print(f"  Step 9: 删除 {n} 个文本框注释")
                    data = xml.encode('utf-8')
                zout.writestr(item, data)
    shutil.move(tmp, out_path)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python make.py <原始模板.docx> [输出路径]")
        sys.exit(1)
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else "template.docx"
    print(f"制作模板: {src} -> {out}")
    make(src, out)
    print(f"完成: {out}")
