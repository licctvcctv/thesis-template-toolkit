"""
上海健康医学院本科毕业论文模板制作（2023版样表）。
用法: cd thesis_project && python -m templates.shjkyxy.make <样表.docx>

从学校样表(副本上海健康医学院毕业论文（设计）样表-正文（2023版）.docx)生成 docxtpl 模板。

样表结构 (paragraph index 仅供参考，以 style + 内容匹配为准):
- 封面页: 校徽/院徽 + 中英文题目 + 6行封面信息表格
- 原创性声明 + 使用授权说明
- 中文摘要 + 关键词 (heading 1 "摘  要")
- 英文摘要 + 关键词 (heading 1 "ABSTRACT")
- 目录 (heading 1 "目  录")
- 正文 (heading 1/2/3 样例章节)
- 参考文献 (heading 1)
- 致谢 (heading 1)

样表的 style 名:
  heading 1 — 所有一级标题(摘要/ABSTRACT/目录/章标题/参考文献/致谢)
  heading 2 — 节标题 (1.1 xxx)
  heading 3 — 小节标题 (1.1.1 xxx)
  Normal    — 正文
"""
import os
import sys
import copy
import re

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../.."))
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _ensure_rPr_font(rPr, east_asia=None, ascii_font=None, sz=None):
    """Ensure rPr element has explicit font and size attributes."""
    fonts = rPr.find(qn('w:rFonts'))
    if fonts is None:
        fonts = OxmlElement('w:rFonts')
        rPr.insert(0, fonts)
    if east_asia and not fonts.get(qn('w:eastAsia')):
        fonts.set(qn('w:eastAsia'), east_asia)
    if ascii_font and not fonts.get(qn('w:ascii')):
        fonts.set(qn('w:ascii'), ascii_font)
    if sz:
        sz_el = rPr.find(qn('w:sz'))
        if sz_el is None:
            sz_el = OxmlElement('w:sz')
            rPr.append(sz_el)
        if not sz_el.get(qn('w:val')):
            sz_el.set(qn('w:val'), sz)
        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs')
            rPr.append(szCs)
        if not szCs.get(qn('w:val')):
            szCs.set(qn('w:val'), sz)


def _is_colored(run):
    """Check if a run has a non-black font color (red, blue, etc.)."""
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None:
        return False
    color = rPr.find(qn('w:color'))
    if color is None:
        return False
    val = color.get(qn('w:val'), '').upper()
    # 黑色/自动色不算彩色
    if val in ('', '000000', 'AUTO', '333333'):
        return False
    return True


def _is_red(run):
    """Check if a run's font color is red (FF0000)."""
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None:
        return False
    color = rPr.find(qn('w:color'))
    if color is not None and color.get(qn('w:val'), '').upper() in ('FF0000', 'RED'):
        return True
    return False


def _all_runs_red(para):
    """Check if ALL non-empty runs in a paragraph are colored (red/blue/etc.)."""
    non_empty = [r for r in para.runs if r.text.strip()]
    if not non_empty:
        return False
    return all(_is_colored(r) for r in non_empty)


def _has_sect_pr(para):
    """Check if paragraph contains a section break."""
    pPr = para._p.pPr
    return pPr is not None and pPr.find(qn('w:sectPr')) is not None


def _find_heading1(paras, keyword, start=0):
    """Find heading 1 paragraph containing keyword."""
    for i in range(start, len(paras)):
        p = paras[i]
        s = p.style.name if p.style else ''
        t = (p.text or '').strip()
        if 'heading 1' in s.lower() and keyword in t:
            return i
    return None


def _insert_jinja_loop(anchor_para, loop_var, item_var, saved_pPr, saved_rPr):
    """Insert a Jinja2 for-loop after anchor paragraph.

    Creates:
      anchor: {%p for <item_var> in <loop_var> %}
      new p:  {{ <item_var> }}
      new p:  {%p endfor %}
    """
    # Set anchor text to for-loop start
    if anchor_para.runs:
        anchor_para.runs[0].text = f"{{%p for {item_var} in {loop_var} %}}"
        for r in anchor_para.runs[1:]:
            r.text = ""
    else:
        anchor_para.add_run(f"{{%p for {item_var} in {loop_var} %}}")

    # {{ item_var }} paragraph with correct formatting
    p_body = OxmlElement('w:p')
    if saved_pPr:
        p_body.append(copy.deepcopy(saved_pPr))
    r_body = OxmlElement('w:r')
    if saved_rPr:
        r_body.append(copy.deepcopy(saved_rPr))
    t_el = OxmlElement('w:t')
    t_el.set(qn('xml:space'), 'preserve')
    t_el.text = f'{{{{ {item_var} }}}}'
    r_body.append(t_el)
    p_body.append(r_body)
    anchor_para._p.addnext(p_body)

    # {%p endfor %} paragraph
    p_end = OxmlElement('w:p')
    r_end = OxmlElement('w:r')
    t_end = OxmlElement('w:t')
    t_end.text = '{%p endfor %}'
    r_end.append(t_end)
    p_end.append(r_end)
    p_body.addnext(p_end)


def _clear_runs(para):
    """Clear all text in paragraph."""
    for r in para.runs:
        r.text = ""


def make(src_path, out_path):
    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc = Document(src_path)
    paras = doc.paragraphs

    print(f"原始: {len(paras)} 段落, {len(doc.tables)} 表格")

    # ========== Step 1: 删除红色格式说明段落 ==========
    removed_red = 0
    cleaned_runs = 0
    for p in list(paras):
        if _all_runs_red(p) and not _has_sect_pr(p):
            # 全彩色段落：清空文字但保留空行（维持页面间距布局）
            for r in p.runs:
                r.text = ""
            removed_red += 1
        else:
            # 混合段落：删除彩色runs但保留非彩色内容
            for r in list(p.runs):
                if _is_colored(r) and r.text.strip():
                    r.text = ""
                    cleaned_runs += 1
    paras = doc.paragraphs
    print(f"  Step 1: 清空 {removed_red} 个彩色段落, 清理 {cleaned_runs} 个彩色run")

    # ========== Step 2: 封面 ==========
    # 2a: 中文题目 — 找到封面区域中较长的中文文本段
    for i, p in enumerate(paras[:40]):
        t = (p.text or '').strip()
        # 样表中文题目: "论文题目(中文标题)"
        if '论文题目' in t or '中文标题' in t:
            if p.runs:
                p.runs[0].text = "{{ title_zh }}"
                for r in p.runs[1:]:
                    r.text = ""
                print(f"  封面中文题目: P{i}")
            break

    # 2b: 英文题目 — 找 "Guideline" 或 "外文标题" 或纯英文长文本
    paras = doc.paragraphs
    for i, p in enumerate(paras[:40]):
        t = (p.text or '').strip()
        if 'Guideline' in t or '外文标题' in t or ('title' in t.lower() and len(t) > 20):
            if p.runs:
                p.runs[0].text = "{{ title_en }}"
                for r in p.runs[1:]:
                    r.text = ""
                print(f"  封面英文题目: P{i}")
            break

    # 2c: 封面表格 — 6行2列, 替换右列为Jinja2变量
    if doc.tables:
        cover_table = doc.tables[0]  # 第一个表格就是封面信息表
        # 按匹配优先级排列(长的先匹配,避免'学'吃掉'学生'和'学号')
        field_rules = [
            ('学生', '{{ name }}'),
            ('指导', '{{ advisor }}'),
            ('完成', '{{ year }} 年 {{ month }} 月'),
            ('专', '{{ major }}'),
        ]
        for row in cover_table.rows:
            label_text = (row.cells[0].text or '').strip()
            # 去掉全角空格后匹配
            label_clean = label_text.replace('\u3000', '')
            jinja_var = None

            # 学院 vs 学号 用 clean label 精确区分
            if label_clean.startswith('学院') or label_clean == '学院':
                jinja_var = '{{ department }}'
            elif label_clean.startswith('学号') or label_clean == '学号':
                jinja_var = '{{ student_id }}'
            else:
                for prefix, var in field_rules:
                    if label_text.startswith(prefix):
                        jinja_var = var
                        break

            if jinja_var:
                cell = row.cells[1]
                for p in cell.paragraphs:
                    if p.runs:
                        p.runs[0].text = jinja_var
                        for r in p.runs[1:]:
                            r.text = ""
                    else:
                        p.add_run(jinja_var)

        print("  Step 2: 封面字段替换完成")
    else:
        # 没有表格，用旧方式处理封面
        print("  Step 2: 未找到封面表格，跳过")

    # ========== Step 3: 保存格式样本 ==========
    paras = doc.paragraphs
    body_pPr = None
    body_rPr = None
    h1_pPr = None
    h1_rPr = None
    h2_pPr = None
    h2_rPr = None
    h3_pPr = None
    h3_rPr = None
    ref_pPr = None
    ref_rPr = None

    # 从正文区域(目录之后)采集格式样本
    toc_end = 0
    for i, p in enumerate(paras):
        s = (p.style.name or '').lower()
        t = (p.text or '').strip()
        if 'heading 1' in s and '目' in t and '录' in t:
            toc_end = i
            break

    for i, p in enumerate(paras):
        if i <= toc_end:
            continue
        s = (p.style.name or '').lower()
        t = (p.text or '').strip()
        if not t:
            continue

        if 'heading 1' in s and h1_pPr is None:
            # 章标题样本 — 取第一个正文章标题(如 "第1章")
            if re.match(r'^第\d+章', t):
                h1_pPr = copy.deepcopy(p._p.pPr) if p._p.pPr else None
                h1_rPr = copy.deepcopy(p.runs[0]._r.find(qn('w:rPr'))) if p.runs else None
        elif 'heading 2' in s and h2_pPr is None:
            h2_pPr = copy.deepcopy(p._p.pPr) if p._p.pPr else None
            h2_rPr = copy.deepcopy(p.runs[0]._r.find(qn('w:rPr'))) if p.runs else None
        elif 'heading 3' in s and h3_pPr is None:
            h3_pPr = copy.deepcopy(p._p.pPr) if p._p.pPr else None
            h3_rPr = copy.deepcopy(p.runs[0]._r.find(qn('w:rPr'))) if p.runs else None
        elif s == 'normal' and body_pPr is None and len(t) > 20:
            pPr = p._p.pPr
            if pPr is not None:
                ind = pPr.find(qn('w:ind'))
                if ind is not None and ind.get(qn('w:firstLine'), '0') != '0':
                    body_pPr = copy.deepcopy(pPr)
                    body_rPr = copy.deepcopy(p.runs[0]._r.find(qn('w:rPr'))) if p.runs else None

    # 如果 body_rPr 没有 sz，再扫一遍正文段找有显式 sz 的样本
    if body_rPr is not None:
        sz_el = body_rPr.find(qn('w:sz'))
        if sz_el is None or not sz_el.get(qn('w:val')):
            for j in range(toc_end + 1, len(paras)):
                pj = paras[j]
                sj = (pj.style.name or '').lower()
                tj = (pj.text or '').strip()
                if sj != 'normal' or len(tj) < 20 or not pj.runs:
                    continue
                jr = pj.runs[0]._r.find(qn('w:rPr'))
                if jr is None:
                    continue
                jsz = jr.find(qn('w:sz'))
                if jsz is not None and jsz.get(qn('w:val')) == '24':
                    # 优先选没有 ascii 覆盖的(继承 Normal 的 Times New Roman)
                    jfonts = jr.find(qn('w:rFonts'))
                    jascii = jfonts.get(qn('w:ascii')) if jfonts is not None else None
                    if jascii is None:
                        body_rPr = copy.deepcopy(jr)
                        break

    # 参考文献格式样本 — 通常无首行缩进
    ref_idx = _find_heading1(paras, '参考文献')
    if ref_idx is not None:
        for j in range(ref_idx + 1, min(ref_idx + 20, len(paras))):
            pj = paras[j]
            sj = (pj.style.name or '').lower()
            tj = (pj.text or '').strip()
            if 'heading 1' in sj:
                break
            if tj and sj == 'normal' and tj.startswith('['):
                ref_pPr = copy.deepcopy(pj._p.pPr) if pj._p.pPr else None
                ref_rPr = copy.deepcopy(pj.runs[0]._r.find(qn('w:rPr'))) if pj.runs else None
                break

    # ref_rPr 去掉 ascii/hAnsi 覆盖，让英文继承 Normal 的 Times New Roman
    if ref_rPr is not None:
        fonts = ref_rPr.find(qn('w:rFonts'))
        if fonts is not None:
            for attr in (qn('w:ascii'), qn('w:hAnsi')):
                if fonts.get(attr) == '宋体':
                    del fonts.attrib[attr]

    print("  Step 3: 格式样本保存完成")
    print(f"    h1={'✓' if h1_pPr else '✗'} h2={'✓' if h2_pPr else '✗'} "
          f"h3={'✓' if h3_pPr else '✗'} body={'✓' if body_pPr else '✗'} "
          f"ref={'✓' if ref_pPr else '✗'}")

    # ========== Step 4: 中文摘要 → Jinja2 ==========
    paras = doc.paragraphs
    abs_zh_idx = _find_heading1(paras, '摘')
    if abs_zh_idx is not None:
        # 找第一个正文段落
        for j in range(abs_zh_idx + 1, min(abs_zh_idx + 10, len(paras))):
            pj = paras[j]
            tj = (pj.text or '').strip()
            sj = (pj.style.name or '').lower()
            if 'heading 1' in sj:
                break
            if tj and '关键词' not in tj and 'normal' in sj:
                _insert_jinja_loop(pj, 'abstract_zh_list', 'para', body_pPr, body_rPr)
                # 清除剩余正文段
                for k in range(j + 1, min(j + 10, len(paras))):
                    pk = paras[k]
                    tk = (pk.text or '').strip()
                    sk = (pk.style.name or '').lower()
                    if 'heading 1' in sk or '关键词' in tk:
                        break
                    if tk and 'normal' in sk:
                        _clear_runs(pk)
                break

        # 关键词 — 替换第一个，清除后续重复
        paras = doc.paragraphs
        kw_found = False
        for j in range(abs_zh_idx + 1, min(abs_zh_idx + 15, len(paras))):
            pj = paras[j]
            tj = (pj.text or '').strip()
            sj = (pj.style.name or '').lower()
            if 'heading 1' in sj:
                break
            if '关键词' in tj:
                if not kw_found:
                    if pj.runs:
                        pj.runs[0].text = "关键词：{{ keywords_zh }}"
                        for r in pj.runs[1:]:
                            r.text = ""
                    kw_found = True
                else:
                    # 后续重复的关键词段落删除
                    _clear_runs(pj)
        print("  Step 4: 中文摘要")

    # ========== Step 5: 英文摘要 → Jinja2 ==========
    paras = doc.paragraphs
    abs_en_idx = _find_heading1(paras, 'ABSTRACT')
    if abs_en_idx is not None:
        for j in range(abs_en_idx + 1, min(abs_en_idx + 10, len(paras))):
            pj = paras[j]
            tj = (pj.text or '').strip()
            sj = (pj.style.name or '').lower()
            if 'heading 1' in sj:
                break
            if tj and 'KEY WORDS' not in tj.upper() and 'normal' in sj:
                _insert_jinja_loop(pj, 'abstract_en_list', 'para', body_pPr, body_rPr)
                for k in range(j + 1, min(j + 10, len(paras))):
                    pk = paras[k]
                    tk = (pk.text or '').strip()
                    if 'KEY WORDS' in tk.upper() or 'heading 1' in (pk.style.name or '').lower():
                        break
                    if tk and 'normal' in (pk.style.name or '').lower():
                        _clear_runs(pk)
                break

        paras = doc.paragraphs
        kw_en_found = False
        for j in range(abs_en_idx + 1, min(abs_en_idx + 15, len(paras))):
            pj = paras[j]
            tj = (pj.text or '').strip()
            sj = (pj.style.name or '').lower()
            if 'heading 1' in sj:
                break
            if 'KEY WORDS' in tj.upper() or 'Key words' in tj:
                if not kw_en_found:
                    if pj.runs:
                        pj.runs[0].text = "KEY WORDS: {{ keywords_en }}"
                        for r in pj.runs[1:]:
                            r.text = ""
                    kw_en_found = True
                else:
                    _clear_runs(pj)
        print("  Step 5: 英文摘要")

    # ========== Step 6: 目录 - 清空 ==========
    paras = doc.paragraphs
    toc_idx = _find_heading1(paras, '目')
    if toc_idx is not None:
        for j in range(toc_idx + 1, min(toc_idx + 60, len(paras))):
            pj = paras[j]
            sj = (pj.style.name or '').lower()
            if 'heading 1' in sj:
                break
            _clear_runs(pj)
        print("  Step 6: 目录清空")

    # ========== Step 7: 正文 → Jinja2 循环 ==========
    paras = doc.paragraphs
    body_start = None
    ref_heading = None

    # 找正文起点(第一个 "第N章" heading 1) 和 参考文献 heading 1
    for i, p in enumerate(paras):
        s = (p.style.name or '').lower()
        t = (p.text or '').strip()
        if 'heading 1' in s and re.match(r'^第\d+章', t) and body_start is None:
            body_start = i
        if 'heading 1' in s and '参考文献' in t:
            ref_heading = i
            break

    if body_start is None or ref_heading is None:
        print("  ERROR: 找不到正文范围")
        return

    print(f"  正文范围: P{body_start} ~ P{ref_heading} ({ref_heading - body_start} 段)")

    # 删除正文段落(保留body_start位置的14个段用于Jinja2), 保留sectPr段
    # 先移除图片/drawing/表格
    for i in range(body_start, ref_heading):
        p = paras[i]
        for drawing in p._p.findall('.//' + qn('w:drawing')):
            drawing.getparent().remove(drawing)
        for pict in p._p.findall('.//' + qn('w:pict')):
            pict.getparent().remove(pict)

    # 删除所有表格(样表里的示例表格)
    while len(doc.tables) > 1:
        # 保留第一个表格(封面)
        tbl = doc.tables[-1]
        tbl._element.getparent().remove(tbl._element)

    # 删除正文段落, 保留前18个位置给Jinja2标签
    paras = doc.paragraphs
    # 重新定位 ref_heading (段落可能已变)
    ref_heading = None
    for i, p in enumerate(paras):
        s = (p.style.name or '').lower()
        t = (p.text or '').strip()
        if 'heading 1' in s and '参考文献' in t:
            ref_heading = i
            break

    body_start = None
    for i, p in enumerate(paras):
        s = (p.style.name or '').lower()
        t = (p.text or '').strip()
        if 'heading 1' in s and re.match(r'^第\d+章', t) and body_start is None:
            body_start = i

    if body_start is None or ref_heading is None:
        print("  ERROR: 重新定位失败")
        return

    # 需要保留的段数: 18 (for/ch.title/for/item/endfor/for/sec.title/for/item/endfor/for/sub.title/for/item/endfor/endfor/endfor/endfor + 结论部分)
    keep_count = 18
    # 删除多余的正文段(从后往前删)
    for i in range(ref_heading - 1, body_start + keep_count - 1, -1):
        p = paras[i]
        if not _has_sect_pr(p):
            p._p.getparent().remove(p._p)

    print(f"  Step 7a: 删除正文，剩余 {len(doc.paragraphs)} 段")

    # 重新取段落
    paras = doc.paragraphs
    # 重新定位
    ref_heading = None
    body_start = None
    for i, p in enumerate(paras):
        s = (p.style.name or '').lower()
        t = (p.text or '').strip()
        if 'heading 1' in s and re.match(r'^第\d+章', t) and body_start is None:
            body_start = i
        if 'heading 1' in s and '参考文献' in t:
            ref_heading = i
            break

    if body_start is None or ref_heading is None:
        print("  ERROR: 定位失败（删除后）")
        return

    available = ref_heading - body_start
    print(f"  可用段: {available} (body_start={body_start})")

    # 设置Jinja2标签
    # 章循环 (heading 1 for ch.title, heading 2 for sec.title, heading 3 for sub.title)
    tags = [
        (body_start + 0,  "{%p for ch in chapters %}",           None, None),
        (body_start + 1,  "{{ ch.title }}",                      h1_pPr, h1_rPr),
        (body_start + 2,  "{%p for item in ch.content %}",       None, None),
        (body_start + 3,  "{{ item }}",                          body_pPr, body_rPr),
        (body_start + 4,  "{%p endfor %}",                       None, None),
        (body_start + 5,  "{%p for sec in ch.sections %}",       None, None),
        (body_start + 6,  "{{ sec.title }}",                     h2_pPr, h2_rPr),
        (body_start + 7,  "{%p for item in sec.content %}",      None, None),
        (body_start + 8,  "{{ item }}",                          body_pPr, body_rPr),
        (body_start + 9,  "{%p endfor %}",                       None, None),
        (body_start + 10, "{%p for sub in sec.subsections %}",   None, None),
        (body_start + 11, "{{ sub.title }}",                     h3_pPr, h3_rPr),
        (body_start + 12, "{%p for item in sub.content %}",      None, None),
        (body_start + 13, "{{ item }}",                          body_pPr, body_rPr),
    ]

    # 需要额外插入 4 个 endfor 段
    last_p = paras[body_start + min(13, available - 1)]
    endfor_tags = ["{%p endfor %}", "{%p endfor %}", "{%p endfor %}", "{%p endfor %}"]
    for tag_text in endfor_tags:
        p_new = OxmlElement('w:p')
        r_new = OxmlElement('w:r')
        t_new = OxmlElement('w:t')
        t_new.text = tag_text
        r_new.append(t_new)
        p_new.append(r_new)
        last_p._p.addnext(p_new)
        last_p = type('P', (), {'_p': p_new})()

    paras = doc.paragraphs
    for idx, text, saved_pPr, saved_rPr in tags:
        if idx >= len(paras):
            break
        p = paras[idx]
        if p.runs:
            p.runs[0].text = text
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(text)

        # Apply formatting
        if saved_pPr is not None:
            old = p._p.pPr
            if old is not None:
                p._p.remove(old)
            p._p.insert(0, copy.deepcopy(saved_pPr))
        if saved_rPr is not None and p.runs:
            r = p.runs[0]._r
            old_rPr = r.find(qn('w:rPr'))
            if old_rPr is not None:
                r.remove(old_rPr)
            r.insert(0, copy.deepcopy(saved_rPr))

    # 控制/循环标签用 body 格式(避免被当作标题)
    for offset in [0, 2, 4, 5, 7, 9, 10, 12]:
        idx = body_start + offset
        if idx < len(paras):
            p = paras[idx]
            if p.style and 'heading' in (p.style.name or '').lower():
                if body_pPr:
                    old = p._p.pPr
                    if old is not None:
                        p._p.remove(old)
                    p._p.insert(0, copy.deepcopy(body_pPr))

    print("  Step 7: 正文Jinja2循环")

    # ========== Step 7b: 清除章循环和参考文献之间的残留段落 ==========
    paras = doc.paragraphs
    # 找到最后一个 {%p endfor %} (属于章循环)
    last_endfor = None
    ref_heading2 = None
    for i, p in enumerate(paras):
        t = (p.text or '').strip()
        s = (p.style.name or '').lower()
        if t == '{%p endfor %}':
            last_endfor = i
        if 'heading 1' in s and '参考文献' in t:
            ref_heading2 = i
            break

    if last_endfor is not None and ref_heading2 is not None:
        removed_residual = 0
        for i in range(ref_heading2 - 1, last_endfor, -1):
            p = paras[i]
            t = (p.text or '').strip()
            if _has_sect_pr(p):
                continue
            if t and not t.startswith('{%') and not t.startswith('{{'):
                p._p.getparent().remove(p._p)
                removed_residual += 1
        if removed_residual:
            print(f"  Step 7b: 清除 {removed_residual} 个残留段落")

    # ========== Step 8: 结论 → Jinja2 (在参考文献之前插入) ==========
    paras = doc.paragraphs
    ref_heading = None
    for i, p in enumerate(paras):
        s = (p.style.name or '').lower()
        t = (p.text or '').strip()
        if 'heading 1' in s and '参考文献' in t:
            ref_heading = i
            break

    if ref_heading is not None:
        # 在参考文献之前插入: 结论标题 + 结论内容循环
        ref_p = paras[ref_heading]._p

        # 结论标题 heading 1
        p_conc_title = OxmlElement('w:p')
        if h1_pPr:
            p_conc_title.append(copy.deepcopy(h1_pPr))
        r_conc = OxmlElement('w:r')
        if h1_rPr:
            r_conc.append(copy.deepcopy(h1_rPr))
        t_conc = OxmlElement('w:t')
        t_conc.set(qn('xml:space'), 'preserve')
        t_conc.text = '结  论'
        r_conc.append(t_conc)
        p_conc_title.append(r_conc)
        ref_p.addprevious(p_conc_title)

        # {%p for para in conclusion_list %}
        p_for = OxmlElement('w:p')
        r_for = OxmlElement('w:r')
        t_for = OxmlElement('w:t')
        t_for.text = '{%p for para in conclusion_list %}'
        r_for.append(t_for)
        p_for.append(r_for)
        ref_p.addprevious(p_for)

        # {{ para }}
        p_body = OxmlElement('w:p')
        if body_pPr:
            p_body.append(copy.deepcopy(body_pPr))
        r_body = OxmlElement('w:r')
        if body_rPr:
            r_body.append(copy.deepcopy(body_rPr))
        t_body = OxmlElement('w:t')
        t_body.set(qn('xml:space'), 'preserve')
        t_body.text = '{{ para }}'
        r_body.append(t_body)
        p_body.append(r_body)
        ref_p.addprevious(p_body)

        # {%p endfor %}
        p_end = OxmlElement('w:p')
        r_end = OxmlElement('w:r')
        t_end = OxmlElement('w:t')
        t_end.text = '{%p endfor %}'
        r_end.append(t_end)
        p_end.append(r_end)
        ref_p.addprevious(p_end)

        print("  Step 8: 结论")

    # ========== Step 9: 参考文献 → Jinja2 ==========
    paras = doc.paragraphs
    ref_heading = None
    for i, p in enumerate(paras):
        s = (p.style.name or '').lower()
        t = (p.text or '').strip()
        if 'heading 1' in s and '参考文献' in t:
            ref_heading = i
            break

    if ref_heading is not None:
        first_ref = None
        for j in range(ref_heading + 1, min(ref_heading + 25, len(paras))):
            pj = paras[j]
            tj = (pj.text or '').strip()
            sj = (pj.style.name or '').lower()
            if 'heading 1' in sj:
                break
            if tj and 'normal' in sj:
                if first_ref is None:
                    first_ref = j
                    _insert_jinja_loop(pj, 'references', 'ref',
                                       ref_pPr or body_pPr, ref_rPr or body_rPr)
                else:
                    _clear_runs(pj)

        print("  Step 9: 参考文献")

    # ========== Step 10: 致谢 → Jinja2 ==========
    paras = doc.paragraphs
    ack_idx = _find_heading1(paras, '致')
    if ack_idx is not None:
        for j in range(ack_idx + 1, min(ack_idx + 10, len(paras))):
            pj = paras[j]
            tj = (pj.text or '').strip()
            sj = (pj.style.name or '').lower()
            if tj and 'normal' in sj:
                _insert_jinja_loop(pj, 'acknowledgement_list', 'para',
                                   body_pPr, body_rPr)
                for k in range(j + 1, min(j + 8, len(paras))):
                    pk = paras[k]
                    if (pk.text or '').strip() and 'normal' in (pk.style.name or '').lower():
                        _clear_runs(pk)
                break
        print("  Step 10: 致谢")

    # ========== Step 11: 清理空段落(摘要之后) ==========
    paras = doc.paragraphs
    clean_start = 0
    for i, p in enumerate(paras):
        s = (p.style.name or '').lower()
        t = (p.text or '').strip()
        if 'heading 1' in s and '摘' in t:
            clean_start = i
            break

    removed = 0
    for i in range(len(paras) - 1, clean_start - 1, -1):
        p = paras[i]
        text = (p.text or "").strip()
        if _has_sect_pr(p):
            continue
        if text:
            continue
        s = (p.style.name or '').lower()
        if 'heading' in s:
            continue
        has_drawing = bool(p._p.findall('.//' + qn('w:drawing')))
        if not has_drawing:
            p._p.getparent().remove(p._p)
            removed += 1

    print(f"  Step 11: 删除 {removed} 个空段落")

    # ========== Save ==========
    doc.save(out_path)

    # Verify
    doc2 = Document(out_path)
    p_count = len(doc2.paragraphs)
    t_count = len(doc2.tables)
    print(f"\n结果: {p_count} 段落, {t_count} 表格")
    for i, p in enumerate(doc2.paragraphs):
        t = (p.text or "").strip()
        if t:
            s = p.style.name if p.style else 'None'
            print(f"  P{i:3d} [{s:25s}] {t[:90]}")

    full = "\n".join(p.text or "" for p in doc2.paragraphs)
    for tag in ["{%", "%}", "{{", "}}"]:
        count = full.count(tag)
        if count:
            print(f"  Jinja2 '{tag}': {count}")


def main():
    if len(sys.argv) < 2:
        print("用法: python -m templates.shjkyxy.make <样表.docx>")
        sys.exit(1)

    src = sys.argv[1]
    here = os.path.dirname(os.path.abspath(__file__))
    out = os.path.join(here, "template.docx")

    print(f"输入: {src}")
    print(f"输出: {out}")
    make(src, out)
    print("完成!")


if __name__ == "__main__":
    main()
