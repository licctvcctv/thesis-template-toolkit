"""
信息工程学院（电子信息、计算机类）毕业设计模版制作器。
将原始论文 .docx 转换为 docxtpl 模版。

用法: python make.py <原始.docx> [输出模版.docx]

原始文档结构（分析结果）:
  封面（Section 0）:
    p2:  本科毕业设计
    p4:  论文题目（Body Text，黑体26pt）
    p11: 日期 [SECTION_BREAK]
    Table 0: 封面信息表（院部、姓名、学号、专业、届别、导师）

  中文摘要页（Section 1）:
    p12: 中文标题（黑体16pt 居中）
    p13: 摘要：+ 正文（"摘要：" 加粗）
    p14: 关键词：+ 关键词 [SECTION_BREAK]
         注意：关键词/：/内容 分在不同 run 里

  英文摘要页（Section 2）:
    p15: 英文标题（Times New Roman 16pt 居中）
    p16: Abstract: + 正文
    p17: Keywords: + 关键词 [SECTION_BREAK]
         注意：Keyword/s/:/ 分在不同 run 里

  正文（Section 3, 从 p18 开始）:
    章标题: x■□□ — 只有 p18 是 Heading 1，其余（p50/69/88）用 Normal+黑体16pt
    节标题: x.x■□□ — 只有 p19 是 Heading 2，其余用 Normal+黑体15pt
    小节标题: x.x.x■□□ — 全部 Normal+黑体14pt
    正文: Normal 12pt / Body Text
    图标注: Body Text 居中（如 "图2-1  管理员用户用例图"）
    表标注: _Style 1 居中（如 "表3-1  管理员信息表"）+ Table 1

  结论: p123 Normal 黑体16pt 居中
  参考文献: p132 "参考文献", p133-147 各条 Body Text
  致谢: p150 HTML Preformatted 黑体16pt
  附录: p154 Heading 1
"""
import os
import sys
import re
import copy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def main():
    if len(sys.argv) < 2:
        print("用法: python make.py <原始.docx> [输出.docx]")
        sys.exit(1)

    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else "template.docx"

    doc = Document(src)
    paras = doc.paragraphs

    # ========== Step 1: 封面字段替换 ==========
    print("Step 1: 封面字段替换...")

    # p4: 论文题目
    _replace_para_text(paras[4], '{{ title_zh }}')
    print("  p4: title_zh")

    # p11: 日期
    _replace_para_text(paras[11], '{{ date }}')
    print("  p11: date")

    # Table 0: 封面信息表
    table = doc.tables[0]
    cover_vars = [
        (0, 'college', '院部'),
        (1, 'name', '姓名'),
        (2, 'student_id', '学号'),
        (3, 'major', '专业'),
        (4, 'class_year', '届别'),
        (5, 'advisor', '导师'),
    ]
    for ri, var, desc in cover_vars:
        cell = table.rows[ri].cells[1]
        for p in cell.paragraphs:
            _replace_para_text(p, '{{ ' + var + ' }}')
        print(f"  table row{ri}: {desc} -> {{{{{var}}}}}")

    # ========== Step 2: 中文摘要 ==========
    print("Step 2: 中文摘要...")

    # p12: 摘要页标题
    _replace_para_text(paras[12], '{{ title_zh }}')
    print("  p12: title_zh (abstract page)")

    # p13: 摘要正文 — "摘要：" 加粗，后面是内容
    _replace_label_value(paras[13], '摘要', '：', 'abstract_zh')
    print("  p13: abstract_zh")

    # p14: 关键词 — "关键词" 和 "：" 可能在不同 run 里
    _replace_label_value(paras[14], '关键词', '：', 'keywords_zh')
    print("  p14: keywords_zh")

    # ========== Step 3: 英文摘要 ==========
    print("Step 3: 英文摘要...")

    # p15: 英文标题
    _replace_para_text(paras[15], '{{ title_en }}')
    print("  p15: title_en")

    # p16: Abstract 正文（冒号后要加空格）
    _replace_label_value(paras[16], 'Abstract', ':', 'abstract_en',
                         space_after_colon=' ')
    print("  p16: abstract_en")

    # p17: Keywords — "Keyword"+"s"+":" 分在3个 run
    _replace_label_value(paras[17], 'Keyword', ':', 'keywords_en',
                         space_after_colon=' ')
    print("  p17: keywords_en")

    # ========== Step 4: 定位正文区域 ==========
    print("Step 4: 定位正文区域...")

    body_start = None
    conclusion_idx = None
    refs_idx = None
    ack_idx = None
    appendix_idx = None

    for i, p in enumerate(paras):
        style = p.style.name if p.style else ''
        text = p.text.strip()

        if style == 'Heading 1' and body_start is None:
            body_start = i
        if text == '结论' and i > 100:
            conclusion_idx = i
        if text == '参考文献' and i > 100:
            refs_idx = i
        if text == '致谢' and i > 100:
            ack_idx = i
        if '附录' in text and style == 'Heading 1':
            appendix_idx = i

    body_end = conclusion_idx - 1 if conclusion_idx else refs_idx - 1
    while body_end > body_start and not paras[body_end].text.strip():
        body_end -= 1

    print(f"  body: p{body_start} - p{body_end}")
    print(f"  conclusion: p{conclusion_idx}")
    print(f"  refs: p{refs_idx}")
    print(f"  ack: p{ack_idx}")
    print(f"  appendix: p{appendix_idx}")

    # ========== Step 5: 正文 Jinja2 循环 ==========
    print("Step 5: 正文循环...")

    # 这个模版样式很乱：只有第1章用 Heading 1, 第1节用 Heading 2
    # 其余章/节/小节全用 Normal/Body Text + 黑体字体
    # 策略：保留 p18(H1), p19(H2), p20(subsection) 作为样本
    sample_h1 = body_start  # p18, Heading 1
    sample_h2 = None
    sample_h3 = None

    for i in range(body_start + 1, body_end + 1):
        p = paras[i]
        style = p.style.name if p.style else ''
        text = p.text.strip()
        if style == 'Heading 2' and sample_h2 is None:
            sample_h2 = i
        elif re.match(r'\d+\.\d+\.\d+', text) and sample_h3 is None:
            sample_h3 = i

    print(f"  samples: H1=p{sample_h1}, H2=p{sample_h2}, H3=p{sample_h3}")

    # 清空正文区域（保留样本）
    keep = {sample_h1, sample_h2, sample_h3}
    for i in range(body_end, body_start - 1, -1):
        if i not in keep:
            _remove_para(paras[i])

    # 删除正文区域内残留的表格（Table 1 等）
    body_xml = doc.element.body
    keep_elements = set()
    paras = doc.paragraphs
    for p in paras:
        style = p.style.name if p.style else ''
        if style.startswith('Heading') and p.text.strip():
            keep_elements.add(p._p)

    in_body = False
    tables_to_remove = []
    for elem in body_xml:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            if elem in keep_elements and not in_body:
                in_body = True
            texts = [t.text for t in elem.findall('.//{%s}t' % WNS) if t.text]
            text = ''.join(texts).strip()
            if text == '结论':
                in_body = False
        elif tag == 'tbl' and in_body:
            tables_to_remove.append(elem)

    for tbl in tables_to_remove:
        body_xml.remove(tbl)
    print(f"  删除 {len(tables_to_remove)} 个残留表格")

    # 重新扫描定位样本
    paras = doc.paragraphs
    p_h1 = p_h2 = p_h3 = None
    for p in paras:
        style = p.style.name if p.style else ''
        text = p.text.strip()
        if style == 'Heading 1' and text and '附录' not in text and p_h1 is None:
            p_h1 = p
        elif style == 'Heading 2' and text and p_h2 is None:
            p_h2 = p
        elif re.match(r'\d+\.\d+\.\d+', text) and p_h3 is None:
            p_h3 = p

    if not p_h1 or not p_h2 or not p_h3:
        print(f"  ERROR: 找不到样本段落 H1={p_h1} H2={p_h2} H3={p_h3}")
        sys.exit(1)

    # 插入 Jinja2 循环
    _insert_jinja_para_before(p_h1, '{%p for ch in chapters %}')
    _replace_para_text(p_h1, '{{ ch.title }}')
    _insert_jinja_para_after(p_h1, '{%p for sec in ch.sections %}')

    _replace_para_text(p_h2, '{{ sec.title }}')
    _insert_jinja_para_after(p_h2, '{%p for item in sec.content %}')
    paras = doc.paragraphs
    for p in paras:
        if p.text.strip() == '{%p for item in sec.content %}':
            _insert_jinja_para_after(p, '{{ item }}')
            break
    paras = doc.paragraphs
    for p in paras:
        if p.text.strip() == '{{ item }}':
            _insert_jinja_para_after(p, '{%p endfor %}')
            break

    paras = doc.paragraphs
    first_endfor = None
    for p in paras:
        if p.text.strip() == '{%p endfor %}':
            first_endfor = p
            break
    _insert_jinja_para_after(first_endfor, '{%p for sub in sec.subsections %}')

    _replace_para_text(p_h3, '{{ sub.title }}')
    _insert_jinja_para_after(p_h3, '{%p for item in sub.content %}')
    paras = doc.paragraphs
    for p in paras:
        if p.text.strip() == '{%p for item in sub.content %}':
            _insert_jinja_para_after(p, '{{ item }}')
            break
    paras = doc.paragraphs
    found_sub_item = False
    for p in paras:
        if 'for item in sub.content' in p.text:
            found_sub_item = True
            continue
        if found_sub_item and p.text.strip() == '{{ item }}':
            _insert_jinja_para_after(p, '{%p endfor %}')
            break

    # endfor sub
    paras = doc.paragraphs
    in_sub_content = False
    sub_content_endfor = None
    for p in paras:
        if 'for item in sub.content' in p.text:
            in_sub_content = True
        if in_sub_content and p.text.strip() == '{%p endfor %}':
            sub_content_endfor = p
            break
    _insert_jinja_para_after(sub_content_endfor, '{%p endfor %}')

    # endfor sec
    paras = doc.paragraphs
    endfor_list = [p for p in paras if p.text.strip() == '{%p endfor %}']
    _insert_jinja_para_after(endfor_list[-1], '{%p endfor %}')

    # endfor ch
    paras = doc.paragraphs
    endfor_list = [p for p in paras if p.text.strip() == '{%p endfor %}']
    _insert_jinja_para_after(endfor_list[-1], '{%p endfor %}')

    print("  循环结构已插入")

    # 清除 {{ item }} 段落的字体格式，让正文回归 Normal 默认样式（宋体12pt）
    # 否则会继承标题的黑体字
    paras = doc.paragraphs
    for p in paras:
        if p.text.strip() == '{{ item }}':
            for r in p.runs:
                rPr = r._r.rPr
                if rPr is not None:
                    # 清除字体名、字号、加粗等，让 Normal 样式生效
                    for child in list(rPr):
                        rPr.remove(child)
    print("  清除 {{ item }} 的标题字体继承")

    # ========== Step 6: 结论 ==========
    print("Step 6: 结论...")
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if p.text.strip() == '结论':
            for j in range(i + 1, min(i + 10, len(paras))):
                if paras[j].text.strip():
                    _replace_para_text(paras[j], '{{ conclusion }}')
                    for k in range(j + 1, len(paras)):
                        text_k = paras[k].text.strip()
                        if text_k == '参考文献':
                            break
                        if text_k:
                            _clear_para(paras[k])
                    print(f"  p{j}: conclusion")
                    break
            break

    # ========== Step 7: 参考文献 ==========
    print("Step 7: 参考文献...")
    paras = doc.paragraphs
    refs_start = None
    for i, p in enumerate(paras):
        if p.text.strip() == '参考文献' and i > 20:
            refs_start = i
            break

    if refs_start:
        first_ref = None
        for i in range(refs_start + 1, len(paras)):
            text = paras[i].text.strip()
            if text and text.startswith('['):
                first_ref = i
                break

        if first_ref:
            _insert_jinja_para_before(paras[first_ref], '{%p for ref in references %}')
            paras = doc.paragraphs
            for i, p in enumerate(paras):
                if p.text.strip() == '{%p for ref in references %}':
                    ref_para = paras[i + 1]
                    _replace_para_text(ref_para, '{{ ref }}')
                    paras = doc.paragraphs
                    in_refs = False
                    to_remove = []
                    for j, p2 in enumerate(paras):
                        if '{{ ref }}' in p2.text:
                            in_refs = True
                            continue
                        if in_refs:
                            text2 = p2.text.strip()
                            if text2 == '致谢' or '附录' in text2:
                                break
                            if text2:
                                to_remove.append(p2)
                    for p2 in to_remove:
                        _remove_para(p2)
                    paras = doc.paragraphs
                    for p2 in paras:
                        if '{{ ref }}' in p2.text:
                            _insert_jinja_para_after(p2, '{%p endfor %}')
                            break
                    break
        print("  refs loop set up")

    # ========== Step 8: 致谢 ==========
    print("Step 8: 致谢...")
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        text = p.text.strip()
        if text == '致谢' and i > 20:
            for j in range(i + 1, min(i + 10, len(paras))):
                if paras[j].text.strip():
                    _replace_para_text(paras[j], '{{ acknowledgement }}')
                    for k in range(j + 1, len(paras)):
                        text_k = paras[k].text.strip()
                        style_k = paras[k].style.name if paras[k].style else ''
                        if '附录' in text_k or style_k == 'Heading 1':
                            break
                        if text_k:
                            _clear_para(paras[k])
                    print(f"  p{j}: acknowledgement")
                    break
            break

    # ========== Step 9: 附录 ==========
    print("Step 9: 附录...")
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if '附录' in p.text.strip() and p.style and p.style.name == 'Heading 1':
            _replace_para_text(p, '{{ appendix_title }}')
            for j in range(i + 1, min(i + 5, len(paras))):
                if paras[j].text.strip():
                    _replace_para_text(paras[j], '{{ appendix_content }}')
                    for k in range(j + 1, len(paras)):
                        if paras[k].text.strip():
                            _clear_para(paras[k])
                    break
            print(f"  p{i}: appendix")
            break

    # ========== Step 10: 清理残留正文 ==========
    print("Step 10: 清理残留正文...")
    paras = doc.paragraphs
    in_body_loop = False
    body_loop_done = False
    endfor_count = 0
    for p in paras:
        text = p.text.strip()
        if 'for ch in chapters' in text:
            in_body_loop = True
            continue
        if in_body_loop and not body_loop_done:
            if text == '{%p endfor %}':
                endfor_count += 1
                if endfor_count >= 5:
                    body_loop_done = True
                continue
            style = p.style.name if p.style else ''
            if style in ('Heading 1', 'Heading 2', 'Heading 3'):
                continue
            if '{%' in text or '{{' in text:
                continue
            if text:
                _clear_para(p)

    # ========== Step 11: 修复模板样式和图片格式 ==========
    print("Step 11: 修复样式...")

    # Normal 样式自带 firstLineChars="200" 会导致图片/表标注也被缩进
    # 把它去掉，正文缩进改在 {{ item }} 段落里单独设置
    from docx.oxml.ns import qn as _qn
    for style in doc.styles.element.findall(_qn('w:style')):
        name_el = style.find(_qn('w:name'))
        if name_el is None:
            continue
        sname = name_el.get(_qn('w:val'))
        if sname == 'Normal':
            pPr = style.find(_qn('w:pPr'))
            if pPr is not None:
                ind = pPr.find(_qn('w:ind'))
                if ind is not None:
                    pPr.remove(ind)
                    print("  Normal 样式: 移除 firstLineChars 缩进")

    # 给 {{ item }} 段落加回首行缩进（正文需要）
    paras = doc.paragraphs
    for p in paras:
        text = p.text.strip()
        if text == '{{ item }}':
            from docx.oxml import OxmlElement
            pPr = p._p.find(_qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p._p.insert(0, pPr)
            ind = OxmlElement('w:ind')
            ind.set(_qn('w:firstLine'), '480')
            ind.set(_qn('w:firstLineChars'), '200')
            pPr.append(ind)
    print("  {{ item }} 段落: 添加首行缩进")

    # ========== 保存 ==========
    doc.save(out)
    print(f"\n完成: {out}")

    # ========== 验证 ==========
    _verify(out)


def _replace_label_value(para, label_text, colon_char, var_name,
                         space_after_colon=''):
    """
    替换 "标签+冒号+值" 格式的段落。
    处理标签/冒号/值分散在多个 run 的情况。

    例如: 关键词(run0) + 空(run1) + ：(run2) + 值(run3+)
    → 保留 run0~run2（标签+冒号），run3 替换为变量，后续清空
    """
    runs = para.runs
    if not runs:
        return

    # 第一步：找到标签 run
    label_found = False
    colon_found = False
    value_set = False

    for ri, r in enumerate(runs):
        text = r.text

        # 还没找到标签
        if not label_found:
            if label_text in text:
                label_found = True
                # 检查这个 run 里是否也包含冒号
                if colon_char in text:
                    colon_found = True
                    # 冒号后可能还有值
                    idx = text.index(colon_char) + len(colon_char)
                    if idx < len(text):
                        # 冒号后面有内容，截断
                        r.text = text[:idx] + space_after_colon
                        # 标记下一个 run 设变量
                    else:
                        # 冒号就是最后，加空格后值在后续 run
                        r.text = text + space_after_colon
            continue

        # 已找到标签，寻找冒号
        if not colon_found:
            if colon_char in text:
                colon_found = True
                idx = text.index(colon_char) + len(colon_char)
                r.text = text[:idx] + space_after_colon
            elif text.strip() == '':
                # 空 run（如 bold 切换），保留
                pass
            else:
                # 标签的一部分（如 Keywords 的 "s"），保留
                pass
            continue

        # 已找到冒号，设置变量
        if not value_set:
            r.text = '{{ ' + var_name + ' }}'
            value_set = True
        else:
            r.text = ''

    # 如果冒号在标签同一 run 但值没找到 run，追加
    if colon_found and not value_set:
        # 找最后一个非空 run 后追加
        for r in runs:
            if r.text and not value_set:
                pass
        # 简单处理：在最后一个 run 后面
        if runs:
            runs[-1].text = '{{ ' + var_name + ' }}'


def _replace_para_text(para, new_text):
    """替换段落文本，保留第一个 run 的格式"""
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ''
    else:
        para.text = new_text


def _clear_para(para):
    """清空段落文本"""
    for r in para.runs:
        r.text = ''


def _remove_para(para):
    """从文档中删除段落"""
    p = para._p
    p.getparent().remove(p)


def _insert_jinja_para_before(para, text):
    """在段落前插入一个 Jinja2 标记段落"""
    new_p = OxmlElement('w:p')
    new_r = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set(qn('xml:space'), 'preserve')
    if para.runs and para.runs[0]._r.rPr is not None:
        new_r.append(copy.deepcopy(para.runs[0]._r.rPr))
    new_r.append(new_t)
    new_p.append(new_r)
    para._p.addprevious(new_p)


def _insert_jinja_para_after(para, text):
    """在段落后插入一个 Jinja2 标记段落"""
    new_p = OxmlElement('w:p')
    new_r = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set(qn('xml:space'), 'preserve')
    if para.runs and para.runs[0]._r.rPr is not None:
        new_r.append(copy.deepcopy(para.runs[0]._r.rPr))
    new_r.append(new_t)
    new_p.append(new_r)
    para._p.addnext(new_p)


def _verify(docx_path):
    """验证模版关键变量"""
    doc = Document(docx_path)
    print("\n=== 模版验证 ===")
    expected = [
        'title_zh', 'title_en', 'date',
        'college', 'name', 'student_id', 'major', 'class_year', 'advisor',
        'abstract_zh', 'keywords_zh', 'abstract_en', 'keywords_en',
        'ch.title', 'sec.title', 'sub.title',
        'conclusion', 'acknowledgement',
        'ref',
    ]
    found = set()
    for p in doc.paragraphs:
        text = p.text
        for var in expected:
            if '{{ ' + var + ' }}' in text:
                found.add(var)
    # 也检查表格
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for var in expected:
                        if '{{ ' + var + ' }}' in text:
                            found.add(var)
                        if '{{ ' + var + ' }}' in p.text:
                            found.add(var)

    for var in expected:
        status = '✓' if var in found else '✗ MISSING'
        print(f"  {status} {{{{{var}}}}}")

    # 检查关键词格式
    for p in doc.paragraphs:
        if '关键词' in p.text:
            print(f"  关键词段落: [{p.text[:50]}]")
        if 'Keyword' in p.text:
            print(f"  Keywords段落: [{p.text[:50]}]")


if __name__ == '__main__':
    main()
