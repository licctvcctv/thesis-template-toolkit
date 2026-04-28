"""
闽江学院理工类本科毕业论文模板清理脚本。

用法：
    cd /Users/a136/vs/45425/thesis_project
    python templates/mjxy_antenna/make.py \
        templates/mjxy_antenna/source.docx \
        templates/mjxy_antenna/template.docx

脚本保留 source.docx 的封面、诚信声明、摘要、页眉页脚、页面设置和
原有段落格式，只把可变字段替换成 docxtpl 占位符，并删除旧目录、
旧正文和旧参考文献。论文内容由 papers/<paper>/JSON + build.py 注入。
"""
from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def para_key(p):
    return re.sub(r"\s+", "", p.text or "")


def find_para(paragraphs, predicate, label):
    for i, p in enumerate(paragraphs):
        if predicate(p.text or "", p):
            return i, p
    raise RuntimeError(f"未找到段落: {label}")


def remove_para(p):
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def remove_block_range(start_el, end_el=None, include_start=True, include_end=False):
    body = start_el.getparent()
    children = list(body)
    start = children.index(start_el)
    if end_el is None:
        end = len(children)
    else:
        end = children.index(end_el)
        if include_end:
            end += 1
    if not include_start:
        start += 1
    for el in children[start:end]:
        if el.tag == qn("w:sectPr"):
            continue
        body.remove(el)


def clear_para(p, text):
    for r in list(p.runs):
        p._p.remove(r._r)
    p.add_run(text)


def copy_font(dst_run, src_run):
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    dst_run.font.size = src_run.font.size
    dst_run.font.name = src_run.font.name
    if src_run._r.rPr is not None and dst_run._r.rPr is not None:
        src_fonts = src_run._r.rPr.rFonts
        dst_fonts = dst_run._r.rPr.rFonts
        if src_fonts is not None and dst_fonts is not None:
            for attr in ("eastAsia", "ascii", "hAnsi"):
                val = src_fonts.get(qn(f"w:{attr}"))
                if val:
                    dst_fonts.set(qn(f"w:{attr}"), val)


def set_single_run(p, text, size=None, bold=None):
    template = p.runs[0] if p.runs else None
    for r in list(p.runs):
        p._p.remove(r._r)
    run = p.add_run(text)
    if template is not None:
        copy_font(run, template)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    return run


def set_keywords_zh(p):
    templates = list(p.runs)
    for r in list(p.runs):
        p._p.remove(r._r)
    r0 = p.add_run("关键词")
    if len(templates) > 0:
        copy_font(r0, templates[0])
    r1 = p.add_run("：")
    if len(templates) > 1:
        copy_font(r1, templates[1])
    r2 = p.add_run("{{ keywords_zh }}")
    if len(templates) > 2:
        copy_font(r2, templates[2])


def set_keywords_en(p):
    templates = list(p.runs)
    for r in list(p.runs):
        p._p.remove(r._r)
    r0 = p.add_run("Key words")
    if len(templates) > 0:
        copy_font(r0, templates[0])
    r1 = p.add_run(": ")
    if len(templates) > 1:
        copy_font(r1, templates[1])
    r2 = p.add_run("{{ keywords_en }}")
    if len(templates) > 2:
        copy_font(r2, templates[2])


def set_cover_field(p, label, value_tag, value_spaces=18):
    """Rebuild one source cover line while preserving the paragraph itself."""
    label_run = p.runs[0] if p.runs else None
    value_template = None
    for run in p.runs:
        if run.underline:
            value_template = run
            break
    for r in list(p.runs):
        p._p.remove(r._r)

    r = p.add_run(label)
    if label_run is not None:
        copy_font(r, label_run)
    r.underline = False

    spacer = p.add_run("    ")
    if value_template is not None:
        copy_font(spacer, value_template)
    spacer.underline = True

    v = p.add_run(value_tag)
    if value_template is not None:
        copy_font(v, value_template)
    v.underline = True

    tail = p.add_run(" " * value_spaces)
    if value_template is not None:
        copy_font(tail, value_template)
    tail.underline = True


def normalize_cover(doc: Document):
    paragraphs = list(doc.paragraphs)
    cover_map = {
        "题目": ("题    目", "{{ title_zh }}", 4),
        "学生姓名": ("学生姓名", "{{ author }}", 12),
        "学号": ("学    号", "{{ student_id }}", 10),
        "系别": ("系    别", "{{ college }}", 2),
        "年级": ("年    级", "{{ grade }}", 18),
        "专业": ("专    业", "{{ major }}", 14),
        "指导教师": ("指导教师", "{{ advisor }}", 12),
        "职称": ("职    称", "{{ advisor_title }}", 14),
        "完成日期": ("完成日期", "{{ finish_date }}", 10),
    }
    for p in paragraphs[:20]:
        key = para_key(p)
        for compact, args in cover_map.items():
            if key.startswith(compact):
                set_cover_field(p, *args)
                break


def strip_shading(p):
    for shd in list(p._p.xpath(".//w:shd")):
        parent = shd.getparent()
        if parent is not None:
            parent.remove(shd)


def ensure_page_break_before(p):
    p.paragraph_format.page_break_before = True


def set_body_placeholder_style(p):
    p.paragraph_format.page_break_before = True
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def patch_update_fields(docx_path: Path):
    tmp = Path(str(docx_path) + ".tmp")
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "word/settings.xml":
                    xml = data.decode("utf-8")
                    if "w:updateFields" not in xml:
                        xml = xml.replace("</w:settings>", '<w:updateFields w:val="true"/></w:settings>')
                    data = xml.encode("utf-8")
                zout.writestr(info, data)
    tmp.replace(docx_path)


def normalize_declaration(paragraphs):
    _, body = find_para(paragraphs, lambda t, p: "兹提交的毕业论文" in t, "诚信声明正文")
    set_single_run(
        body,
        "兹提交的毕业论文（设计）《{{ title_zh }}》，是本人在指导老师{{ advisor }}的指导下独立研究、撰写的成果；论文（设计）未剽窃、抄袭他人的学术观点、思想和成果，未篡改研究数据，论文（设计）中所引用的文字、研究成果均已在论文（设计）中以明确的方式标明；在毕业论文（设计）工作过程中，本人恪守学术规范，遵守学校有关规定，依法享有和承担由此论文（设计）产生的权利和责任。",
    )
    _, signature = find_para(paragraphs, lambda t, p: t.strip().startswith("声明人（签名）"), "声明签名")
    set_single_run(signature, "声明人（签名）：{{ author }}")
    _, date = find_para(paragraphs, lambda t, p: "年" in t and "月" in t and "日" in t and para_key(p).startswith("2021"), "声明日期")
    set_single_run(date, "{{ finish_date_cn }}")


def ensure_summary_placeholders(paragraphs, title_idx, key_idx, prefix, keep_count=3):
    body = [p for p in paragraphs[title_idx + 1:key_idx] if (p.text or "").strip()]
    for i, p in enumerate(body):
        if i < keep_count:
            set_single_run(p, "{{ %s_%d }}" % (prefix, i + 1))
            strip_shading(p)
        else:
            remove_para(p)
    if len(body) < keep_count:
        key_p = paragraphs[key_idx]
        for i in range(len(body), keep_count):
            inserted = key_p.insert_paragraph_before("{{ %s_%d }}" % (prefix, i + 1))
            inserted.style = body[0].style if body else key_p.style
            strip_shading(inserted)


def add_center_title(doc: Document, text: str, page_break=True):
    p = doc.add_paragraph()
    if page_break:
        p.paragraph_format.page_break_before = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    return p


def main():
    if len(sys.argv) < 2:
        raise SystemExit("用法: python make.py <source.docx> [template.docx]")

    source = Path(sys.argv[1])
    output = Path(sys.argv[2]) if len(sys.argv) > 2 else Path(__file__).with_name("template.docx")

    doc = Document(source)
    normalize_cover(doc)
    paragraphs = list(doc.paragraphs)
    normalize_declaration(paragraphs)

    paragraphs = list(doc.paragraphs)
    zh_title_idx, zh_title = find_para(paragraphs, lambda t, p: para_key(p) == "摘要", "中文摘要标题")
    zh_key_idx, zh_key = find_para(paragraphs, lambda t, p: t.strip().startswith("关键词"), "中文关键词")
    en_title_idx, en_title = find_para(paragraphs, lambda t, p: para_key(p) == "Abstract", "英文摘要标题")
    en_key_idx, en_key = find_para(paragraphs, lambda t, p: t.strip().startswith("Key words"), "英文关键词")
    _, toc_title = find_para(paragraphs, lambda t, p: para_key(p) == "目录", "目录标题")
    _, body_start = find_para(paragraphs, lambda t, p: para_key(p).startswith("1.绪论"), "正文起点")

    print(f"摘要: P{zh_title_idx}, Abstract: P{en_title_idx}, 正文: P{paragraphs.index(body_start)}")

    ensure_summary_placeholders(paragraphs, zh_title_idx, zh_key_idx, "abstract_zh", 3)
    set_keywords_zh(zh_key)
    strip_shading(zh_key)

    paragraphs = list(doc.paragraphs)
    _, zh_key = find_para(paragraphs, lambda t, p: t.strip().startswith("关键词"), "中文关键词")
    _, en_title = find_para(paragraphs, lambda t, p: para_key(p) == "Abstract", "英文摘要标题")
    remove_block_range(zh_key._element, en_title._element, include_start=False, include_end=False)
    ensure_page_break_before(en_title)

    paragraphs = list(doc.paragraphs)
    en_title_idx, en_title = find_para(paragraphs, lambda t, p: para_key(p) == "Abstract", "英文摘要标题")
    en_key_idx, en_key = find_para(paragraphs, lambda t, p: t.strip().startswith("Key words"), "英文关键词")
    ensure_summary_placeholders(paragraphs, en_title_idx, en_key_idx, "abstract_en", 3)
    set_keywords_en(en_key)
    strip_shading(en_key)

    paragraphs = list(doc.paragraphs)
    _, en_key = find_para(paragraphs, lambda t, p: t.strip().startswith("Key words"), "英文关键词")
    _, toc_title = find_para(paragraphs, lambda t, p: para_key(p) == "目录", "目录标题")
    _, body_start = find_para(paragraphs, lambda t, p: para_key(p).startswith("1.绪论"), "正文起点")

    remove_block_range(en_key._element, toc_title._element, include_start=False, include_end=False)
    paragraphs = list(doc.paragraphs)
    _, toc_title = find_para(paragraphs, lambda t, p: para_key(p) == "目录", "目录标题")
    _, body_start = find_para(paragraphs, lambda t, p: para_key(p).startswith("1.绪论"), "正文起点")
    remove_block_range(toc_title._element, body_start._element, include_start=False, include_end=False)
    remove_block_range(body_start._element)

    ensure_page_break_before(toc_title)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_p = doc.add_paragraph("{{p toc_doc}}")
    toc_p.paragraph_format.line_spacing = Pt(18)
    toc_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    body_p = doc.add_paragraph("{{p body}}")
    set_body_placeholder_style(body_p)
    add_center_title(doc, "结论", page_break=True)
    doc.add_paragraph("{{p conclusion_doc}}")
    add_center_title(doc, "致谢", page_break=True)
    doc.add_paragraph("{{p acknowledgement_doc}}")
    add_center_title(doc, "参考文献", page_break=True)
    doc.add_paragraph("{{p references_doc}}")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    patch_update_fields(output)

    text = "\n".join(p.text for p in Document(str(output)).paragraphs)
    for tag in ("{{p toc_doc}}", "{{p body}}", "{{p conclusion_doc}}", "{{p references_doc}}"):
        if tag not in text:
            raise RuntimeError(f"模板缺少占位符: {tag}")
    print(f"模板已生成: {output}")


if __name__ == "__main__":
    main()
