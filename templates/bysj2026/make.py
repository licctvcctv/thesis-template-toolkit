"""
毕业设计说明书 2026版 模版制作器。

用法: python make.py <原始.docx> [输出.docx]
"""
import os, sys, re, copy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def main():
    if len(sys.argv) < 2:
        print("用法: python make.py <原始.docx> [输出.docx]")
        sys.exit(1)

    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else "template.docx"
    doc = Document(src)
    paras = doc.paragraphs

    # ====== Step 0: 定位所有关键段落 ======
    print("Step 0: 定位关键段落...")
    idx = {}
    for i, p in enumerate(paras):
        text = p.text.strip()
        norm = text.replace(' ', '').replace('\u3000', '')
        if re.match(r'^1\s+引言', text) or re.match(r'^1\s+\S', text) and 'H1' not in idx:
            idx['H1'] = i
        if re.match(r'^1[．.]\s*1', text) and 'H2' not in idx:
            idx['H2'] = i
        if re.match(r'^1\.1\.1', text) and 'H3' not in idx:
            idx['H3'] = i
        if re.match(r'^2\s+×', text) and 'CH2' not in idx:
            idx['CH2'] = i
        if norm.startswith('结论') and i > 50:
            idx['conclusion'] = i
        if norm.startswith('致谢') and i > 50:
            idx['ack'] = i
        if '参考文献' in norm and i > 50 and len(norm) < 15:
            idx['refs'] = i

    for k, v in idx.items():
        print(f"  {k}: p{v} = {paras[v].text.strip()[:50]}")

    # ====== Step 1: 封面 ======
    print("Step 1: 封面...")
    # 按内容定位封面字段（不依赖硬编码索引）
    for i, p in enumerate(paras[:20]):
        text = p.text.strip()
        if text.startswith('题') and (':' in text or '：' in text):
            # 题目行：保留 "题    目:" 标签，后面替换
            colon_found = False
            for r in p.runs:
                if ':' in r.text or '：' in r.text:
                    sep = ':' if ':' in r.text else '：'
                    ic = r.text.index(sep)
                    r.text = r.text[:ic + 1] + ' '
                    colon_found = True
                    continue
                if colon_found:
                    r.text = ''
            if p.runs:
                p.runs[-1].text = '{{ title_zh }}'
            # 清除下一个非空段落（可能是标题第二行）
            for j in range(i + 1, min(i + 3, 20)):
                if paras[j].text.strip():
                    _clear_para(paras[j])
                    break
        elif '院系名称' in text:
            _replace_inline_fields(p, {
                '院系名称': 'college', '专业班级': 'class_name'})
        elif '学生姓名' in text:
            _replace_inline_fields(p, {
                '学生姓名': 'name', '学号': 'student_id'})
        elif '指导教师' in text:
            _replace_inline_fields(p, {
                '指导教师': 'advisor', '教师职称': 'advisor_title'})
        elif re.match(r'^\d{4}年', text):
            _replace_para_text(p, '{{ date }}')
    print("  封面变量已替换")

    # ====== Step 2: 摘要表格 ======
    print("Step 2: 摘要...")
    table = doc.tables[0]
    # 中文摘要 (row 0)
    for p in table.rows[0].cells[0].paragraphs:
        text = p.text.strip()
        if '×' * 3 in text and '关键' not in text:
            _replace_para_text(p, '{{ abstract_zh }}')
        elif '关键词' in text:
            _replace_label_value(p, '关键词', '：', 'keywords_zh')
        elif text.startswith('（') and ('黑体' in text or '不要' in text):
            _clear_para(p)

    # 英文摘要 (row 1)
    for p in table.rows[1].cells[0].paragraphs:
        text = p.text.strip()
        if text.startswith('Title'):
            _replace_para_text(p, '{{ title_en }}')
        elif '×' * 3 in text and 'Keyword' not in text:
            _replace_para_text(p, '{{ abstract_en }}')
        elif 'Keyword' in text:
            _replace_label_value(p, 'Keyword', ' ', 'keywords_en',
                                 merge_label_parts=True)
        elif text.startswith('（') and 'Times' in text:
            _clear_para(p)
    print("  摘要变量已替换")

    # ====== Step 3: 清除所有格式说明 ======
    print("Step 3: 清除格式说明...")

    # 先定位正文样本段落（H2后、H3后的 ×× 段落），保护它们不被清除
    h1_idx = idx['H1']
    h2_idx = idx.get('H2', h1_idx)
    h3_idx = idx.get('H3', h2_idx)
    conclusion_idx = idx.get('conclusion', len(paras))

    body_sec_idx = None  # H2后的正文样本
    for i in range(h2_idx + 1, h3_idx):
        if paras[i].text.strip() and '×' in paras[i].text:
            body_sec_idx = i; break

    body_sub_idx = None  # H3后的正文样本
    for i in range(h3_idx + 1, min(h3_idx + 5, len(paras))):
        if paras[i].text.strip() and '×' in paras[i].text:
            body_sub_idx = i; break

    protected = set(idx.values())
    if body_sec_idx: protected.add(body_sec_idx)
    if body_sub_idx: protected.add(body_sub_idx)
    removed = 0

    # 3a: 目次区域 (p20 到 H1 之前)
    for i in range(20, h1_idx):
        if paras[i].text.strip():
            _clear_para(paras[i])
            removed += 1

    # 3b: 正文区域格式说明 (H3+正文样本之后 到 结论前)
    clean_start = (body_sub_idx + 1) if body_sub_idx else (h3_idx + 1)
    for i in range(clean_start, conclusion_idx):
        if i in protected:
            continue
        text = paras[i].text.strip()
        norm = text.replace(' ', '').replace('\u3000', '')
        if norm.startswith('结论') or norm.startswith('致谢') or '参考文献' in norm:
            continue
        if text:
            _clear_para(paras[i])
            removed += 1

    # 3c: 结论后到文末的格式说明
    refs_idx = idx.get('refs', len(paras))
    for i in range(refs_idx + 1, len(paras)):
        text = paras[i].text.strip()
        if '{%' in text or '{{' in text:
            continue
        if text and ('不少于' in text or '实事求是' in text or
                     '作者' in text and '超过' in text or
                     '著录' in text or '编排规范' in text or
                     '陈建军' in text or text == '例如：' or
                     '连续出版物' in text or text.startswith('*')):
            _clear_para(paras[i])
            removed += 1

    # 3d: 删除示例表格 (Table 1, 2, 3)
    body_xml = doc.element.body
    keep_tbl = table._tbl
    for tbl in body_xml.findall(f'{{{WNS}}}tbl'):
        if tbl is not keep_tbl:
            tbl.getparent().remove(tbl)
            removed += 1

    print(f"  清除 {removed} 个元素")

    # ====== Step 4: 正文循环 ======
    print("Step 4: 正文循环...")

    paras = doc.paragraphs
    # 重新定位样本
    p_h1 = p_h2 = p_h3 = None
    for p in paras:
        text = p.text.strip()
        if re.match(r'^1\s+引言', text) or re.match(r'^1\s+\S', text) and p_h1 is None:
            p_h1 = p
        if re.match(r'^1[．.]\s*1', text) and p_h2 is None:
            p_h2 = p
        if re.match(r'^1\.1\.1', text) and p_h3 is None:
            p_h3 = p

    if not p_h1:
        print("  ERROR: 找不到H1"); sys.exit(1)
    if not p_h2:
        p_h2 = p_h1
        print("  WARNING: H2 fallback to H1")
    if not p_h3:
        p_h3 = p_h2
        print("  WARNING: H3 fallback to H2")

    # 清除样本标题中的格式说明
    for sample in [p_h1, p_h2, p_h3]:
        _strip_format_hint(sample)

    # 定位正文样本段落（H1后面的 ×× 段落、H3后面的 ×× 段落）
    paras_list = list(paras)
    h1_i = paras_list.index(p_h1)
    h2_i = paras_list.index(p_h2) if p_h2 is not p_h1 else h1_i
    h3_i = paras_list.index(p_h3) if p_h3 is not p_h2 else h2_i

    # 找 H2 后面第一个正文样本 → sec.content 的 {{ item }}
    p_body_sec = None
    for i in range(h2_i + 1, h3_i):
        text = paras_list[i].text.strip()
        if text and '×' in text:
            p_body_sec = paras_list[i]
            break

    # 找 H3 后面第一个正文样本 → sub.content 的 {{ item }}
    p_body_sub = None
    for i in range(h3_i + 1, min(h3_i + 5, len(paras_list))):
        text = paras_list[i].text.strip()
        if text and '×' in text:
            p_body_sub = paras_list[i]
            break

    # 清除其余占位正文（不是用作 {{ item }} 的）
    for i in range(h1_i + 1, h3_i + 5):
        if i >= len(paras_list):
            break
        p = paras_list[i]
        if p is p_h2 or p is p_h3 or p is p_body_sec or p is p_body_sub:
            continue
        text = p.text.strip()
        if text and '×' in text:
            _clear_para(p)

    # 构建 Jinja2 循环，用原始正文段落做 {{ item }}
    _insert_jinja_para_before(p_h1, '{%p for ch in chapters %}')
    # H1 有多个 run（run0 可能是图片/分页标记用宋体小号，run1 才是标题黑体15pt）
    # 把文本放到正确的标题格式 run 里
    _replace_para_text_in_title_run(p_h1, '{{ ch.title }}')
    _insert_jinja_para_after(p_h1, '{%p for sec in ch.sections %}')

    _replace_para_text(p_h2, '{{ sec.title }}')
    _insert_jinja_para_after(p_h2, '{%p for item in sec.content %}')

    # 把 p_body_sec 变成 {{ item }}（保留原始段落的 pPr 和 rPr）
    _strip_format_hint(p_body_sec)
    _replace_para_text(p_body_sec, '{{ item }}')
    _insert_jinja_para_after(p_body_sec, '{%p endfor %}')

    # subsections loop
    paras = doc.paragraphs
    for p in paras:
        if p.text.strip() == '{%p endfor %}':
            _insert_jinja_para_after(p, '{%p for sub in sec.subsections %}')
            break

    _replace_para_text(p_h3, '{{ sub.title }}')
    _insert_jinja_para_after(p_h3, '{%p for item in sub.content %}')

    # 把 p_body_sub 变成 {{ item }}（保留原始段落格式）
    _strip_format_hint(p_body_sub)
    _replace_para_text(p_body_sub, '{{ item }}')
    _insert_jinja_para_after(p_body_sub, '{%p endfor %}')

    # Close: endfor sub, endfor sec, endfor ch
    paras = doc.paragraphs
    endfor_list = [p for p in paras if p.text.strip() == '{%p endfor %}']
    _insert_jinja_para_after(endfor_list[-1], '{%p endfor %}')
    paras = doc.paragraphs
    endfor_list = [p for p in paras if p.text.strip() == '{%p endfor %}']
    _insert_jinja_para_after(endfor_list[-1], '{%p endfor %}')
    paras = doc.paragraphs
    endfor_list = [p for p in paras if p.text.strip() == '{%p endfor %}']
    _insert_jinja_para_after(endfor_list[-1], '{%p endfor %}')

    print("  循环已插入")

    # ====== Step 5: 结论 ======
    print("Step 5: 结论/致谢/参考文献...")
    paras = doc.paragraphs

    # 结论
    for i, p in enumerate(paras):
        norm = p.text.strip().replace(' ', '').replace('\u3000', '')
        if norm.startswith('结论') and i > 30:
            _strip_format_hint(p)
            _clear_hint_paras(paras, i + 1, '致谢', 'conclusion')
            break

    # 致谢
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        norm = p.text.strip().replace(' ', '').replace('\u3000', '')
        if norm.startswith('致谢') and i > 30:
            _strip_format_hint(p)
            _clear_hint_paras(paras, i + 1, '参考文献', 'acknowledgement')
            break

    # 参考文献
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        norm = p.text.strip().replace(' ', '').replace('\u3000', '')
        if '参考文献' in norm and i > 30 and len(norm) < 15:
            _strip_format_hint(p)
            # 找到第一个 ×× 段落，设为 ref 循环
            for j in range(i + 1, min(i + 8, len(paras))):
                t = paras[j].text.strip()
                if t.startswith('（空'):
                    _clear_para(paras[j])
                elif '×' in t:
                    _insert_jinja_para_before(paras[j], '{%p for ref in references %}')
                    _replace_para_text(paras[j], '{{ ref }}')
                    _insert_jinja_para_after(paras[j], '{%p endfor %}')
                    # 清除后续的 ×× 和 …… 段落
                    paras2 = doc.paragraphs
                    after_ref = False
                    for p2 in paras2:
                        if '{%p endfor %}' in p2.text and after_ref:
                            break
                        if '{{ ref }}' in p2.text:
                            after_ref = True
                            continue
                        if after_ref:
                            pass  # endfor is next
                    # 清除 endfor 之后的残留
                    paras2 = doc.paragraphs
                    past_endfor = False
                    for p2 in paras2:
                        if past_endfor:
                            t2 = p2.text.strip()
                            if t2 and ('×' in t2 or '…' in t2):
                                _clear_para(p2)
                        if '{{ ref }}' in p2.text:
                            pass
                        if '{%p endfor %}' in p2.text:
                            # find the one after ref
                            pass
                    break
            break

    # 最终清理所有残留的 ×× 和格式说明
    paras = doc.paragraphs
    in_loop = False
    for p in paras:
        text = p.text.strip()
        if 'for ch in chapters' in text:
            in_loop = True
        if in_loop and '{%p endfor %}' in text:
            pass
        # 清除循环外的残留
        if text and '×' in text and '{%' not in text and '{{' not in text:
            _clear_para(p)
        if text == '（空1行）' or text == '（空2行）':
            _clear_para(p)
        if text == '…………':
            _clear_para(p)

    # 章标题加分页（每章新起一页）
    from docx.oxml.ns import qn as _qn
    paras = doc.paragraphs
    for p in paras:
        if p.text.strip() == '{{ ch.title }}':
            pPr = p._p.find(_qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p._p.insert(0, pPr)
            pb = OxmlElement('w:pageBreakBefore')
            pPr.append(pb)
            break

    # ====== Step 6: 删除文本框 ======
    print("Step 6: 清除文本框...")
    body_xml = doc.element.body
    ns_wps = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    ns_mc = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
    ns_wp = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    ns_wp14 = 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing'
    ns_v = 'urn:schemas-microsoft-com:vml'
    ns_w = WNS

    # 删除包含文本框的 drawing 和 pict 元素
    removed_tb = 0
    # 封面段落的 XML 元素，需要保护
    cover_p = doc.paragraphs[0]._p

    # 方法1: 找 wps:txbx 的祖先 w:r 并删除（跳过封面）
    for txbx in body_xml.findall(f'.//{{{ns_wps}}}txbx'):
        parent = txbx.getparent()
        while parent is not None:
            tag = parent.tag.split('}')[-1]
            if tag == 'r':
                # 跳过封面段落里的 run
                if parent.getparent() is cover_p:
                    break
                parent.getparent().remove(parent)
                removed_tb += 1
                break
            if tag == 'body':
                break
            parent = parent.getparent()

    # 方法2: 找 v:textbox 的祖先 w:r 并删除（跳过封面）
    for txbx in body_xml.findall(f'.//{{{ns_v}}}textbox'):
        parent = txbx.getparent()
        while parent is not None:
            tag = parent.tag.split('}')[-1]
            if tag == 'r':
                if parent.getparent() is cover_p:
                    break
                parent.getparent().remove(parent)
                removed_tb += 1
                break
            if tag == 'body':
                break
            parent = parent.getparent()

    print(f"  删除 {removed_tb} 个文本框")

    # ====== Step 7: 清除摘要表格内的格式说明 ======
    print("Step 7: 清除摘要表格格式说明...")
    table = doc.tables[0]
    cleaned_t = 0
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                text = p.text.strip()
                if text.startswith('（空'):
                    _clear_para(p)
                    cleaned_t += 1
                # "摘  要（小三黑体，居中）" → 保留 "摘  要"，删格式说明
                elif '摘' in text and '要' in text and '（' in text:
                    _strip_format_hint(p)
                    cleaned_t += 1
    print(f"  清除 {cleaned_t} 个格式说明")

    # ====== Step 8: 清除示例图片 ======
    print("Step 8: 清除示例图片...")
    ns_drawing = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    removed_img = 0
    # 段落中的图片（跳过 p0 校徽）
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if i == 0:
            continue  # 保留封面校徽
        drawings = p._p.findall(f'.//{{{ns_drawing}}}drawing')
        if drawings:
            for d in drawings:
                d.getparent().remove(d)
                removed_img += 1
    # 表格内的图片（如 Keywords 行的箭头）
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    drawings = p._p.findall(f'.//{{{ns_drawing}}}drawing')
                    if drawings:
                        for d in drawings:
                            d.getparent().remove(d)
                            removed_img += 1
    print(f"  删除 {removed_img} 个示例图片")

    doc.save(out)
    print(f"\n完成: {out}")
    _verify(out)


def _clear_hint_paras(paras, start, stop_keyword, var_name):
    """清除 start 到 stop_keyword 之间的段落，把第一个正文段替换为变量"""
    replaced = False
    for j in range(start, len(paras)):
        text = paras[j].text.strip()
        norm = text.replace(' ', '').replace('\u3000', '')
        if stop_keyword in norm:
            break
        if text.startswith('（空'):
            _clear_para(paras[j])
        elif text and '×' in text and not replaced:
            _replace_para_text(paras[j], '{{ ' + var_name + ' }}')
            replaced = True
        elif text and '×' in text:
            _clear_para(paras[j])


def _strip_format_hint(para):
    """清除段落中的 '（小3号黑体，居中）' 等格式说明"""
    in_hint = False
    for r in para.runs:
        if '（' in r.text and ('小' in r.text or '号' in r.text or '作为' in r.text or '可作为' in r.text):
            idx = r.text.index('（')
            r.text = r.text[:idx]
            in_hint = True
        elif in_hint:
            if '）' in r.text:
                idx = r.text.index('）')
                r.text = r.text[idx + 1:]
                in_hint = False
            else:
                r.text = ''


def _replace_inline_fields(para, field_map):
    """替换 '标签1：值  标签2：值' 格式的封面行"""
    full_text = para.text
    runs = para.runs
    if not runs:
        return

    fields = list(field_map.items())
    field_idx = 0
    waiting_value = False

    for r in runs:
        if field_idx >= len(fields):
            r.text = ''
            continue

        label, var = fields[field_idx]
        text = r.text

        # 检查是否包含当前标签
        if label in text:
            for sep in ['：', ':']:
                if sep in text:
                    idx = text.index(sep)
                    r.text = text[:idx + 1] + ' '
                    waiting_value = True
                    break
            continue

        # 检查是否包含冒号（标签和冒号分开）
        if not waiting_value and ('：' in text or ':' in text):
            for sep in ['：', ':']:
                if sep in text:
                    idx = text.index(sep)
                    r.text = text[:idx + 1] + ' '
                    waiting_value = True
                    break
            continue

        if waiting_value:
            r.text = '{{ ' + var + ' }}'
            field_idx += 1
            waiting_value = False
            continue

        # 可能是标签的一部分（如 "号" in "学    号："）
        # 跳过
        if text.strip() in ['号', '称', '级']:
            continue


def _replace_label_value(para, label_text, colon_char, var_name,
                         space_after_colon='', merge_label_parts=False):
    runs = para.runs
    if not runs: return
    label_found = colon_found = value_set = False
    for r in runs:
        text = r.text
        if not label_found:
            if label_text in text:
                label_found = True
                if colon_char in text:
                    colon_found = True
                    idx = text.index(colon_char) + len(colon_char)
                    r.text = text[:idx] + space_after_colon
            continue
        if not colon_found:
            if colon_char in text:
                colon_found = True
                idx = text.index(colon_char) + len(colon_char)
                r.text = text[:idx] + space_after_colon
            continue
        if not value_set:
            r.text = '{{ ' + var_name + ' }}'
            value_set = True
        else:
            r.text = ''
    if colon_found and not value_set and runs:
        runs[-1].text = '{{ ' + var_name + ' }}'


def _replace_para_text(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]: r.text = ''
    else:
        para.text = new_text


def _replace_para_text_in_title_run(para, new_text):
    """替换段落文本，把文本放到字号最大的 run 里（标题 run），其余清空"""
    runs = para.runs
    if not runs:
        para.text = new_text
        return
    # 找字号最大的 run 索引
    best_idx = 0
    best_size = 0
    for i, r in enumerate(runs):
        sz = r.font.size or 0
        if sz > best_size:
            best_size = sz
            best_idx = i
    for i, r in enumerate(runs):
        if i == best_idx:
            r.text = new_text
        else:
            r.text = ''


def _clear_para(para):
    for r in para.runs: r.text = ''


def _remove_para(para):
    para._p.getparent().remove(para._p)


def _insert_jinja_para_before(para, text):
    new_p = OxmlElement('w:p')
    new_r = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set(qn('xml:space'), 'preserve')
    if para.runs and para.runs[0]._r.rPr is not None:
        new_r.append(copy.deepcopy(para.runs[0]._r.rPr))
    new_r.append(new_t)
    new_p.append(new_r)
    para._p.addprevious(new_p)


def _insert_jinja_para_after(para, text):
    new_p = OxmlElement('w:p')
    new_r = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set(qn('xml:space'), 'preserve')
    if para.runs and para.runs[0]._r.rPr is not None:
        new_r.append(copy.deepcopy(para.runs[0]._r.rPr))
    new_r.append(new_t)
    new_p.append(new_r)
    para._p.addnext(new_p)


def _verify(docx_path):
    doc = Document(docx_path)
    print("\n=== 模版验证 ===")
    expected = [
        'title_zh', 'title_en', 'date',
        'college', 'name', 'student_id', 'advisor',
        'abstract_zh', 'keywords_zh', 'abstract_en', 'keywords_en',
        'ch.title', 'sec.title', 'sub.title',
        'conclusion', 'acknowledgement', 'ref',
    ]
    found = set()
    for p in doc.paragraphs:
        for var in expected:
            if '{{ ' + var + ' }}' in p.text:
                found.add(var)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for var in expected:
                        if '{{ ' + var + ' }}' in p.text:
                            found.add(var)
    for var in expected:
        status = '✓' if var in found else '✗ MISSING'
        print(f"  {status} {{{{{var}}}}}")
    # 检查残留
    dirty = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if text and '×' in text:
            dirty += 1
            print(f"  残留: {text[:60]}")
    if dirty == 0:
        print("  ✓ 无残留")


if __name__ == '__main__':
    main()
