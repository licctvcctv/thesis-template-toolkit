"""
Convert an already-filled thesis body document into a project docxtpl template.

Usage:
    python make.py <source.docx> [output.docx]

The output follows this repository's thesis template contract:
meta fields + chapters/sections/subsections + conclusion + acknowledgement
+ references.
"""
import copy
import os
import re
import sys

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def main():
    if len(sys.argv) < 2:
        print("Usage: python make.py <source.docx> [output.docx]")
        sys.exit(1)

    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else "template.docx"
    doc = Document(src)

    anchors = _find_anchors(doc)
    _replace_toc(doc, anchors)
    anchors = _find_anchors(doc)
    _replace_front_matter(doc, anchors)
    anchors = _find_anchors(doc)
    _replace_body(doc, anchors)
    _replace_tail(doc)
    _clear_headers(doc)
    _remove_nonessential_notes(doc)
    _remove_all_tables(doc)
    _set_page_break_for_chapter_loop(doc)

    os.makedirs(os.path.dirname(os.path.abspath(out)) or ".", exist_ok=True)
    doc.save(out)
    _mark_fields_update_on_open(out)
    print(f"Saved: {out}")
    _verify(out)


def _find_anchors(doc):
    anchors = {}
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().replace("\t", " ")
        norm = text.replace(" ", "").replace("\u3000", "")
        if norm == "摘要" and "abstract_zh_title" not in anchors:
            anchors["abstract_zh_title"] = i
        elif norm == "Abstract" and "abstract_en_title" not in anchors:
            anchors["abstract_en_title"] = i
        elif re.match(r"^前言$", norm) and "body_start" not in anchors:
            anchors["body_start"] = i
        elif norm == "参考文献" and "refs" not in anchors:
            anchors["refs"] = i
        elif norm == "致谢" and "ack" not in anchors:
            anchors["ack"] = i

    required = ["abstract_zh_title", "abstract_en_title", "body_start", "refs", "ack"]
    missing = [k for k in required if k not in anchors]
    if missing:
        raise RuntimeError(f"Missing anchors: {missing}")
    return anchors


def _replace_front_matter(doc, anchors):
    paras = doc.paragraphs

    zh_title = paras[anchors["abstract_zh_title"]]
    en_title = paras[anchors["abstract_en_title"]]
    refs = paras[anchors["refs"]]
    ack = paras[anchors["ack"]]

    # Chinese abstract: keep title paragraph, replace content up to English title.
    _replace_range_with_blocks(
        doc,
        zh_title,
        en_title,
        [
            _clone_with_text(paras[anchors["abstract_zh_title"] + 1], "{{ abstract_zh }}"),
            _clone_with_text(paras[anchors["abstract_zh_title"] + 3], "关键词：{{ keywords_zh }}"),
        ],
    )

    # English abstract: replace content up to body start.
    paras = doc.paragraphs
    anchors = _find_anchors(doc)
    en_title = paras[anchors["abstract_en_title"]]
    body_start = paras[anchors["body_start"]]
    _replace_range_with_blocks(
        doc,
        en_title,
        body_start,
        [
            _clone_with_text(paras[anchors["abstract_en_title"] + 1], "{{ abstract_en }}"),
            _clone_with_text(paras[anchors["abstract_en_title"] + 3], "Keywords: {{ keywords_en }}"),
        ],
    )


def _replace_toc(doc, anchors):
    paras = doc.paragraphs
    toc_title = paras[0]
    zh_title = paras[anchors["abstract_zh_title"]]
    toc1 = paras[1]
    toc2 = paras[4]
    toc3 = paras[5]
    blocks = [
        _clone_with_text(toc1, "{%p for ch in chapters %}"),
        _clone_with_text(toc1, "{{ ch.title }}"),
        _clone_with_text(toc2, "{%p for sec in ch.sections %}"),
        _clone_with_text(toc2, "{{ sec.title }}"),
        _clone_with_text(toc3, "{%p for sub in sec.subsections %}"),
        _clone_with_text(toc3, "{{ sub.title }}"),
        _clone_with_text(toc3, "{%p endfor %}"),
        _clone_with_text(toc2, "{%p endfor %}"),
        _clone_with_text(toc1, "{%p endfor %}"),
    ]
    _replace_range_with_blocks(doc, toc_title, zh_title, blocks)


def _replace_body(doc, anchors):
    paras = doc.paragraphs
    p_h1 = paras[anchors["body_start"]]
    p_h2 = _first_after(doc, anchors["body_start"], "Heading 2")
    p_h3 = _first_after(doc, anchors["body_start"], "Heading 3")
    p_body = _first_after(doc, anchors["body_start"], "Normal")

    p_refs = paras[anchors["refs"]]
    blocks = [
        _clone_with_text(p_body, "{%p for ch in chapters %}"),
        _clone_with_text(p_h1, "{{ ch.title }}", drop_style=True),
        _clone_with_text(p_body, "{%p for item in ch.content %}"),
        _clone_with_text(p_body, "{{ item }}"),
        _clone_with_text(p_body, "{%p endfor %}"),
        _clone_with_text(p_body, "{%p for sec in ch.sections %}"),
        _clone_with_text(p_h2, "{{ sec.title }}", drop_style=True),
        _clone_with_text(p_body, "{%p for item in sec.content %}"),
        _clone_with_text(p_body, "{{ item }}"),
        _clone_with_text(p_body, "{%p endfor %}"),
        _clone_with_text(p_body, "{%p for sub in sec.subsections %}"),
        _clone_with_text(p_h3, "{{ sub.title }}", drop_style=True),
        _clone_with_text(p_body, "{%p for item in sub.content %}"),
        _clone_with_text(p_body, "{{ item }}"),
        _clone_with_text(p_body, "{%p endfor %}"),
        _clone_with_text(p_body, "{%p endfor %}"),
        _clone_with_text(p_body, "{%p endfor %}"),
        _clone_with_text(p_body, "{%p endfor %}"),
        _clone_with_text(p_h1, "结论", drop_style=True),
        _clone_with_text(p_body, "{%p if conclusion %}"),
        _clone_with_text(p_body, "{{ conclusion }}"),
        _clone_with_text(p_body, "{%p endif %}"),
        _clone_with_text(p_body, "{%p for para in conclusion_list %}"),
        _clone_with_text(p_body, "{{ para }}"),
        _clone_with_text(p_body, "{%p endfor %}"),
    ]
    _replace_range_with_blocks(doc, p_h1, p_refs, blocks, include_start=True)


def _replace_tail(doc):
    paras = doc.paragraphs
    refs = next(p for p in paras if p.text.strip() == "参考文献")
    ack = next(p for p in paras if p.text.strip() == "致谢")
    ref_sample = _first_between(doc, refs, ack) or _first_style(doc, "endnote text") or _first_style(doc, "Normal")

    _replace_range_with_blocks(
        doc,
        refs,
        ack,
        [
            _clone_with_text(ref_sample, "{%p for ref in references %}"),
            _clone_with_text(ref_sample, "{{ ref }}"),
            _clone_with_text(ref_sample, "{%p endfor %}"),
        ],
    )

    paras = doc.paragraphs
    ack = next(p for p in paras if p.text.strip() == "致谢")
    ack_sample = _first_after_index(doc, _para_index(doc, ack), "Normal") or _first_style(doc, "Normal")
    # Replace everything after acknowledgement with one acknowledgement placeholder.
    body = doc.element.body
    children = list(body)
    start = children.index(ack._p) + 1
    end = len(children)
    tail_sect = children[-1] if children and children[-1].tag == qn("w:sectPr") else None
    insert_at = start
    for elem in children[start:end]:
        if elem is tail_sect:
            continue
        body.remove(elem)
    body.insert(insert_at, _clone_with_text(ack_sample, "{{ acknowledgement }}"))


def _replace_range_with_blocks(doc, start_para, stop_para, blocks, include_start=False):
    """Replace content between two paragraphs; keep stop paragraph.

    By default the start paragraph is kept. When include_start=True, the
    start paragraph is replaced as part of the range.
    """
    body = doc.element.body
    children = list(body)
    start = children.index(start_para._p) + (0 if include_start else 1)
    stop = children.index(stop_para._p)
    for elem in children[start:stop]:
        body.remove(elem)
    for offset, elem in enumerate(blocks):
        body.insert(start + offset, elem)


def _clone_with_text(sample_para, text, drop_style=False):
    elem = OxmlElement("w:p")
    ppr = sample_para._p.pPr
    if ppr is not None:
        ppr2 = copy.deepcopy(ppr)
        for tag in ["w:numPr", "w:pageBreakBefore"]:
            node = ppr2.find(qn(tag))
            if node is not None:
                ppr2.remove(node)
        if drop_style:
            pstyle = ppr2.find(qn("w:pStyle"))
            if pstyle is not None:
                ppr2.remove(pstyle)
        elem.append(ppr2)

    run = OxmlElement("w:r")
    if sample_para.runs:
        rpr = sample_para.runs[0]._r.rPr
        if rpr is not None:
            run.append(copy.deepcopy(rpr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    elem.append(run)
    return elem


def _replace_para_text(p, text):
    if not p.runs:
        p.add_run("")
    p.runs[0].text = text
    for r in p.runs[1:]:
        r.text = ""


def _first_after(doc, idx, style_name):
    for p in doc.paragraphs[idx + 1:]:
        if p.style and p.style.name == style_name and p.text.strip():
            return p
    raise RuntimeError(f"No paragraph with style {style_name} after {idx}")


def _first_after_index(doc, idx, style_name):
    for p in doc.paragraphs[idx + 1:]:
        if p.style and p.style.name == style_name and p.text.strip():
            return p
    return None


def _first_style(doc, style_name):
    for p in doc.paragraphs:
        if p.style and p.style.name == style_name and p.text.strip():
            return p
    return None


def _first_between(doc, start_para, stop_para):
    paras = doc.paragraphs
    start = _para_index(doc, start_para)
    stop = _para_index(doc, stop_para)
    for p in paras[start + 1:stop]:
        if p.text.strip():
            return p
    return None


def _para_index(doc, target_para):
    for i, p in enumerate(doc.paragraphs):
        if p._p is target_para._p:
            return i
    raise ValueError("paragraph not found")


def _remove_all_tables(doc):
    for tbl in list(doc.element.body.findall(qn("w:tbl"))):
        parent = tbl.getparent()
        if parent is not None:
            parent.remove(tbl)


def _clear_headers(doc):
    for section in doc.sections:
        section.header.is_linked_to_previous = False
        for p in section.header.paragraphs:
            _replace_para_text(p, "")


def _remove_nonessential_notes(doc):
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if "灯泡注释" in style:
            _replace_para_text(p, "")


def _set_page_break_for_chapter_loop(doc):
    title_paras = [p for p in doc.paragraphs if p.text.strip() == "{{ ch.title }}"]
    if title_paras:
        p = title_paras[-1]
        pPr = p._p.get_or_add_pPr()
        if pPr.find(qn("w:pageBreakBefore")) is None:
            pPr.append(OxmlElement("w:pageBreakBefore"))
    for title in ["结论", "参考文献", "致谢"]:
        for p in doc.paragraphs:
            if p.text.strip() == title:
                p.paragraph_format.page_break_before = True
                break


def _verify(path):
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    required = [
        "{{ abstract_zh }}",
        "{{ keywords_zh }}",
        "{{ abstract_en }}",
        "{{ keywords_en }}",
        "{%p for ch in chapters %}",
        "{{ ch.title }}",
        "{%p for sec in ch.sections %}",
        "{{ sec.title }}",
        "{%p for sub in sec.subsections %}",
        "{{ sub.title }}",
        "{%p for ref in references %}",
        "{{ ref }}",
        "{{ acknowledgement }}",
    ]
    missing = [tag for tag in required if tag not in text]
    if missing:
        raise RuntimeError(f"Template verification failed, missing: {missing}")
    print("Verification passed")


def _mark_fields_update_on_open(path):
    import shutil
    import tempfile
    import zipfile
    from lxml import etree

    w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    tmp = tempfile.mktemp(suffix=".docx")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml":
                root = etree.fromstring(data)
                update = root.find(f"{{{w}}}updateFields")
                if update is None:
                    update = etree.Element(f"{{{w}}}updateFields")
                    root.insert(0, update)
                update.set(f"{{{w}}}val", "true")
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            elif item.filename == "word/document.xml":
                root = etree.fromstring(data)
                for fld in root.findall(f".//{{{w}}}fldChar"):
                    fld.set(f"{{{w}}}dirty", "true")
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(item, data)
    shutil.move(tmp, path)


if __name__ == "__main__":
    main()
