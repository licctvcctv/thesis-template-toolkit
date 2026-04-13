"""
重庆移通学院 毕业论文模版制作器。
将原始论文 .docx 转换为 docxtpl 模版。

用法: python make.py <原始.docx> [输出模版.docx]

封面结构: 纯段落+下划线填充（非表格）
  p6:  2026届毕业论文（设计）
  p8:  论文（设计）题目：__xxxx__
  p11: 学        院：____xxxx__
  p12: 专        业：____xxxx__
  p13: 班        级：____xxxx__
  p14: 学 生  姓 名：____xxxx__
  p15: 学 生  学 号：____xxxx__
  p16: 指 导  教 师：____xxxx__
  p17: 答辩组负责人：____xxxx__

摘要: p40 "摘    要", p42 正文, p44 关键词
英文摘要: p53 "ABSTRACT", p55 正文, p57 关键词
正文: Heading 1/2/3, 从 p134 开始
致谢: p359 "致    谢"
参考文献: p368 "参考文献", p370+ 各条
"""
import os
import sys
import re
import copy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def main():
    if len(sys.argv) < 2:
        print("用法: python make.py <原始.docx> [输出.docx]")
        sys.exit(1)

    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else "template.docx"

    doc = Document(src)
    paras = doc.paragraphs

    # ========== Step 1: 封面字段替换 ==========
    print("Step 1: 封面字段替换...")
    cover_fields = [
        (8,  '题目', 'title_zh'),
        (11, '学院', 'college'),
        (12, '专业', 'major'),
        (13, '班级', 'class_name'),
        (14, '姓名', 'name'),
        (15, '学号', 'student_id'),
        (16, '导师', 'advisor'),
    ]
    for pi, desc, var in cover_fields:
        p = paras[pi]
        # 找到冒号后的 run，替换值部分
        found_colon = False
        for r in p.runs:
            if '：' in r.text:
                found_colon = True
                # 冒号后面可能有值，保留冒号
                idx = r.text.index('：')
                r.text = r.text[:idx + 1]
                continue
            if found_colon:
                # 这些是下划线+值的 run
                if any(c not in '_  \t' for c in r.text):
                    # 含有实际内容的 run，替换为变量
                    r.text = '{{ ' + var + ' }}'
                    found_colon = False  # 后续 run 清空
                else:
                    # 纯下划线/空格 run，保留
                    pass
        print(f"  p{pi}: {desc} -> {{{{{var}}}}}")

    # ========== Step 2: 中文摘要 ==========
    print("Step 2: 中文摘要...")
    # p42 是摘要正文 — 可能跨多段到 p43
    # 找到摘要正文段落（p40 是标题，往后找到关键词前的段落）
    abstract_start = None
    keywords_zh_idx = None
    for i in range(40, 50):
        text = paras[i].text.strip()
        if text and '摘' not in text and '关键词' not in text and abstract_start is None:
            abstract_start = i
        if '关键词' in text:
            keywords_zh_idx = i
            break

    if abstract_start:
        # 保留第一段，替换内容
        _replace_para_text(paras[abstract_start], '{{ abstract_zh }}')
        # 清空中间段落（如果有多段摘要）
        for i in range(abstract_start + 1, keywords_zh_idx):
            _clear_para(paras[i])
        print(f"  p{abstract_start}: abstract_zh")

    if keywords_zh_idx:
        # 关键词格式：关键词：xxx；xxx → 关键词：{{ keywords_zh }}
        p = paras[keywords_zh_idx]
        found_label = False
        set_var = False
        for r in p.runs:
            if '关键词' in r.text:
                # 保留"关键词："，清掉后面的值
                idx = r.text.index('：') + 1 if '：' in r.text else len(r.text)
                r.text = r.text[:idx]
                found_label = True
                continue
            if found_label and not set_var:
                r.text = '{{ keywords_zh }}'
                set_var = True
                continue
            if set_var:
                r.text = ''
        print(f"  p{keywords_zh_idx}: keywords_zh")

    # ========== Step 3: 英文摘要 ==========
    print("Step 3: 英文摘要...")
    abstract_en_start = None
    keywords_en_idx = None
    for i in range(52, 60):
        text = paras[i].text.strip()
        if text and 'ABSTRACT' not in text and 'Key' not in text and abstract_en_start is None:
            abstract_en_start = i
        if text.startswith('Key'):
            keywords_en_idx = i
            break

    if abstract_en_start:
        _replace_para_text(paras[abstract_en_start], '{{ abstract_en }}')
        for i in range(abstract_en_start + 1, keywords_en_idx):
            _clear_para(paras[i])
        print(f"  p{abstract_en_start}: abstract_en")

    if keywords_en_idx:
        p = paras[keywords_en_idx]
        found_label = False
        set_var = False
        for r in p.runs:
            if 'Key' in r.text:
                idx = r.text.index(':') + 1 if ':' in r.text else len(r.text)
                r.text = r.text[:idx] + ' '
                found_label = True
                continue
            if found_label and not set_var:
                r.text = '{{ keywords_en }}'
                set_var = True
                continue
            if set_var:
                r.text = ''
        print(f"  p{keywords_en_idx}: keywords_en")

    # ========== Step 4: 找到正文区域 ==========
    print("Step 4: 定位正文区域...")
    body_start = None
    body_end = None
    ack_idx = None
    refs_idx = None

    for i, p in enumerate(paras):
        style = p.style.name if p.style else ''
        text = p.text.strip()
        if style == 'Heading 1' and body_start is None:
            body_start = i
        if '致' in text and '谢' in text and style == 'Normal' and i > 300:
            ack_idx = i
        if text == '参考文献' and style == 'Normal' and i > 300:
            refs_idx = i

    # body_end = last Heading or content paragraph before acknowledgment
    if ack_idx:
        body_end = ack_idx - 1
        # skip empty paragraphs, but don't go past致谢 title
        while body_end > body_start:
            text = paras[body_end].text.strip()
            if text and '致' not in text:
                break
            if '致' in text and '谢' in text:
                body_end -= 1  # 致谢标题不算正文
                break
            body_end -= 1

    print(f"  body: p{body_start} - p{body_end}")
    print(f"  ack: p{ack_idx}")
    print(f"  refs: p{refs_idx}")

    # ========== Step 5: 设置正文 Jinja2 循环 ==========
    print("Step 5: 正文循环...")

    # 保存样本段落（各 Heading 和 Normal 的第一个）
    sample_h1 = None
    sample_h2 = None
    sample_h3 = None
    sample_normal = None

    for i in range(body_start, body_end + 1):
        p = paras[i]
        style = p.style.name if p.style else ''
        if style == 'Heading 1' and sample_h1 is None:
            sample_h1 = i
        elif style == 'Heading 2' and sample_h2 is None:
            sample_h2 = i
        elif style == 'Heading 3' and sample_h3 is None:
            sample_h3 = i
        elif style == 'Normal' and p.text.strip() and sample_normal is None:
            sample_normal = i

    print(f"  samples: H1=p{sample_h1}, H2=p{sample_h2}, H3=p{sample_h3}, Normal=p{sample_normal}")

    # 清空正文区域（除了样本段落），从后往前
    keep = {sample_h1, sample_h2, sample_h3, sample_normal}
    for i in range(body_end, body_start - 1, -1):
        if i not in keep:
            _remove_para(paras[i])

    # 删除正文区域内残留的表格
    # 表格在 XML 中是 w:tbl 元素，不属于段落，需要单独清理
    body_xml = doc.element.body
    # 找到正文循环的位置范围（在第一个 Heading 附近）
    # 先定位 keep 的段落 XML 元素
    keep_elements = set()
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ''
        if style.startswith('Heading') and p.text.strip():
            keep_elements.add(p._p)

    # 遍历 body 子元素，删除位于正文区域内的 w:tbl
    # 正文区域 = Heading 1 样本到致谢之间
    in_body = False
    tables_to_remove = []
    for elem in body_xml:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            # 检查是否是 Heading 1 样本
            texts = [t.text for t in elem.findall('.//{%s}t' % WNS) if t.text]
            text = ''.join(texts).strip()
            if elem in keep_elements and not in_body:
                in_body = True
            if '致' in text and '谢' in text and len(text) < 10:
                in_body = False
        elif tag == 'tbl' and in_body:
            tables_to_remove.append(elem)

    for tbl in tables_to_remove:
        body_xml.remove(tbl)
    print(f"  删除 {len(tables_to_remove)} 个残留表格")

    # 现在重新定位样本段落（因为删除了其他段落）
    # 重新扫描
    paras = doc.paragraphs  # refresh
    h1_idx = h2_idx = h3_idx = norm_idx = None
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ''
        text = p.text.strip()
        if style == 'Heading 1' and text and h1_idx is None:
            h1_idx = i
        elif style == 'Heading 2' and text and h2_idx is None:
            h2_idx = i
        elif style == 'Heading 3' and text and h3_idx is None:
            h3_idx = i

    # 在 Heading 1 前插入章节循环开始标记
    # 用 H1 做章标题，H2 做节标题，H3 做小节标题
    # Normal 段落做正文内容

    # 方案：在 H1 样本前插入 for 循环
    p_h1 = None
    p_h2 = None
    p_h3 = None
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ''
        if style == 'Heading 1' and p.text.strip() and p_h1 is None:
            p_h1 = p
        elif style == 'Heading 2' and p.text.strip() and p_h2 is None:
            p_h2 = p
        elif style == 'Heading 3' and p.text.strip() and p_h3 is None:
            p_h3 = p

    # 插入 Jinja2 循环结构
    # {%p for ch in chapters %}
    # H1: {{ ch.title }}
    # {%p for sec in ch.sections %}
    # H2: {{ sec.title }}
    # {%p for item in sec.content %}{{ item }}{%p endfor %}
    # {%p for sub in sec.subsections %}
    # H3: {{ sub.title }}
    # {%p for item in sub.content %}{{ item }}{%p endfor %}
    # {%p endfor %}
    # {%p endfor %}
    # {%p endfor %}

    # 在 H1 前插入 for-ch
    _insert_jinja_para_before(p_h1, '{%p for ch in chapters %}')
    # H1 内容改为 {{ ch.title }}
    _replace_para_text(p_h1, '{{ ch.title }}')
    # H1 后插入 for-sec
    _insert_jinja_para_after(p_h1, '{%p for sec in ch.sections %}')

    # H2 内容改为 {{ sec.title }}
    _replace_para_text(p_h2, '{{ sec.title }}')
    # H2 后插入内容循环（三个段落：for / item / endfor）
    _insert_jinja_para_after(p_h2, '{%p for item in sec.content %}')
    paras = doc.paragraphs
    for p in paras:
        if p.text.strip() == '{%p for item in sec.content %}':
            _insert_jinja_para_after(p, '{{ item }}')
            break
    paras = doc.paragraphs
    for p in paras:
        if p.text.strip() == '{{ item }}':
            _insert_jinja_para_after(p, '{%p endfor %}')
            break
    # 然后插入 for-sub
    paras = doc.paragraphs
    # 找第一个 endfor（属于 sec.content 的）
    first_endfor = None
    for p in paras:
        if p.text.strip() == '{%p endfor %}':
            first_endfor = p
            break
    _insert_jinja_para_after(first_endfor, '{%p for sub in sec.subsections %}')

    # H3 内容改为 {{ sub.title }}
    _replace_para_text(p_h3, '{{ sub.title }}')
    # H3 后插入子内容循环（三个段落）
    _insert_jinja_para_after(p_h3, '{%p for item in sub.content %}')
    paras = doc.paragraphs
    for p in paras:
        if p.text.strip() == '{%p for item in sub.content %}':
            _insert_jinja_para_after(p, '{{ item }}')
            break
    paras = doc.paragraphs
    # 找 sub.content 的 {{ item }}
    found_sub_item = False
    for p in paras:
        if 'for item in sub.content' in p.text:
            found_sub_item = True
            continue
        if found_sub_item and p.text.strip() == '{{ item }}':
            _insert_jinja_para_after(p, '{%p endfor %}')
            break

    # 现在结构是:
    # {%p for ch %}
    #   H1: {{ ch.title }}
    #   {%p for sec %}
    #     H2: {{ sec.title }}
    #     {%p for item in sec.content %}
    #     {{ item }}
    #     {%p endfor %}          ← sec.content endfor
    #     {%p for sub %}
    #       H3: {{ sub.title }}
    #       {%p for item in sub.content %}
    #       {{ item }}
    #       {%p endfor %}        ← sub.content endfor
    #     {%p endfor %}          ← sub endfor (need to add)
    #   {%p endfor %}            ← sec endfor (need to add)
    # {%p endfor %}              ← ch endfor (need to add)

    # 找 sub.content 的 endfor，在其后加 endfor sub
    paras = doc.paragraphs
    in_sub_content = False
    sub_content_endfor = None
    for p in paras:
        if 'for item in sub.content' in p.text:
            in_sub_content = True
        if in_sub_content and p.text.strip() == '{%p endfor %}':
            sub_content_endfor = p
            break
    _insert_jinja_para_after(sub_content_endfor, '{%p endfor %}')  # endfor sub

    # 找刚加的 endfor sub，在其后加 endfor sec
    paras = doc.paragraphs
    endfor_list = [p for p in paras if p.text.strip() == '{%p endfor %}']
    # endfor_list: [sec.content, sub.content, sub]
    _insert_jinja_para_after(endfor_list[-1], '{%p endfor %}')  # endfor sec

    # 再加 endfor ch
    paras = doc.paragraphs
    endfor_list = [p for p in paras if p.text.strip() == '{%p endfor %}']
    _insert_jinja_para_after(endfor_list[-1], '{%p endfor %}')  # endfor ch

    # 删除多余的 Normal 样本段落
    paras = doc.paragraphs
    for p in paras:
        style = p.style.name if p.style else ''
        text = p.text.strip()
        if style == 'Normal' and text and 'endfor' not in text and 'for ' not in text and '{{' not in text:
            # 这可能是残留的样本正文
            # 检查是否在 body 区域内（在 H1 附近）
            if p_h1 and p_h2 and p_h3:
                # 检查位置：在 h1 到 endfor 之间
                pass  # 先不管，后面清理

    print("  循环结构已插入")

    # ========== Step 6: 致谢 ==========
    print("Step 6: 致谢...")
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        text = p.text.strip()
        style = p.style.name if p.style else ''
        # 跳过 TOC 条目，只找正文区域的致谢标题
        if '致' in text and '谢' in text and len(text) < 10 and 'toc' not in style.lower():
            # 下一个非空段落就是致谢内容
            for j in range(i + 1, min(i + 10, len(paras))):
                if paras[j].text.strip():
                    _replace_para_text(paras[j], '{{ acknowledgement }}')
                    # 清空后续致谢段落
                    for k in range(j + 1, len(paras)):
                        if paras[k].text.strip() == '参考文献':
                            break
                        if paras[k].text.strip():
                            _clear_para(paras[k])
                    print(f"  p{j}: acknowledgement")
                    break
            break

    # ========== Step 7: 参考文献 ==========
    print("Step 7: 参考文献...")
    paras = doc.paragraphs
    refs_start = None
    for i, p in enumerate(paras):
        if p.text.strip() == '参考文献':
            refs_start = i
            break

    if refs_start:
        # 找到第一条参考文献
        first_ref = None
        for i in range(refs_start + 1, len(paras)):
            text = paras[i].text.strip()
            if text:
                first_ref = i
                break

        if first_ref:
            # 在第一条参考文献前插入 for 循环
            _insert_jinja_para_before(paras[first_ref], '{%p for ref in references %}')
            # 重新获取段落列表
            paras = doc.paragraphs
            # 找到第一条参考文献（现在位置偏移了1）
            for i, p in enumerate(paras):
                if p.text.strip() == '{%p for ref in references %}':
                    # 下一段就是第一条引用
                    ref_para = paras[i + 1]
                    _replace_para_text(ref_para, '{{ ref }}')
                    # 删除其余参考文献
                    paras = doc.paragraphs
                    in_refs = False
                    for j, p in enumerate(paras):
                        if '{{ ref }}' in p.text:
                            in_refs = True
                            continue
                        if in_refs:
                            if p.text.strip():
                                _remove_para(p)
                    # 在 {{ ref }} 后插入 endfor
                    paras = doc.paragraphs
                    for p in paras:
                        if '{{ ref }}' in p.text:
                            _insert_jinja_para_after(p, '{%p endfor %}')
                            break
                    break
        print(f"  refs loop set up")

    # ========== Step 8: 清理残留正文 ==========
    print("Step 8: 清理残留正文...")
    paras = doc.paragraphs
    in_body_loop = False
    body_loop_done = False
    endfor_count = 0
    for p in paras:
        text = p.text.strip()
        if 'for ch in chapters' in text:
            in_body_loop = True
            continue
        if in_body_loop and not body_loop_done:
            if text == '{%p endfor %}':
                endfor_count += 1
                if endfor_count >= 3:
                    body_loop_done = True
                continue
            style = p.style.name if p.style else ''
            # 保留 Heading 和 Jinja2 标记，删除其他
            if style in ('Heading 1', 'Heading 2', 'Heading 3'):
                continue
            if '{%' in text or '{{' in text:
                continue
            if text:
                _clear_para(p)

    # ========== 保存 ==========
    doc.save(out)
    print(f"\n完成: {out}")


def _replace_para_text(para, new_text):
    """替换段落文本，保留第一个 run 的格式"""
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ''
    else:
        para.text = new_text


def _clear_para(para):
    """清空段落文本"""
    for r in para.runs:
        r.text = ''


def _remove_para(para):
    """从文档中删除段落"""
    p = para._p
    p.getparent().remove(p)


def _insert_jinja_para_before(para, text):
    """在段落前插入一个 Jinja2 标记段落"""
    new_p = OxmlElement('w:p')
    new_r = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set(qn('xml:space'), 'preserve')
    # 复制 rPr（字体等）
    if para.runs and para.runs[0]._r.rPr is not None:
        new_r.append(copy.deepcopy(para.runs[0]._r.rPr))
    new_r.append(new_t)
    new_p.append(new_r)
    para._p.addprevious(new_p)


def _insert_jinja_para_after(para, text):
    """在段落后插入一个 Jinja2 标记段落"""
    new_p = OxmlElement('w:p')
    new_r = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set(qn('xml:space'), 'preserve')
    if para.runs and para.runs[0]._r.rPr is not None:
        new_r.append(copy.deepcopy(para.runs[0]._r.rPr))
    new_r.append(new_t)
    new_p.append(new_r)
    para._p.addnext(new_p)


if __name__ == '__main__':
    main()
