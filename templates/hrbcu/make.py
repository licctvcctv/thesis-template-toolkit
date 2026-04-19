"""
哈尔滨商业大学本科毕业设计（论文）模板制作。
用法: python make.py <原始模板.docx> [输出路径]

模板特点:
- 中文封面: P0-23, 标题P5, 姓名P14/导师P15-16/专业P17/学号P18/学院P19, 日期P23
- 英文封面: P26-44, 标题P31, Table0(5行2列: Student/Supervisor/Specialty/ID/School), 日期P44
- 中文摘要: P45 "摘  要"(Heading1), P46-49正文(Plain Text/Normal), P51关键词
- 英文摘要: P58 "Abstract"(Heading1), P59-61正文, P63 Keywords
- 目录: P69-108, toc 1/2/3
- 正文: P116起, Heading 1/2/3 + Plain Text/Normal
- 结论: P238 "结  论"(Heading1)
- 参考文献: P247 (Heading1), P252起 [1][2]条目
- 致谢: P263 (Heading1), P264正文
- 附录: P268起
"""
import os
import sys
import re
import copy
import zipfile
import shutil

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../.."))
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def make(src_path, out_path):
    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc = Document(src_path)
    paras = doc.paragraphs

    # ========== Step 1: 中文封面 ==========
    # P5: 论文题目
    p5 = paras[5]
    if p5.runs:
        p5.runs[0].text = "{{ title_zh }}"
        for r in p5.runs[1:]:
            r.text = ""

    # P14-19: 姓名/导师/专业/学号/学院
    # runs[0]=label, runs[1]=underlined spaces (value area)
    cover_fields = {
        14: "{{ name }}",           # 学 生 姓 名：
        15: "{{ advisor }}",        # 指 导 教 师：（校内）
        17: "{{ major_class }}",    # 专 业 班 级：
        18: "{{ student_id }}",     # 学       号：
        19: "{{ college }}",        # 学       院：
    }
    for idx, var in cover_fields.items():
        p = paras[idx]
        if len(p.runs) >= 2:
            p.runs[1].text = var
            for r in p.runs[2:]:
                r.text = ""

    # P16: 校外导师（可选），清空即可
    p16 = paras[16]
    for r in p16.runs:
        r.text = ""

    # P23: 日期
    p23 = paras[23]
    if p23.runs:
        p23.runs[0].text = "{{ finish_date }}"
        for r in p23.runs[1:]:
            r.text = ""

    print("  Step 1: 中文封面")

    # ========== Step 2: 英文封面 ==========
    # P31: English title
    p31 = paras[31]
    if p31.runs:
        p31.runs[0].text = "{{ title_en }}"
        for r in p31.runs[1:]:
            r.text = ""

    # Table 0: English cover fields
    tbl = doc.tables[0]
    en_fields = {
        0: "{{ name_en }}",           # Student
        1: "{{ advisor_en }}",        # Supervisor (On-campus)
        2: "{{ major_class_en }}",    # Specialty and Class
        3: "{{ student_id }}",        # Student ID
        4: "{{ college_en }}",        # School
    }
    for row_idx, var in en_fields.items():
        cell = tbl.rows[row_idx].cells[1]
        for p in cell.paragraphs:
            if p.runs:
                p.runs[0].text = var
                for r in p.runs[1:]:
                    r.text = ""

    # P44: English date
    p44 = paras[44]
    if p44.runs:
        p44.runs[0].text = "{{ finish_date_en }}"
        for r in p44.runs[1:]:
            r.text = ""

    print("  Step 2: 英文封面")

    # ========== Step 3: 清除注释段落 ==========
    # 这是模板撰写规范文件，有很多"注释："开头的说明段落需要清除
    note_indices = []
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        if text.startswith("注释") or text.startswith("注释：") or text.startswith("格式要求见"):
            note_indices.append(i)
            for r in p.runs:
                r.text = ""
    # Also clear P8 (title format note), P11 (document structure note)
    for idx in [8, 11, 34]:
        for r in paras[idx].runs:
            r.text = ""

    print(f"  Step 3: 清除 {len(note_indices) + 3} 个注释段落")

    # ========== Step 4: 中文摘要 ==========
    # P46: 第一个摘要正文段落 → 循环
    p46 = paras[46]
    p46_p = p46._p
    p46_p.addprevious(_mk_ctrl("{%p for abs_p in abstract_zh_list %}"))
    if p46.runs:
        p46.runs[0].text = "{{ abs_p }}"
        for r in p46.runs[1:]:
            r.text = ""
    p46_p.addnext(_mk_ctrl("{%p endfor %}"))
    # Clear remaining abstract paragraphs (P47-50)
    for idx in range(47, 51):
        for r in paras[idx].runs:
            r.text = ""

    # P51: 关键词 → runs[0]="关键词", runs[1]="：", runs[2]=values
    p51 = paras[51]
    if p51.runs and len(p51.runs) >= 3:
        p51.runs[2].text = "{{ keywords_zh }}"
        for r in p51.runs[3:]:
            r.text = ""

    print("  Step 4: 中文摘要")

    # ========== Step 5: 英文摘要 ==========
    # P59: English abstract body → loop
    p59 = paras[59]
    p59_p = p59._p
    p59_p.addprevious(_mk_ctrl("{%p for abs_p in abstract_en_list %}"))
    if p59.runs:
        p59.runs[0].text = "{{ abs_p }}"
        for r in p59.runs[1:]:
            r.text = ""
    p59_p.addnext(_mk_ctrl("{%p endfor %}"))
    # Clear P60-62
    for idx in range(60, 63):
        for r in paras[idx].runs:
            r.text = ""

    # P63: Keywords → runs[0]="Keywords", runs[1]=": ", runs[2+]=values
    p63 = paras[63]
    if p63.runs and len(p63.runs) >= 3:
        p63.runs[2].text = "{{ keywords_en }}"
        for r in p63.runs[3:]:
            r.text = ""

    # Clear note after keywords (P64-65)
    for idx in range(64, 68):
        if idx < len(paras):
            for r in paras[idx].runs:
                r.text = ""

    # 清除关键词说明段落 (P53-56区域)
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        if "关键词是供检索用的主题词条" in text or "关键词应选取3-6个" in text:
            for r in p.runs:
                r.text = ""
        if "英文摘要中的关键词小写" in text:
            for r in p.runs:
                r.text = ""

    print("  Step 5: 英文摘要")

    # ========== Step 5.5: 结论 ==========
    concl_idx = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and "结" in text and "论" in text and not re.match(r'^\d', text):
            concl_idx = i
            break

    if concl_idx:
        for j in range(concl_idx + 1, min(concl_idx + 20, len(paras))):
            pj = paras[j]
            text = (pj.text or "").strip()
            style = pj.style.name if pj.style else ""
            if "Heading" in style:
                break
            if pj.runs and text:
                pj_p = pj._p
                pj_p.addprevious(_mk_ctrl("{%p for cp in conclusion_list %}"))
                pj.runs[0].text = "{{ cp }}"
                for r in pj.runs[1:]:
                    r.text = ""
                pj_p.addnext(_mk_ctrl("{%p endfor %}"))
                for k in range(j + 1, min(j + 20, len(paras))):
                    pk = paras[k]
                    pk_style = pk.style.name if pk.style else ""
                    if "Heading" in pk_style:
                        break
                    for r in pk.runs:
                        r.text = ""
                break
        print("  Step 5.5: 结论")

    # ========== Step 6: 致谢 ==========
    ack_idx = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and "致" in text and "谢" in text:
            ack_idx = i
            break

    if ack_idx:
        for j in range(ack_idx + 1, min(ack_idx + 20, len(paras))):
            pj = paras[j]
            text = (pj.text or "").strip()
            style = pj.style.name if pj.style else ""
            if "Heading" in style:
                break
            if pj.runs and text:
                pj.runs[0].text = "{{ acknowledgement }}"
                for r in pj.runs[1:]:
                    r.text = ""
                for k in range(j + 1, min(j + 20, len(paras))):
                    pk = paras[k]
                    pk_style = pk.style.name if pk.style else ""
                    if "Heading" in pk_style:
                        break
                    for r in pk.runs:
                        r.text = ""
                break

    print("  Step 6: 致谢")

    # ========== Step 7: 参考文献 ==========
    ref_idx = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and "参考文献" in text:
            ref_idx = i
            break

    if ref_idx:
        # Clear description paragraphs between heading and first [N]
        first_ref = None
        for j in range(ref_idx + 1, min(ref_idx + 15, len(paras))):
            pt = (paras[j].text or "").strip()
            if re.match(r'\[\d+\]', pt):
                first_ref = j
                break
            # Clear description text
            if pt:
                for r in paras[j].runs:
                    r.text = ""

        # Also clear empty Heading 1 before 参考文献
        if ref_idx > 0:
            prev_p = paras[ref_idx - 1]
            prev_style = prev_p.style.name if prev_p.style else ""
            if prev_style == "Heading 1" and not (prev_p.text or "").strip():
                prev_p.style = doc.styles['Normal']

        if first_ref:
            pj = paras[first_ref]
            ed_p = pj._p
            ed_p.addprevious(_mk_ctrl("{%p for ref in references %}"))
            if pj.runs:
                pj.runs[0].text = "{{ ref }}"
                for r in pj.runs[1:]:
                    r.text = ""
            ed_p.addnext(_mk_ctrl("{%p endfor %}"))

            # Clear all remaining reference content until 致谢
            for k in range(first_ref + 1, len(paras)):
                pk = paras[k]
                pk_style = pk.style.name if pk.style else ""
                if "Heading" in pk_style:
                    break
                for r in pk.runs:
                    r.text = ""

    print("  Step 7: 参考文献")

    # ========== Step 8: 删除附录内容 ==========
    app_idx = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and "附" in text and "录" in text:
            app_idx = i
            break

    if app_idx:
        for k in range(app_idx + 1, len(paras)):
            pk = paras[k]
            for r in pk.runs:
                r.text = ""
            style = pk.style.name if pk.style else ""
            if style.startswith("Heading"):
                pk.style = doc.styles['Normal']
        # 删除附录区域的表格（保留 Table 0 英文封面表）
        body_elem = doc.element.body
        # 找附录起始位置（段落计数）
        app_tables = []
        pc = 0
        for elem in body_elem:
            tag = elem.tag.split('}')[-1]
            if tag == 'p':
                pc += 1
            elif tag == 'tbl' and pc > app_idx:
                app_tables.append(elem)
        for tbl_elem in app_tables:
            body_elem.remove(tbl_elem)
        if app_tables:
            print(f"  Step 8: 清除附录内容 + {len(app_tables)} 个附录表格")
        else:
            print("  Step 8: 清除附录内容")

    # ========== Step 8.5: 清除目录说明文字 ==========
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        if "目录应包括论文中全部章" in text or "目录中章标题用宋体小四号字" in text:
            for r in p.runs:
                r.text = ""

    # ========== Step 8.6: 删除封面和参考文献区域的多余空段落 ==========
    paras = doc.paragraphs  # refresh
    to_remove = []

    # 中文封面：删除P8-P13区域的空段（原注释文字清空后留下的空行）
    for i in range(8, 14):
        if i < len(paras):
            p = paras[i]
            if not (p.text or "").strip() and p._p.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr') is None:
                to_remove.append(p._p)

    # 英文封面：删除P31之后到P44之前的空段（原注释区域）
    # 找英文标题和日期之间的空段
    in_en_cover = False
    for i in range(32, 44):
        if i < len(paras):
            p = paras[i]
            text = (p.text or "").strip()
            has_sect = p._p.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr') is not None
            if not text and not has_sect:
                # 保留最多2个空行（英文标题和表格之间需要一些间距）
                to_remove.append(p._p)

    # 参考文献：删除heading和{%p for之间的空段 + endfor之后到下一个heading之间的空段
    ref_heading_idx = None
    ref_for_idx = None
    ref_endfor_idx = None
    next_heading_after_ref = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and "参考文献" in text:
            ref_heading_idx = i
        if ref_heading_idx and "{%p for ref" in text:
            ref_for_idx = i
        if ref_heading_idx and ref_for_idx and "{%p endfor" in text:
            ref_endfor_idx = i
        if ref_endfor_idx and style == "Heading 1" and i > ref_endfor_idx:
            next_heading_after_ref = i
            break

    if ref_heading_idx and ref_for_idx:
        for i in range(ref_heading_idx + 1, ref_for_idx):
            p = paras[i]
            if not (p.text or "").strip():
                has_sect = p._p.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr') is not None
                if not has_sect:
                    to_remove.append(p._p)

    if ref_endfor_idx and next_heading_after_ref:
        for i in range(ref_endfor_idx + 1, next_heading_after_ref):
            p = paras[i]
            if not (p.text or "").strip():
                has_sect = p._p.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr') is not None
                if not has_sect:
                    to_remove.append(p._p)

    # 通用清理：结论→参考文献→致谢→附录→末尾，删除多余空段落
    # 策略：每个Heading 1之前只保留1个SECT空段，其余空段全删
    paras = doc.paragraphs  # refresh
    _NS_W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    # 找结论的位置（正文循环之后的第一个Heading 1）
    concl_start = None
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        if style == "Heading 1" and "结" in text and "论" in text:
            concl_start = i
            break

    if concl_start:
        # 从结论开始到文档末尾，合并连续SECT段，删除多余空段
        sect_seen_before_heading = False
        i = concl_start
        while i < len(paras):
            p = paras[i]
            text = (p.text or "").strip()
            style = p.style.name if p.style else ""
            has_sect = p._p.find(f'.//{_NS_W}sectPr') is not None

            if "Heading" in style and text:
                # Heading段本身不删
                sect_seen_before_heading = False
                i += 1
                continue

            if text and "{%" not in text:
                # 有内容的段不删
                sect_seen_before_heading = False
                i += 1
                continue

            if not text and has_sect:
                if sect_seen_before_heading:
                    # 已经有一个SECT了，这个多余，删
                    to_remove.append(p._p)
                else:
                    sect_seen_before_heading = True
                i += 1
                continue

            if not text and not has_sect:
                # 空段无SECT，检查是否在内容段和Heading之间
                # 如果前面已经有SECT，这个空段是多余的
                to_remove.append(p._p)
                i += 1
                continue

            i += 1

    removed = 0
    for elem in to_remove:
        if elem.getparent() is not None:
            elem.getparent().remove(elem)
            removed += 1
    if removed:
        print(f"  Step 8.6: 删除 {removed} 个多余空段落")

    doc.save(out_path)

    # ========== Step 9: 正文循环 ==========
    _setup_body(out_path)
    print("  Step 9: 正文循环")

    # ========== Step 10: 删除文本框/AlternateContent + 图片 ==========
    tmp = out_path + ".tmp"
    with zipfile.ZipFile(out_path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w') as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    xml = data.decode('utf-8')
                    n = xml.count('<mc:AlternateContent')
                    if n > 0:
                        xml = re.sub(
                            r'<mc:AlternateContent>.*?</mc:AlternateContent>',
                            '', xml, flags=re.DOTALL)
                        print(f"  Step 10: 删除 {n} 个文本框")
                    # 删除内嵌图片 (w:drawing)
                    n2 = xml.count('<w:drawing>')
                    if n2 > 0:
                        xml = re.sub(
                            r'<w:drawing>.*?</w:drawing>',
                            '', xml, flags=re.DOTALL)
                        print(f"  Step 10: 删除 {n2} 个内嵌图片")
                    data = xml.encode('utf-8')
                zout.writestr(item, data)
    shutil.move(tmp, out_path)


def _setup_body(doc_path):
    """设置正文 Jinja2 循环。Heading 1/2/3 三级标题。"""
    doc = Document(doc_path)
    paras = doc.paragraphs

    body_start = None
    body_end = len(paras)

    for i, p in enumerate(paras):
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()

        if style == "Heading 1" and re.match(r'^\d+\s', text):
            if body_start is None:
                body_start = i

        if style == "Heading 1" and "结" in text and "论" in text and not re.match(r'^\d', text):
            body_end = i
            break

    if not body_start:
        print("  警告: 未找到正文起始")
        return

    h1_idx = body_start
    h2_idx = None
    h3_idx = None
    body_idx = None

    for i in range(body_start + 1, body_end):
        p = paras[i]
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        if not text:
            continue

        if style == "Heading 1":
            continue
        if style == "Heading 2" and h2_idx is None:
            h2_idx = i
        if style == "Heading 3" and h3_idx is None:
            h3_idx = i
        if body_idx is None and style in ("Plain Text", "Normal") and len(text) > 20:
            body_idx = i

    if not body_idx:
        print(f"  警告: 样本不足 h1={h1_idx} h2={h2_idx} h3={h3_idx} body={body_idx}")
        return

    print(f"    样本: h1=[{h1_idx}] h2=[{h2_idx}] h3=[{h3_idx}] body=[{body_idx}]")

    keep = {h1_idx, h2_idx, h3_idx, body_idx} - {None}

    for i in range(body_start, body_end):
        if i in keep:
            continue
        p = paras[i]
        if p._p.find('.//w:sectPr', _NS) is not None:
            continue
        for r in p.runs:
            r.text = ""
        style = p.style.name if p.style else ""
        if style.startswith("Heading") or style.startswith("标题"):
            p.style = doc.styles['Normal']

    # 删除正文区域表格
    body_elem = doc.element.body
    tables_rm = []
    pc = 0
    for elem in body_elem:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            pc += 1
        elif tag == 'tbl' and body_start <= pc <= body_end:
            tables_rm.append(elem)
    for tbl in tables_rm:
        body_elem.remove(tbl)
    if tables_rm:
        print(f"    删除 {len(tables_rm)} 个示例表格")

    def _set_text(para, text):
        if para.runs:
            para.runs[0].text = text
            for r in para.runs[1:]:
                r.text = ""

    paras[h1_idx].paragraph_format.page_break_before = True
    _set_text(paras[h1_idx], "{{ ch.title }}")
    if h2_idx:
        _set_text(paras[h2_idx], "{{ sec.title }}")
    if h3_idx:
        _set_text(paras[h3_idx], "{{ sub.title }}")
    _set_text(paras[body_idx], "{{ para }}")

    h1_p = paras[h1_idx]._p
    h2_p = paras[h2_idx]._p if h2_idx else None
    h3_p = paras[h3_idx]._p if h3_idx else None
    body_p = paras[body_idx]._p
    body_copy = copy.deepcopy(body_p)

    def _mk_body(text):
        p = copy.deepcopy(body_copy)
        for t_elem in p.findall('.//w:t', _NS):
            t_elem.text = ""
        t_elems = p.findall('.//w:t', _NS)
        if t_elems:
            t_elems[0].text = text
        return p

    parent = h1_p.getparent()
    anchor = h1_p.getprevious()

    for elem in [h1_p, h2_p, h3_p, body_p]:
        if elem is not None and elem.getparent() is not None:
            parent.remove(elem)

    cursor = anchor

    def _after(cursor, elem):
        cursor.addnext(elem)
        return elem

    cursor = _after(cursor, _mk_ctrl("{%p for ch in chapters %}"))
    cursor = _after(cursor, h1_p)

    if h2_p is not None:
        cursor = _after(cursor, _mk_ctrl("{%p for sec in ch.sections %}"))
        cursor = _after(cursor, h2_p)
        cursor = _after(cursor, _mk_ctrl("{%p for para in sec.content %}"))
        cursor = _after(cursor, body_p)
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))

        if h3_p is not None:
            cursor = _after(cursor, _mk_ctrl("{%p for sub in sec.subsections %}"))
            cursor = _after(cursor, h3_p)
            cursor = _after(cursor, _mk_ctrl("{%p for p2 in sub.content %}"))
            cursor = _after(cursor, _mk_body("{{ p2 }}"))
            cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))
            cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))

        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))
    else:
        cursor = _after(cursor, _mk_ctrl("{%p for sec in ch.sections %}"))
        cursor = _after(cursor, _mk_ctrl("{%p for para in sec.content %}"))
        cursor = _after(cursor, body_p)
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))

    cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))

    _remove_empty_between(doc, body_end)
    doc.save(doc_path)


def _remove_empty_between(doc, body_end):
    paras = doc.paragraphs
    last_endfor = None
    target_idx = None

    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        if "{%p endfor %}" in text:
            last_endfor = i
        style = p.style.name if p.style else ""
        if style == "Heading 1" and i >= body_end:
            target_idx = i
            break

    if last_endfor is None or target_idx is None:
        return

    to_remove = []
    for i in range(last_endfor + 1, target_idx):
        p = paras[i]
        text = (p.text or "").strip()
        has_sect = p._p.find('.//w:sectPr', _NS) is not None
        if has_sect:
            if text:
                for r in p.runs:
                    r.text = ""
            continue
        if not text:
            to_remove.append(p._p)

    for elem in to_remove:
        if elem.getparent() is not None:
            elem.getparent().remove(elem)
    if to_remove:
        print(f"    删除 {len(to_remove)} 个空段落")


def _mk_ctrl(text):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python make.py <原始模板.docx> [输出路径]")
        sys.exit(1)
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else \
        os.path.join(os.path.dirname(__file__), "template.docx")
    print(f"制作模板: {src} -> {out}")
    make(src, out)
    print(f"完成: {out}")
