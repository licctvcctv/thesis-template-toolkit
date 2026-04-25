"""
天津职业技术师范大学毕业设计模板清理脚本。

用法:
    cd /Users/a136/vs/45425/thesis_project
    python templates/tjut_traffic/make.py \
        templates/tjut_traffic/source.docx \
        templates/tjut_traffic/template.docx

脚本保留原模板封面、声明、页眉页脚、页码与页面设置，清理示例正文、
旧目录和摘要说明文字，并插入 docxtpl 占位符。
"""
from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def para_key(p):
    return re.sub(r"\s+", "", p.text or "")


def find_para(paragraphs, predicate, label):
    for i, p in enumerate(paragraphs):
        if predicate(p.text or "", p):
            return i, p
    raise RuntimeError(f"未找到段落: {label}")


def find_all(paragraphs, predicate):
    return [(i, p) for i, p in enumerate(paragraphs) if predicate(p.text or "", p)]


def remove_para(p):
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def remove_block_range(start_el, end_el=None, include_start=True, include_end=False):
    body = start_el.getparent()
    children = list(body)
    start = children.index(start_el)
    if end_el is None:
        end = len(children)
    else:
        end = children.index(end_el)
        if include_end:
            end += 1
    if not include_start:
        start += 1
    for el in children[start:end]:
        if el.tag == qn("w:sectPr"):
            continue
        body.remove(el)


def clear_para(p, text):
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def set_single_run(p, text, size=None, bold=None):
    clear_para(p, text)
    for r in p.runs:
        if size is not None:
            r.font.size = Pt(size)
        if bold is not None:
            r.bold = bold


def strip_shading(p):
    for shd in list(p._p.xpath(".//w:shd")):
        parent = shd.getparent()
        if parent is not None:
            parent.remove(shd)


def ensure_page_break_before(p):
    p.paragraph_format.page_break_before = True


def remove_template_callouts(doc: Document):
    """Remove instructional callout shapes left in the source Word template."""
    for tag in ("w:drawing", "w:pict"):
        for el in list(doc.element.xpath(f".//{tag}")):
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)


def add_toc_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(instr)

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run._r.append(sep)

    t = OxmlElement("w:t")
    t.text = "目录更新后显示"
    run._r.append(t)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def patch_update_fields(docx_path: Path):
    with zipfile.ZipFile(docx_path, "a") as zf:
        try:
            xml = zf.read("word/settings.xml").decode("utf-8")
        except KeyError:
            xml = (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                f'<w:settings xmlns:w="{W_NS}"></w:settings>'
            )
        if "w:updateFields" not in xml:
            xml = xml.replace(
                "</w:settings>",
                '<w:updateFields w:val="true"/></w:settings>',
            )
        zf.writestr("word/settings.xml", xml)


def normalize_cover(doc: Document):
    paragraphs = list(doc.paragraphs)
    title_paras = find_all(paragraphs, lambda t, p: para_key(p) == "中文题目")
    en_title_paras = find_all(paragraphs, lambda t, p: para_key(p).lower() == "englishtitle")
    if len(title_paras) < 2 or len(en_title_paras) < 2:
        raise RuntimeError("封面题目占位段落数量异常")

    for _, p in title_paras[:2]:
        set_single_run(p, "{{ title_zh }}")
    for _, p in en_title_paras[:2]:
        set_single_run(p, "{{ title_en }}")

    replacements = {
        "专业：": "专    业： {{ major }}",
        "班级学号：": "班级学号： {{ class_id }} - {{ student_no }}",
        "学生姓名：": "学生姓名： {{ student_name }}",
        "指导教师：": "指导教师： {{ advisor }} {{ advisor_title }}",
        "专业班级：": "专业班级：{{ class_id }}",
        "学 院：": "学    院：{{ college }}",
        "二〇": "{{ year_cn }}年{{ month_cn }}",
        "20年月": "{{ year }} 年 {{ month }} 月",
    }
    for p in paragraphs:
        key = para_key(p)
        raw = p.text or ""
        if raw.strip().startswith("专业班级"):
            set_single_run(p, replacements["专业班级："])
        elif key.startswith("专业："):
            set_single_run(p, replacements["专业："])
        elif raw.strip().startswith("班级学号"):
            set_single_run(p, replacements["班级学号："])
        elif raw.strip().startswith("学生姓名"):
            set_single_run(p, replacements["学生姓名："])
        elif raw.strip().startswith("指导教师"):
            set_single_run(p, replacements["指导教师："])
        elif raw.strip().startswith("学 院") or key.startswith("学院："):
            set_single_run(p, replacements["学 院："])
        elif "二〇" in raw and "年" in raw and "月" in raw:
            set_single_run(p, replacements["二〇"])
        elif key == "20年月":
            set_single_run(p, replacements["20年月"])


def ensure_summary_placeholders(paragraphs, title_idx, key_idx, placeholder_prefix, keep_count=3):
    body = [p for p in paragraphs[title_idx + 1:key_idx] if (p.text or "").strip()]
    for i, p in enumerate(body):
        if i < keep_count:
            set_single_run(p, "{{ %s_%d }}" % (placeholder_prefix, i + 1))
            strip_shading(p)
        else:
            remove_para(p)
    if len(body) < keep_count:
        key_p = paragraphs[key_idx]
        for i in range(len(body), keep_count):
            inserted = key_p.insert_paragraph_before("{{ %s_%d }}" % (placeholder_prefix, i + 1))
            inserted.style = body[0].style if body else key_p.style
            strip_shading(inserted)


def main():
    if len(sys.argv) < 2:
        raise SystemExit("用法: python make.py <source.docx> [template.docx]")

    source = Path(sys.argv[1])
    output = Path(sys.argv[2]) if len(sys.argv) > 2 else Path(__file__).with_name("template.docx")

    doc = Document(source)
    remove_template_callouts(doc)
    normalize_cover(doc)
    paragraphs = list(doc.paragraphs)

    zh_title_idx, zh_title = find_para(paragraphs, lambda t, p: para_key(p) == "摘要", "中文摘要标题")
    zh_key_idx, zh_key = find_para(paragraphs, lambda t, p: t.strip().startswith("关键词"), "中文关键词")
    en_title_idx, en_title = find_para(paragraphs, lambda t, p: para_key(p).upper() == "ABSTRACT", "英文摘要标题")
    en_key_idx, en_key = find_para(paragraphs, lambda t, p: t.strip().startswith("Key Words") or t.strip().startswith("Keywords"), "英文关键词")
    toc_idx, toc_title = find_para(paragraphs, lambda t, p: para_key(p) == "目录", "目录标题")
    body_idx, body_start = find_para(paragraphs, lambda t, p: para_key(p).startswith("1毕业设计的结构"), "正文起点")

    print(f"摘要: P{zh_title_idx}, Abstract: P{en_title_idx}, 目录: P{toc_idx}, 正文: P{body_idx}")

    ensure_page_break_before(zh_title)
    ensure_summary_placeholders(paragraphs, zh_title_idx, zh_key_idx, "abstract_zh", 3)
    set_single_run(zh_key, "关键词：{{ keywords_zh }}")
    strip_shading(zh_key)

    paragraphs = list(doc.paragraphs)
    _, zh_key = find_para(paragraphs, lambda t, p: t.strip().startswith("关键词"), "中文关键词")
    _, en_title = find_para(paragraphs, lambda t, p: para_key(p).upper() == "ABSTRACT", "英文摘要标题")
    remove_block_range(zh_key._element, en_title._element, include_start=False, include_end=False)
    ensure_page_break_before(en_title)

    paragraphs = list(doc.paragraphs)
    en_title_idx, en_title = find_para(paragraphs, lambda t, p: para_key(p).upper() == "ABSTRACT", "英文摘要标题")
    en_key_idx, en_key = find_para(paragraphs, lambda t, p: t.strip().startswith("Key Words") or t.strip().startswith("Keywords"), "英文关键词")
    ensure_summary_placeholders(paragraphs, en_title_idx, en_key_idx, "abstract_en", 3)
    set_single_run(en_key, "Key Words: {{ keywords_en }}")
    strip_shading(en_key)

    paragraphs = list(doc.paragraphs)
    _, en_key = find_para(paragraphs, lambda t, p: t.strip().startswith("Key Words") or t.strip().startswith("Keywords"), "英文关键词")
    _, toc_title = find_para(paragraphs, lambda t, p: para_key(p) == "目录", "目录标题")
    _, body_start = find_para(paragraphs, lambda t, p: para_key(p).startswith("1毕业设计的结构"), "正文起点")

    # 删除英文关键词和目录之间的空白/提示，清除旧目录项和示例正文。
    remove_block_range(en_key._element, toc_title._element, include_start=False, include_end=False)
    paragraphs = list(doc.paragraphs)
    _, toc_title = find_para(paragraphs, lambda t, p: para_key(p) == "目录", "目录标题")
    _, body_start = find_para(paragraphs, lambda t, p: para_key(p).startswith("1毕业设计的结构"), "正文起点")
    remove_block_range(toc_title._element, body_start._element, include_start=False, include_end=False)
    remove_block_range(body_start._element)

    ensure_page_break_before(toc_title)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in toc_title.runs:
        r.bold = True
        r.font.size = Pt(16)

    toc_field = doc.add_paragraph()
    add_toc_field(toc_field)

    body_p = doc.add_paragraph("{{p body}}")
    ensure_page_break_before(body_p)

    ref_title = doc.add_paragraph()
    ensure_page_break_before(ref_title)
    ref_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = ref_title.add_run("参考文献")
    rr.bold = True
    rr.font.size = Pt(16)

    doc.add_paragraph("{{p references_doc}}")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    patch_update_fields(output)
    print(f"模板已生成: {output}")


if __name__ == "__main__":
    main()
