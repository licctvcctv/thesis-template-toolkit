"""
上海海洋大学模板制作。
用法: cd thesis_project && python -m templates.shou.make <原始模板.docx>
"""
import os
import sys
import re
import zipfile
import shutil

# 确保能 import 项目根目录的模块
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../.."))
from editor import TemplateEditor
from body_maker import setup_body_template
from refs_maker import setup_refs_template
from docx import Document


def make(src_path, out_path):
    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    # Step 1: 封面 + 摘要
    ed = TemplateEditor(src_path)
    ed.replace_run(6, 3, "{{ title_zh }}")
    ed.replace_run(8, 4, "{{ title_en }}")
    ed.replace_run(10, 1, "学 院：{{ college }}")
    ed.replace_run(11, 0, "   班 级：{{ class_name }}")
    ed.replace_run(12, 0, "   姓 名：{{ name }}")
    ed.replace_run(13, 0, "   学 号：{{ student_id }}")
    ed.replace_run(14, 0, "指导教师：{{ advisor }}")
    ed.replace_run(16, 0, "{{ year }}年")
    ed.replace_run(16, 2, "{{ month }}月")
    ed.clear_para(40)
    ed.replace_run(41, 1, "{{ abstract_zh }}")
    ed.clear_runs(41, [2, 3, 4, 5, 6, 7, 8, 9, 10, 11])
    ed.clear_para(42)
    ed.replace_run(43, 1, "：{{ keywords_zh }}")
    # 删除多余空 run
    for r in reversed(ed.doc.paragraphs[43].runs[2:]):
        r._r.getparent().remove(r._r)
    ed.clear_para(45)
    ed.clear_para(47)
    ed.replace_run(48, 0, "{{ abstract_en }}")
    ed.clear_runs(48, [1, 2, 3, 4])
    ed.replace_run(49, 2, "{{ keywords_en }}")
    for r in reversed(ed.doc.paragraphs[49].runs[3:]):
        r._r.getparent().remove(r._r)
    ed.clear_para(51)
    ed.clear_para(53)
    ed.save(out_path)
    print("  Step 1: 封面/摘要")

    # Step 2: 正文循环
    setup_body_template(out_path, out_path)
    print("  Step 2: 正文循环")

    # Step 3: 致谢/参考文献
    doc = Document(out_path)
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if text in ["（空一行）", "（空两行）"] and i > 60:
            for r in p.runs: r.text = ""
        if "宋体小四号" in text and "行间距" in text:
            for r in p.runs: r.text = ""
        if "小四号宋体" in text and "行间距" in text:
            for r in p.runs: r.text = ""
        if "正文中公式" in text:
            for r in p.runs: r.text = ""
        if text == "参考文献" and i > 60:
            setup_refs_template(doc, i)
        if "致谢内容致谢内容" in text:
            if p.runs:
                p.runs[0].text = "    {{ acknowledgement }}"
                for r in p.runs[1:]: r.text = ""
    doc.save(out_path)
    print("  Step 3: 致谢/参考文献")

    # Step 4: 删除文本框注释
    tmp = out_path + ".tmp"
    with zipfile.ZipFile(out_path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w') as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    xml = data.decode('utf-8')
                    n = xml.count('<mc:AlternateContent')
                    xml = re.sub(
                        r'<mc:AlternateContent>.*?</mc:AlternateContent>',
                        '', xml, flags=re.DOTALL)
                    print(f"  Step 4: 删除 {n} 个文本框注释")
                    data = xml.encode('utf-8')
                zout.writestr(item, data)
    shutil.move(tmp, out_path)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python make.py <原始模板.docx> [输出路径]")
        sys.exit(1)
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else "template.docx"
    print(f"制作模板: {src} -> {out}")
    make(src, out)
    print(f"完成: {out}")
