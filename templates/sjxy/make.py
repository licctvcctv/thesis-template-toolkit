"""
数计学院（人工智能与信息工程学院）大数据分析类毕业设计模板制作。
用法: cd thesis_project && python -m templates.sjxy.make <原始模板.docx>

模板特点（无封面页、无目录页、无声明页）:
- P0-40: 书写式样说明（全部删除）
- P41: 论文题目 (style=论文题目)
- P42: 专业+姓名 (style=论文正文, run[1]=专业, run[3]=姓名)
- P43: 指导老师 (style=论文正文, run[1]=教师姓名)
- P44: 中文摘要 (style=Normal, run[0-1]="摘要：" bold, run[2+]=正文)
- P45: 关键词 (style=Normal, run[0-1]="关键词：" bold, run[2+]=值)
- P46: Abstract (style=Normal, run[2-3]="Abstract:" bold, run[4+]=正文)
- P47: Keywords (style=Normal, run[1]="Keywords:" bold, run[2+]=值)
- P48-51: 空段
- P52起: 正文 (一级标题/二级标题/三级标题/论文正文/图标题/code block)
- 结论: 最后一个一级标题章节（如"6 总结与展望"），属于正文循环内
- 参考文献: style=一级标题, text="参考文献"
- 致谢: style=一级标题, text="致谢"
"""
import os
import sys
import re
import copy
import zipfile
import shutil

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../.."))
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def make(src_path, out_path):
    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc = Document(src_path)
    paras = doc.paragraphs

    # ========== Step 1: 删除书写式样说明 (P0-P40) ==========
    # 用文本匹配找"二、书写样式"作为说明区结束
    style_guide_end = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        if text == "二、书写样式":
            style_guide_end = i
            break

    removed_guide = 0
    if style_guide_end:
        for i in range(0, style_guide_end + 1):
            for r in paras[i].runs:
                r.text = ""
            removed_guide += 1
        print(f"  Step 1: 清空 {removed_guide} 个书写说明段落")
    else:
        print("  Step 1: 未找到书写式样说明")

    # ========== Step 2: 论文题目 ==========
    # 有两个"论文题目"样式段落，P39是说明区空段，P41���真正的题目
    title_idx = None
    for i, p in enumerate(paras):
        if p.style and p.style.name == "论文题目" and (p.text or "").strip():
            title_idx = i
            break

    if title_idx:
        p = paras[title_idx]
        if p.runs:
            p.runs[0].text = "{{ title_zh }}"
            for r in p.runs[1:]:
                r.text = ""
        print(f"  Step 2: 论文题目 [{title_idx}]")

    # ========== Step 3: 专业+姓名、指导老师 ==========
    # 题目之后的两段: 专业+姓名, 指导老师
    if title_idx:
        # P42: "    软件工程专业  董雲炜（学生姓名）"
        p_info = paras[title_idx + 1]
        if p_info.runs and len(p_info.runs) >= 4:
            p_info.runs[1].text = "{{ major }}"
            p_info.runs[2].text = "专业  "
            p_info.runs[3].text = "{{ name }}"
            for r in p_info.runs[4:]:
                r.text = ""

        # P43: "指导老师：吴奇 "
        p_adv = paras[title_idx + 2]
        if p_adv.runs and len(p_adv.runs) >= 2:
            p_adv.runs[1].text = "{{ advisor }}"
            for r in p_adv.runs[2:]:
                r.text = ""

        print("  Step 3: 专业+姓名、指导老师")

    # ========== Step 4: 中文摘要 ==========
    # 找"摘要："开头的段落
    abs_zh_idx = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        if text.startswith("摘要") and "：" in text[:5]:
            abs_zh_idx = i
            break

    if abs_zh_idx:
        p = paras[abs_zh_idx]
        # run[0-1] = "摘要：" (bold, 保留), run[2+] = 正文内容 → 替换
        if p.runs and len(p.runs) >= 3:
            p.runs[2].text = "{{ abstract_zh }}"
            for r in p.runs[3:]:
                r.text = ""
        print(f"  Step 4: 中文摘要 [{abs_zh_idx}]")

    # ========== Step 5: 关键词 ==========
    kw_zh_idx = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        if text.startswith("关键词") and "：" in text[:5]:
            kw_zh_idx = i
            break

    if kw_zh_idx:
        p = paras[kw_zh_idx]
        if p.runs and len(p.runs) >= 3:
            p.runs[2].text = "{{ keywords_zh }}"
            for r in p.runs[3:]:
                r.text = ""
        print(f"  Step 5: 关键词 [{kw_zh_idx}]")

    # ========== Step 6: 英文摘要 ==========
    abs_en_idx = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        if "Abstract" in text and ":" in text[:15]:
            abs_en_idx = i
            break

    if abs_en_idx:
        p = paras[abs_en_idx]
        # run[0-1] 空, run[2-3]="Abstract:" bold, run[4+]=正文
        if p.runs and len(p.runs) >= 5:
            p.runs[4].text = "{{ abstract_en }}"
            for r in p.runs[5:]:
                r.text = ""
        print(f"  Step 6: 英文摘要 [{abs_en_idx}]")

    # ========== Step 7: 英文关键词 ==========
    kw_en_idx = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        if "Keywords:" in text or "Keywords：" in text:
            kw_en_idx = i
            break

    if kw_en_idx:
        p = paras[kw_en_idx]
        # run[1]="Keywords:" bold, run[2]=空格, run[3+]=值
        if p.runs and len(p.runs) >= 4:
            p.runs[2].text = " "
            p.runs[3].text = "{{ keywords_en }}"
            for r in p.runs[4:]:
                r.text = ""
        print(f"  Step 7: 英文关键词 [{kw_en_idx}]")

    # ========== Step 8: 结论（循环） ==========
    # 找"总结"/"结论"一级标题，将其正文段设为循环
    concl_idx = None
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        if style == "一级标题" and ("总结" in text or "结论" in text):
            concl_idx = i
            break

    if concl_idx:
        # 找结论后第一个正文段 → 设循环
        for j in range(concl_idx + 1, min(concl_idx + 30, len(paras))):
            pj = paras[j]
            text = (pj.text or "").strip()
            style = pj.style.name if pj.style else ""
            if style == "一级标题":
                break
            if style == "论文正文" and text and pj.runs:
                pj_p = pj._p
                pj_p.addprevious(_mk_ctrl("{%p for cp in conclusion_list %}"))
                pj.runs[0].text = "{{ cp }}"
                for r in pj.runs[1:]:
                    r.text = ""
                pj_p.addnext(_mk_ctrl("{%p endfor %}"))
                # 清后续正文段
                for k in range(j + 1, min(j + 30, len(paras))):
                    pk = paras[k]
                    pk_style = pk.style.name if pk.style else ""
                    if pk_style == "一级标题":
                        break
                    for r in pk.runs:
                        r.text = ""
                break
        print(f"  Step 8: 结论 [{concl_idx}]")

    # ========== Step 9: 参考文献 ==========
    ref_idx = None
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        if style == "一级标题" and text == "参考文献":
            ref_idx = i
            break

    if ref_idx:
        first_ref = None
        for j in range(ref_idx + 1, min(ref_idx + 5, len(paras))):
            pj = paras[j]
            pt = (pj.text or "").strip()
            if pt and pj.runs:
                first_ref = j
                break

        if first_ref:
            pj = paras[first_ref]
            ed_p = pj._p
            ed_p.addprevious(_mk_ctrl("{%p for ref in references %}"))
            if pj.runs:
                pj.runs[0].text = "{{ ref }}"
                for r in pj.runs[1:]:
                    r.text = ""
            ed_p.addnext(_mk_ctrl("{%p endfor %}"))

            # 清后续参考文献条目直到致谢
            for k in range(first_ref + 1, len(paras)):
                pk = paras[k]
                pk_style = pk.style.name if pk.style else ""
                pk_text = (pk.text or "").strip()
                if pk_style == "一级标题":
                    break
                for r in pk.runs:
                    r.text = ""

        print(f"  Step 9: 参考文献 [{ref_idx}]")

    # ========== Step 10: 致谢 ==========
    ack_idx = None
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        if style == "一级标题" and text == "致谢":
            ack_idx = i
            break

    if ack_idx:
        for j in range(ack_idx + 1, min(ack_idx + 20, len(paras))):
            pj = paras[j]
            text = (pj.text or "").strip()
            if not text:
                continue
            if pj.runs:
                pj_p = pj._p
                pj_p.addprevious(_mk_ctrl("{%p for ack_p in acknowledgement_list %}"))
                pj.runs[0].text = "{{ ack_p }}"
                for r in pj.runs[1:]:
                    r.text = ""
                pj_p.addnext(_mk_ctrl("{%p endfor %}"))
                # 清后续致谢段
                for k in range(j + 1, len(paras)):
                    pk = paras[k]
                    pk_style = pk.style.name if pk.style else ""
                    if pk_style == "一级标题":
                        break
                    for r in pk.runs:
                        r.text = ""
                break

        print(f"  Step 10: 致谢 [{ack_idx}]")

    # ========== Step 11: 删除空段落和残留内容 ==========
    to_remove = []

    # 1) 删除书写说明区域所有段落（文字已清空，含空run残留段）
    if style_guide_end:
        for i in range(0, style_guide_end + 1):
            p = paras[i]
            has_sect = p._p.find('.//w:sectPr', _NS) is not None
            if not has_sect:
                to_remove.append(p._p)

    # 2) 删除英文关键词 → 正文之间的空段落
    if kw_en_idx:
        for i in range(kw_en_idx + 1, len(paras)):
            p = paras[i]
            style = p.style.name if p.style else ""
            text = (p.text or "").strip()
            if style == "一级标题" or text:
                break
            has_sect = p._p.find('.//w:sectPr', _NS) is not None
            if not has_sect:
                to_remove.append(p._p)

    # 3) 删除参考文献 → 致谢之间的空段落
    if ref_idx and ack_idx:
        for i in range(ref_idx + 1, ack_idx):
            p = paras[i]
            if not (p.text or "").strip():
                has_sect = p._p.find('.//w:sectPr', _NS) is not None
                if not has_sect:
                    to_remove.append(p._p)

    # 4) 删除致谢之后到文档末尾的空段落
    if ack_idx:
        in_ack = False
        for i in range(ack_idx + 1, len(paras)):
            p = paras[i]
            text = (p.text or "").strip()
            if not text:
                has_sect = p._p.find('.//w:sectPr', _NS) is not None
                if not has_sect:
                    to_remove.append(p._p)

    # 5) 删除说明区内的表格（正文之前的表格都是说明用的）
    body_elem = doc.element.body
    first_h1_pc = None
    pc = 0
    pre_tables = []
    for elem in body_elem:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            pc += 1
            if first_h1_pc is None:
                p = paras[pc - 1] if pc - 1 < len(paras) else None
                if p and p.style and p.style.name == "一级标题" and (p.text or "").strip():
                    first_h1_pc = pc
        elif tag == 'tbl' and (first_h1_pc is None or pc < first_h1_pc):
            pre_tables.append(elem)
    for tbl_elem in pre_tables:
        body_elem.remove(tbl_elem)
    if pre_tables:
        print(f"  Step 11: 删除 {len(pre_tables)} 个说明区表格")

    removed = 0
    for elem in to_remove:
        if elem.getparent() is not None:
            elem.getparent().remove(elem)
            removed += 1
    if removed:
        print(f"  Step 11: 删除 {removed} 个空段落")

    doc.save(out_path)

    # ========== Step 12: 正文循环 ==========
    _setup_body(out_path, concl_idx)
    print("  Step 12: 正文循环")

    # ========== Step 13: 删除文本框/AlternateContent + 内嵌图片 ==========
    tmp = out_path + ".tmp"
    with zipfile.ZipFile(out_path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w') as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    xml = data.decode('utf-8')
                    n = xml.count('<mc:AlternateContent')
                    if n > 0:
                        xml = re.sub(
                            r'<mc:AlternateContent>.*?</mc:AlternateContent>',
                            '', xml, flags=re.DOTALL)
                        print(f"  Step 13: 删除 {n} 个文本框")
                    n2 = xml.count('<w:drawing>')
                    if n2 > 0:
                        xml = re.sub(
                            r'<w:drawing>.*?</w:drawing>',
                            '', xml, flags=re.DOTALL)
                        print(f"  Step 13: 删除 {n2} 个内嵌图片")
                    data = xml.encode('utf-8')
                zout.writestr(item, data)
    shutil.move(tmp, out_path)


def _setup_body(doc_path, concl_idx_hint=None):
    """设置正文 Jinja2 循环。

    此模板特点:
    - 章标题: 一级标题 (e.g. "1 绪论")
    - 二级标题: 二级标题 (e.g. "1.1 项目背景与意义")
    - 三级标题: 三级标题 (e.g. "2.3.1 缺失值处理")
    - 正文: 论文正文
    - 图表标题: 图标题
    - 代码块: code block
    """
    doc = Document(doc_path)
    paras = doc.paragraphs

    # 找正文范围：第一个一级标题 ~ "参考文献"一级标题
    body_start = None
    body_end = len(paras)

    for i, p in enumerate(paras):
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()

        if style == "一级标题" and re.match(r'^\d+\s', text):
            if body_start is None:
                body_start = i

        if style == "一级标题" and text == "参考文献":
            body_end = i
            break

    if not body_start:
        print("  警告: 未找到正文起始")
        return

    # 找样本段落
    h1_idx = body_start
    h2_idx = None
    h3_idx = None
    body_idx = None

    for i in range(body_start + 1, body_end):
        p = paras[i]
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        if not text:
            continue

        if style == "一级标题":
            continue
        if style == "二级标题" and h2_idx is None:
            h2_idx = i
        if style == "三级标题" and h3_idx is None:
            h3_idx = i
        if style == "论文正文" and body_idx is None and len(text) > 20:
            body_idx = i

    if not body_idx:
        print(f"  警告: 样本不足 h1={h1_idx} h2={h2_idx} h3={h3_idx} body={body_idx}")
        return

    print(f"    样本: h1=[{h1_idx}] h2=[{h2_idx}] h3=[{h3_idx}] body=[{body_idx}]")

    keep = {h1_idx, h2_idx, h3_idx, body_idx} - {None}

    # 清空非样本段落
    for i in range(body_start, body_end):
        if i in keep:
            continue
        p = paras[i]
        if p._p.find('.//w:sectPr', _NS) is not None:
            continue
        for r in p.runs:
            r.text = ""
        style = p.style.name if p.style else ""
        if style in ("一级标题", "二级标题", "三级标题", "图标题", "code block"):
            try:
                p.style = doc.styles['论文正文']
            except KeyError:
                p.style = doc.styles['Normal']

    # 删除正文区域表格
    body_elem = doc.element.body
    tables_rm = []
    pc = 0
    for elem in body_elem:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            pc += 1
        elif tag == 'tbl' and body_start <= pc <= body_end:
            tables_rm.append(elem)
    for tbl in tables_rm:
        body_elem.remove(tbl)
    if tables_rm:
        print(f"    删除 {len(tables_rm)} 个示例表格")

    # 替换样本文本
    def _set_text(para, text):
        if para.runs:
            para.runs[0].text = text
            for r in para.runs[1:]:
                r.text = ""

    paras[h1_idx].paragraph_format.page_break_before = True
    _set_text(paras[h1_idx], "{{ ch.title }}")
    if h2_idx:
        _set_text(paras[h2_idx], "{{ sec.title }}")
    if h3_idx:
        _set_text(paras[h3_idx], "{{ sub.title }}")
    _set_text(paras[body_idx], "{{ para }}")

    h1_p = paras[h1_idx]._p
    h2_p = paras[h2_idx]._p if h2_idx else None
    h3_p = paras[h3_idx]._p if h3_idx else None
    body_p = paras[body_idx]._p
    body_copy = copy.deepcopy(body_p)

    def _mk_body(text):
        p = copy.deepcopy(body_copy)
        for t_elem in p.findall('.//w:t', _NS):
            t_elem.text = ""
        t_elems = p.findall('.//w:t', _NS)
        if t_elems:
            t_elems[0].text = text
        return p

    # 移除样本，按正确顺序重新插入
    parent = h1_p.getparent()
    anchor = h1_p.getprevious()

    for elem in [h1_p, h2_p, h3_p, body_p]:
        if elem is not None and elem.getparent() is not None:
            parent.remove(elem)

    cursor = anchor

    def _after(cursor, elem):
        cursor.addnext(elem)
        return elem

    cursor = _after(cursor, _mk_ctrl("{%p for ch in chapters %}"))
    cursor = _after(cursor, h1_p)

    if h2_p is not None:
        cursor = _after(cursor, _mk_ctrl("{%p for sec in ch.sections %}"))
        cursor = _after(cursor, h2_p)
        cursor = _after(cursor, _mk_ctrl("{%p for para in sec.content %}"))
        cursor = _after(cursor, body_p)
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))  # content

        if h3_p is not None:
            cursor = _after(cursor, _mk_ctrl("{%p for sub in sec.subsections %}"))
            cursor = _after(cursor, h3_p)
            cursor = _after(cursor, _mk_ctrl("{%p for p2 in sub.content %}"))
            cursor = _after(cursor, _mk_body("{{ p2 }}"))
            cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))  # sub.content
            cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))  # subsections

        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))  # sections
    else:
        cursor = _after(cursor, _mk_ctrl("{%p for sec in ch.sections %}"))
        cursor = _after(cursor, _mk_ctrl("{%p for para in sec.content %}"))
        cursor = _after(cursor, body_p)
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))

    cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))  # chapters

    # 清正文末尾到参考文献之间的残留空段落
    _remove_empty_between(doc, body_end)

    doc.save(doc_path)


def _remove_empty_between(doc, body_end):
    """清除正文循环 endfor 和参考文献之间的空段落"""
    paras = doc.paragraphs
    last_endfor = None
    target_idx = None

    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        if "{%p endfor %}" in text:
            last_endfor = i
        style = p.style.name if p.style else ""
        if style == "一级标题" and text == "参考文献":
            target_idx = i
            break

    if last_endfor is None or target_idx is None:
        return

    to_remove = []
    for i in range(last_endfor + 1, target_idx):
        p = paras[i]
        text = (p.text or "").strip()
        if p._p.find('.//w:sectPr', _NS) is not None:
            if text:
                for r in p.runs:
                    r.text = ""
            continue
        if not text:
            to_remove.append(p._p)

    for elem in to_remove:
        if elem.getparent() is not None:
            elem.getparent().remove(elem)
    if to_remove:
        print(f"    删除 {len(to_remove)} 个空段落 (endfor→参考文献)")


def _mk_ctrl(text):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python make.py <原始模板.docx> [输出路径]")
        sys.exit(1)
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else \
        os.path.join(os.path.dirname(__file__), "template.docx")
    print(f"制作模板: {src} -> {out}")
    make(src, out)
    print(f"完成: {out}")
