"""
青岛科技大学高等学历继续教育本科毕业论文模板制作。
用法: cd thesis_project && python -m templates.qdkjdx.make <原始模板.docx>

模板结构:
- P0-P17:  封面页 (校名、题目下划线、表格[指导教师/专业/层次/姓名/学号]、日期) [sectPr@P17]
- P18:     中文论文题目 (style=中文论文题目)
- P19:     "摘 要" 标题 (style=论文正文章标题)
- P20-P21: 中文摘要正文 (style=中文摘要正文)
- P22:     关键词 (style=中文关键词内容, 含IMG装饰)
- P23-P37: 英文摘要区域 (含IMG装饰) [sectPr@P37]
- P38-P59: 目录页 [sectPr@P59]
- P60-P61: 前言 [sectPr@P62]
- P63-P72: 第1章示例 (章标题/节标题/小节/正文/图/四级标题) [sectPr@P73]
- P74-P80: 第2章示例 (含示例表格) [sectPr@P81]
- P82-P84: 参考文献 [sectPr@P86]
- P87-P88: 致谢

关键 style 名称:
  中文论文题目, 论文正文章标题, 论文正文,
  中文摘要正文, 中文关键词内容,
  参考文献标题, 中文参考文献, 英文参考文献,
  致谢标题,
  目录章标题, 目录节标题, 目录小节标题, 目录小节子标题
"""
import os
import sys
import re
import copy
import zipfile
import shutil

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../.."))
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def make(src_path, out_path):
    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc = Document(src_path)
    paras = doc.paragraphs

    print(f"原始: {len(paras)} 段落, {len(doc.tables)} 表格")

    # ========== Step 1: 封面 - 替换题目、表格字段、日期 ==========
    # P9-P10: 题目下划线 → Jinja2
    for i in (9, 10):
        p = paras[i]
        if p.runs:
            if i == 9:
                # "题 目 ____" → "题 目 {{ title_zh }}"
                p.runs[0].text = "题 目 "
                if len(p.runs) > 1:
                    p.runs[1].text = "{{ title_zh }}"
                    for r in p.runs[2:]:
                        r.text = ""
                else:
                    p.runs[0].text = "题 目 {{ title_zh }}"
            else:
                # P10: second line of title (for long titles)
                for r in p.runs:
                    r.text = ""

    # Table 0: 指导教师/专业名称/层次/学生姓名/学生学号
    if doc.tables:
        tbl = doc.tables[0]
        field_map = {
            '指导教师': '{{ advisor }}',
            '专业名称': '{{ major }}',
            '层    次': '{{ level }}',
            '层次': '{{ level }}',
            '学生姓名': '{{ name }}',
            '学生学号': '{{ student_id }}',
        }
        for row in tbl.rows:
            label = row.cells[0].text.strip()
            if label in field_map:
                cell = row.cells[1]
                for p in cell.paragraphs:
                    if p.runs:
                        p.runs[0].text = field_map[label]
                        for r in p.runs[1:]:
                            r.text = ""
                    else:
                        run = p.add_run(field_map[label])

    # P17: 日期 "2025年 5月"
    p17 = paras[17]
    if p17.runs:
        p17.runs[0].text = "{{ year }}"
        if len(p17.runs) > 1:
            p17.runs[1].text = "年 "
        if len(p17.runs) > 2:
            p17.runs[2].text = "{{ month }}"
            for r in p17.runs[3:]:
                r.text = ""

    print("  Step 1: 封面字段替换完成")

    # ========== Step 2: 中文论文题目 ==========
    p18 = paras[18]
    if p18.runs:
        p18.runs[0].text = "{{ title_zh }}"
        for r in p18.runs[1:]:
            r.text = ""
    print("  Step 2: 中文论文题目")

    # ========== Step 3: 中文摘要 ==========
    # P19 "摘 要" - keep as is (heading)
    # P20-P21: 摘要正文 → single Jinja2 var
    p20 = paras[20]
    if p20.runs:
        p20.runs[0].text = "{{ abstract_zh }}"
        for r in p20.runs[1:]:
            r.text = ""
    # P21: clear second paragraph
    p21 = paras[21]
    for r in p21.runs:
        r.text = ""

    print("  Step 3: 中文摘要")

    # ========== Step 4: 关键词 ==========
    p22 = paras[22]
    if p22.runs:
        # Keep "关键词：" prefix, replace rest
        found_kw = False
        for ri, r in enumerate(p22.runs):
            if '关键词' in r.text:
                r.text = "关键词："
                found_kw = True
                # Clear remaining runs, put value in next
                for r2 in p22.runs[ri+1:]:
                    r2.text = ""
                if ri + 1 < len(p22.runs):
                    p22.runs[ri+1].text = "{{ keywords_zh }}"
                else:
                    r.text = "关键词：{{ keywords_zh }}"
                break
        if not found_kw and p22.runs:
            p22.runs[0].text = "关键词：{{ keywords_zh }}"
            for r in p22.runs[1:]:
                r.text = ""

    print("  Step 4: 关键词")

    # ========== Step 5: 英文摘要区域 (P23-P37) → 清空但保留sectPr ==========
    # These contain decorative images and empty paragraphs
    # We don't need English abstract in this template variant
    for i in range(23, 38):
        p = paras[i]
        # Remove images (drawing elements)
        for drawing in p._p.findall('.//' + qn('w:drawing')):
            drawing.getparent().remove(drawing)
        for r in p.runs:
            r.text = ""
    print("  Step 5: 清空英文摘要区域")

    # ========== Step 6: 目录页 (P38-P59) → 清空但保留sectPr ==========
    for i in range(38, 60):
        p = paras[i]
        for r in p.runs:
            r.text = ""
    print("  Step 6: 清空目录页")

    # ========== Step 7: 前言(P60-P61) → 改为 Jinja2 前言循环 ==========
    # P60: "前言" 章标题 - keep
    # P61: 正文 → preface placeholder
    p61 = paras[61]
    if p61.runs:
        p61.runs[0].text = "{% for para in preface %}{{ para }}\n{% endfor %}"
        for r in p61.runs[1:]:
            r.text = ""
    print("  Step 7: 前言")

    # ========== Step 8: 正文章节区域 → Jinja2 循环 ==========
    # 关键: {%p for/endfor %} 必须独占一段，不能与 {{ }} 同段
    # 需要18个段落: P63-P80 刚好够用
    #
    #   P63: {%p for ch in chapters %}
    #   P64: {{ ch.title }}              ← 章标题样式
    #   P65: {%p for item in ch.content %}
    #   P66: {{ item }}                  ← 正文样式
    #   P67: {%p endfor %}               -- end content
    #   P68: {%p for sec in ch.sections %}
    #   P69: {{ sec.title }}             ← 节标题样式
    #   P70: {%p for item in sec.content %}
    #   P71: {{ item }}                  ← 正文样式
    #   P72: {%p endfor %}               -- end sec.content
    #   P73: {%p for sub in sec.subsections %}
    #   P74: {{ sub.title }}             ← 小节标题样式
    #   P75: {%p for item in sub.content %}
    #   P76: {{ item }}                  ← 正文样式
    #   P77: {%p endfor %}               -- end sub.content
    #   P78: {%p endfor %}               -- end subsections
    #   P79: {%p endfor %}               -- end sections
    #   P80: {%p endfor %}               -- end chapters

    # Save original formatting BEFORE overwriting any content
    # P64 = body text (论文正文), P65 = section heading, P67 = subsection heading
    body_pPr = copy.deepcopy(paras[64]._p.pPr) if paras[64]._p.pPr is not None else None
    body_rPr = copy.deepcopy(paras[64].runs[0]._r.find(qn('w:rPr'))) if paras[64].runs else None
    sec_heading_pPr = copy.deepcopy(paras[65]._p.pPr) if paras[65]._p.pPr is not None else None
    sec_heading_rPr = copy.deepcopy(paras[65].runs[0]._r.find(qn('w:rPr'))) if paras[65].runs else None
    sub_heading_pPr = copy.deepcopy(paras[67]._p.pPr) if paras[67]._p.pPr is not None else None
    sub_heading_rPr = copy.deepcopy(paras[67].runs[0]._r.find(qn('w:rPr'))) if paras[67].runs else None

    tags = [
        (63, "{%p for ch in chapters %}"),
        (64, "{{ ch.title }}"),
        (65, "{%p for item in ch.content %}"),
        (66, "{{ item }}"),
        (67, "{%p endfor %}"),
        (68, "{%p for sec in ch.sections %}"),
        (69, "{{ sec.title }}"),
        (70, "{%p for item in sec.content %}"),
        (71, "{{ item }}"),
        (72, "{%p endfor %}"),
        (73, "{%p for sub in sec.subsections %}"),
        (74, "{{ sub.title }}"),
        (75, "{%p for item in sub.content %}"),
        (76, "{{ item }}"),
        (77, "{%p endfor %}"),
        (78, "{%p endfor %}"),
        (79, "{%p endfor %}"),
        (80, "{%p endfor %}"),
    ]

    for idx, text in tags:
        p = paras[idx]
        # Remove images
        for drawing in p._p.findall('.//' + qn('w:drawing')):
            drawing.getparent().remove(drawing)
        if p.runs:
            p.runs[0].text = text
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(text)

    # Fix paragraph styles: chapter title → 论文正文章标题, no indent
    paras[64].style = doc.styles['论文正文章标题']
    ind = paras[64]._p.pPr.find(qn('w:ind')) if paras[64]._p.pPr is not None else None
    if ind is not None:
        paras[64]._p.pPr.remove(ind)

    # Fix section title (P69): apply original P65 pPr and run rPr (黑体 15pt)
    def _apply_heading_fmt(para, saved_pPr, saved_rPr):
        """Replace paragraph pPr and first run rPr with saved heading formatting."""
        if saved_pPr is not None:
            old = para._p.pPr
            if old is not None:
                para._p.remove(old)
            para._p.insert(0, copy.deepcopy(saved_pPr))
        if saved_rPr is not None and para.runs:
            r = para.runs[0]._r
            old = r.find(qn('w:rPr'))
            if old is not None:
                r.remove(old)
            r.insert(0, copy.deepcopy(saved_rPr))

    _apply_heading_fmt(paras[69], sec_heading_pPr, sec_heading_rPr)
    _apply_heading_fmt(paras[74], sub_heading_pPr, sub_heading_rPr)

    # Fix content paragraphs → apply saved body text formatting (from original P64)
    for idx in [66, 71, 76]:
        _apply_heading_fmt(paras[idx], body_pPr, body_rPr)

    # P81: has sectPr - clear text only
    p81 = paras[81]
    for r in p81.runs:
        r.text = ""

    # Remove Table 1 (sample data table)
    if len(doc.tables) > 1:
        tbl1 = doc.tables[1]
        tbl1._element.getparent().remove(tbl1._element)
        print("  Step 8a: 删除示例数据表格")

    print("  Step 8: 正文章节循环")

    # ========== Step 9: 参考文献 ==========
    # P82: "参考文献：" heading - keep
    # P83: 中文参考文献 → loop
    p83 = paras[83]
    if p83.runs:
        p83.runs[0].text = "{% for ref in references %}{{ ref }}\n{% endfor %}"
        for r in p83.runs[1:]:
            r.text = ""
    # P84: 英文参考文献 → clear
    p84 = paras[84]
    for r in p84.runs:
        r.text = ""

    print("  Step 9: 参考文献")

    # ========== Step 10: 致谢 ==========
    # P87: "致谢" heading - keep
    # P88: 致谢正文
    p88 = paras[88]
    if p88.runs:
        p88.runs[0].text = "{% for para in acknowledgement_list %}{{ para }}\n{% endfor %}"
        for r in p88.runs[1:]:
            r.text = ""
    print("  Step 10: 致谢")

    # ========== Step 11: 清理残余图片 ==========
    # Remove decorative images from cover (P8, P22, P23, P32, P33)
    for i in [8]:
        p = paras[i]
        for drawing in p._p.findall('.//' + qn('w:drawing')):
            drawing.getparent().remove(drawing)

    # Remove images from P22 (keywords paragraph has decorative image)
    for drawing in paras[22]._p.findall('.//' + qn('w:drawing')):
        drawing.getparent().remove(drawing)

    print("  Step 11: 清理装饰图片")

    # ========== Step 12: 删除空段落 ==========
    removed = 0
    for i in range(len(paras) - 1, -1, -1):
        p = paras[i]
        text = (p.text or "").strip()
        style = p.style.name if p.style else ""

        # Skip paragraphs with sectPr
        pPr = p._p.pPr
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            continue

        # Skip non-empty and styled paragraphs
        if text:
            continue

        # Skip paragraphs in cover area (P0-P17)
        if i <= 17:
            continue

        # Skip key structural paragraphs
        if style in ('论文正文章标题', '中文论文题目', '参考文献标题', '致谢标题'):
            continue

        # Remove truly empty paragraphs in cleared sections
        if 23 <= i <= 36:  # English abstract area
            p._p.getparent().remove(p._p)
            removed += 1
        elif 38 <= i <= 58:  # TOC area
            p._p.getparent().remove(p._p)
            removed += 1
        elif 69 <= i <= 72:  # Cleared example area
            p._p.getparent().remove(p._p)
            removed += 1
        elif 74 <= i <= 81:  # Cleared chapter 2 area
            if not text:
                p._p.getparent().remove(p._p)
                removed += 1

    print(f"  Step 12: 删除 {removed} 个空段落")

    # ========== Save ==========
    doc.save(out_path)

    # Verify
    doc2 = Document(out_path)
    p_count = len(doc2.paragraphs)
    t_count = len(doc2.tables)
    texts = [(p.text or "").strip() for p in doc2.paragraphs if (p.text or "").strip()]
    print(f"\n结果: {p_count} 段落, {t_count} 表格")
    for i, p in enumerate(doc2.paragraphs):
        t = (p.text or "").strip()
        if t:
            print(f"  P{i:2d} [{p.style.name:25s}] {t[:80]}")

    # Check for residual template issues
    full = "\n".join(p.text or "" for p in doc2.paragraphs)
    for tag in ["{%", "%}", "{{", "}}"]:
        count = full.count(tag)
        if count:
            print(f"  Jinja2 '{tag}': {count}")


def main():
    if len(sys.argv) < 2:
        print("用法: python -m templates.qdkjdx.make <原始模板.docx>")
        sys.exit(1)

    src = sys.argv[1]
    here = os.path.dirname(os.path.abspath(__file__))
    out = os.path.join(here, "template.docx")

    print(f"输入: {src}")
    print(f"输出: {out}")
    make(src, out)
    print("完成!")


if __name__ == "__main__":
    main()
