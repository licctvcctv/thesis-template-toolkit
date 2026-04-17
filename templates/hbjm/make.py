"""
河北经贸大学本科毕业论文模板制作。
用法: cd thesis_project && python -m templates.hbjm.make <原始模板.docx>

模板特点:
- 封面: P0-17, Logo图片, "本科毕业论文（设计）", 标题P5, 学院P10/专业P11/姓名P12/学号P13/导师P14, 日期P17
- 原创性声明: P19-27, P22含论文标题引用
- 版权授权: P34-41, P41有SECT分节符
- 中文摘要: P42 "摘　　要"(1正文), P43-45正文, P47关键词
- 英文摘要: P58 "ABSTRACT"(1正文), P59-61正文, P63 Keywords
- 目录: P65-105, toc 1/2/3, P105有SECT分节符
- 正文: P106起, Heading 1/2 + 标题3
- 结论: P526 "结　　论" (Heading 1)
- 参考文献: P532 (Heading 1), P533-554条目(Normal, 无numPr)
- 致谢: P555 (Heading 1), P556-559正文
"""
import os
import sys
import re
import copy
import zipfile
import shutil

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../.."))
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def make(src_path, out_path):
    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc = Document(src_path)
    paras = doc.paragraphs

    # ========== Step 1: 封面 ==========
    # P5: 论文标题 "健身管理小程序"
    p5 = paras[5]
    if p5.runs:
        p5.runs[0].text = "{{ title_zh }}"
        for r in p5.runs[1:]:
            r.text = ""

    # P10-P14: 学院/专业/姓名/学号/导师
    # These have complex multi-run structure with labels and values separated
    # Pattern: label runs + spacing + value run + trailing spacing
    cover_fields = {
        10: ("国际教育学院", "{{ college }}"),
        11: ("计算机科学与技术", "{{ major }}"),
        12: ("赵恩", "{{ name }}"),
        13: ("202162310501", "{{ student_id }}"),
        14: ("曾文献", "{{ advisor }}"),
    }
    for idx, (orig_value, var) in cover_fields.items():
        p = paras[idx]
        found = False
        for j, r in enumerate(p.runs):
            if orig_value in r.text:
                r.text = var
                # Clear trailing whitespace runs
                for k in range(j + 1, len(p.runs)):
                    p.runs[k].text = ""
                found = True
                break
        if not found:
            # Fallback: find the run containing the actual value by looking for
            # non-whitespace, non-label content
            full = p.text
            colon_pos = full.find("：")
            if colon_pos >= 0:
                # Find the run that starts the value after the colon
                char_count = 0
                for j, r in enumerate(p.runs):
                    char_count += len(r.text)
                    if char_count > colon_pos + 1 and r.text.strip():
                        r.text = var
                        for k in range(j + 1, len(p.runs)):
                            p.runs[k].text = ""
                        break

    # P17: 提交日期 "提交日期： 二〇二五  年  四  月"
    # runs[0]="提交日期：", runs[2]="二〇二五", runs[5]="年", runs[8]="四", runs[11]="月"
    # Replace the date portion with {{ finish_date }}
    p17 = paras[17]
    if p17.runs and len(p17.runs) >= 3:
        # Keep "提交日期：" in runs[0], put variable in runs[1]
        p17.runs[1].text = "{{ finish_date }}"
        for r in p17.runs[2:]:
            r.text = ""

    print("  Step 1: 封面")

    # ========== Step 2: 原创性声明 ==========
    # P22: 本人所提交的学位论文"健身管理小程序"，...
    # runs[0]="本人所提交的学位论文", runs[1]=""", runs[2]="健身管理小程序", runs[3]=""", runs[4]=rest
    p22 = paras[22]
    if p22.runs and len(p22.runs) >= 3:
        p22.runs[2].text = "{{ title_zh }}"

    print("  Step 2: 原创性声明")

    # ========== Step 3: 中文摘要 ==========
    # P43-45: 摘要正文段落 → 用循环替换
    p43 = paras[43]
    p43_p = p43._p
    p43_p.addprevious(_mk_ctrl("{%p for abs_p in abstract_zh_list %}"))
    if p43.runs:
        p43.runs[0].text = "{{ abs_p }}"
        for r in p43.runs[1:]:
            r.text = ""
    p43_p.addnext(_mk_ctrl("{%p endfor %}"))
    # Clear remaining abstract paragraphs
    for idx in [44, 45]:
        for r in paras[idx].runs:
            r.text = ""

    # P47: 关键词 → runs[0]="关键词：", runs[1]以后是值
    p47 = paras[47]
    if p47.runs and len(p47.runs) >= 2:
        p47.runs[1].text = "{{ keywords_zh }}"
        for r in p47.runs[2:]:
            r.text = ""

    print("  Step 3: 中文摘要")

    # ========== Step 4: 英文摘要 ==========
    # P59-61: 英文摘要正文 → 用循环替换
    p59 = paras[59]
    p59_p = p59._p
    p59_p.addprevious(_mk_ctrl("{%p for abs_p in abstract_en_list %}"))
    if p59.runs:
        p59.runs[0].text = "{{ abs_p }}"
        for r in p59.runs[1:]:
            r.text = ""
    p59_p.addnext(_mk_ctrl("{%p endfor %}"))
    for idx in [60, 61]:
        for r in paras[idx].runs:
            r.text = ""

    # P63: Key words → runs[0]="Key words", runs[1]="：", runs[2+] values
    p63 = paras[63]
    if p63.runs and len(p63.runs) >= 3:
        # Keep "Key words" + "：", replace value runs
        p63.runs[2].text = "{{ keywords_en }}"
        for r in p63.runs[3:]:
            r.text = ""

    print("  Step 4: 英文摘要")

    # ========== Step 4.5: 结论 ==========
    concl_idx = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and "结" in text and "论" in text and not re.match(r'^第\d', text):
            concl_idx = i
            break

    if concl_idx:
        for j in range(concl_idx + 1, min(concl_idx + 20, len(paras))):
            pj = paras[j]
            text = (pj.text or "").strip()
            style = pj.style.name if pj.style else ""
            if "Heading" in style:
                break
            if pj.runs and text:
                pj_p = pj._p
                pj_p.addprevious(_mk_ctrl("{%p for cp in conclusion_list %}"))
                pj.runs[0].text = "{{ cp }}"
                for r in pj.runs[1:]:
                    r.text = ""
                pj_p.addnext(_mk_ctrl("{%p endfor %}"))
                # Clear subsequent conclusion paragraphs
                for k in range(j + 1, min(j + 20, len(paras))):
                    pk = paras[k]
                    pk_style = pk.style.name if pk.style else ""
                    if "Heading" in pk_style:
                        break
                    for r in pk.runs:
                        r.text = ""
                break
        print("  Step 4.5: 结论")

    # ========== Step 5: 致谢 ==========
    ack_idx = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and "致" in text and "谢" in text:
            ack_idx = i
            break

    if ack_idx:
        for j in range(ack_idx + 1, min(ack_idx + 20, len(paras))):
            pj = paras[j]
            text = (pj.text or "").strip()
            style = pj.style.name if pj.style else ""
            if "Heading" in style:
                break
            if pj.runs and text:
                pj.runs[0].text = "{{ acknowledgement }}"
                for r in pj.runs[1:]:
                    r.text = ""
                for k in range(j + 1, min(j + 20, len(paras))):
                    pk = paras[k]
                    pk_style = pk.style.name if pk.style else ""
                    if "Heading" in pk_style:
                        break
                    for r in pk.runs:
                        r.text = ""
                break

    print("  Step 5: 致谢")

    # ========== Step 6: 参考文献 ==========
    ref_idx = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and "参考文献" in text:
            ref_idx = i
            break

    if ref_idx:
        # References in this doc are Normal style without numPr, starting with [1], [2]...
        first_ref = None
        for j in range(ref_idx + 1, min(ref_idx + 5, len(paras))):
            pj = paras[j]
            pt = (pj.text or "").strip()
            if pt and re.match(r'\[\d+\]', pt):
                first_ref = j
                break

        if first_ref:
            pj = paras[first_ref]
            ed_p = pj._p
            ed_p.addprevious(_mk_ctrl("{%p for ref in references %}"))
            if pj.runs:
                pj.runs[0].text = "{{ ref }}"
                for r in pj.runs[1:]:
                    r.text = ""
            ed_p.addnext(_mk_ctrl("{%p endfor %}"))

            # Clear remaining reference entries
            for k in range(first_ref + 1, len(paras)):
                pk = paras[k]
                pk_text = (pk.text or "").strip()
                pk_style = pk.style.name if pk.style else ""
                if "Heading" in pk_style:
                    break
                if pk_text and re.match(r'\[\d+\]', pk_text):
                    for r in pk.runs:
                        r.text = ""

    print("  Step 6: 参考文献")

    doc.save(out_path)

    # ========== Step 7: 正文循环 ==========
    _setup_body(out_path)
    print("  Step 7: 正文循环")

    # ========== Step 8: 删除文本框/AlternateContent ==========
    tmp = out_path + ".tmp"
    with zipfile.ZipFile(out_path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w') as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    xml = data.decode('utf-8')
                    n = xml.count('<mc:AlternateContent')
                    if n > 0:
                        xml = re.sub(
                            r'<mc:AlternateContent>.*?</mc:AlternateContent>',
                            '', xml, flags=re.DOTALL)
                        print(f"  Step 8: 删除 {n} 个文本框")
                    else:
                        print("  Step 8: 无文本框")
                    data = xml.encode('utf-8')
                zout.writestr(item, data)
    shutil.move(tmp, out_path)


def _setup_body(doc_path):
    """设置正文 Jinja2 循环。Heading 1/2 + 标题3 三级标题。"""
    doc = Document(doc_path)
    paras = doc.paragraphs

    # 找正文范围: 第1章 → 结论
    body_start = None
    body_end = len(paras)

    for i, p in enumerate(paras):
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()

        if style == "Heading 1" and re.match(r'^第\d+章', text):
            if body_start is None:
                body_start = i

        # Stop at 结论 or 致谢
        if style == "Heading 1" and "结" in text and "论" in text and not re.match(r'^第\d', text):
            body_end = i
            break
        if style == "Heading 1" and "致" in text and "谢" in text:
            body_end = i
            break

    if not body_start:
        print("  警告: 未找到正文起始")
        return

    # 找样本段落
    h1_idx = body_start
    h2_idx = None
    h3_idx = None
    body_idx = None

    for i in range(body_start + 1, body_end):
        p = paras[i]
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        if not text:
            continue

        if style == "Heading 1":
            continue
        if style == "Heading 2" and h2_idx is None:
            h2_idx = i
        if style in ("标题3", "Heading 3") and h3_idx is None:
            h3_idx = i
        if body_idx is None and style == "Normal" and len(text) > 20:
            body_idx = i

    if not body_idx:
        print(f"  警告: 样本不足 h1={h1_idx} h2={h2_idx} h3={h3_idx} body={body_idx}")
        return

    print(f"    样本: h1=[{h1_idx}] h2=[{h2_idx}] h3=[{h3_idx}] body=[{body_idx}]")

    keep = {h1_idx, h2_idx, h3_idx, body_idx} - {None}

    # 清空非样本段落
    for i in range(body_start, body_end):
        if i in keep:
            continue
        p = paras[i]
        if p._p.find('.//w:sectPr', _NS) is not None:
            continue
        for r in p.runs:
            r.text = ""
        style = p.style.name if p.style else ""
        if style.startswith("Heading") or style.startswith("标题"):
            p.style = doc.styles['Normal']
        if style == "Caption":
            p.style = doc.styles['Normal']
        # Handle 论文正文 style (figure/table captions)
        if style == "论文正文":
            p.style = doc.styles['Normal']

    # 删除正文区域表格
    body_elem = doc.element.body
    tables_rm = []
    pc = 0
    for elem in body_elem:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            pc += 1
        elif tag == 'tbl' and body_start <= pc <= body_end:
            tables_rm.append(elem)
    for tbl in tables_rm:
        body_elem.remove(tbl)
    if tables_rm:
        print(f"    删除 {len(tables_rm)} 个示例表格")

    # 替换样本文本
    def _set_text(para, text):
        if para.runs:
            para.runs[0].text = text
            for r in para.runs[1:]:
                r.text = ""

    paras[h1_idx].paragraph_format.page_break_before = True
    _set_text(paras[h1_idx], "{{ ch.title }}")
    if h2_idx:
        _set_text(paras[h2_idx], "{{ sec.title }}")
    if h3_idx:
        _set_text(paras[h3_idx], "{{ sub.title }}")
    _set_text(paras[body_idx], "{{ para }}")

    h1_p = paras[h1_idx]._p
    h2_p = paras[h2_idx]._p if h2_idx else None
    h3_p = paras[h3_idx]._p if h3_idx else None
    body_p = paras[body_idx]._p
    body_copy = copy.deepcopy(body_p)

    def _mk_body(text):
        p = copy.deepcopy(body_copy)
        for t_elem in p.findall('.//w:t', _NS):
            t_elem.text = ""
        t_elems = p.findall('.//w:t', _NS)
        if t_elems:
            t_elems[0].text = text
        return p

    parent = h1_p.getparent()
    anchor = h1_p.getprevious()

    for elem in [h1_p, h2_p, h3_p, body_p]:
        if elem is not None and elem.getparent() is not None:
            parent.remove(elem)

    cursor = anchor

    def _after(cursor, elem):
        cursor.addnext(elem)
        return elem

    cursor = _after(cursor, _mk_ctrl("{%p for ch in chapters %}"))
    cursor = _after(cursor, h1_p)

    if h2_p is not None:
        cursor = _after(cursor, _mk_ctrl("{%p for sec in ch.sections %}"))
        cursor = _after(cursor, h2_p)
        cursor = _after(cursor, _mk_ctrl("{%p for para in sec.content %}"))
        cursor = _after(cursor, body_p)
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))

        if h3_p is not None:
            cursor = _after(cursor, _mk_ctrl("{%p for sub in sec.subsections %}"))
            cursor = _after(cursor, h3_p)
            cursor = _after(cursor, _mk_ctrl("{%p for p2 in sub.content %}"))
            cursor = _after(cursor, _mk_body("{{ p2 }}"))
            cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))
            cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))

        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))
    else:
        cursor = _after(cursor, _mk_ctrl("{%p for sec in ch.sections %}"))
        cursor = _after(cursor, _mk_ctrl("{%p for para in sec.content %}"))
        cursor = _after(cursor, body_p)
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))

    cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))

    # 清空正文到结论/致谢之间的残留空段落
    _remove_empty_between(doc, body_end)

    doc.save(doc_path)


def _remove_empty_between(doc, body_end):
    """清除正文循环 endfor 和结论/致谢之间的空段落"""
    paras = doc.paragraphs
    last_endfor = None
    target_idx = None

    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        if "{%p endfor %}" in text:
            last_endfor = i
        style = p.style.name if p.style else ""
        if style == "Heading 1" and i >= body_end:
            target_idx = i
            break

    if last_endfor is None or target_idx is None:
        return

    to_remove = []
    for i in range(last_endfor + 1, target_idx):
        p = paras[i]
        text = (p.text or "").strip()
        has_sect = p._p.find('.//w:sectPr', _NS) is not None
        if has_sect:
            if text:
                for r in p.runs:
                    r.text = ""
            continue
        if not text:
            to_remove.append(p._p)

    for elem in to_remove:
        if elem.getparent() is not None:
            elem.getparent().remove(elem)
    if to_remove:
        print(f"    删除 {len(to_remove)} 个空段落")


def _mk_ctrl(text):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python make.py <原始模板.docx> [输出路径]")
        sys.exit(1)
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else \
        os.path.join(os.path.dirname(__file__), "template.docx")
    print(f"制作模板: {src} -> {out}")
    make(src, out)
    print(f"完成: {out}")
