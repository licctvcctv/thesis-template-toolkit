"""
心理学类论文模板制作。
用法: python make.py <原始论文.docx> [输出路径]

模板特点:
- 无封面、无诚信承诺页
- P0: 摘要标题
- P1-P3: 中文摘要正文 → {{ abstract_zh }}
- P4: 关键词：... → 关键词：{{ keywords_zh }}
- P6: Abstract 标题
- P7-P8: 英文摘要正文 → {{ abstract_en }}
- P9: Keywords: ... → Keywords: {{ keywords_en }}
- P11: 目录标题 (TOC Title)
- P12-P63: 目录条目
- P64+: 正文 (Thesis Heading 1 / Thesis Heading 2 / None)
- 致谢 (Thesis Heading 1) → {{ acknowledgement }}
- 参考文献 (Thesis Heading 1) → for/endfor loop
- 只有两级标题: Thesis Heading 1 (章) + Thesis Heading 2 (节/小节)
"""
import os
import sys
import re
import copy
import zipfile
import shutil

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../.."))
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def make(src_path, out_path):
    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc = Document(src_path)

    # ========== Step 1: 中文摘要 ==========
    # P1-P3: 摘要正文 → 合并到 P1
    p1 = doc.paragraphs[1]
    if p1.runs:
        p1.runs[0].text = "{{ abstract_zh }}"
        for r in p1.runs[1:]:
            r.text = ""
    # 清空 P2, P3
    for idx in [2, 3]:
        for r in doc.paragraphs[idx].runs:
            r.text = ""

    # P4: "关键词：..." → 保留 "关键词：" 标签，替换值
    # runs: [关键词：] [] [value]
    p4 = doc.paragraphs[4]
    if p4.runs:
        # Keep run[0] (关键词：), replace run with value onwards
        # Find the run that has the actual keywords value
        for i, r in enumerate(p4.runs):
            if r.text.strip() and '关键词' not in r.text and '：' not in r.text:
                r.text = "{{ keywords_zh }}"
                for rr in p4.runs[i + 1:]:
                    rr.text = ""
                break
        else:
            # Fallback: put it all in run[0]
            full = p4.text
            if "：" in full:
                prefix = full[:full.index("：") + 1]
                p4.runs[0].text = prefix + "{{ keywords_zh }}"
                for r in p4.runs[1:]:
                    r.text = ""

    print("  Step 1: 中文摘要")

    # ========== Step 2: 英文摘要 ==========
    # P7-P8: 英文摘要正文 → 合并到 P7
    p7 = doc.paragraphs[7]
    if p7.runs:
        p7.runs[0].text = "{{ abstract_en }}"
        for r in p7.runs[1:]:
            r.text = ""
    # 清空 P8
    for r in doc.paragraphs[8].runs:
        r.text = ""

    # P9: "Keywords: ..." → 保留 "Keywords: " 标签，替换值
    # runs: [Keywords:] [ ] [value]
    p9 = doc.paragraphs[9]
    if p9.runs:
        for i, r in enumerate(p9.runs):
            if r.text.strip() and 'Keywords' not in r.text and 'keywords' not in r.text:
                r.text = "{{ keywords_en }}"
                for rr in p9.runs[i + 1:]:
                    rr.text = ""
                break
        else:
            p9.runs[0].text = "Keywords: {{ keywords_en }}"
            for r in p9.runs[1:]:
                r.text = ""

    print("  Step 2: 英文摘要")

    # ========== Step 3: 致谢 ==========
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        style = _style_name(p)
        if style == "Thesis Heading 1" and text == "致谢":
            for j in range(i + 1, min(i + 15, len(doc.paragraphs))):
                pj = doc.paragraphs[j]
                sj = _style_name(pj)
                if sj == "Thesis Heading 1":
                    break
                if pj.runs and pj.text.strip():
                    pj.runs[0].text = "    {{ acknowledgement }}"
                    for r in pj.runs[1:]:
                        r.text = ""
                    # 清后续段落到下一个 Heading
                    for k in range(j + 1, min(j + 20, len(doc.paragraphs))):
                        pk = doc.paragraphs[k]
                        pk_style = _style_name(pk)
                        if pk_style == "Thesis Heading 1":
                            break
                        for r in pk.runs:
                            r.text = ""
                    break
            break

    print("  Step 3: 致谢")

    # ========== Step 4: 参考文献 ==========
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        style = _style_name(p)
        if style == "Thesis Heading 1" and "参考文献" in text:
            for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                pj = doc.paragraphs[j]
                pt = (pj.text or "").strip()
                if pj.runs and pt and "[" in pt:
                    ed_p = pj._p
                    # 清除 bookmarkStart/bookmarkEnd/hyperlink 等交叉引用元素
                    _strip_field_refs(ed_p)
                    # for 标签
                    ed_p.addprevious(_mk_ctrl("{%p for ref in references %}"))
                    pj.runs[0].text = "{{ ref }}"
                    for r in pj.runs[1:]:
                        r.text = ""
                    ed_p.addnext(_mk_ctrl("{%p endfor %}"))
                    # 清掉后续参考文献条目 (扩大搜索范围到 25 条)
                    cleaning = False
                    for k in range(j + 1, min(j + 25, len(doc.paragraphs))):
                        pk = doc.paragraphs[k]
                        pk_style = _style_name(pk)
                        if pk_style == "Thesis Heading 1":
                            break
                        pk_text = (pk.text or "").strip()
                        if "{%p endfor" in pk_text:
                            cleaning = True
                            continue
                        if cleaning:
                            _strip_field_refs(pk._p)
                            for r in pk.runs:
                                r.text = ""
                    break
            break

    print("  Step 4: 参考文献")

    doc.save(out_path)

    # ========== Step 5: 正文循环 ==========
    _setup_body(out_path)
    print("  Step 5: 正文循环")

    # ========== Step 6: 删除文本框注释 ==========
    tmp = out_path + ".tmp"
    with zipfile.ZipFile(out_path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w') as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    xml = data.decode('utf-8')
                    n = xml.count('<mc:AlternateContent')
                    xml = re.sub(
                        r'<mc:AlternateContent>.*?</mc:AlternateContent>',
                        '', xml, flags=re.DOTALL)
                    print(f"  Step 6: 删除 {n} 个文本框注释")
                    data = xml.encode('utf-8')
                zout.writestr(item, data)
    shutil.move(tmp, out_path)


def _style_name(para):
    """安全获取段落样式名"""
    if para.style:
        return para.style.name
    return "None"


def _strip_field_refs(p_elem):
    """移除段落中的 bookmarkStart/bookmarkEnd/hyperlink 等交叉引用元素。
    这些元素是文档中的引用标记 (如 [1])，在模板中不需要。"""
    to_remove = []
    for child in p_elem:
        tag = child.tag.split('}')[-1]
        if tag in ('bookmarkStart', 'bookmarkEnd', 'hyperlink'):
            to_remove.append(child)
    for elem in to_remove:
        p_elem.remove(elem)


def _setup_body(doc_path):
    """设置正文循环 (两级: Thesis Heading 1 + Thesis Heading 2)"""
    doc = Document(doc_path)
    paras = doc.paragraphs

    # 定位样本段落
    h1_idx = h2_idx = body_idx = None
    body_end = len(paras)

    # 找章标题和 body_end
    for i, p in enumerate(paras):
        style = _style_name(p)
        text = (p.text or "").strip()
        if style == "Thesis Heading 1":
            if any(kw in text for kw in ["致谢", "致  谢", "参考", "附录"]):
                body_end = i
                break
            if h1_idx is None:
                h1_idx = i

    if not h1_idx:
        print("  警告: 未找到 Thesis Heading 1")
        return

    # 从 h1_idx 之后找 h2/body 样本
    for i in range(h1_idx + 1, body_end):
        p = paras[i]
        style = _style_name(p)
        text = (p.text or "").strip()

        if style == "Thesis Heading 2" and h2_idx is None:
            h2_idx = i

        # 正文段落: None style with substantial text
        if body_idx is None and text and len(text) > 20:
            if style == "None" or style == "Normal":
                body_idx = i

    if not h1_idx or not body_idx:
        print(f"  警告: 样本不足 h1={h1_idx} body={body_idx}")
        return

    keep = {h1_idx, h2_idx, body_idx} - {None}

    # 清空非样本段落 (包括交叉引用元素)
    for i in range(h1_idx, body_end):
        if i in keep:
            continue
        p = paras[i]
        if p._p.find('.//w:sectPr', _NS) is not None:
            continue
        _strip_field_refs(p._p)
        for r in p.runs:
            r.text = ""
        style = _style_name(p)
        if style.startswith("Thesis Heading"):
            # Reset heading to body style (None)
            try:
                p.style = doc.styles['Normal']
            except KeyError:
                pass

    # 删除正文区域示例表格
    body_elem = doc.element.body
    tables_rm = []
    pc = 0
    for elem in body_elem:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            pc += 1
        elif tag == 'tbl' and h1_idx <= pc <= body_end:
            tables_rm.append(elem)
    for tbl in tables_rm:
        body_elem.remove(tbl)

    # 替换样本文本
    def _set_text(para, text):
        if para.runs:
            para.runs[0].text = text
            for r in para.runs[1:]:
                r.text = ""

    paras[h1_idx].paragraph_format.page_break_before = True
    _set_text(paras[h1_idx], "{{ ch.title }}")
    if h2_idx:
        _set_text(paras[h2_idx], "{{ sec.title }}")
    # 清除 body 样本的交叉引用元素
    _strip_field_refs(paras[body_idx]._p)
    _set_text(paras[body_idx], "{{ para }}")

    h1_p = paras[h1_idx]._p
    h2_p = paras[h2_idx]._p if h2_idx else None
    body_p = paras[body_idx]._p

    # 移除样本，按正确顺序重新插入
    parent = h1_p.getparent()
    anchor = h1_p.getprevious()

    for elem in [h1_p, h2_p, body_p]:
        if elem is not None and elem.getparent() is not None:
            parent.remove(elem)

    cursor = anchor

    def _after(cursor, elem):
        cursor.addnext(elem)
        return elem

    # {%p for ch in chapters %}
    #   {{ ch.title }}           ← Thesis Heading 1
    #   {%p for sec in ch.sections %}
    #     {{ sec.title }}        ← Thesis Heading 2
    #     {%p for para in sec.content %}
    #       {{ para }}           ← None (body text)
    #     {%p endfor %}
    #   {%p endfor %}
    # {%p endfor %}
    cursor = _after(cursor, _mk_ctrl("{%p for ch in chapters %}"))
    cursor = _after(cursor, h1_p)

    if h2_p is not None:
        cursor = _after(cursor, _mk_ctrl("{%p for sec in ch.sections %}"))
        cursor = _after(cursor, h2_p)
        cursor = _after(cursor, _mk_ctrl("{%p for para in sec.content %}"))
        cursor = _after(cursor, body_p)
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))  # content
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))  # sections
    else:
        cursor = _after(cursor, _mk_ctrl("{%p for sec in ch.sections %}"))
        cursor = _after(cursor, _mk_ctrl("{%p for para in sec.content %}"))
        cursor = _after(cursor, body_p)
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))

    cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))  # chapters

    # 清正文末尾到文档末尾之间的残留段落 (致谢后、参考文献后)
    for i in range(body_end, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        style = _style_name(p)
        text = (p.text or "").strip()
        if style == "Thesis Heading 1":
            continue
        if "acknowledgement" in text:
            continue
        if "{{ ref" in text or "{%p" in text:
            continue
        _strip_field_refs(p._p)
        for r in p.runs:
            r.text = ""

    # 删除循环区域和致谢之间的空段落
    _remove_empty_between(doc, "{%p endfor %}", "致谢")
    # 删除参考文献 endfor 之后的残留空段落
    _remove_trailing_empty(doc, "参考文献")

    doc.save(doc_path)


def _remove_empty_between(doc, start_marker, end_keyword, start_from_last=False):
    """删除 start_marker 和 end_keyword(Thesis Heading 1) 之间的空段落"""
    paras = doc.paragraphs
    marker_idx = None
    for i, p in enumerate(paras):
        if start_marker in (p.text or ""):
            if start_from_last:
                marker_idx = i
            elif marker_idx is None:
                marker_idx = i

    end_idx = None
    for i, p in enumerate(paras):
        style = _style_name(p)
        text = (p.text or "").strip()
        if style == "Thesis Heading 1" and end_keyword in text and i > (marker_idx or 0):
            end_idx = i
            break

    if marker_idx is None or end_idx is None:
        return

    to_remove = []
    for i in range(marker_idx + 1, end_idx):
        p = paras[i]
        text = (p.text or "").strip()
        style = _style_name(p)
        if style.startswith("Thesis Heading"):
            continue
        if "{%p" in text or "{{" in text:
            continue
        if p._p.find('.//w:sectPr', _NS) is not None:
            continue
        if not text:
            to_remove.append(p._p)

    for elem in to_remove:
        if elem.getparent() is not None:
            elem.getparent().remove(elem)
    if to_remove:
        print(f"    删除 {len(to_remove)} 个空段落 ({start_marker[:20]}→{end_keyword})")


def _remove_trailing_empty(doc, heading_keyword):
    """删除指定 Heading 之后、endfor 之后的所有空段落"""
    paras = doc.paragraphs
    # 找到该 heading
    heading_idx = None
    for i, p in enumerate(paras):
        style = _style_name(p)
        text = (p.text or "").strip()
        if style == "Thesis Heading 1" and heading_keyword in text:
            heading_idx = i
            break
    if heading_idx is None:
        return

    # 找到 endfor 之后
    endfor_idx = None
    for i in range(heading_idx + 1, len(paras)):
        if "{%p endfor" in (paras[i].text or ""):
            endfor_idx = i
            break
    if endfor_idx is None:
        return

    to_remove = []
    for i in range(endfor_idx + 1, len(paras)):
        p = paras[i]
        text = (p.text or "").strip()
        style = _style_name(p)
        if style.startswith("Thesis Heading"):
            break
        if p._p.find('.//w:sectPr', _NS) is not None:
            continue
        if not text:
            to_remove.append(p._p)

    for elem in to_remove:
        if elem.getparent() is not None:
            elem.getparent().remove(elem)
    if to_remove:
        print(f"    删除 {len(to_remove)} 个残留空段落 ({heading_keyword} 之后)")


def _mk_ctrl(text):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python make.py <原始论文.docx> [输出路径]")
        sys.exit(1)
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else "template.docx"
    print(f"制作模板: {src} -> {out}")
    make(src, out)
    print(f"完成: {out}")
