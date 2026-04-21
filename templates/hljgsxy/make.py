"""
黑龙江工商学院软件工程专业毕业设计模板制作。
用法: cd thesis_project && python -m templates.hljgsxy.make <原始模板.docx>

模板结构:
- P0-P21:   说明页 (删除)
- P22-P42:  封面页 (毕业论文题目、姓名、学号等)
- P45-P55:  诚信声明
- P58:      sectPr (目录标题)
- P59-P64:  中文摘要+关键词 [sectPr@P64]
- P65-P70:  英文摘要+关键词
- P71-P114: 目录页 [sectPr@P114]
- P115-P234: 正文五章示例内容
- P235-P239: 结论
- P240-P246: 参考文献
- P247-P250: 致谢
- P252-P254: 附录

关键 style 名称:
  论-一级标题（章）, 论-二级标题（节）, 论-三级标题（条）,
  论-摘要题目, 论-摘要正文, 论-摘要关键字,
  论-结论, 论-参考文献正文, 论-致谢标题,
  论-封面填写信息, 论-封面论文题目
"""
import os
import sys
import copy

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../.."))
from docx import Document
from docx.oxml.ns import qn


def make(src_path, out_path):
    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc = Document(src_path)
    paras = doc.paragraphs

    print(f"原始: {len(paras)} 段落, {len(doc.tables)} 表格")

    # ========== Step 1: 删除说明页 (P0-P21) ==========
    # P20 是校徽图片（论-图片居中样式），需要保留
    removed = 0
    for i in range(21, -1, -1):
        p = paras[i]
        # 保留有sectPr的段落
        pPr = p._p.pPr
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            continue
        # 保留有图片的段落（校徽logo）
        has_drawing = bool(p._p.findall('.//' + qn('w:drawing')))
        if has_drawing:
            continue
        p._p.getparent().remove(p._p)
        removed += 1
    print(f"  Step 1: 删除说明页 {removed} 段")

    # Re-read paragraphs after deletion
    paras = doc.paragraphs
    # 删除第一段（校徽图片段）里残留的分页符（原说明页→封面之间的分页）
    if paras:
        for br in paras[0]._p.findall('.//' + qn('w:br')):
            if br.get(qn('w:type')) == 'page':
                br.getparent().remove(br)
    print(f"  删除后: {len(paras)} 段落")

    # ========== Step 2: 封面 - 替换字段 ==========
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ''
        t = (p.text or '').strip()

        # 封面论文题目 (style=论-封面论文题目)
        if '面向对象设计' in t or ('论文题目' in style):
            if p.runs:
                p.runs[0].text = "{{ title_zh }}"
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.add_run("{{ title_zh }}")

        # 封面填写信息
        if style == '论-封面填写信息':
            if '姓' in t and '名' in t:
                _set_field(p, '姓    名：', '{{ name }}')
            elif '学' in t and '号' in t:
                _set_field(p, '学    号：', '{{ student_id }}')
            elif '学院' in t:
                _set_field(p, '学院(系)：', '{{ department }}')
            elif '专' in t and '业' in t:
                _set_field(p, '专    业：', '{{ major }}')
            elif '年' in t and '级' in t:
                _set_field(p, '年    级：', '{{ grade }}')
            elif '指导教师' in t:
                _set_field(p, '指导教师：', '{{ advisor }}')
            elif '职' in t and '称' in t:
                _set_field(p, '职    称：', '{{ advisor_title }}')

        # 封面日期
        if '年   月   日' in t:
            if p.runs:
                p.runs[0].text = "{{ year }}年 {{ month }}月"
                for r in p.runs[1:]:
                    r.text = ""

    print("  Step 2: 封面字段替换完成")

    # ========== Step 3: 中文摘要 + 关键词 ==========
    # Save 论-摘要正文 formatting for later use
    abstract_body_pPr = None
    abstract_body_rPr = None
    for p in paras:
        style = p.style.name if p.style else ''
        if style == '论-摘要正文' and abstract_body_pPr is None:
            abstract_body_pPr = copy.deepcopy(p._p.pPr) if p._p.pPr is not None else None
            abstract_body_rPr = copy.deepcopy(p.runs[0]._r.find(qn('w:rPr'))) if p.runs else None

    for i, p in enumerate(paras):
        style = p.style.name if p.style else ''
        t = (p.text or '').strip()

        if style == '论-摘要题目' and '摘' in t and 'Abstract' not in t:
            # Keep title as is
            pass

        # 摘要正文 → 段落级循环 + 论-摘要正文 样式
        if style == 'Normal' and '摘要是对毕业设计内容' in t:
            p.runs[0].text = "{%p for para in abstract_zh_list %}"
            for r in p.runs[1:]:
                r.text = ""
            # Insert {{ para }} and endfor after this paragraph
            from docx.oxml import OxmlElement
            # {{ para }} paragraph with 论-摘要正文 style
            p_para = OxmlElement('w:p')
            if abstract_body_pPr is not None:
                p_para.append(copy.deepcopy(abstract_body_pPr))
            r_para = OxmlElement('w:r')
            if abstract_body_rPr is not None:
                r_para.append(copy.deepcopy(abstract_body_rPr))
            t_para = OxmlElement('w:t')
            t_para.set(qn('xml:space'), 'preserve')
            t_para.text = '{{ para }}'
            r_para.append(t_para)
            p_para.append(r_para)
            p._p.addnext(p_para)
            # {%p endfor %} paragraph
            p_end = OxmlElement('w:p')
            r_end = OxmlElement('w:r')
            t_end = OxmlElement('w:t')
            t_end.text = '{%p endfor %}'
            r_end.append(t_end)
            p_end.append(r_end)
            p_para.addnext(p_end)

        # Clear remaining摘要说明段
        if style == 'Normal' and ('摘要内容应在300' in t or '关键词是供检索使用' in t):
            for r in p.runs:
                r.text = ""

        # 关键词
        if style == '论-摘要关键字' and '关键词' in t and 'Key' not in t:
            if p.runs:
                for ri, r in enumerate(p.runs):
                    if '关键词' in r.text:
                        r.text = "关键词  "
                        for r2 in p.runs[ri+1:]:
                            r2.text = ""
                        if ri + 1 < len(p.runs):
                            p.runs[ri+1].text = "{{ keywords_zh }}"
                        else:
                            r.text = "关键词  {{ keywords_zh }}"
                        break

    print("  Step 3: 中文摘要")

    # ========== Step 4: 英文摘要 → 段落级循环 ==========
    # Find the Abstract title first, then look for 论-摘要正文 AFTER it
    paras = doc.paragraphs  # refresh after Step 3 inserted paragraphs
    abstract_title_idx = None
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ''
        t = (p.text or '').strip()
        if style == '论-摘要题目' and 'Abstract' in t:
            abstract_title_idx = i
            break

    en_abstract_done = False
    if abstract_title_idx is not None:
        for i in range(abstract_title_idx + 1, min(abstract_title_idx + 10, len(paras))):
            p = paras[i]
            style = p.style.name if p.style else ''
            t = (p.text or '').strip()
            if style == '论-摘要正文' and not en_abstract_done and t and '{%' not in t:
                # Replace first real 论-摘要正文 with paragraph-level loop
                p.runs[0].text = "{%p for para in abstract_en_list %}"
                for r in p.runs[1:]:
                    r.text = ""
                # Insert {{ para }} and endfor
                from docx.oxml import OxmlElement
                p_para = OxmlElement('w:p')
                if abstract_body_pPr is not None:
                    p_para.append(copy.deepcopy(abstract_body_pPr))
                r_para = OxmlElement('w:r')
                if abstract_body_rPr is not None:
                    r_para.append(copy.deepcopy(abstract_body_rPr))
                t_para = OxmlElement('w:t')
                t_para.set(qn('xml:space'), 'preserve')
                t_para.text = '{{ para }}'
                r_para.append(t_para)
                p_para.append(r_para)
                p._p.addnext(p_para)
                p_end = OxmlElement('w:p')
                r_end = OxmlElement('w:r')
                t_end = OxmlElement('w:t')
                t_end.text = '{%p endfor %}'
                r_end.append(t_end)
                p_end.append(r_end)
                p_para.addnext(p_end)
                en_abstract_done = True
            elif style == '论-摘要正文' and en_abstract_done:
                for r in p.runs:
                    r.text = ""
            if style == '论-摘要关键字' and 'Key' in t:
                if p.runs:
                    p.runs[0].text = "Key words  {{ keywords_en }}"
                    for r in p.runs[1:]:
                        r.text = ""
                # Clear any following 论-摘要关键字 with just punctuation
                for k in range(i + 1, min(i + 3, len(paras))):
                    pk = paras[k]
                    sk = pk.style.name if pk.style else ''
                    tk = (pk.text or '').strip()
                    if sk == '论-摘要关键字' and len(tk) <= 2:
                        for r in pk.runs:
                            r.text = ""
                break

    print("  Step 4: 英文摘要")

    # ========== Step 5: 目录页 - 清空 ==========
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ''
        if style.startswith('toc ') or (style == '论-目录标题' and '目' in (p.text or '')):
            # Keep sectPr paragraphs
            pPr = p._p.pPr
            has_sect = pPr is not None and pPr.find(qn('w:sectPr')) is not None
            if not has_sect:
                for r in p.runs:
                    r.text = ""

    print("  Step 5: 清空目录页")

    # ========== Step 6: 正文区域 → Jinja2循环 ==========
    # 找到正文起始和结束位置
    paras = doc.paragraphs  # refresh

    # Save formatting from original styled paragraphs before overwriting
    ch_heading_pPr = None
    ch_heading_rPr = None
    sec_heading_pPr = None
    sec_heading_rPr = None
    sub_heading_pPr = None
    sub_heading_rPr = None
    body_pPr = None
    body_rPr = None

    for p in paras:
        style = p.style.name if p.style else ''
        t = (p.text or '').strip()
        if style == '论-一级标题（章）' and ch_heading_pPr is None:
            ch_heading_pPr = copy.deepcopy(p._p.pPr) if p._p.pPr is not None else None
            ch_heading_rPr = copy.deepcopy(p.runs[0]._r.find(qn('w:rPr'))) if p.runs else None
        elif style == '论-二级标题（节）' and sec_heading_pPr is None:
            sec_heading_pPr = copy.deepcopy(p._p.pPr) if p._p.pPr is not None else None
            sec_heading_rPr = copy.deepcopy(p.runs[0]._r.find(qn('w:rPr'))) if p.runs else None
        elif style == '论-三级标题（条）' and sub_heading_pPr is None:
            sub_heading_pPr = copy.deepcopy(p._p.pPr) if p._p.pPr is not None else None
            sub_heading_rPr = copy.deepcopy(p.runs[0]._r.find(qn('w:rPr'))) if p.runs else None
        elif style == 'Normal' and body_pPr is None and len(t) > 20:
            # Find a real body text paragraph: must have firstLine indent, not centered
            pPr = p._p.pPr
            if pPr is not None:
                # Skip TOC, centered, colored paragraphs
                pStyle = pPr.find(qn('w:pStyle'))
                jc = pPr.find(qn('w:jc'))
                rPr_node = pPr.find(qn('w:rPr'))
                is_centered = jc is not None and jc.get(qn('w:val')) == 'center'
                is_colored = False
                if rPr_node is not None:
                    color = rPr_node.find(qn('w:color'))
                    if color is not None and color.get(qn('w:val'), '000000') != '000000':
                        is_colored = True
                has_indent = pPr.find(qn('w:ind')) is not None
                if not is_centered and not is_colored and has_indent:
                    body_pPr = copy.deepcopy(pPr)
                    body_rPr = copy.deepcopy(p.runs[0]._r.find(qn('w:rPr'))) if p.runs else None

    # Find the range of body content paragraphs (from first 一级标题 to 结论)
    body_start = None
    body_end = None
    conclusion_idx = None

    for i, p in enumerate(paras):
        style = p.style.name if p.style else ''
        t = (p.text or '').strip()
        if style == '论-一级标题（章）' and body_start is None:
            body_start = i
        if style == '论-结论' and '结' in t and '论' in t and conclusion_idx is None:
            conclusion_idx = i
            body_end = i

    print(f"  正文范围: P{body_start} ~ P{body_end} ({body_end - body_start} 段)")

    # We need 18 paragraphs for the Jinja2 loop. Use body_start to body_start+17.
    # Clear all body content paragraphs except the first 18.

    # First, remove all paragraphs between body_start+18 and body_end (reverse order)
    for i in range(body_end - 1, body_start + 17, -1):
        p = paras[i]
        pPr = p._p.pPr
        has_sect = pPr is not None and pPr.find(qn('w:sectPr')) is not None
        if not has_sect:
            p._p.getparent().remove(p._p)

    # Remove images from remaining paragraphs
    paras = doc.paragraphs
    for i in range(body_start, min(body_start + 18, len(paras))):
        p = paras[i]
        for drawing in p._p.findall('.//' + qn('w:drawing')):
            drawing.getparent().remove(drawing)

    # Remove ALL example tables (they get duplicated by Jinja2 loop)
    while len(doc.tables) > 0:
        tbl = doc.tables[-1]
        tbl._element.getparent().remove(tbl._element)
    print(f"  Step 6a: 删除所有示例表格，剩余 {len(doc.paragraphs)} 段")

    # Now set the 18 Jinja2 tags on body_start to body_start+17
    paras = doc.paragraphs
    tags = [
        (body_start + 0,  "{%p for ch in chapters %}"),
        (body_start + 1,  "{{ ch.title }}"),
        (body_start + 2,  "{%p for item in ch.content %}"),
        (body_start + 3,  "{{ item }}"),
        (body_start + 4,  "{%p endfor %}"),
        (body_start + 5,  "{%p for sec in ch.sections %}"),
        (body_start + 6,  "{{ sec.title }}"),
        (body_start + 7,  "{%p for item in sec.content %}"),
        (body_start + 8,  "{{ item }}"),
        (body_start + 9,  "{%p endfor %}"),
        (body_start + 10, "{%p for sub in sec.subsections %}"),
        (body_start + 11, "{{ sub.title }}"),
        (body_start + 12, "{%p for item in sub.content %}"),
        (body_start + 13, "{{ item }}"),
        (body_start + 14, "{%p endfor %}"),
        (body_start + 15, "{%p endfor %}"),
        (body_start + 16, "{%p endfor %}"),
        (body_start + 17, "{%p endfor %}"),
    ]

    for idx, text in tags:
        if idx >= len(paras):
            break
        p = paras[idx]
        for drawing in p._p.findall('.//' + qn('w:drawing')):
            drawing.getparent().remove(drawing)
        if p.runs:
            p.runs[0].text = text
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(text)

    # Apply correct styles
    def _apply_fmt(para, saved_pPr, saved_rPr):
        if saved_pPr is not None:
            old = para._p.pPr
            if old is not None:
                para._p.remove(old)
            para._p.insert(0, copy.deepcopy(saved_pPr))
        if saved_rPr is not None and para.runs:
            r = para.runs[0]._r
            old = r.find(qn('w:rPr'))
            if old is not None:
                r.remove(old)
            r.insert(0, copy.deepcopy(saved_rPr))

    # ch.title → 一级标题
    _apply_fmt(paras[body_start + 1], ch_heading_pPr, ch_heading_rPr)
    # sec.title → 二级标题
    _apply_fmt(paras[body_start + 6], sec_heading_pPr, sec_heading_rPr)
    # sub.title → 三级标题
    _apply_fmt(paras[body_start + 11], sub_heading_pPr, sub_heading_rPr)
    # item (content) → body text
    for offset in [3, 8, 13]:
        _apply_fmt(paras[body_start + offset], body_pPr, body_rPr)

    print("  Step 6: 正文Jinja2循环")

    # ========== Step 7: 结论 → 段落级循环 ==========
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ''
        t = (p.text or '').strip()
        if style == '论-结论' and '结' in t and '论' in t and '参考' not in t:
            # Keep heading, replace first body paragraph with loop
            for j in range(i + 1, min(i + 4, len(paras))):
                pj = paras[j]
                sj = pj.style.name if pj.style else ''
                if sj in ('论-结论', '论-参考文献正文', '论-致谢标题'):
                    break
                if pj.runs:
                    # Save body formatting
                    conclusion_body_pPr = copy.deepcopy(pj._p.pPr) if pj._p.pPr is not None else None
                    conclusion_body_rPr = copy.deepcopy(pj.runs[0]._r.find(qn('w:rPr'))) if pj.runs else None
                    pj.runs[0].text = "{%p for para in conclusion_list %}"
                    for r in pj.runs[1:]:
                        r.text = ""
                    # Insert {{ para }} with body style
                    from docx.oxml import OxmlElement
                    p_para = OxmlElement('w:p')
                    if conclusion_body_pPr is not None:
                        p_para.append(copy.deepcopy(conclusion_body_pPr))
                    r_para = OxmlElement('w:r')
                    if conclusion_body_rPr is not None:
                        r_para.append(copy.deepcopy(conclusion_body_rPr))
                    t_para = OxmlElement('w:t')
                    t_para.set(qn('xml:space'), 'preserve')
                    t_para.text = '{{ para }}'
                    r_para.append(t_para)
                    p_para.append(r_para)
                    pj._p.addnext(p_para)
                    p_end = OxmlElement('w:p')
                    r_end = OxmlElement('w:r')
                    t_end = OxmlElement('w:t')
                    t_end.text = '{%p endfor %}'
                    r_end.append(t_end)
                    p_end.append(r_end)
                    p_para.addnext(p_end)
                    break
            break

    print("  Step 7: 结论")

    # ========== Step 8: 参考文献 → 段落级循环 ==========
    paras = doc.paragraphs  # refresh
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ''
        t = (p.text or '').strip()
        if style == '论-结论' and '参考文献' in t:
            # Next paragraph should be 论-参考文献正文
            for j in range(i + 1, min(i + 8, len(paras))):
                pj = paras[j]
                sj = pj.style.name if pj.style else ''
                if sj == '论-参考文献正文':
                    # Save formatting
                    ref_pPr = copy.deepcopy(pj._p.pPr) if pj._p.pPr is not None else None
                    ref_rPr = copy.deepcopy(pj.runs[0]._r.find(qn('w:rPr'))) if pj.runs else None
                    pj.runs[0].text = "{%p for ref in references %}"
                    for r in pj.runs[1:]:
                        r.text = ""
                    # Insert {{ ref }} with reference style
                    from docx.oxml import OxmlElement
                    p_ref = OxmlElement('w:p')
                    if ref_pPr is not None:
                        p_ref.append(copy.deepcopy(ref_pPr))
                    r_ref = OxmlElement('w:r')
                    if ref_rPr is not None:
                        r_ref.append(copy.deepcopy(ref_rPr))
                    t_ref = OxmlElement('w:t')
                    t_ref.set(qn('xml:space'), 'preserve')
                    t_ref.text = '{{ ref }}'
                    r_ref.append(t_ref)
                    p_ref.append(r_ref)
                    pj._p.addnext(p_ref)
                    p_end = OxmlElement('w:p')
                    r_end = OxmlElement('w:r')
                    t_end = OxmlElement('w:t')
                    t_end.text = '{%p endfor %}'
                    r_end.append(t_end)
                    p_end.append(r_end)
                    p_ref.addnext(p_end)
                    # Clear remaining reference paragraphs
                    for k in range(j + 1, min(j + 10, len(paras))):
                        pk = paras[k]
                        sk = pk.style.name if pk.style else ''
                        if sk == '论-参考文献正文':
                            for r in pk.runs:
                                r.text = ""
                        else:
                            break
                    break
            break

    print("  Step 8: 参考文献")

    # ========== Step 9: 致谢 → 段落级循环 ==========
    paras = doc.paragraphs  # refresh
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ''
        t = (p.text or '').strip()
        if style == '论-致谢标题' and '致' in t:
            for j in range(i + 1, min(i + 5, len(paras))):
                pj = paras[j]
                sj = pj.style.name if pj.style else ''
                if sj in ('论-附录标题',):
                    break
                tj = (pj.text or '').strip()
                if tj and '不建议' not in tj and '排版要求' not in tj:
                    # Save formatting
                    ack_pPr = copy.deepcopy(pj._p.pPr) if pj._p.pPr is not None else None
                    ack_rPr = copy.deepcopy(pj.runs[0]._r.find(qn('w:rPr'))) if pj.runs else None
                    pj.runs[0].text = "{%p for para in acknowledgement_list %}"
                    for r in pj.runs[1:]:
                        r.text = ""
                    # Insert {{ para }} and endfor
                    from docx.oxml import OxmlElement
                    p_para = OxmlElement('w:p')
                    if ack_pPr is not None:
                        p_para.append(copy.deepcopy(ack_pPr))
                    r_para = OxmlElement('w:r')
                    if ack_rPr is not None:
                        r_para.append(copy.deepcopy(ack_rPr))
                    t_para = OxmlElement('w:t')
                    t_para.set(qn('xml:space'), 'preserve')
                    t_para.text = '{{ para }}'
                    r_para.append(t_para)
                    p_para.append(r_para)
                    pj._p.addnext(p_para)
                    p_end = OxmlElement('w:p')
                    r_end = OxmlElement('w:r')
                    t_end = OxmlElement('w:t')
                    t_end.text = '{%p endfor %}'
                    r_end.append(t_end)
                    p_end.append(r_end)
                    p_para.addnext(p_end)
                    break
                elif tj:
                    for r in pj.runs:
                        r.text = ""
            break

    # Clear remaining致谢说明段
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ''
        t = (p.text or '').strip()
        if style == '论-致谢标题' and '致' in t:
            for j in range(i + 1, min(i + 8, len(paras))):
                pj = paras[j]
                tj = (pj.text or '').strip()
                if '不建议' in tj or '排版要求' in tj:
                    for r in pj.runs:
                        r.text = ""
            break

    print("  Step 9: 致谢")

    # ========== Step 10: 清理空段落 + 残留说明文字 ==========
    # IMPORTANT: 只清理"摘要题目"之后的空段落。
    # 封面和诚信声明区域的空段落是Word排版间距，不能删。
    paras = doc.paragraphs
    removed = 0
    instruction_keywords = ['毕业设计的结论是', '排版要求', '不建议', '可以没有附录', '如果没有', '结论要单独成页', '概括说明设计']

    # 找到第一个"摘要题目"的位置，之前的段落全部保留
    body_start = len(paras)
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ''
        if style == '论-摘要题目':
            body_start = i
            break

    for i in range(len(paras) - 1, -1, -1):
        # 封面+诚信声明区域的段落不动
        if i < body_start:
            continue

        p = paras[i]
        text = (p.text or "").strip()
        style = p.style.name if p.style else ""
        pPr = p._p.pPr
        has_sect = pPr is not None and pPr.find(qn('w:sectPr')) is not None

        if has_sect:
            continue

        # Remove residual instruction paragraphs
        if text and any(kw in text for kw in instruction_keywords):
            p._p.getparent().remove(p._p)
            removed += 1
            continue

        if text:
            continue
        # Skip structural styles
        if style in ('论-一级标题（章）', '论-摘要题目', '论-结论', '论-致谢标题', '论-附录标题'):
            continue

        # Remove truly empty paragraphs in cleared sections
        has_drawing = bool(p._p.findall('.//' + qn('w:drawing')))
        if not has_drawing:
            p._p.getparent().remove(p._p)
            removed += 1

    # The generated papers in this project do not render appendix content.
    # Keep the thesis clean by removing the empty appendix title paragraph.
    for p in list(doc.paragraphs):
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        if style == '论-附录标题' and text == '附录':
            p._p.getparent().remove(p._p)
            removed += 1

    print(f"  Step 10: 删除 {removed} 个空段落/说明段")

    # ========== Save ==========
    doc.save(out_path)

    # Verify
    doc2 = Document(out_path)
    p_count = len(doc2.paragraphs)
    t_count = len(doc2.tables)
    print(f"\n结果: {p_count} 段落, {t_count} 表格")
    for i, p in enumerate(doc2.paragraphs):
        t = (p.text or "").strip()
        if t:
            print(f"  P{i:2d} [{p.style.name:25s}] {t[:80]}")

    full = "\n".join(p.text or "" for p in doc2.paragraphs)
    for tag in ["{%", "%}", "{{", "}}"]:
        count = full.count(tag)
        if count:
            print(f"  Jinja2 '{tag}': {count}")


def _set_field(p, label, value):
    """Set a cover page field: keep label, replace value."""
    if p.runs:
        # Find the run with the label
        for ri, r in enumerate(p.runs):
            if any(c in r.text for c in label[:2]):
                r.text = label
                if ri + 1 < len(p.runs):
                    p.runs[ri + 1].text = value
                    for r2 in p.runs[ri + 2:]:
                        r2.text = ""
                else:
                    r.text = label + value
                return
        # Fallback: set first run
        p.runs[0].text = label + value
        for r in p.runs[1:]:
            r.text = ""


def main():
    if len(sys.argv) < 2:
        print("用法: python -m templates.hljgsxy.make <原始模板.docx>")
        sys.exit(1)

    src = sys.argv[1]
    here = os.path.dirname(os.path.abspath(__file__))
    out = os.path.join(here, "template.docx")

    print(f"输入: {src}")
    print(f"输出: {out}")
    make(src, out)
    print("完成!")


if __name__ == "__main__":
    main()
