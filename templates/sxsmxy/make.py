"""
陕西国际商贸学院本科毕业论文模板制作。
用法: cd thesis_project && python -m templates.sxsmxy.make <原始模板.docx>

模板特点:
- 封面: 段落式（非表格），P0 附件2, P2 届次, P6 标题, P10-P14 信息, P21 日期
- 声明: P23-P37
- 中文摘要: P38 标题, P39-P41 正文, P43 关键词
- 英文摘要: P54 标题, P55-P57 正文, P59 关键词
- 正文: P75 起，Thesis Heading 1 为章标题，
       二级/三级标题均为 Normal + run-level sz=177800
- 致谢: P297
- 参考文献: P308 起，带 numPr 编号
- 无目录页（或目录在分节符中间被跳过）
"""
import os
import sys
import re
import copy
import zipfile
import shutil

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../.."))
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def make(src_path, out_path):
    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc = Document(src_path)
    paras = doc.paragraphs

    # ========== Step 1: 封面 ==========
    # P2: "2026 届本科生毕业论文" → year
    p2 = paras[2]
    if p2.runs and len(p2.runs) >= 2:
        # runs: ['  ', '2026', ' ', '届本科生毕业论文']
        p2.runs[1].text = "{{ year }}"

    # P6: 标题
    p6 = paras[6]
    if p6.runs:
        p6.runs[0].text = "{{ title_zh }}"
        for r in p6.runs[1:]:
            r.text = ""

    # P10: 学生姓名 - runs: [学生][姓名][：][spaces_underline][space_underline]
    _replace_cover_field(paras[10], "{{ name }}")
    # P11: 学号
    _replace_cover_field(paras[11], "{{ student_id }}")
    # P12: 学院
    _replace_cover_field(paras[12], "{{ college }}")
    # P13: 专业
    _replace_cover_field(paras[13], "{{ major }}")
    # P14: 指导教师 - runs: [指导教师：][校内导师姓名  职称]
    p14 = paras[14]
    if p14.runs and len(p14.runs) >= 2:
        p14.runs[1].text = "{{ advisor }}"
        for r in p14.runs[2:]:
            r.text = ""

    # P21: 日期 "20XX年X月"
    p21 = paras[21]
    if p21.runs:
        p21.runs[0].text = "{{ finish_date }}"
        for r in p21.runs[1:]:
            r.text = ""

    print("  Step 1: 封面")

    # ========== Step 2: 声明页 ==========
    # P25: 声明中标题引用 → {{ title_zh }}
    p25 = paras[25]
    if p25.runs:
        # runs contain: ...《[空白underline]...》...
        # runs[1]-[3] are underlined blanks for title
        for j, r in enumerate(p25.runs):
            if r.text.strip() == '' and r.font.underline:
                r.text = "{{ title_zh }}"
                # clear remaining underline runs
                for k in range(j + 1, len(p25.runs)):
                    if p25.runs[k].text.strip() == '' and p25.runs[k].font.underline:
                        p25.runs[k].text = ""
                    else:
                        break
                break

    print("  Step 2: 声明页")

    # ========== Step 3: 中文摘要 ==========
    # P39: 摘要正文样本 → 改为循环
    # P40-P41: 清空（循环会自动生成多段）
    p39 = paras[39]
    p39_p = p39._p
    p39_p.addprevious(_mk_ctrl("{%p for abs_p in abstract_zh_list %}"))
    if p39.runs:
        p39.runs[0].text = "{{ abs_p }}"
        for r in p39.runs[1:]:
            r.text = ""
    p39_p.addnext(_mk_ctrl("{%p endfor %}"))
    for idx in [40, 41]:
        for r in paras[idx].runs:
            r.text = ""

    # P43: 关键词 → 保留"关键词："标签，替换值
    p43 = paras[43]
    if p43.runs and len(p43.runs) >= 2:
        # runs[0]: "关键词：" (黑体), runs[1]: value (宋体)
        p43.runs[1].text = "{{ keywords_zh }}"
        for r in p43.runs[2:]:
            r.text = ""

    print("  Step 3: 中文摘要")

    # ========== Step 4: 英文摘要 ==========
    # P55: 英文摘要样本 → 改为循环
    # P56-P57: 清空
    p55 = paras[55]
    p55_p = p55._p
    p55_p.addprevious(_mk_ctrl("{%p for abs_p in abstract_en_list %}"))
    if p55.runs:
        p55.runs[0].text = "{{ abs_p }}"
        for r in p55.runs[1:]:
            r.text = ""
    p55_p.addnext(_mk_ctrl("{%p endfor %}"))
    for idx in [56, 57]:
        for r in paras[idx].runs:
            r.text = ""

    # P59: Keywords → runs[0]: "Keywords: " (bold), runs[1]: value
    p59 = paras[59]
    if p59.runs and len(p59.runs) >= 2:
        p59.runs[1].text = "{{ keywords_en }}"
        for r in p59.runs[2:]:
            r.text = ""

    print("  Step 4: 英文摘要")

    # ========== Step 5: 致谢 ==========
    ack_idx = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        if text in ("致  谢", "致 谢", "致谢"):
            ack_idx = i
            break

    if ack_idx:
        # 致谢正文从 ack_idx+1 开始
        for j in range(ack_idx + 1, min(ack_idx + 20, len(paras))):
            pj = paras[j]
            text = (pj.text or "").strip()
            if not text:
                continue
            # 遇到 "参考文献" 停止
            if "参考文献" in text:
                break
            if pj.runs:
                pj.runs[0].text = "{{ acknowledgement }}"
                for r in pj.runs[1:]:
                    r.text = ""
                # 清后续致谢段落
                for k in range(j + 1, min(j + 20, len(paras))):
                    pk = paras[k]
                    pk_text = (pk.text or "").strip()
                    if not pk_text:
                        continue
                    if "参考文献" in pk_text:
                        break
                    if pk._p.find('.//w:sectPr', _NS) is not None:
                        continue
                    for r in pk.runs:
                        r.text = ""
                break

    print("  Step 5: 致谢")

    # ========== Step 6: 参考文献 ==========
    ref_idx = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        if text == "参考文献":
            ref_idx = i
            break

    if ref_idx:
        first_ref = None
        for j in range(ref_idx + 1, min(ref_idx + 5, len(paras))):
            pj = paras[j]
            pt = (pj.text or "").strip()
            has_num = pj._p.find('.//w:numPr', _NS) is not None
            if pt and has_num:
                first_ref = j
                break

        if first_ref:
            pj = paras[first_ref]
            ed_p = pj._p
            ed_p.addprevious(_mk_ctrl("{%p for ref in references %}"))
            # 参考文献条目有多个空 run + 正文 run，找到正文 run
            for r in pj.runs:
                if r.text.strip():
                    r.text = "{{ ref }}"
                    break
            # 清空其余 run
            found_var = False
            for r in pj.runs:
                if "{{ ref }}" in r.text:
                    found_var = True
                    continue
                if found_var:
                    r.text = ""
            ed_p.addnext(_mk_ctrl("{%p endfor %}"))

            # 清掉后续参考文献条目（first_ref+1 到文档末尾的所有带 numPr 的段落）
            for k in range(first_ref + 1, len(paras)):
                pk = paras[k]
                pk_text = (pk.text or "").strip()
                has_num = pk._p.find('.//w:numPr', _NS) is not None
                if has_num and pk_text:
                    for r in pk.runs:
                        r.text = ""

    print("  Step 6: 参考文献")

    doc.save(out_path)

    # ========== Step 7: 正文循环 ==========
    _setup_body(out_path)
    print("  Step 7: 正文循环")

    # ========== Step 8: 删除文本框/AlternateContent ==========
    tmp = out_path + ".tmp"
    with zipfile.ZipFile(out_path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w') as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    xml = data.decode('utf-8')
                    n = xml.count('<mc:AlternateContent')
                    if n > 0:
                        xml = re.sub(
                            r'<mc:AlternateContent>.*?</mc:AlternateContent>',
                            '', xml, flags=re.DOTALL)
                        print(f"  Step 8: 删除 {n} 个文本框注释")
                    else:
                        print("  Step 8: 无文本框注释")
                    data = xml.encode('utf-8')
                zout.writestr(item, data)
    shutil.move(tmp, out_path)


def _replace_cover_field(para, var_text):
    """替换封面字段：保留标签 runs，将下划线空白 runs 替换为模板变量。"""
    runs = para.runs
    if not runs:
        return
    for j, r in enumerate(runs):
        if r.text.strip() == '' and r.font.underline:
            r.text = var_text
            for k in range(j + 1, len(runs)):
                runs[k].text = ""
            return
    # fallback: 替换最后一个 run
    runs[-1].text = var_text


def _setup_body(doc_path):
    """设置正文 Jinja2 循环。

    此模板特点:
    - 章标题: Thesis Heading 1 (e.g. "第1章 绪论")
    - 二级标题: Normal + run sz=177800 (e.g. "1.1 研究背景与意义")
    - 三级标题: Normal + run sz=177800 (e.g. "5.1.1 用户登录与注册")
    - 正文: Normal + run sz=152400
    - 图标题: Normal + run sz=133350
    """
    doc = Document(doc_path)
    paras = doc.paragraphs

    # 找正文范围
    body_start = None
    body_end = len(paras)
    ack_idx = None

    for i, p in enumerate(paras):
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()

        if style == "Thesis Heading 1":
            if re.match(r'^第\d+章', text):
                if body_start is None:
                    body_start = i

        if text in ("致  谢", "致 谢", "致谢"):
            ack_idx = i
            body_end = i
            break

    if not body_start:
        print("  警告: 未找到正文起始")
        return

    # 找样本段落
    h1_idx = body_start  # Thesis Heading 1
    h2_idx = None  # Normal, sz=177800, pattern x.x
    h3_idx = None  # Normal, sz=177800, pattern x.x.x
    body_idx = None  # Normal, sz=152400, body text
    caption_idx = None  # Normal, sz=133350, figure/table caption

    for i in range(body_start + 1, body_end):
        p = paras[i]
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        if not text:
            continue
        if style == "Thesis Heading 1":
            continue

        # Check run-level font size
        sz = None
        if p.runs:
            sz = p.runs[0].font.size

        if sz == 177800:
            # 二级或三级标题
            if re.match(r'^\d+\.\d+\.\d+', text) and h3_idx is None:
                h3_idx = i
            elif re.match(r'^\d+\.\d+\s', text) and h2_idx is None:
                h2_idx = i
        elif body_idx is None and len(text) > 20:
            if style == "Normal" and sz and sz <= 152400:
                body_idx = i

    if not body_idx:
        print(f"  警告: 样本不足 h1={h1_idx} h2={h2_idx} body={body_idx}")
        return

    print(f"    样本: h1=[{h1_idx}] h2=[{h2_idx}] h3=[{h3_idx}] body=[{body_idx}]")

    keep = {h1_idx, h2_idx, h3_idx, body_idx} - {None}

    # 清空非样本段落
    for i in range(body_start, body_end):
        if i in keep:
            continue
        p = paras[i]
        if p._p.find('.//w:sectPr', _NS) is not None:
            continue
        for r in p.runs:
            r.text = ""
        style = p.style.name if p.style else ""
        if style == "Thesis Heading 1":
            p.style = doc.styles['Normal']
        if style == "Thesis Heading 2":
            p.style = doc.styles['Normal']

    # 删除正文区域表格
    body_elem = doc.element.body
    tables_rm = []
    pc = 0
    for elem in body_elem:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            pc += 1
        elif tag == 'tbl' and body_start <= pc <= body_end:
            tables_rm.append(elem)
    for tbl in tables_rm:
        body_elem.remove(tbl)
    if tables_rm:
        print(f"    删除 {len(tables_rm)} 个示例表格")

    # 替换样本文本
    def _set_text(para, text):
        if para.runs:
            para.runs[0].text = text
            for r in para.runs[1:]:
                r.text = ""

    paras[h1_idx].paragraph_format.page_break_before = True
    _set_text(paras[h1_idx], "{{ ch.title }}")
    if h2_idx:
        _set_text(paras[h2_idx], "{{ sec.title }}")
    if h3_idx:
        _set_text(paras[h3_idx], "{{ sub.title }}")
    _set_text(paras[body_idx], "{{ para }}")

    h1_p = paras[h1_idx]._p
    h2_p = paras[h2_idx]._p if h2_idx else None
    h3_p = paras[h3_idx]._p if h3_idx else None
    body_p = paras[body_idx]._p
    body_copy = copy.deepcopy(body_p)

    def _mk_body(text):
        p = copy.deepcopy(body_copy)
        for t_elem in p.findall('.//w:t', _NS):
            t_elem.text = ""
        t_elems = p.findall('.//w:t', _NS)
        if t_elems:
            t_elems[0].text = text
        return p

    # 移除样本，按正确顺序重新插入
    parent = h1_p.getparent()
    anchor = h1_p.getprevious()

    for elem in [h1_p, h2_p, h3_p, body_p]:
        if elem is not None and elem.getparent() is not None:
            parent.remove(elem)

    cursor = anchor

    def _after(cursor, elem):
        cursor.addnext(elem)
        return elem

    cursor = _after(cursor, _mk_ctrl("{%p for ch in chapters %}"))
    cursor = _after(cursor, h1_p)

    if h2_p is not None:
        cursor = _after(cursor, _mk_ctrl("{%p for sec in ch.sections %}"))
        cursor = _after(cursor, h2_p)
        cursor = _after(cursor, _mk_ctrl("{%p for para in sec.content %}"))
        cursor = _after(cursor, body_p)
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))  # content

        if h3_p is not None:
            cursor = _after(cursor, _mk_ctrl("{%p for sub in sec.subsections %}"))
            cursor = _after(cursor, h3_p)
            cursor = _after(cursor, _mk_ctrl("{%p for p2 in sub.content %}"))
            cursor = _after(cursor, _mk_body("{{ p2 }}"))
            cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))  # sub.content
            cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))  # subsections

        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))  # sections
    else:
        cursor = _after(cursor, _mk_ctrl("{%p for sec in ch.sections %}"))
        cursor = _after(cursor, _mk_ctrl("{%p for para in sec.content %}"))
        cursor = _after(cursor, body_p)
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))

    cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))  # chapters

    # 清正文末尾到致谢之间的残留空段落
    _remove_empty_between(doc, body_end)

    doc.save(doc_path)


def _remove_empty_between(doc, body_end):
    """清除正文循环 endfor 和致谢之间的空段落"""
    paras = doc.paragraphs
    last_endfor = None
    ack_idx = None

    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        if "{%p endfor %}" in text:
            last_endfor = i
        if text in ("致  谢", "致 谢", "致谢"):
            ack_idx = i
            break

    if last_endfor is None or ack_idx is None:
        return

    to_remove = []
    for i in range(last_endfor + 1, ack_idx):
        p = paras[i]
        text = (p.text or "").strip()
        if p._p.find('.//w:sectPr', _NS) is not None:
            continue
        if not text:
            to_remove.append(p._p)

    for elem in to_remove:
        if elem.getparent() is not None:
            elem.getparent().remove(elem)
    if to_remove:
        print(f"    删除 {len(to_remove)} 个空段落 (endfor→致谢)")


def _mk_ctrl(text):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python make.py <原始模板.docx> [输出路径]")
        sys.exit(1)
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else \
        os.path.join(os.path.dirname(__file__), "template.docx")
    print(f"制作模板: {src} -> {out}")
    make(src, out)
    print(f"完成: {out}")
