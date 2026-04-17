"""
电商会员数据化运营系统论文模板制作（大理学院/定会风格）。
用法: cd thesis_project && python -m templates.dslg.make <原始模板.docx>

模板特点:
- 封面1(中文): P17-18 标题, P20-24 学院/班级/姓名/学号/导师, P27 日期
- 封面2(英文): P36-39 英文标题
- 摘要: Heading 1 "摘 要", P54-55 正文, P57 关键词
- Abstract: Heading 1, P72-73 正文, P75 Key words
- 正文: Heading 1/2/3 三级标题
- 致谢: Heading 1 "致 谢"
- 参考文献: Heading 1, List Paragraph + numPr
"""
import os
import sys
import re
import copy
import zipfile
import shutil

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../.."))
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def make(src_path, out_path):
    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc = Document(src_path)
    paras = doc.paragraphs

    # ========== Step 1: 封面(中文) ==========
    # P17: 标题行1 + P18: 标题行2 → 合并为 {{ title_zh }}
    p17 = paras[17]
    if p17.runs:
        p17.runs[0].text = "{{ title_zh }}"
        for r in p17.runs[1:]:
            r.text = ""
    for r in paras[18].runs:
        r.text = ""

    # P20-P24: 学院/班级/姓名/学号/导师
    cover_fields = {
        20: ("学    院：", "{{ college }}"),
        21: ("专业班级：", "{{ class_name }}"),
        22: ("姓    名：", "{{ name }}"),
        23: ("学    号：", "{{ student_id }}"),
        24: ("指导教师：", "{{ advisor }}"),
    }
    for idx, (label, var) in cover_fields.items():
        p = paras[idx]
        if p.runs:
            # 找到冒号后面的内容替换
            full = p.text
            colon_pos = full.find("：")
            if colon_pos >= 0:
                p.runs[0].text = full[:colon_pos + 1] + var
                for r in p.runs[1:]:
                    r.text = ""

    # P27: 日期 "2023年5月"
    p27 = paras[27]
    if p27.runs:
        p27.runs[0].text = "{{ finish_date }}"
        for r in p27.runs[1:]:
            r.text = ""

    print("  Step 1: 封面(中文)")

    # ========== Step 2: 封面(英文) ==========
    # P36-37: 英文标题 → P38-39: 英文标题第二部分
    p36 = paras[36]
    if p36.runs:
        p36.runs[0].text = "{{ title_en_line1 }}"
        for r in p36.runs[1:]:
            r.text = ""
    for r in paras[37].runs:
        r.text = ""
    p38 = paras[38]
    if p38.runs:
        p38.runs[0].text = "{{ title_en_line2 }}"
        for r in p38.runs[1:]:
            r.text = ""
    for r in paras[39].runs:
        r.text = ""

    print("  Step 2: 封面(英文)")

    # ========== Step 3: 中文摘要 ==========
    # P54-55: 摘要正文 → 循环
    p54 = paras[54]
    p54_p = p54._p
    p54_p.addprevious(_mk_ctrl("{%p for abs_p in abstract_zh_list %}"))
    if p54.runs:
        p54.runs[0].text = "{{ abs_p }}"
        for r in p54.runs[1:]:
            r.text = ""
    p54_p.addnext(_mk_ctrl("{%p endfor %}"))
    for r in paras[55].runs:
        r.text = ""

    # P57: 关键词
    p57 = paras[57]
    if p57.runs and len(p57.runs) >= 2:
        p57.runs[1].text = "{{ keywords_zh }}"
        for r in p57.runs[2:]:
            r.text = ""

    print("  Step 3: 中文摘要")

    # ========== Step 4: 英文摘要 ==========
    # P72-73: 英文摘要正文 → 循环
    p72 = paras[72]
    p72_p = p72._p
    p72_p.addprevious(_mk_ctrl("{%p for abs_p in abstract_en_list %}"))
    if p72.runs:
        p72.runs[0].text = "{{ abs_p }}"
        for r in p72.runs[1:]:
            r.text = ""
    p72_p.addnext(_mk_ctrl("{%p endfor %}"))
    for r in paras[73].runs:
        r.text = ""

    # P75: Key words
    p75 = paras[75]
    if p75.runs and len(p75.runs) >= 2:
        p75.runs[1].text = " {{ keywords_en }}"
        for r in p75.runs[2:]:
            r.text = ""

    print("  Step 4: 英文摘要")

    # ========== Step 4.5: 结论 ==========
    concl_idx = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and "结" in text and "论" in text and not re.match(r'^\d', text):
            concl_idx = i
            break

    if concl_idx:
        # 找结论正文段落 → 用循环替换
        for j in range(concl_idx + 1, min(concl_idx + 20, len(paras))):
            pj = paras[j]
            text = (pj.text or "").strip()
            style = pj.style.name if pj.style else ""
            if "Heading" in style:
                break
            if pj.runs and text:
                pj_p = pj._p
                pj_p.addprevious(_mk_ctrl("{%p for cp in conclusion_list %}"))
                pj.runs[0].text = "{{ cp }}"
                for r in pj.runs[1:]:
                    r.text = ""
                pj_p.addnext(_mk_ctrl("{%p endfor %}"))
                # 清后续结论段落
                for k in range(j + 1, min(j + 20, len(paras))):
                    pk = paras[k]
                    pk_style = pk.style.name if pk.style else ""
                    if "Heading" in pk_style:
                        break
                    for r in pk.runs:
                        r.text = ""
                break
        print("  Step 4.5: 结论")

    # ========== Step 5: 致谢 ==========
    ack_idx = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and "致" in text and "谢" in text:
            ack_idx = i
            break

    if ack_idx:
        for j in range(ack_idx + 1, min(ack_idx + 20, len(paras))):
            pj = paras[j]
            text = (pj.text or "").strip()
            style = pj.style.name if pj.style else ""
            if "Heading" in style:
                break
            if pj.runs and text:
                pj.runs[0].text = "{{ acknowledgement }}"
                for r in pj.runs[1:]:
                    r.text = ""
                for k in range(j + 1, min(j + 20, len(paras))):
                    pk = paras[k]
                    pk_style = pk.style.name if pk.style else ""
                    if "Heading" in pk_style:
                        break
                    for r in pk.runs:
                        r.text = ""
                break

    print("  Step 5: 致谢")

    # ========== Step 6: 参考文献 ==========
    ref_idx = None
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        style = p.style.name if p.style else ""
        if style == "Heading 1" and "参考文献" in text:
            ref_idx = i
            break

    if ref_idx:
        first_ref = None
        for j in range(ref_idx + 1, min(ref_idx + 5, len(paras))):
            pj = paras[j]
            pt = (pj.text or "").strip()
            has_num = pj._p.find('.//w:numPr', _NS) is not None
            if pt and has_num:
                first_ref = j
                break

        if first_ref:
            pj = paras[first_ref]
            ed_p = pj._p
            ed_p.addprevious(_mk_ctrl("{%p for ref in references %}"))
            if pj.runs:
                pj.runs[0].text = "{{ ref }}"
                for r in pj.runs[1:]:
                    r.text = ""
            ed_p.addnext(_mk_ctrl("{%p endfor %}"))

            # 清掉后续参考文献条目
            for k in range(first_ref + 1, len(paras)):
                pk = paras[k]
                has_num = pk._p.find('.//w:numPr', _NS) is not None
                pk_text = (pk.text or "").strip()
                if has_num and pk_text:
                    for r in pk.runs:
                        r.text = ""
                elif pk_text and not has_num:
                    break

    print("  Step 6: 参考文献")

    doc.save(out_path)

    # ========== Step 7: 正文循环 ==========
    _setup_body(out_path)
    print("  Step 7: 正文循环")

    # ========== Step 8: 删除文本框 ==========
    tmp = out_path + ".tmp"
    with zipfile.ZipFile(out_path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w') as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    xml = data.decode('utf-8')
                    n = xml.count('<mc:AlternateContent')
                    if n > 0:
                        xml = re.sub(
                            r'<mc:AlternateContent>.*?</mc:AlternateContent>',
                            '', xml, flags=re.DOTALL)
                        print(f"  Step 8: 删除 {n} 个文本框")
                    else:
                        print("  Step 8: 无文本框")
                    data = xml.encode('utf-8')
                zout.writestr(item, data)
    shutil.move(tmp, out_path)


def _setup_body(doc_path):
    """设置正文 Jinja2 循环。Heading 1/2/3 三级标题。"""
    doc = Document(doc_path)
    paras = doc.paragraphs

    # 找正文范围
    body_start = None
    body_end = len(paras)

    for i, p in enumerate(paras):
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()

        if style == "Heading 1" and re.match(r'^\d+\s', text):
            if body_start is None:
                body_start = i

        if style == "Heading 1" and ("致" in text and "谢" in text):
            body_end = i
            break
        if style == "Heading 1" and "结" in text and "论" in text and not re.match(r'^\d', text):
            body_end = i
            break

    if not body_start:
        print("  警告: 未找到正文起始")
        return

    # 找样本段落
    h1_idx = body_start
    h2_idx = None
    h3_idx = None
    body_idx = None

    for i in range(body_start + 1, body_end):
        p = paras[i]
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        if not text:
            continue

        if style == "Heading 1":
            continue
        if style == "Heading 2" and h2_idx is None:
            h2_idx = i
        if style == "Heading 3" and h3_idx is None:
            h3_idx = i
        if body_idx is None and style in ("Normal", "List Paragraph") and len(text) > 20:
            body_idx = i

    if not body_idx:
        print(f"  警告: 样本不足 h1={h1_idx} h2={h2_idx} body={body_idx}")
        return

    print(f"    样本: h1=[{h1_idx}] h2=[{h2_idx}] h3=[{h3_idx}] body=[{body_idx}]")

    keep = {h1_idx, h2_idx, h3_idx, body_idx} - {None}

    # 清空非样本段落
    for i in range(body_start, body_end):
        if i in keep:
            continue
        p = paras[i]
        if p._p.find('.//w:sectPr', _NS) is not None:
            continue
        for r in p.runs:
            r.text = ""
        style = p.style.name if p.style else ""
        if style.startswith("Heading") or style.startswith("标题"):
            p.style = doc.styles['Normal']
        if style == "Caption":
            p.style = doc.styles['Normal']

    # 删除正文区域表格
    body_elem = doc.element.body
    tables_rm = []
    pc = 0
    for elem in body_elem:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            pc += 1
        elif tag == 'tbl' and body_start <= pc <= body_end:
            tables_rm.append(elem)
    for tbl in tables_rm:
        body_elem.remove(tbl)
    if tables_rm:
        print(f"    删除 {len(tables_rm)} 个示例表格")

    # 替换样本文本
    def _set_text(para, text):
        if para.runs:
            para.runs[0].text = text
            for r in para.runs[1:]:
                r.text = ""

    paras[h1_idx].paragraph_format.page_break_before = True
    _set_text(paras[h1_idx], "{{ ch.title }}")
    if h2_idx:
        _set_text(paras[h2_idx], "{{ sec.title }}")
    if h3_idx:
        _set_text(paras[h3_idx], "{{ sub.title }}")
    _set_text(paras[body_idx], "{{ para }}")

    h1_p = paras[h1_idx]._p
    h2_p = paras[h2_idx]._p if h2_idx else None
    h3_p = paras[h3_idx]._p if h3_idx else None
    body_p = paras[body_idx]._p
    body_copy = copy.deepcopy(body_p)

    def _mk_body(text):
        p = copy.deepcopy(body_copy)
        for t_elem in p.findall('.//w:t', _NS):
            t_elem.text = ""
        t_elems = p.findall('.//w:t', _NS)
        if t_elems:
            t_elems[0].text = text
        return p

    parent = h1_p.getparent()
    anchor = h1_p.getprevious()

    for elem in [h1_p, h2_p, h3_p, body_p]:
        if elem is not None and elem.getparent() is not None:
            parent.remove(elem)

    cursor = anchor

    def _after(cursor, elem):
        cursor.addnext(elem)
        return elem

    cursor = _after(cursor, _mk_ctrl("{%p for ch in chapters %}"))
    cursor = _after(cursor, h1_p)

    if h2_p is not None:
        cursor = _after(cursor, _mk_ctrl("{%p for sec in ch.sections %}"))
        cursor = _after(cursor, h2_p)
        cursor = _after(cursor, _mk_ctrl("{%p for para in sec.content %}"))
        cursor = _after(cursor, body_p)
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))

        if h3_p is not None:
            cursor = _after(cursor, _mk_ctrl("{%p for sub in sec.subsections %}"))
            cursor = _after(cursor, h3_p)
            cursor = _after(cursor, _mk_ctrl("{%p for p2 in sub.content %}"))
            cursor = _after(cursor, _mk_body("{{ p2 }}"))
            cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))
            cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))

        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))
    else:
        cursor = _after(cursor, _mk_ctrl("{%p for sec in ch.sections %}"))
        cursor = _after(cursor, _mk_ctrl("{%p for para in sec.content %}"))
        cursor = _after(cursor, body_p)
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))
        cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))

    cursor = _after(cursor, _mk_ctrl("{%p endfor %}"))

    # 清空正文到结论/致谢之间的残留空段落
    _remove_empty_between(doc, body_end)

    doc.save(doc_path)


def _remove_empty_between(doc, body_end):
    """清除正文循环 endfor 和致谢/结论之间的空段落"""
    paras = doc.paragraphs
    last_endfor = None
    target_idx = None

    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        if "{%p endfor %}" in text:
            last_endfor = i
        style = p.style.name if p.style else ""
        if style == "Heading 1" and i >= body_end:
            target_idx = i
            break

    if last_endfor is None or target_idx is None:
        return

    to_remove = []
    for i in range(last_endfor + 1, target_idx):
        p = paras[i]
        text = (p.text or "").strip()
        has_sect = p._p.find('.//w:sectPr', _NS) is not None
        if has_sect:
            # 有分节符的段落不能删，但清空文字
            if text:
                for r in p.runs:
                    r.text = ""
            continue
        if not text:
            to_remove.append(p._p)

    for elem in to_remove:
        if elem.getparent() is not None:
            elem.getparent().remove(elem)
    if to_remove:
        print(f"    删除 {len(to_remove)} 个空段落")


def _mk_ctrl(text):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python make.py <原始模板.docx> [输出路径]")
        sys.exit(1)
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else \
        os.path.join(os.path.dirname(__file__), "template.docx")
    print(f"制作模板: {src} -> {out}")
    make(src, out)
    print(f"完成: {out}")
