"""
景德镇学院本科毕业论文模板制作。

用法:
    cd thesis_project
    python templates/jdzxy/make.py <原始定稿.docx> [输出template.docx]

本脚本只做模板清理和占位符设置，不写具体论文内容。
生成论文内容请使用 papers/<project>/build.py + JSON 数据。
"""
from __future__ import annotations

import copy
import re
import sys
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": WNS}


def _clear_para(p):
    for r in p.runs:
        r.text = ""


def _replace_para_text(p, text):
    if not p.runs:
        p.add_run(text)
        return
    target = 0
    # 标题段落经常有多个 run，优先放进字号最大的 run 里。
    best_size = 0
    for i, r in enumerate(p.runs):
        if r.font.size and r.font.size.pt > best_size:
            best_size = r.font.size.pt
            target = i
    for i, r in enumerate(p.runs):
        r.text = text if i == target else ""


def _make_p(text, like=None):
    p = OxmlElement("w:p")
    if like is not None and like.pPr is not None:
        p.append(copy.deepcopy(like.pPr))
    r = OxmlElement("w:r")
    if like is not None:
        first_r = like.find(qn("w:r"))
        if first_r is not None and first_r.rPr is not None:
            r.append(copy.deepcopy(first_r.rPr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def _replace_xml_text(p_xml, text):
    texts = p_xml.findall(".//w:t", NS)
    if not texts:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        p_xml.append(r)
        return
    texts[0].text = text
    for t in texts[1:]:
        t.text = ""


def _add_page_break_before(p_xml):
    ppr = p_xml.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        p_xml.insert(0, ppr)
    pbr = ppr.find(qn("w:pageBreakBefore"))
    if pbr is None:
        pbr = OxmlElement("w:pageBreakBefore")
        ppr.append(pbr)
    pbr.set(qn("w:val"), "1")


def _remove_page_break_before(p_xml):
    ppr = p_xml.find(qn("w:pPr"))
    if ppr is None:
        return
    for node in list(ppr.findall(qn("w:pageBreakBefore"))):
        ppr.remove(node)


def _remove_between(start_xml, stop_xml, keep_section_markers=True):
    cur = start_xml.getnext()
    while cur is not None and cur is not stop_xml:
        nxt = cur.getnext()
        if keep_section_markers and cur.find(".//w:sectPr", NS) is not None:
            cur = nxt
            continue
        cur.getparent().remove(cur)
        cur = nxt


def _find_para(paras, predicate, start=0):
    for i, p in enumerate(paras[start:], start):
        if predicate((p.text or "").strip(), p):
            return i, p
    raise RuntimeError("paragraph not found")


def _cleanup_update_fields(path):
    with ZipFile(path) as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    document = etree.fromstring(files["word/document.xml"])
    _remove_ref_fields(document)
    _collapse_redundant_page_breaks(document)
    _collapse_redundant_section_breaks(document)
    _remove_front_page_break_before(document)
    _normalize_page_number_sections(document)
    files["word/document.xml"] = etree.tostring(
        document, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    settings = etree.fromstring(files["word/settings.xml"])
    if not settings.xpath("./w:updateFields", namespaces=NS):
        node = etree.SubElement(settings, f"{{{WNS}}}updateFields")
        node.set(f"{{{WNS}}}val", "true")
    files["word/settings.xml"] = etree.tostring(
        settings, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    # 清理页脚里模板遗留文本，只保留 PAGE 字段。
    for name, data in list(files.items()):
        if not (name.startswith("word/footer") and name.endswith(".xml")):
            continue
        root = etree.fromstring(data)
        if not root.xpath(".//w:instrText[contains(., 'PAGE')]", namespaces=NS):
            continue
        for p in root.xpath("./w:p", namespaces=NS):
            root.remove(p)
        p = etree.SubElement(root, f"{{{WNS}}}p")
        ppr = etree.SubElement(p, f"{{{WNS}}}pPr")
        jc = etree.SubElement(ppr, f"{{{WNS}}}jc")
        jc.set(f"{{{WNS}}}val", "center")
        r = etree.SubElement(p, f"{{{WNS}}}r")
        begin = etree.SubElement(r, f"{{{WNS}}}fldChar")
        begin.set(f"{{{WNS}}}fldCharType", "begin")
        instr = etree.SubElement(r, f"{{{WNS}}}instrText")
        instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        instr.text = " PAGE  \\* MERGEFORMAT "
        sep = etree.SubElement(r, f"{{{WNS}}}fldChar")
        sep.set(f"{{{WNS}}}fldCharType", "separate")
        t = etree.SubElement(r, f"{{{WNS}}}t")
        t.text = "1"
        end = etree.SubElement(r, f"{{{WNS}}}fldChar")
        end.set(f"{{{WNS}}}fldCharType", "end")
        files[name] = etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    with ZipFile(path, "w", ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)


def _remove_ref_fields(root):
    """Remove broken REF cross-reference fields while preserving TOC fields."""
    for p in root.xpath(".//w:p", namespaces=NS):
        runs = list(p.xpath("./w:r", namespaces=NS))
        i = 0
        while i < len(runs):
            run = runs[i]
            begins = run.xpath(".//w:fldChar[@w:fldCharType='begin']", namespaces=NS)
            if not begins:
                i += 1
                continue
            field_runs = [run]
            is_ref = bool(run.xpath(".//w:instrText[contains(., 'REF')]", namespaces=NS))
            j = i + 1
            ended = bool(run.xpath(".//w:fldChar[@w:fldCharType='end']", namespaces=NS))
            while j < len(runs) and not ended:
                field_runs.append(runs[j])
                if runs[j].xpath(".//w:instrText[contains(., 'REF')]", namespaces=NS):
                    is_ref = True
                if runs[j].xpath(".//w:fldChar[@w:fldCharType='end']", namespaces=NS):
                    ended = True
                j += 1
            if is_ref:
                for fr in field_runs:
                    if fr.getparent() is not None:
                        fr.getparent().remove(fr)
                runs = list(p.xpath("./w:r", namespaces=NS))
                continue
            i = j


def _paragraph_text(p):
    return "".join(p.xpath(".//w:t/text()", namespaces=NS)).strip()


def _has_page_break(p):
    return bool(p.xpath(".//w:br[@w:type='page']", namespaces=NS))


def _has_section_break(p):
    return p.find(".//w:sectPr", NS) is not None


def _collapse_redundant_page_breaks(root):
    """Drop duplicated blank page-break paragraphs from the source template."""
    body = root.find(".//w:body", namespaces=NS)
    if body is None:
        return
    children = list(body)
    keep_break_seen = False
    for el in children:
        if el.tag != f"{{{WNS}}}p":
            continue
        text = _paragraph_text(el)
        has_break = _has_page_break(el)
        if text:
            keep_break_seen = False
            continue
        if has_break:
            if keep_break_seen and not _has_section_break(el):
                body.remove(el)
                continue
            keep_break_seen = True
        elif _has_section_break(el):
            keep_break_seen = False


def _collapse_redundant_section_breaks(root):
    """Remove empty one-page sections left by the source print layout."""
    body = root.find(".//w:body", namespaces=NS)
    if body is None:
        return
    prev_empty_sect = False
    for el in list(body):
        if el.tag != f"{{{WNS}}}p":
            prev_empty_sect = False
            continue
        empty = not _paragraph_text(el)
        has_sect = _has_section_break(el)
        if empty and has_sect and prev_empty_sect:
            body.remove(el)
            continue
        prev_empty_sect = empty and has_sect


def _remove_front_page_break_before(root):
    """The old document used pageBreakBefore heavily; keep only chapter breaks."""
    body = root.find(".//w:body", namespaces=NS)
    if body is None:
        return
    for p in body.xpath("./w:p", namespaces=NS):
        text = _paragraph_text(p)
        if text == "{{ ch.title }}":
            continue
        ppr = p.find(qn("w:pPr"))
        if ppr is None:
            continue
        for node in list(ppr.findall(qn("w:pageBreakBefore"))):
            ppr.remove(node)


def _normalize_page_number_sections(root):
    """Keep roman front matter and restart decimal numbering for the body."""
    sects = root.xpath(".//w:sectPr", namespaces=NS)
    if not sects:
        return

    # The final sectPr applies to the generated body/references/appendix section.
    body_sect = sects[-1]
    pg = body_sect.find(qn("w:pgNumType"))
    if pg is None:
        pg = etree.SubElement(body_sect, f"{{{WNS}}}pgNumType")
    pg.set(f"{{{WNS}}}fmt", "decimal")
    pg.set(f"{{{WNS}}}start", "1")


def make(src, out):
    doc = Document(src)
    paras = doc.paragraphs

    print("Step 1: 封面与摘要占位符...")
    replacements = {
        1: "学号{{ student_id }}",
        2: "密级 {{ secrecy_level }}",
        10: "{{ title_zh }}",
        11: "",
        15: "学 院 名 称：{{ college }}",
        16: "专 业 名 称：{{ major }}",
        17: "学 生 姓 名：{{ name }}",
        18: "指 导 教 师：{{ advisor }}",
        22: "{{ date_zh }}",
        29: "{{ title_en }}",
        30: "",
        32: "School ：{{ school_en }}",
        33: "Major ：{{ major_en }}",
        34: "Name ：{{ name_en }}",
        35: "Supervisor：{{ advisor_en }}",
        45: "{{ date_en }}",
        67: "{{ abstract_zh_1 }}",
        68: "{{ abstract_zh_2 }}",
        70: "关键词：{{ keywords_zh }}",
        74: "{{ abstract_en_1 }}",
        75: "{{ abstract_en_2 }}",
        76: "Key words: {{ keywords_en }}",
    }
    for idx, text in replacements.items():
        _replace_para_text(paras[idx], text)

    print("Step 2: 清理目录结果，保留可更新目录占位...")
    _, toc_title = _find_para(
        paras, lambda text, _p: text.replace(" ", "").replace("　", "") == "目录"
    )
    _, body_start = _find_para(paras, lambda text, _p: text == "1 绪论")
    # 保留目录标题后面的 section marker，其他指纹目录结果清掉。
    _remove_between(toc_title._p, body_start._p, keep_section_markers=True)
    toc_title._p.addnext(_make_p("__TOC_PLACEHOLDER__", toc_title._p))

    print("Step 3: 正文区域改为 Jinja2 循环...")
    paras = doc.paragraphs
    _, body_start = _find_para(paras, lambda text, _p: text == "1 绪论")
    _, refs_heading = _find_para(paras, lambda text, _p: text == "参考文献", start=80)
    _, ack_heading = _find_para(
        paras, lambda text, _p: text.replace(" ", "").replace("　", "") == "致谢", start=80
    )
    _, appendix_heading = _find_para(
        paras, lambda text, _p: text.replace(" ", "").replace("　", "") in {"附件", "附录"},
        start=80,
    )
    _, h2_sample = _find_para(paras, lambda text, p: p.style and p.style.name == "Heading 2", start=body_start._p.getparent().index(body_start._p))
    body_sample = None
    for p in paras:
        if p._p.getparent().index(p._p) > h2_sample._p.getparent().index(h2_sample._p):
            if (p.text or "").strip() and p.style and p.style.name == "Normal":
                body_sample = p
                break
    if body_sample is None:
        raise RuntimeError("找不到正文样本段落")

    h1_xml = copy.deepcopy(body_start._p)
    h2_xml = copy.deepcopy(h2_sample._p)
    h3_xml = copy.deepcopy(h2_sample._p)
    body_xml = copy.deepcopy(body_sample._p)
    body2_xml = copy.deepcopy(body_sample._p)
    _replace_xml_text(h1_xml, "{{ ch.title }}")
    _replace_xml_text(h2_xml, "{{ sec.title }}")
    _replace_xml_text(h3_xml, "{{ sub.title }}")
    _replace_xml_text(body_xml, "{{ item }}")
    _replace_xml_text(body2_xml, "{{ item }}")
    for p_xml in (h2_xml, h3_xml, body_xml, body2_xml):
        _remove_page_break_before(p_xml)
    _add_page_break_before(h1_xml)

    _remove_between(body_start._p.getprevious(), refs_heading._p, keep_section_markers=False)
    anchor = refs_heading._p
    blocks = [
        _make_p("{%p for ch in chapters %}"),
        h1_xml,
        _make_p("{%p for sec in ch.sections %}"),
        h2_xml,
        _make_p("{%p for item in sec.content %}"),
        body_xml,
        _make_p("{%p endfor %}"),
        _make_p("{%p for sub in sec.subsections %}"),
        h3_xml,
        _make_p("{%p for item in sub.content %}"),
        body2_xml,
        _make_p("{%p endfor %}"),
        _make_p("{%p endfor %}"),
        _make_p("{%p endfor %}"),
        _make_p("{%p endfor %}"),
    ]
    for block in blocks:
        anchor.addprevious(block)

    print("Step 4: 参考文献、致谢、附件占位符...")
    paras = doc.paragraphs
    _, refs_heading = _find_para(paras, lambda text, _p: text == "参考文献", start=80)
    _, ack_heading = _find_para(
        paras, lambda text, _p: text.replace(" ", "").replace("　", "") == "致谢", start=80
    )
    _, appendix_heading = _find_para(
        paras, lambda text, _p: text.replace(" ", "").replace("　", "") in {"附件", "附录"},
        start=80,
    )
    _replace_para_text(refs_heading, "参考文献")
    _replace_para_text(ack_heading, "致 谢")
    _replace_para_text(appendix_heading, "附 件")

    _add_page_break_before(refs_heading._p)
    ref_sample = _make_p("[{{ ref.idx }}] {{ ref.text }}", body_sample._p)
    _remove_between(refs_heading._p, ack_heading._p, keep_section_markers=True)
    refs_heading._p.addnext(_make_p("{%p endfor %}"))
    refs_heading._p.addnext(ref_sample)
    refs_heading._p.addnext(_make_p("{%p for ref in references %}"))

    body = doc.element.body
    final_sect = body.find(qn("w:sectPr"))
    cur = ack_heading._p
    while cur is not None and cur is not final_sect:
        nxt = cur.getnext()
        cur.getparent().remove(cur)
        cur = nxt

    print("Step 5: 清理残留图片和指纹论文文本...")
    body_loop_idx = None
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip() == "{%p for ch in chapters %}":
            body_loop_idx = i
            break
    if body_loop_idx is None:
        raise RuntimeError("找不到正文循环占位符，无法安全清理旧图片")

    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if "指纹" in text or "打卡" in text or "AS608" in text or "DS1307" in text:
            if "{{" not in text and "{%" not in text:
                _clear_para(p)
        if i < body_loop_idx:
            continue
        for drawing in list(p._p.findall(f".//{{{WNS}}}drawing")):
            parent = drawing.getparent()
            while parent is not None and parent.tag != qn("w:r"):
                parent = parent.getparent()
            if parent is not None:
                parent.getparent().remove(parent)

    doc.save(out)
    _cleanup_update_fields(out)
    _verify(out)
    print(f"完成: {out}")


def _verify(path):
    doc = Document(path)
    needed = [
        "title_zh", "student_id", "abstract_zh_1", "keywords_zh",
        "abstract_en_1", "keywords_en", "ch.title", "sec.title",
        "item", "ref",
    ]
    text = "\n".join(p.text for p in doc.paragraphs)
    print("\n=== 模板验证 ===")
    for key in needed:
        print(f"  {'✓' if key in text else '✗'} {key}")
    dirty = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t and re.search(r"指纹|打卡|AS608|DS1307", t):
            dirty.append(t[:60])
    if dirty:
        for t in dirty[:5]:
            print(f"  残留: {t}")
    else:
        print("  ✓ 无旧论文关键词残留")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python make.py <原始定稿.docx> [输出template.docx]")
        sys.exit(1)
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else str(Path(__file__).with_name("template.docx"))
    print(f"制作模板: {src} -> {out}")
    make(src, out)
