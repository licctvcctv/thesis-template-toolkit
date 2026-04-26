"""
哈尔滨理工大学软件工程专业论文模板制作。

用法:
  cd /Users/a136/vs/45425/thesis_project
  python -m templates.hrbust_software.make \
    "/Users/a136/vs/45425/要求426/软件工程专业-本科毕业设计（论文）模板-20250422(3).docx"

本脚本只负责把学校 Word 模板清理成 docxtpl 模板：
- 保留原模板的标题、摘要、目录、页眉页脚、章节样式和页码结构；
- 把正文、结论、致谢、参考文献、附录替换为子文档占位符；
- 不写具体论文内容，论文内容由 papers/study-buddy/build.py + JSON 渲染。
"""
from __future__ import annotations

import os
import re
import sys
import zipfile
import copy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


HERE = Path(__file__).resolve().parent
DEFAULT_OUT = HERE / "template.docx"


def _set_para_text(p, text: str):
    if p.runs:
        p.runs[0].text = text
        _clear_run_emphasis(p.runs[0])
        for r in p.runs[1:]:
            r.text = ""
    else:
        r = p.add_run(text)
        _clear_run_emphasis(r)


def _clear_run_emphasis(run):
    """Remove sample-template highlight/red-font residue from replaced runs."""
    r_pr = run._element.get_or_add_rPr()
    for tag in ("w:highlight", "w:color", "w:shd"):
        for node in list(r_pr.findall(qn(tag))):
            r_pr.remove(node)


def _clear_para(p):
    for r in p.runs:
        r.text = ""


def _copy_run_style(src_run, dst_run):
    if src_run is None:
        return
    r_pr = src_run._element.rPr
    if r_pr is not None:
        dst_run._element.insert(0, copy.deepcopy(r_pr))


def _replace_runs_with_templates(p, specs):
    old_runs = list(p.runs)
    for run in old_runs:
        p._p.remove(run._r)
    for text, template_run in specs:
        r = p.add_run(text)
        _copy_run_style(template_run, r)
        _clear_run_emphasis(r)


def _set_keyword_line(p, label: str, placeholder: str):
    """Preserve the template's mixed keyword styling: label != content."""
    old_runs = list(p.runs)
    label_run = old_runs[0] if old_runs else None
    sep_run = old_runs[1] if len(old_runs) > 1 else label_run
    content_run = None
    for run in old_runs[2:]:
        if (run.text or "").strip():
            content_run = run
            break
    if content_run is None:
        content_run = old_runs[-1] if old_runs else label_run
    _replace_runs_with_templates(
        p,
        [
            (label, label_run),
            ("\u3000", sep_run),
            (placeholder, content_run),
        ],
    )


def _make_placeholder_para(text: str, like=None):
    p = OxmlElement("w:p")
    if like is not None and like.pPr is not None:
        import copy

        p.append(copy.deepcopy(like.pPr))
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def _remove_blocks_between(doc: Document, start_idx: int, end_idx: int):
    """Remove paragraphs and tables between paragraph indexes [start, end)."""
    body = doc.element.body
    para_pos = {}
    p_count = -1
    for elem in body:
        if elem.tag == qn("w:p"):
            p_count += 1
            para_pos[elem] = p_count

    to_remove = []
    current_para = -1
    for elem in body:
        if elem.tag == qn("w:p"):
            current_para = para_pos[elem]
            if start_idx <= current_para < end_idx:
                pPr = elem.find(qn("w:pPr"))
                has_sect = pPr is not None and pPr.find(qn("w:sectPr")) is not None
                if not has_sect:
                    to_remove.append(elem)
        elif elem.tag == qn("w:tbl") and start_idx <= current_para < end_idx:
            to_remove.append(elem)

    for elem in to_remove:
        if elem.getparent() is not None:
            elem.getparent().remove(elem)
    return len(to_remove)


def _find_heading(paras, keyword: str, start=0):
    for i in range(start, len(paras)):
        p = paras[i]
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        compact = text.replace(" ", "").replace("　", "")
        if (style.startswith("Heading") or style == "Title") and keyword in compact:
            return i
    return None


def _patch_update_fields(docx_path: Path):
    tmp = Path(str(docx_path) + ".tmp")
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "word/settings.xml":
                    xml = data.decode("utf-8")
                    if "w:updateFields" not in xml:
                        xml = xml.replace(
                            "</w:settings>",
                            '<w:updateFields w:val="true"/></w:settings>',
                        )
                    data = xml.encode("utf-8")
                zout.writestr(info, data)
    tmp.replace(docx_path)


def _clean_template_residue(docx_path: Path):
    """Remove old comment balloons/text boxes left by sample templates."""
    tmp = Path(str(docx_path) + ".tmp")
    comment_parts = {
        "word/comments.xml",
        "word/commentsExtended.xml",
        "word/commentsIds.xml",
        "word/people.xml",
    }
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                if info.filename in comment_parts:
                    continue
                data = zin.read(info.filename)
                if info.filename in {
                    "word/document.xml",
                    "word/header1.xml",
                    "word/header2.xml",
                    "word/header3.xml",
                    "word/footer1.xml",
                    "word/footer2.xml",
                    "word/footer3.xml",
                }:
                    xml = data.decode("utf-8")
                    xml = re.sub(
                        r"<mc:AlternateContent>.*?</mc:AlternateContent>",
                        "",
                        xml,
                        flags=re.DOTALL,
                    )
                    xml = re.sub(r"<w:commentRangeStart\b[^>]*/>", "", xml)
                    xml = re.sub(r"<w:commentRangeEnd\b[^>]*/>", "", xml)
                    xml = re.sub(r"<w:commentReference\b[^>]*/>", "", xml)
                    data = xml.encode("utf-8")
                elif info.filename == "word/_rels/document.xml.rels":
                    xml = data.decode("utf-8")
                    xml = re.sub(
                        r'<Relationship\b[^>]*(?:comments|commentsExtended|commentsIds|people)[^>]*/>',
                        "",
                        xml,
                    )
                    data = xml.encode("utf-8")
                elif info.filename == "[Content_Types].xml":
                    xml = data.decode("utf-8")
                    xml = re.sub(
                        r'<Override\b[^>]*(?:comments|commentsExtended|commentsIds|people)[^>]*/>',
                        "",
                        xml,
                    )
                    data = xml.encode("utf-8")
                zout.writestr(info, data)
    tmp.replace(docx_path)


def make(src_path: str | os.PathLike, out_path: str | os.PathLike = DEFAULT_OUT):
    src_path = Path(src_path)
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(src_path))
    paras = doc.paragraphs
    print(f"原始: {len(paras)} 段落, {len(doc.tables)} 表格")

    # Front matter. The supplied software template has stable paragraph order:
    # title, Chinese abstract title/body/keywords, English abstract/body/keywords, TOC.
    _set_para_text(paras[0], "{{ title_zh }}")
    _set_para_text(paras[2], "{{ abstract_zh_1 }}")
    _set_para_text(paras[3], "{{ abstract_zh_2 }}")
    _set_para_text(paras[4], "{{ abstract_zh_3 }}")
    _clear_para(paras[5])
    _clear_para(paras[6])
    _set_keyword_line(paras[8], "关键词", "{{ keywords_zh }}")

    _set_para_text(paras[9], "{{ title_en }}")
    _set_para_text(paras[11], "{{ abstract_en_1 }}")
    _set_para_text(paras[12], "{{ abstract_en_2 }}")
    _set_para_text(paras[13], "{{ abstract_en_3 }}")
    _clear_para(paras[14])
    _set_keyword_line(paras[15], "Keywords", "{{ keywords_en }}")
    print("  Step 1: 摘要与关键词占位完成")

    body_start = None
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        if style == "Heading 1" and i > 40 and "结论" not in text:
            body_start = i
            break
    toc_title_idx = None
    for i, p in enumerate(paras):
        style = p.style.name if p.style else ""
        compact = (p.text or "").replace(" ", "").replace("　", "")
        if "章标题" in style and compact == "目录":
            toc_title_idx = i
            break
    conclusion_idx = _find_heading(paras, "结论", start=body_start or 0)
    ack_idx = _find_heading(paras, "致谢", start=conclusion_idx or 0)
    refs_idx = _find_heading(paras, "参考文献", start=ack_idx or 0)
    appendix_idx = _find_heading(paras, "附录", start=refs_idx or 0)

    if body_start is None or conclusion_idx is None or ack_idx is None or refs_idx is None:
        raise RuntimeError(
            f"无法定位正文区域: body={body_start}, conclusion={conclusion_idx}, "
            f"ack={ack_idx}, refs={refs_idx}, appendix={appendix_idx}"
        )

    if toc_title_idx is not None and toc_title_idx + 1 < body_start:
        toc_anchor = paras[toc_title_idx]._p
        removed = _remove_blocks_between(doc, toc_title_idx + 1, body_start)
        toc_anchor.addnext(_make_placeholder_para("{{p toc_doc}}"))
        print(f"  Step 1b: 目录结果清理 {removed} 个块")
        paras = doc.paragraphs
        # The body start index changes after removing the old TOC result.
        body_start = None
        for i, p in enumerate(paras):
            style = p.style.name if p.style else ""
            text = (p.text or "").strip()
            if style == "Heading 1" and i > 20 and "结论" not in text:
                body_start = i
                break
        conclusion_idx = _find_heading(paras, "结论", start=body_start or 0)

    # Body: replace all sample chapters with {{ body }} while preserving TOC/front matter.
    anchor = paras[body_start]._p.getprevious()
    removed = _remove_blocks_between(doc, body_start, conclusion_idx)
    anchor.addnext(_make_placeholder_para("{{p body}}"))
    print(f"  Step 2: 正文样例清理 {removed} 个块")

    # Refresh after DOM edits.
    paras = doc.paragraphs
    conclusion_idx = _find_heading(paras, "结论")
    ack_idx = _find_heading(paras, "致谢", start=conclusion_idx or 0)
    refs_idx = _find_heading(paras, "参考文献", start=ack_idx or 0)
    appendix_idx = _find_heading(paras, "附录", start=refs_idx or 0)

    # Conclusion.
    c_anchor = paras[conclusion_idx]._p
    removed = _remove_blocks_between(doc, conclusion_idx + 1, ack_idx)
    c_anchor.addnext(_make_placeholder_para("{{p conclusion_doc}}"))
    print(f"  Step 3: 结论样例清理 {removed} 个块")

    # Acknowledgement.
    paras = doc.paragraphs
    ack_idx = _find_heading(paras, "致谢")
    refs_idx = _find_heading(paras, "参考文献", start=ack_idx or 0)
    a_anchor = paras[ack_idx]._p
    removed = _remove_blocks_between(doc, ack_idx + 1, refs_idx)
    a_anchor.addnext(_make_placeholder_para("{{p acknowledgement_doc}}"))
    print(f"  Step 4: 致谢样例清理 {removed} 个块")

    # References.
    paras = doc.paragraphs
    refs_idx = _find_heading(paras, "参考文献")
    appendix_idx = _find_heading(paras, "附录", start=refs_idx or 0)
    r_anchor = paras[refs_idx]._p
    if appendix_idx is None:
        appendix_idx = len(paras)
    removed = _remove_blocks_between(doc, refs_idx + 1, appendix_idx)
    r_anchor.addnext(_make_placeholder_para("{{p references_doc}}"))
    print(f"  Step 5: 参考文献样例清理 {removed} 个块")

    # Appendix, if the source contains one.
    paras = doc.paragraphs
    appendix_idx = _find_heading(paras, "附录")
    if appendix_idx is not None:
        ap_anchor = paras[appendix_idx]._p
        removed = _remove_blocks_between(doc, appendix_idx + 1, len(paras))
        ap_anchor.addnext(_make_placeholder_para("{{p appendix_doc}}"))
        print(f"  Step 6: 附录样例清理 {removed} 个块")

    doc.save(out_path)
    _patch_update_fields(out_path)
    _clean_template_residue(out_path)

    out_doc = Document(str(out_path))
    full = "\n".join(p.text or "" for p in out_doc.paragraphs)
    for tag in ["{{p body}}", "{{p conclusion_doc}}", "{{p references_doc}}"]:
        if tag not in full:
            raise RuntimeError(f"模板缺少占位符: {tag}")
    print(f"完成: {out_path} ({len(out_doc.paragraphs)} 段, {len(out_doc.tables)} 表)")


def main():
    if len(sys.argv) < 2:
        print("用法: python -m templates.hrbust_software.make <原始模板.docx> [输出.docx]")
        sys.exit(1)
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else str(DEFAULT_OUT)
    make(src, out)


if __name__ == "__main__":
    main()
