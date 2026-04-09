"""
核心构建逻辑：将章节数据和参考文献构建为 Word 子文档。
支持多段落、图片插入、三线表、引用上标。
"""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ================================================================
# 书签 & 引用超链接
# ================================================================

_bookmark_id_counter = [100]


def _next_bookmark_id():
    _bookmark_id_counter[0] += 1
    return str(_bookmark_id_counter[0])


def add_bookmark(paragraph, bookmark_name):
    """给段落添加书签"""
    bid = _next_bookmark_id()
    p = paragraph._p
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), bid)
    start.set(qn('w:name'), bookmark_name)
    p.insert(0, start)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), bid)
    p.append(end)


def _cite_bookmark(n):
    return f"_Ref_cite_{n}"


def add_citation(paragraph, cite_num):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), _cite_bookmark(cite_num))
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    vertAlign = OxmlElement('w:vertAlign')
    vertAlign.set(qn('w:val'), 'superscript')
    rPr.append(vertAlign)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = f"[{cite_num}]"
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


# ================================================================
# 三线表
# ================================================================

def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), val.get('sz', '4'))
            el.set(qn('w:color'), val.get('color', '000000'))
            el.set(qn('w:space'), '0')
            borders.append(el)
    tcPr.append(borders)


def _add_table_block(subdoc, table_data, line_spacing=1.5):
    """添加三线表：表题 + 表头 + 数据行"""
    caption = table_data.get("caption", "")
    headers = table_data.get("headers", [])
    rows = table_data.get("rows", [])
    ncols = len(headers)

    # 表题段落（表上方，五号宋体居中）
    p_cap = subdoc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run(caption)
    run.font.size = Pt(10.5)

    # 创建表格
    table = subdoc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 表头
    thick = {'val': 'single', 'sz': '12', 'color': '000000'}
    thin = {'val': 'single', 'sz': '4', 'color': '000000'}
    none = {'val': 'nil'}

    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(10.5)
        # 三线表：表头上粗线、下细线、左右无线
        _set_cell_border(cell, top=thick, bottom=thin,
                         left=none, right=none)

    # 数据行
    for i, row in enumerate(rows):
        is_last = (i == len(rows) - 1)
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(10.5)
            # 最后一行下方粗线，其余无线
            _set_cell_border(
                cell,
                top=none,
                bottom=thick if is_last else none,
                left=none, right=none
            )

    return table


# ================================================================
# 段落构建
# ================================================================

def _add_text_paragraph(subdoc, text, citations=None,
                        line_spacing=1.5):
    """添加一个正文段落"""
    p = subdoc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.85)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    run.font.size = Pt(12)
    for n in (citations or []):
        add_citation(p, n)
    return p


def _add_image_block(subdoc, image_path, caption, width_cm=12):
    """添加图片 + 图注"""
    p_img = subdoc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(image_path, width=Cm(width_cm))

    p_cap = subdoc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = p_cap.add_run(caption)
    run_cap.font.size = Pt(10.5)
    return p_img, p_cap


# ================================================================
# 内容块处理
# ================================================================

def _build_content_blocks(subdoc, content_list, line_spacing=1.5):
    """
    处理 content 列表，支持四种内容块：
      - str: 纯文本段落
      - {"text": ..., "cite": [...]}: 带引用的段落
      - {"image": ..., "caption": ..., "width_cm": ...}: 图片
      - {"table": {"caption": ..., "headers": [...], "rows": [...]}}: 三线表
    """
    for block in content_list:
        if isinstance(block, str):
            _add_text_paragraph(subdoc, block,
                                line_spacing=line_spacing)
        elif isinstance(block, dict):
            if "image" in block:
                _add_image_block(
                    subdoc, block["image"],
                    block.get("caption", ""),
                    block.get("width_cm", 12),
                )
            elif "table" in block:
                _add_table_block(subdoc, block["table"],
                                 line_spacing=line_spacing)
            elif "text" in block:
                _add_text_paragraph(
                    subdoc, block["text"], block.get("cite", []),
                    line_spacing=line_spacing,
                )


# ================================================================
# 公开接口
# ================================================================

def _add_heading(subdoc, text, level=1, page_break=False,
                 line_spacing=1.5):
    """添加标题段落，兼容无 Heading 样式的模板"""
    try:
        p = subdoc.add_heading(text, level=level)
    except (KeyError, ValueError):
        # 模板没有 Heading 样式，手动格式化
        p = subdoc.add_paragraph()
        p.paragraph_format.line_spacing = line_spacing
        run = p.add_run(text)
        run.bold = True
        if level == 1:
            run.font.size = Pt(15)  # 小三号
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif level == 2:
            run.font.size = Pt(14)  # 四号
        else:
            run.font.size = Pt(12)  # 小四号
    if page_break:
        p.paragraph_format.page_break_before = True
    return p


def build_chapters(doc, chapters, line_spacing=1.5):
    """构建正文子文档"""
    sd = doc.new_subdoc()

    for i, ch in enumerate(chapters):
        _add_heading(sd, ch["title"], level=1,
                     page_break=(i > 0), line_spacing=line_spacing)

        for sec in ch["sections"]:
            _add_heading(sd, sec["title"], level=2,
                         line_spacing=line_spacing)
            _build_content_blocks(sd, sec.get("content", []),
                                  line_spacing)

            for subsec in sec.get("subsections", []):
                _add_heading(sd, subsec["title"], level=3,
                             line_spacing=line_spacing)
                _build_content_blocks(sd, subsec.get("content", []),
                                      line_spacing)
    return sd


def build_references(doc, references, line_spacing=1.5):
    """构建参考文献子文档（带书签）"""
    sd = doc.new_subdoc()

    for i, ref in enumerate(references, 1):
        p = sd.add_paragraph()
        p.paragraph_format.line_spacing = line_spacing
        run = p.add_run(f"[{i}] {ref}")
        run.font.size = Pt(12)
        add_bookmark(p, _cite_bookmark(i))

    return sd
